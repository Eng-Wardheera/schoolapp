import base64
from collections import defaultdict
import csv
from decimal import Decimal
from email import parser
import io
import platform
import socket
import string
import sys
import traceback
from xml.dom.minidom import Document
from flask_wtf import FlaskForm
from googletrans import Translator
import imgkit
from pandas import read_csv 
import pandas as pd
from flask_mail import Message
from math import ceil
import os
import random
import re
from sqlite3 import IntegrityError
import uuid
import psutil
import pycountry
import pyotp
import qrcode
from datetime import date, datetime, timedelta
import pytz
from flask import Blueprint, Response, after_this_request, current_app, flash, g, json, jsonify, make_response, render_template, request, redirect, send_file, send_from_directory, session, url_for
from flask_login import login_user, logout_user, login_required, current_user, user_logged_in
import requests
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
from sqlalchemy.orm import aliased, joinedload
from sqlalchemy import Integer, and_, case, cast, exists, extract, func, or_
from wtforms import DecimalField, SubmitField, TextAreaField
from app import EAT, now_eat, UPLOAD_FOLDER, db, ALLOWED_EXTENSIONS
from app import mail
from app import google 
from app import github 
import phonenumbers
from phonenumbers import NumberParseException, PhoneMetadata, parse, is_valid_number, format_number, PhoneNumberFormat
from app.modal import (AcademicYear, Branch, Class, ClassLevel, ClassSubject, Exam, ExamHall, ExamHallAssignment, ExamPaper, ExamQuestion, ExamSubject, ExamTicket, ExamTimetable, FeeInvoice, Parent, Permission, Role, RolePermission, School, SchoolSiteSettings, Section, SettingsData, SomaliaLocation, Student, StudentAttendance, StudentExamMark, StudentExamResult, StudentFeeCollection, StudentPromotion, Subject, Teacher, TeacherAssignment, Term, TimeSlot, Timetable,User, UserLog, UserPermission, UserRole, UserSession)
from app.utils import get_academic_year
from app.view import AcademicYearForm, AttendanceForm, BranchForm, ChangePasswordForm, ClassForm, ClassLevelForm, ClassSubjectForm, ExamForm, ExamHallAssignmentForm, ExamHallForm, ExamMultiPublishForm, ExamQuestionForm, ExamSubjectForm, ExamTimetableForm, ExamTimetableForm, ForgotPasswordChangeForm, ForgotPasswordForm, LoginForm, ParentForm, RegisterForm, SchoolForm, SchoolSiteSettingsForm, SectionForm, SettingsDataForm, SomaliaLocationForm, StudentExamMarkForm, StudentExamResultForm, StudentFeeCollectionForm, StudentForm, StudentPromotionForm, SubjectForm, TeacherAssignmentForm, TeacherForm, TermForm, TimeSlotForm, TimetableForm, TwoFactorForm, UserForm, UserProfileForm, VerifyOTPForm 
from user_agents import parse as parse_ua  # install: pip install pyyaml user-agents
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from app import csrf   # ✅ NOW THIS WORKS




bp = Blueprint('main', __name__)

#------------------------------------------
#---- Function: 1 | Func Allowed Files  ---
#------------------------------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
 

@bp.route('/', methods=['GET', 'POST'])
def index():
    settings = SettingsData.query.first()

    form = LoginForm()  # create an instance of your login form

    if form.validate_on_submit():
        # Example login logic
        login_id = form.login_id.data
        password = form.password.data
        remember = form.remember.data

        # TODO: Add your authentication logic here
        # e.g., user = authenticate_user(login_id, password)

        flash('Logged in successfully!', 'success')
        return redirect(url_for('main.dashboard'))  # redirect to dashboard

    # **Pass the form to the template**
    return render_template("backend/auth/auth-login.html", form=form,  settings=settings,)




@bp.context_processor
def inject_settings():
    settings = SettingsData.query.first()
    return dict(settings=settings)


# Get real IP
# Get real client IP

def get_ip():
    """Return the real client IP, handling proxies."""
    headers_to_check = ['X-Forwarded-For', 'X-Real-IP', 'CF-Connecting-IP', 'Forwarded']
    for header in headers_to_check:
        ip = request.headers.get(header)
        if ip:
            ip = ip.split(',')[0].strip()
            if ip:
                return ip
    return request.remote_addr or "Unknown"

def get_active_network_interface():
    if_addrs = psutil.net_if_addrs()
    if_stats = psutil.net_if_stats()

    for iface, stats in if_stats.items():
        if stats.isup:
            name_lower = iface.lower()
            if "wi-fi" in name_lower or "wlan" in name_lower or "wireless" in name_lower:
                return "Wi-Fi"
            elif "eth" in name_lower or "en" in name_lower:
                return "Ethernet"
            elif "loopback" in name_lower:
                continue  # skip loopback
            else:
                return iface
    return "Unknown"


def get_device_name_from_ua(ua_string, user_id=""):
    """Return detailed device info from User-Agent."""
    ua = parse_ua(ua_string)
    device_type = "Desktop"
    os_name = ua.os.family
    os_version = ua.os.version_string or "Unknown"
    browser_name = ua.browser.family or "Unknown"
    manufacturer = ""
    model = ""

    if ua.is_mobile or ua.is_tablet:
        device_type = "Mobile" if ua.is_mobile else "Tablet"

        if "Android" in ua_string:
            match = re.search(r'Android [\d\.]+; ([^;)\]]+)', ua_string)
            if match:
                raw_model = match.group(1).strip()
                parts = raw_model.split(" ")
                if len(parts) > 1:
                    manufacturer = parts[0]
                    model = " ".join(parts[1:])
                else:
                    model = raw_model
            os_name = "Android"
            ver_match = re.search(r'Android ([\d\.]+)', ua_string)
            if ver_match:
                os_version = ver_match.group(1)

        elif "iPhone" in ua_string or "iPad" in ua_string:
            manufacturer = "Apple"
            model = "iPhone" if "iPhone" in ua_string else "iPad"
            os_name = "iOS"
            ver_match = re.search(r'OS ([\d_]+)', ua_string)
            if ver_match:
                os_version = ver_match.group(1).replace('_', '.')

    device_name = f"{user_id} | {os_name} {os_version} | {device_type} | {manufacturer} {model}".strip()
    return device_name, os_name, os_version, device_type, browser_name, manufacturer, model

def create_user_log(user_id, action, extra_info="", status="info"):
    """Logs user action and updates user's device info fields."""
    ua_string = request.headers.get('User-Agent', '')
    device_name, os_name, os_version, device_type, browser_name, manufacturer, model = get_device_name_from_ua(
        ua_string, user_id=str(user_id)
    )

    # Server-side system info
    system_info = {
        "architecture": platform.architecture()[0],
        "processor": platform.processor() or "Unknown",
        "ram_gb": round(psutil.virtual_memory().total / (1024**3), 2)
    }

    full_device_name = f"{device_name} | {system_info['architecture']} | {system_info['processor']} | RAM {system_info['ram_gb']} GB"

    # Get network interface (optional, demo placeholder)
    interface_name = get_active_network_interface()
    # Update password
    africa_time = datetime.now(pytz.timezone("Africa/Nairobi"))
    
    # Create log
    log = UserLog(
        user_id=user_id,
        login_time=africa_time,
        ip_address=get_ip(),
        device=device_type,         # Desktop / Mobile / Tablet
        browser=browser_name,
        platform=os_name,
        device_name=full_device_name,
        interface_name=interface_name,
        action=action,
        status=status,
        extra_info=extra_info
    )

    # Update the user record
    user = User.query.get(user_id)
    if user:
        user.device = device_type
        user.browser = browser_name
        user.ip_address=get_ip()
        user.platform = os_name
        user.device_name = full_device_name
        user.interface_name = interface_name
        user.extra_info = extra_info
        user.last_login_ip = get_ip()
        user.login_time = africa_time
        db.session.add(user)

    db.session.add(log)
    db.session.commit()

    
@bp.route("/api/device_info", methods=["POST"])
def device_info():
    data = request.json

    log = UserLog(
        user_id=data.get("user_id"),
        ip_address=data.get("ip"),
        subnet_mask=data.get("subnet"),
        gateway=data.get("gateway"),
        mac_address=data.get("mac"),
        device_name=data.get("device_name"),
        interface_name=data.get("interface_name"),
        platform=data.get("platform"),
        device="Desktop Agent",
        browser="N/A",
        action="device_info",
        status="info",
        extra_info="Device info from agent"
    )

    db.session.add(log)
    db.session.commit()
    return {"status": "success"}



#-----------------------------------------------
#---- Route: 1 | Home - Frontend Template ------
#-----------------------------------------------





#-----------------------------------------------
#---- Route: 2 | Login - Auth Template ------
#-----------------------------------------------
@bp.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        flash("You are already logged in!", "info")
        return redirect(url_for("main.dashboard"))

    form = LoginForm()

    if form.validate_on_submit():
        login_id = form.login_id.data.strip()
        password = form.password.data

        # 🔍 Fetch user
        user = User.query.filter(
            (User.email == login_id) |
            (User.username == login_id) |
            (User.phone == login_id)
        ).first()

        if not user or not check_password_hash(user.password, password):
            flash("Invalid login credentials", "danger")
            return redirect(url_for("main.login"))

        if not bool(user.status):
            flash("Your account is inactive", "danger")
            return redirect(url_for("main.login"))

        # 🔹 Clear old local session
        session.clear()

        now = now_eat()
        ua_string = request.headers.get('User-Agent', '')
        ip_address = get_ip()
        interface_name = get_active_network_interface()

        # 🔹 Detect device details
        device_name, os_name, os_version, device_type, browser_name, manufacturer, model = get_device_name_from_ua(
            ua_string,
            user_id=user.username
        )

        # ---------------------------------------------------------
        # 🔥 UNIQUE SESSION LOGIC (Hal Qalab = Hal Token)
        # ---------------------------------------------------------
        # Waxaan fiirinaynaa haddii uu jiro session hore oo isku mid ka ah:
        # User ID + User Agent (Browser-ka) + IP Address
        existing_session = UserSession.query.filter_by(
            user_id=user.id,
            user_agent=ua_string,
            ip_address=ip_address
        ).first()

        import uuid

        if existing_session:
            # Qofkaas waa la yaqaan, isla Token-kiisii hore u daay
            session_entry = existing_session
            session_entry.last_activity = now
            session_token = session_entry.session_token
            extra_msg = "Existing device & browser session reused."
        else:
            # Haddii uu yahay qalab cusub ama browser kale: Samey Token cusub
            session_token = str(uuid.uuid4())
            session_entry = UserSession(
                id=uuid.uuid4().hex,
                user_id=user.id,
                session_token=session_token,  # Token-kan wuxuu noqonayaa aqoonsiga qalabkan
                ip_address=ip_address,
                user_agent=ua_string,
                device=device_type,
                browser=browser_name,
                platform=os_name,
                payload=None,
                last_activity=now
            )
            db.session.add(session_entry)
            extra_msg = "New session created (New device/browser)."

        # -----------------------------
        # ✅ Create UserLog
        # -----------------------------
        log = UserLog(
            user_id=user.id,
            login_time=now,
            ip_address=ip_address,
            device=device_type,
            browser=browser_name,
            platform=os_name,
            device_name=device_name,
            interface_name=interface_name,
            extra_info=f"{extra_msg} | Manufacturer: {manufacturer}, Model: {model}",
            status="login",
            action="login"
        )
        db.session.add(log)

        # -----------------------------
        # ✅ Update USER Table
        # -----------------------------
        user.device = device_type
        user.browser = browser_name
        user.platform = os_name
        user.device_name = device_name
        user.interface_name = interface_name
        user.last_login_ip = ip_address
        user.login_time = now
        user.last_active = now
        user.auth_status = "login"

        db.session.add(user)
        db.session.commit()

        # 🔥 SAVE TO FLASK SESSION
        session["session_id"] = session_entry.id
        session["session_token"] = session_token # Token-kan waa kan loo isticmaali doono hubinta qalabka
        session["log_id"] = log.id

        # 🔥 LOGIN USER
        login_user(user)

        flash(f"Welcome back, {user.fullname}!", "success")
        return redirect(url_for("main.dashboard"))

    return render_template("backend/auth/auth-login.html", form=form)



@bp.before_request
def check_session_validity():

    if current_user.is_authenticated:

        session_token = session.get("session_token")

        user_session = UserSession.query.filter_by(
            user_id=current_user.id,
            session_token=session_token
        ).first()

        # ❌ haddii session la delete gareeyay
        if not user_session:
            logout_user()
            session.clear()


@bp.route('/logout-other-session', methods=['POST'])
@login_required
def logout_other_session():

    data = request.get_json()
    session_id = data.get('session_id')
    password = data.get('password')

    # 🔐 check password
    if not check_password_hash(current_user.password, password):
        return jsonify({'success': False, 'message': 'Password is incorrect'})

    # 🔍 find session
    session_to_delete = UserSession.query.filter_by(
        id=session_id,
        user_id=current_user.id
    ).first()

    if not session_to_delete:
        return jsonify({'success': False, 'message': 'Session not found'})

    try:
        # 🔥 IMPORTANT: invalidate session token (force logout)
        session_to_delete.session_token = "invalid_" + str(session_to_delete.id)

        db.session.delete(session_to_delete)
        db.session.commit()

        return jsonify({
            'success': True,
            'message': 'Session logged out successfully'
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({
            'success': False,
            'message': str(e)
        })





@bp.route('/ping_session', methods=['POST'])
@login_required
def ping_session():
    session_id = session.get('session_id')
    if not session_id:
        return jsonify({'success': False})

    user_session = UserSession.query.get(session_id)
    if user_session:
        # Update last_activity timestamp
        user_session.last_activity = now_eat
        db.session.commit()
    return jsonify({'success': True})

# ----- Single User Last Active -----
# ----- Helper Function -----
def humanize_time_diff(past_time, now=None):
    if not past_time:
        return "Never"

    # Use EAT timezone to get current time
    now = now or now_eat()

    # If past_time is naive (no timezone), localize it to EAT
    if past_time.tzinfo is None or past_time.tzinfo.utcoffset(past_time) is None:
        past_time = EAT.localize(past_time)
    else:
        past_time = past_time.astimezone(EAT)

    if past_time > now:
        return "Just now"

    diff = now - past_time
    seconds = int(diff.total_seconds())
    minutes = seconds // 60
    hours = minutes // 60
    days = diff.days

    if seconds < 60:
        return f"{seconds} sec{'s' if seconds != 1 else ''} ago"
    elif minutes < 60:
        return f"{minutes} min{'s' if minutes != 1 else ''} ago"
    elif hours < 24:
        return f"{hours} hr{'s' if hours != 1 else ''} ago"
    elif days < 7:
        return f"{days} day{'s' if days != 1 else ''} ago"
    elif days < 30:
        weeks = days // 7
        return f"{weeks} week{'s' if weeks != 1 else ''} ago"
    elif days < 365:
        months = days // 30
        return f"{months} month{'s' if months != 1 else ''} ago"
    else:
        years = days // 365
        return f"{years} year{'s' if years != 1 else ''} ago"



    
# Heartbeat (update last_active)
@bp.route("/api/heartbeat", methods=["POST"])
@csrf.exempt    # ✅ correct
@login_required
def heartbeat():
    try:
        now = now_eat()  # 🔹 current datetime

        # 🔹 Update USER
        current_user.last_active = now
        current_user.auth_status = "login"

        # 🔹 Get session ID safely
        session_id = session.get("session_id")
        if not session_id:
            return jsonify({
                "success": False,
                "error": "No session_id found"
            }), 400

        # 🔹 Update CURRENT SESSION (SAFE QUERY)
        user_session = UserSession.query.filter_by(
            id=session_id,
            user_id=current_user.id
        ).first()

        if not user_session:
            return jsonify({
                "success": False,
                "error": "Session not found"
            }), 404

        # ✅ SAFE: simply update last_activity
        user_session.last_activity = now

        # optional column
        if hasattr(user_session, "is_active"):
            user_session.is_active = True

        db.session.commit()

        return jsonify({
            "success": True,
            "last_active": now.strftime('%Y-%m-%d %H:%M:%S'),
            "session_id": session_id
        })

    except Exception as e:
        db.session.rollback()
        print("🔥 HEARTBEAT ERROR:")
        traceback.print_exc()   # 🔹 full traceback for debugging

        return jsonify({
            "success": False,
            "error": "Server error"
        }), 500


@bp.route("/api/online-status")
@login_required
def api_online_status():
    eat = pytz.timezone("Africa/Nairobi")
    now = now_eat()
    
    users = User.query.order_by(User.login_time.desc()).limit(25).all()
    data = []
    for user in users:
        last_active = None
        if user.last_active:
            last_active = (
                eat.localize(user.last_active) if user.last_active.tzinfo is None
                else user.last_active.astimezone(eat)
            )
        is_online = user.auth_status == 'login' and user.session_token and last_active and (now - last_active).total_seconds() < 60
        data.append({
            "id": user.id,
            "username": user.username,
            "is_online": is_online,
            "last_seen_ago": humanize_time_diff(last_active, now) if last_active else "Never"
        })
    return jsonify(data)


@bp.route("/api/user-last-active/<int:user_id>")
@login_required
def get_user_last_active(user_id):
    eat = pytz.timezone("Africa/Nairobi")
    now = datetime.now(eat)
    user = User.query.get_or_404(user_id)

    if user.last_active:
        last_active = (
            eat.localize(user.last_active) if user.last_active.tzinfo is None
            else user.last_active.astimezone(eat)
        )
        last_seen_ago = humanize_time_diff(last_active, now)
        formatted_time = last_active.strftime('%d %b %Y, %I:%M %p')
    else:
        last_active = None
        last_seen_ago = "Never"
        formatted_time = "Unknown"

    is_online = user.auth_status == 'login' and user.session_token and last_active and (now - last_active).total_seconds() < 30

    return jsonify({
        "id": user.id,
        "username": user.username,
        "is_online": is_online,
        "last_seen_ago": last_seen_ago,
        "formatted_time": formatted_time
    })



#----------------------------------------------------------
#---- Route: 39 |  STEP 1: Request Reset (Enter Email) ----
#----------------------------------------------------------
@bp.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    form = ForgotPasswordForm()

    if form.validate_on_submit():
        email = form.email.data
        session['forgot_password_email'] = email

        user = User.query.filter_by(email=email).first()
        if not user:
            flash("If this email is registered, you will receive password reset instructions shortly.", "info")
            return redirect(url_for('main.forgot_password'))

        # Generate OTP
        otp_code = str(random.randint(100000, 999999))
        session['forgot_password_otp'] = otp_code
        session['otp_created_at'] = now_eat().isoformat()

        # Send OTP Email
        try:
            reset_link = url_for('main.forgot_password_verify_otp', _external=True)
            send_otp_email(
                user_email=email,
                otp_code=otp_code,
                username=user.username,
                reset_link=reset_link
            )
            flash("OTP sent to your email. Please check your inbox.", "success")
        except Exception as e:
            flash(f"Failed to send OTP email. Please try again later. ({str(e)})", "danger")

        return redirect(url_for('main.forgot_password_verify_otp'))

    site_data = SettingsData.query.first()
    return render_template(
        'backend/auth/auth-reset-creative.html',
        site_data=site_data,
        form=form
    )


#----------------------------------------------------------------
#---- Route: 40 |  STEP 2: Verify OTP ---------------------------
#----------------------------------------------------------------
@bp.route('/forgot-password/verify-otp', methods=['GET', 'POST'])
def forgot_password_verify_otp():
    if current_user.is_authenticated:
        flash("You are already logged in.", "info")
        return redirect(url_for("main.dashboard"))

    if 'forgot_password_otp' not in session or 'forgot_password_email' not in session:
        flash("Session expired. Please start the password reset again.", "error")
        return redirect(url_for('main.forgot_password'))

    email = session.get('forgot_password_email')
    saved_otp = session.get('forgot_password_otp')
    otp_created_at = session.get('otp_created_at')

    form = VerifyOTPForm()

    if form.validate_on_submit():
        input_otp = form.otp_code.data

        # OTP expiry check
        if otp_created_at:
            otp_time = datetime.fromisoformat(otp_created_at)
            if now_eat() - otp_time > timedelta(minutes=5):
                session.clear()
                flash("OTP expired. Please request a new password reset.", "error")
                return redirect(url_for('main.forgot_password'))

        if input_otp == saved_otp:
            user = User.query.filter_by(email=email).first()
            if user:
                session['forgot_password_verified_email'] = email
                flash("OTP verified successfully. Please change your password.", "success")
                return redirect(url_for('main.forgot_password_change_password'))
            else:
                flash("User not found for this email.", "error")
                return redirect(url_for('main.forgot_password'))
        else:
            flash("Invalid OTP. Please try again.", "error")

    site_data = SettingsData.query.first()
    return render_template(
        'backend/auth/auth-verify-creative.html',
        site_data=site_data,
        email=email,
        form=form
    )


 
#-------------------------------------------------------------
#---- Route: 41 |  STEP 3: Change Password Section after OTP -
#-------------------------------------------------------------
@bp.route('/forgot-change-password', methods=['GET', 'POST'])
def forgot_password_change_password():
    # 1. Hubi in email-ka uu session-ka ku jiro
    email = session.get('forgot_password_verified_email')
    if not email:
        flash("Session expired. Please restart the password reset process.", "error")
        return redirect(url_for('main.forgot_password'))

    # 2. Hel qofka isticmaalaha ah (User)
    user = User.query.filter_by(email=email).first()
    if not user:
        flash("User not found.", "error")
        return redirect(url_for('main.forgot_password'))

    form = ForgotPasswordChangeForm()

    # 3. Marka foomka la soo gudbiyo (POST)
    if form.validate_on_submit():
        new_password = form.new_password.data
        confirm_password = form.confirm_password.data

        # 4. Xaqiijinta Shuruudaha Password-ka (Server-side Validation)
        if not is_valid_password(new_password):
            flash("Password-ku ma buuxin shuruudaha (8 xaraf, weyn, yar, nambar iyo astaan).", "warning")
            return redirect(request.url)

        # 5. Hubi inay isku mid yihiin
        if new_password != confirm_password:
            flash("Password-ada aad qortay isku mid ma ahan!", "danger")
            return redirect(request.url)

        # 6. Beddel password-ka oo keydi
        try:
            user.password = generate_password_hash(new_password)
            db.session.commit()

            # Clear session si uusan mar kale u isticmaalin link-ga
            session.pop('forgot_password_verified_email', None)

            flash("✅ Password-ka si guul leh ayaa loo beddelay! Hadda waad geli kartaa.", "success")
            return redirect(url_for('main.login'))
            
        except Exception as e:
            db.session.rollback()
            flash("Cilad ayaa dhacday intii password-ka la keydinayay.", "error")

    # 7. Soo bandhig foomka (GET)
    site_data = SettingsData.query.first()
    return render_template(
        'backend/auth/auth-resetting-creative.html',
        form=form,
        site_data=site_data,
        username=user.username
    )

#---------------------------------------------------------------------------------------------
##--------------------------------------------------------------------------------------------
#---- Function: 9 | Send OTP Email using Mailtrap Section ------------------------------------

def send_otp_email(user_email, otp_code, username, reset_link=None,
                   sender_name=None, sender_email=None):
    """
    Send OTP email with full site info, user info, social links, logo, footer,
    and optional reset link. Sender name is visible in inbox.
    """
    try:
        # 🔹 Get site settings
        site_data = SettingsData.query.first()

        # 🔹 Defaults / All fields
        group_name = getattr(site_data, "group_name", "Your Company")
        system_name = getattr(site_data, "system_name", group_name)
        address = getattr(site_data, "address", "")
        phone1 = getattr(site_data, "phone1", "")
        phone2 = getattr(site_data, "phone2", "")
        email_support = getattr(site_data, "email", "support@example.com")
        facebook = getattr(site_data, "facebook", "")
        twitter = getattr(site_data, "twitter", "")
        instagram = getattr(site_data, "instagram", "")
        dribbble = getattr(site_data, "dribbble", "")
        logo = getattr(site_data, "logo", "")

        # 🔹 Current year
        current_year = datetime.now(pytz.timezone("Africa/Nairobi")).year

        # 🔹 Logo URL
        logo_url = url_for('static', filename=logo, _external=True) if logo else None

        # 🔹 Sender email + name visible
        if sender_email is None:
            sender_email = current_app.config.get("MAIL_USERNAME", "liilove668@gmail.com")
        if sender_name is None:
            sender_name = group_name  # Use site group_name as visible sender

        sender_full = f"{sender_name} <{sender_email}>"

        # 🔹 Prepare email
        msg = Message(
            subject=f"{group_name} - OTP Verification Code",
            sender=sender_full,
            recipients=[user_email]
        )
 
        # 🔹 Render HTML template with full data
        msg.html = render_template(
            "backend/auth/auth-sms-verify.html",
            otp_code=otp_code,
            username=username,
            email=user_email,
            reset_link=reset_link,
            current_year=current_year,
            # Site info
            group_name=group_name,
            system_name=system_name,
            address=address,
            phone1=phone1,
            phone2=phone2,
            email_support=email_support,
            facebook=facebook,
            twitter=twitter,
            instagram=instagram,
            dribbble=dribbble,
            logo=logo_url
        )

        # 🔹 Send email
        mail.send(msg)
        print("✅ OTP email sent successfully with full data and visible sender")

    except Exception as e:
        print(f"❌ Error sending OTP email: {str(e)}")
        raise



#------------ save and clean sign in google
@bp.route("/google-login")
def google_login():
    redirect_uri = url_for('main.google_authorize', _external=True)
    return google.authorize_redirect(redirect_uri)

#--------------- google authonzied saved and clean
@bp.route("/google-authorize")
def google_authorize():
    try:
        # 1. Hel xogta Google
        token = google.authorize_access_token()
        user_info = google.userinfo()

        email = user_info.get('email')
        fullname = user_info.get('name')
        google_id = user_info.get('sub')

        if not email:
            flash("Failed to get email from Google.", "danger")
            return redirect(url_for("main.login"))

        # 2. Hubi haddii user-ka uu horey ugu jiro Database-keena
        user = User.query.filter_by(email=email).first()

        # SHARCI: Haddii qofku uusan database-ka ku jirin, ma abuuri karno account cusub halkan
        # Waa inuu account hore u lahaa ama laga soo diwaangaliyay system-ka
        if not user:
            flash("Account-kan kuma diwaangashana system-ka. Fadlan la xiriir maamulka.", "danger")
            return redirect(url_for("main.login"))

        # 3. Hubi haddii account-ku uu xiran yahay (Inactive)
        if not bool(user.status):
            flash("Account-kaaga waa laga joojiyay shaqada.", "danger")
            return redirect(url_for("main.login"))

        # 4. Diyaarinta xogta qalabka (Device & Session)
        session.clear() # Nadiifi session-kii hore
        
        now = now_eat()
        ua_string = request.headers.get('User-Agent', '')
        ip_address = get_ip()
        interface_name = get_active_network_interface()

        # Faahfaahinta qalabka
        device_name, os_name, os_version, device_type, browser_name, manufacturer, model = get_device_name_from_ua(
            ua_string,
            user_id=user.username or user.email
        )

        # 5. 🔥 HUBINTA QALABKA (Isla sharcigii login-ka caadiga ah)
        existing_session = UserSession.query.filter_by(
            user_id=user.id,
            user_agent=ua_string,
            ip_address=ip_address
        ).first()

        if existing_session:
            # Haddii qalabkan la yaqaan, isticmaal session-kiisii hore
            session_entry = existing_session
            session_entry.last_activity = now
            session_token = session_entry.session_token
            extra_msg = "Google Login: Existing device reused."
        else:
            # Haddii uu yahay qalab cusub, u samee mid cusub
            session_token = str(uuid.uuid4())
            session_entry = UserSession(
                id=uuid.uuid4().hex,
                user_id=user.id,
                session_token=session_token,
                ip_address=ip_address,
                user_agent=ua_string,
                device=device_type,
                browser=browser_name,
                platform=os_name,
                payload=None,
                last_activity=now
            )
            db.session.add(session_entry)
            extra_msg = "Google Login: New device detected."

        # 6. ✅ Create UserLog
        log = UserLog(
            user_id=user.id,
            login_time=now,
            ip_address=ip_address,
            device=device_type,
            browser=browser_name,
            platform=os_name,
            device_name=device_name,
            interface_name=interface_name,
            extra_info=f"{extra_msg} | Google Login",
            status="login",
            action="google_auth"
        )
        db.session.add(log)

        # 7. ✅ Update User Table
        user.last_login_ip = ip_address
        user.login_time = now
        user.last_active = now
        user.auth_status = "login"
        # Haddii uu markii ugu horaysay Google ku soo galayo, ku dar Google ID-giisa
        if not user.google:
            user.google = google_id

        db.session.add(user)
        db.session.commit()

        # 8. 🔥 Kaydi Session-ka Flask
        session["session_id"] = session_entry.id
        session["session_token"] = session_token
        session["log_id"] = log.id

        # 9. Login-ka rasmiga ah
        login_user(user)

        flash(f"Welcome back via Google, {user.fullname}!", "success")
        return redirect(url_for("main.dashboard"))

    except Exception as e:
        db.session.rollback()
        flash(f"Google authorization failed: {str(e)}", "danger")
        return redirect(url_for("main.login"))


#------------- githup auth
# -----------------------------
# GitHub login redirect
@bp.route("/github-login")
def github_login():
    redirect_uri = url_for("main.github_authorize", _external=True)
    return github.authorize_redirect(redirect_uri)

# GitHub OAuth callback
@bp.route("/github-authorize")
def github_authorize():
    try:
        # 1. Hel access token iyo xogta user-ka
        token = github.authorize_access_token()
        user_info = github.get("user").json()

        email = user_info.get("email")
        fullname = user_info.get("name")
        github_id = str(user_info.get("id"))

        # 2. Haddii email-ku qarsan yahay GitHub, ka soo saar liiska emails-ka
        if not email:
            emails_resp = github.get("user/emails")
            if emails_resp.status_code == 200:
                emails = emails_resp.json()
                primary_email = next((e['email'] for e in emails if e['primary'] and e['verified']), None)
                email = primary_email

        if not email:
            flash("Failed to get a verified email from GitHub.", "danger")
            return redirect(url_for("main.login"))

        # 3. Hubi haddii user-ku uu horey ugu jiro Database-ka (Email kaliya isticmaal)
        user = User.query.filter_by(email=email).first()

        # SHARCI: User-ka aan diwaangashanayn ma abuuri karo account cusub halkan
        if not user:
            flash("Account-kan kuma diwaangashana system-ka. Fadlan is-diwaangali marka hore.", "danger")
            return redirect(url_for("main.login"))

        # 4. Hubi haddii account-ku firfircoon yahay
        if not bool(user.status):
            flash("Account-kaaga waa mid aan firfirconayn (Inactive).", "danger")
            return redirect(url_for("main.login"))

        # 5. Diyaarinta session-ka iyo xogta qalabka
        session.clear()
        
        now = now_eat()
        ua_string = request.headers.get('User-Agent', '')
        ip_address = get_ip()
        interface_name = get_active_network_interface()

        # Faahfaahinta qalabka
        device_name, os_name, os_version, device_type, browser_name, manufacturer, model = get_device_name_from_ua(
            ua_string,
            user_id=user.username or user.email
        )

        # 6. 🔥 HUBINTA QALABKA (Hal Qalab = Hal Session Token)
        existing_session = UserSession.query.filter_by(
            user_id=user.id,
            user_agent=ua_string,
            ip_address=ip_address
        ).first()

        if existing_session:
            # Haddii isla qalabkii iyo browser-kii yahay, ha abuurin token cusub
            session_entry = existing_session
            session_entry.last_activity = now
            session_token = session_entry.session_token
            extra_msg = "GitHub Login: Existing session reused."
        else:
            # Haddii uu yahay qalab cusub, abuuro mid cusub
            session_token = str(uuid.uuid4())
            session_entry = UserSession(
                id=uuid.uuid4().hex,
                user_id=user.id,
                session_token=session_token,
                ip_address=ip_address,
                user_agent=ua_string,
                device=device_type,
                browser=browser_name,
                platform=os_name,
                payload=None,
                last_activity=now
            )
            db.session.add(session_entry)
            extra_msg = "GitHub Login: New device session created."

        # 7. ✅ Create UserLog
        log = UserLog(
            user_id=user.id,
            login_time=now,
            ip_address=ip_address,
            device=device_type,
            browser=browser_name,
            platform=os_name,
            device_name=device_name,
            interface_name=interface_name,
            extra_info=f"{extra_msg} | GitHub Auth",
            status="login",
            action="github_auth"
        )
        db.session.add(log)

        # 8. ✅ Update User Table
        user.last_login_ip = ip_address
        user.login_time = now
        user.last_active = now
        user.auth_status = "login"
        
        # Ku xir GitHub ID haddii uusan horey ugu xirnayn
        if not user.github:
            user.github = github_id

        db.session.add(user)
        db.session.commit()

        # 9. 🔥 Kaydi Session-ka Flask
        session["session_id"] = session_entry.id
        session["session_token"] = session_token
        session["log_id"] = log.id

        # Login-ka rasmiga ah
        login_user(user)

        flash(f"Welcome back via GitHub, {user.fullname}!", "success")
        return redirect(url_for("main.dashboard"))

    except Exception as e:
        db.session.rollback()
        flash(f"GitHub login failed: {str(e)}", "danger")
        return redirect(url_for("main.login"))


#------------------------------

#-----------------------------------------------
#---- Route: 3 | Register - Auth Template ------
#-----------------------------------------------
@bp.route("/register", methods=["GET", "POST"])
def register():
    if current_user.is_authenticated:
        flash("Logout first to create a new account!", "info")
        return redirect(url_for("main.dashboard"))

    form = RegisterForm()
    
    if form.validate_on_submit():
        fullname = form.fullname.data
        username = form.username.data
        email = form.email.data
        phone = form.phone.data
        password = form.password.data
        confirm_password = form.confirm_password.data

        if password != confirm_password:
            flash("Passwords do not match", "danger")
            return redirect(url_for("main.register"))

        existing_user = User.query.filter(
            (User.username == username) |
            (User.email == email) |
            (User.phone == phone)
        ).first()

        if existing_user:
            flash("Username, email or phone already exists", "danger")
            return redirect(url_for("main.register"))

        user = User(
            fullname=fullname,
            username=username,
            email=email,
            phone=phone,
            password=generate_password_hash(password),
            role="user",
            status=0
        )

        db.session.add(user)
        db.session.commit()

        flash("Account created successfully. Please login!", "success")
        return redirect(url_for("main.login"))

    # ✅ Halkan waa muhiim: ku soo dir 'form' template-ka
    return render_template("backend/auth/auth-register.html", form=form)


# ================= HELPERS =================
def get_user_site_logo(user):
    def safe_path(path):
        if path:
            return path.replace("\\", "/").lstrip("static/").lstrip("/")
        return None

    school_id = getattr(user, "school_id", None)
    branch_id = getattr(user, "branch_id", None)

    # ---------------------------
    # 🔥 SUPERADMIN → GLOBAL ONLY
    # ---------------------------
    if hasattr(user, "role") and user.role.value == "superadmin":
        settings = SettingsData.query.order_by(SettingsData.id.desc()).first()
        return {
            "main_logo": safe_path(settings.logo) if settings and settings.logo else None,
            "sub_logo": safe_path(settings.logo2) if settings and settings.logo2 else None,
            "sign_logo": safe_path(settings.sign_logo) if settings and hasattr(settings, "sign_logo") else None
        }

    # ---------------------------
    # 🔥 BRANCH ADMIN → ONLY ITS BRANCH LOGO
    # ---------------------------
    if hasattr(user, "role") and user.role.value == "branch_admin" and branch_id:
        settings = SchoolSiteSettings.query.filter_by(branch_id=branch_id).order_by(SchoolSiteSettings.id.desc()).first()
        if settings:
            return {
                "main_logo": safe_path(settings.main_logo),
                "sub_logo": safe_path(settings.sub_logo),
                "sign_logo": safe_path(settings.sign_logo)
            }
        return {"main_logo": None, "sub_logo": None, "sign_logo": None}

    # ---------------------------
    # 🔥 SCHOOL ADMIN → SCHOOL LOGO (branch=None)
    # ---------------------------
    if hasattr(user, "role") and user.role.value == "school_admin" and school_id:
        settings = SchoolSiteSettings.query.filter_by(school_id=school_id, branch_id=None).order_by(SchoolSiteSettings.id.desc()).first()
        if settings:
            return {
                "main_logo": safe_path(settings.main_logo),
                "sub_logo": safe_path(settings.sub_logo),
                "sign_logo": safe_path(settings.sign_logo)
            }
        return {"main_logo": None, "sub_logo": None, "sign_logo": None}

    # ---------------------------
    # 🔥 TEACHER / STUDENT / PARENT → ONLY BRANCH LOGO
    # ---------------------------
    if branch_id:
        settings = SchoolSiteSettings.query.filter_by(branch_id=branch_id).order_by(SchoolSiteSettings.id.desc()).first()
        if settings:
            return {
                "main_logo": safe_path(settings.main_logo),
                "sub_logo": safe_path(settings.sub_logo),
                "sign_logo": safe_path(settings.sign_logo)
            }
        else:
            # ❌ No fallback to school/global, strict branch isolation
            return {"main_logo": None, "sub_logo": None, "sign_logo": None}

    # ---------------------------
    # 🔥 FALLBACK → SCHOOL LOGO IF NO BRANCH
    # ---------------------------
    if school_id:
        settings = SchoolSiteSettings.query.filter_by(school_id=school_id, branch_id=None).order_by(SchoolSiteSettings.id.desc()).first()
        if settings:
            return {
                "main_logo": safe_path(settings.main_logo),
                "sub_logo": safe_path(settings.sub_logo),
                "sign_logo": safe_path(settings.sign_logo)
            }

    # ---------------------------
    # 🔥 FINAL FALLBACK → GLOBAL
    # ---------------------------
    global_settings = SettingsData.query.order_by(SettingsData.id.desc()).first()
    return {
        "main_logo": safe_path(global_settings.logo) if global_settings else None,
        "sub_logo": safe_path(global_settings.logo2) if global_settings else None,
        "sign_logo": safe_path(global_settings.sign_logo) if global_settings and hasattr(global_settings, "sign_logo") else None
    }


# ================= BEFORE REQUEST =================
@bp.before_request
def load_site_logo():
    # Tani waxay shaqaynaysaa mar walba, xataa login-ka ka hor
    # dynamic-ally waxay u soo saaraysaa logo-ga qof kasta (Guest ama User)
    g.site_logo = get_user_site_logo(current_user)






@bp.route("/dashboard", methods=["GET", "POST"])
@login_required
def dashboard():
    @after_this_request
    def add_header(response):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        return response

    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for('main.index'))

    # --- 1. Helitaanka Aqoonsiga User-ka ---
    user_role = current_user.role.value 
    s_id = current_user.school_id
    b_id = current_user.branch_id
    current_year =now_eat().year

    # --- 2. Function-ka Shaandhaynta (Global Filter) ---
    # Waxaan ku daray in haddii uu macalin yahayna uu Branch-kiisa kaliya arko
    def apply_filters(model):
        query = model.query
        if user_role != 'superadmin':
            query = query.filter_by(school_id=s_id)
            if b_id:
                query = query.filter_by(branch_id=b_id)
        return query

    # --- 3. Xisaabinta Lacagaha (Financial Analytics) ---
    # Logic-gan wuxuu hadda si otomaatig ah u raacayaa qofka soo galay
    financial_query = db.session.query(
        func.sum(StudentFeeCollection.amount_due).label('total_due'),
        func.sum(StudentFeeCollection.amount_paid).label('total_paid'),
        func.sum(StudentFeeCollection.remaining_balance).label('total_balance')
    )

    if user_role != 'superadmin':
        financial_query = financial_query.filter(StudentFeeCollection.school_id == s_id)
        if b_id:
            financial_query = financial_query.filter(StudentFeeCollection.branch_id == b_id)

    financials = financial_query.first()

    total_due = financials.total_due or 0.0
    total_paid = financials.total_paid or 0.0
    total_balance = financials.total_balance or 0.0
    
    paid_percentage = round((total_paid / total_due * 100), 1) if total_due > 0 else 0
    balance_percentage = round((total_balance / total_due * 100), 1) if total_due > 0 else 0

    # --- 1. Gender Stats (Pie Chart) ---
    gender_query = db.session.query(
        Student.gender, func.count(Student.id)
    ).group_by(Student.gender)

    if user_role != 'superadmin':
        gender_query = gender_query.filter(Student.school_id == s_id)
        if b_id:
            gender_query = gender_query.filter(Student.branch_id == b_id)

    gender_data = dict(gender_query.all())
    male_count = gender_data.get('male', 0) or gender_data.get('Male', 0)
    female_count = gender_data.get('female', 0) or gender_data.get('Female', 0)

    # --- 2. Attendance Stats (Pie Chart) ---
    attendance_base = apply_filters(StudentAttendance)
    present_count = attendance_base.filter_by(status='present').count()
    absent_count = attendance_base.filter_by(status='absent').count()

    # --- 4. Dakhliga Bilaha ee Sanadkan (Monthly Earnings - Current Year) ---
    # --- 4. Dakhliga Bilaha ee Sanadkan ---
    monthly_query = db.session.query(
        extract('month', StudentFeeCollection.payment_date).label('month'),
        func.sum(StudentFeeCollection.amount_paid).label('monthly_total')
    ).filter(extract('year', StudentFeeCollection.payment_date) == current_year)

    if user_role != 'superadmin':
        monthly_query = monthly_query.filter(StudentFeeCollection.school_id == s_id)
        if b_id:
            monthly_query = monthly_query.filter(StudentFeeCollection.branch_id == b_id)

    monthly_stats = monthly_query.group_by('month').order_by('month').all()

    # Diyaarinta xogta 12-ka bilood
    monthly_revenue = [0] * 12
    for month, total in monthly_stats:
        monthly_revenue[int(month) - 1] = float(total)

    # 1. Wadarta Guud ee Sanadka (Yearly Income)
    total_yearly_income = sum(monthly_revenue) 

    # 2. Magacyada bilaha ee Chart-ka
    month_names = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

    # 3. Dakhliga bishan (April)
    current_month_index = datetime.now().month - 1
    current_month_earning = monthly_revenue[current_month_index]

    # --- 5. Wadarta Dakhliga Sanad kasta (Yearly Total Gained) ---
  # routes.py

    # --- 5. Wadarta Dakhliga Sanad kasta ---
    yearly_query = db.session.query(
        extract('year', StudentFeeCollection.payment_date).label('year'),
        func.sum(StudentFeeCollection.amount_paid).label('yearly_total')
    ).filter(StudentFeeCollection.payment_date != None) # Ha ku darin .all() halkan

    # Shaandhaynta School ama Branch (FILTER-KU WAA INUU HALKAN AHAADAA)
    if user_role != 'superadmin':
        yearly_query = yearly_query.filter(StudentFeeCollection.school_id == s_id) # OK
        if b_id:
            yearly_query = yearly_query.filter(StudentFeeCollection.branch_id == b_id) # OK

    # .all() iyo .group_by() waxaa la qoraa ugu dambaynta
    yearly_stats = yearly_query.group_by('year').order_by('year').all() 

    # Hadda inta kale waa caadi
    years_labels = [str(int(y.year)) for y in yearly_stats if y.year]
    years_totals = [float(y.yearly_total or 0) for y in yearly_stats]

   
    # --- 6. Xisaabinta Lacagaha Maqan (Remaining/Balance) ---
    # Waxaan si toos ah uga soo qaadanaynaa Model-ka StudentFeeCollection
    pending_query = db.session.query(
        func.sum(StudentFeeCollection.amount_due).label('total_due'),
        func.sum(StudentFeeCollection.amount_paid).label('total_paid'),
        func.sum(StudentFeeCollection.remaining_balance).label('total_remaining')
    )

    # Shaandhaynta School ama Branch
    if user_role != 'superadmin':
        pending_query = pending_query.filter(StudentFeeCollection.school_id == s_id)
        if b_id:
            pending_query = pending_query.filter(StudentFeeCollection.branch_id == b_id)

    totals = pending_query.one()

    # Variable-yada loogu talagalay Template-ka
    total_due = float(totals.total_due or 0)
    total_paid = float(totals.total_paid or 0)
    total_remaining = float(totals.total_remaining or 0)

    # Boqolleyda (Percentage)
    paid_percent = (total_paid / total_due * 100) if total_due > 0 else 0
    remaining_percent = (total_remaining / total_due * 100) if total_due > 0 else 0


    # --- 7. Ardayda aan bixinin lacagta bishan (Unpaid Students) ---
    page_num = request.args.get('page', 1, type=int)
    search_term = request.args.get('search', '')

    # Variable-kan waa mid unique ah: monthly_debtor_report
    monthly_debtor_base_query = StudentFeeCollection.query.join(Student).filter(
        StudentFeeCollection.payment_status != 'Paid',
        db.extract('month', StudentFeeCollection.payment_date) == datetime.now().month,
        db.extract('year', StudentFeeCollection.payment_date) == datetime.now().year
    )

    # Haddii uu jiro search
    if search_term:
        monthly_debtor_base_query = monthly_debtor_base_query.filter(
            db.or_(
                Student.full_name.ilike(f'%{search_term}%')
            )
        )

    # Global lock check
    if user_role != 'superadmin':
        monthly_debtor_base_query = monthly_debtor_base_query.filter(StudentFeeCollection.school_id == s_id)

    # Pagination-ka rasmiga ah
    monthly_debtor_report = monthly_debtor_base_query.order_by(
        StudentFeeCollection.remaining_balance.desc()
    ).paginate(page=page_num, per_page=10)

    # 8. Waxaan soo saaraynaa 10-ka arday ee ugu maqnaanshaha badan bishan
    # most_absent_students_list
    query = db.session.query(
        Student, 
        func.count(StudentAttendance.id).label('absent_count')
    ).join(StudentAttendance, Student.id == StudentAttendance.student_id).filter(
        and_(
            StudentAttendance.status == 'absent',
            extract('month', StudentAttendance.date) == datetime.now().month,
            extract('year', StudentAttendance.date) == datetime.now().year
        )
    )

    # 2. Role-based Filter (Global Lock Logic) 
    # Halkan wali waa Query, marka filter-ka waa uu aqbalayaa
    if user_role != 'superadmin':
        query = query.filter(StudentAttendance.school_id == s_id)

    # 3. Hadda dhammaystir grouping-ka iyo soo saarista xogta (all)
    most_absent_students_list = query.group_by(Student.id).order_by(
        db.desc('absent_count')
    ).limit(10).all()

    # 9. Helitaanka maalinta maanta (tusaale: 'Monday', 'Tuesday')
    # Fiiro gaar ah: Hubi in maalmaha database-ka ku jira ay u qoran yihiin sidan.
    current_day_name = datetime.now().strftime('%A')

    # Query-ga saldhigga ah
    schedule_query = db.session.query(Timetable).join(TimeSlot).filter(Timetable.day_of_week == current_day_name)

    # Role-based filtering (Invisible Personalization)
    if user_role != 'superadmin':
        schedule_query = schedule_query.filter(Timetable.school_id == s_id)
        if b_id:
            schedule_query = schedule_query.filter(Timetable.branch_id == b_id)

    # Shaandhayn gaar u ah Macallinka ama Ardayga
    if user_role == 'teacher':
        teacher_rec = Teacher.query.filter_by(user_id=current_user.id).first()
        if teacher_rec:
            schedule_query = schedule_query.filter(Timetable.teacher_id == teacher_rec.id)
            
    elif user_role == 'student':
        student_rec = Student.query.filter_by(user_id=current_user.id).first()
        if student_rec:
            schedule_query = schedule_query.filter(
                Timetable.class_id == student_rec.class_id,
                Timetable.section_id == student_rec.section_id
            )

    # Soo saar xogta adoo u kala horaysiinaya waqtiga (TimeSlot)
    upcoming_schedules = schedule_query.order_by(TimeSlot.start_time.asc()).all()

    # --- 3. Financial Stats (Waa kuwii hore u xisaabsanaa) ---
    # Waxaan u baahanahay total_paid iyo total_balance si aan chart ugu samayno

    # --- 4. Xogta Gaarka u ah Macallinka ---
    teacher_stats = {
        'classes_count': 0, 
        'subjects_count': 0, 
        'class_subjects_count': 0,
        'exam_papers_count': 0
    }

    if user_role == 'teacher':
        teacher_record = Teacher.query.filter_by(user_id=current_user.id).first()
        
        if teacher_record:
            # Ka soo saar Assignments-ka (Kaliya kuwa Macallinka iyo Laantiisa)
            assign_query = TeacherAssignment.query.filter_by(teacher_id=teacher_record.id)
            
            # Hubi Branch-ka
            if b_id:
                assign_query = assign_query.filter_by(branch_id=b_id)
            
            assignments = assign_query.all()

            # Unique Counts
            teacher_stats['classes_count'] = len(set([a.class_id for a in assignments]))
            
            sub_ids = []
            for a in assignments:
                if a.subject_ids:
                    sub_ids.extend(a.subject_ids)
            teacher_stats['subjects_count'] = len(set(sub_ids))
            
            teacher_stats['class_subjects_count'] = len(assignments)

            # Exam Papers (Kaliya kuwa Macallinka uu qoray iyo Laantiisa)
            paper_query = ExamPaper.query.filter_by(teacher_id=teacher_record.id)
            if b_id:
                paper_query = paper_query.filter_by(branch_id=b_id)
            
            teacher_stats['exam_papers_count'] = paper_query.count()
   
    # --- Student Specific Logic ---
    student_stats = {
        'attendance_rate': 0,
        'present_days': 0,
        'total_days': 0,
        'total_due': 0.0,
        'total_paid': 0.0,
        'remaining_balance': 0.0
    }

    if user_role == 'student':
        student_record = Student.query.filter_by(user_id=current_user.id).first()
        
        if student_record:
            # A. Xisaabinta Attendance (Kala saarista nooc kasta)
            attendance_query = StudentAttendance.query.filter_by(student_id=student_record.id)
            total_days = attendance_query.count()

            # Helper function si loogu xisaabiyo tirada iyo boqolayda
            def calculate_status(status_val):
                count = attendance_query.filter_by(status=status_val).count()
                rate = round((count / total_days * 100), 1) if total_days > 0 else 0
                return count, rate

            # Helitaanka xogta dhabta ah
            p_count, p_rate = calculate_status('present')
            a_count, a_rate = calculate_status('absent')
            l_count, l_rate = calculate_status('late')
            e_count, e_rate = calculate_status('excused')

            student_stats.update({
                'total_days': total_days,
                'present_count': p_count, 'present_rate': p_rate,
                'absent_count': a_count, 'absent_rate': a_rate,
                'late_count': l_count, 'late_rate': l_rate,
                'excused_count': e_count, 'excused_rate': e_rate
            })

            # B. Xisaabinta Lacagaha (Financials)
            fees = db.session.query(
                func.sum(StudentFeeCollection.amount_due).label('due'),
                func.sum(StudentFeeCollection.amount_paid).label('paid'),
                func.sum(StudentFeeCollection.remaining_balance).label('balance')
            ).filter(StudentFeeCollection.student_id == student_record.id).first()

            student_stats.update({
                'total_due': fees.due or 0.0,
                'total_paid': fees.paid or 0.0,
                'remaining_balance': fees.balance or 0.0
            })
          

    # --- 5. Tirada Guud (Dynamic Counts) ---
    counts = {
        'students': apply_filters(Student).count(),
        'teachers': apply_filters(Teacher).count(),
        'parents': apply_filters(Parent).count(),
        'classes': apply_filters(Class).count(),
        'subjects': apply_filters(Subject).count(),
        'class_subjects': apply_filters(ClassSubject).count(),
        'teacher_assignments': apply_filters(TeacherAssignment).count(),
        'exams': apply_filters(Exam).count(),
    }

    # --- 6. Session & Settings ---
    active_sessions_count = UserSession.query.filter_by(user_id=current_user.id, is_active=True).count()
    current_session = UserSession.query.filter_by(user_id=current_user.id, is_active=True).order_by(UserSession.last_activity.desc()).first()
    current_device_name = f"{current_session.device} {current_session.platform} ({current_session.browser})" if current_session else "Unknown Device"
    settings = SettingsData.query.first()
   
    return render_template(
        "backend/home/dashbaord.html",
        settings=settings,
        user=current_user,
        active_sessions=active_sessions_count,
        current_device=current_device_name,
        counts=counts,
        total_due=total_due,
        total_paid=total_paid,
        total_balance=total_balance,
        paid_percentage=paid_percentage,
        balance_percentage=balance_percentage,
        teacher_stats=teacher_stats,
        student_stats=student_stats,
        male_count=male_count,
        female_count=female_count,
        present_count=present_count,
        absent_count=absent_count,
        monthly_revenue=monthly_revenue,
        month_names=month_names,
        total_yearly_income=total_yearly_income,
        current_month_earning=current_month_earning,
        current_year=current_year,
        years_labels=years_labels, # Tan ku dar
        years_totals=years_totals,  # Tan ku dar
        monthly_debtor_report=monthly_debtor_report, 
        search_term=search_term,
        most_absent_students_list=most_absent_students_list,
        upcoming_schedules=upcoming_schedules,
        current_day_name=current_day_name,
      
        total_remaining_fees=total_remaining, # Lacagta weli maqan
        remaining_percentage=round(remaining_percent, 1),
    )



@bp.route("/get-dashboard-data")
@login_required
def get_dashboard_data():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    try:
        s_id = current_user.school_id
        b_id = current_user.branch_id
        # Hubi in role-ka uu yahay string
        user_role = current_user.role.value if hasattr(current_user.role, 'value') else current_user.role

        # 1. Financials
        financial_query = db.session.query(
            func.sum(StudentFeeCollection.amount_paid).label('paid'),
            func.sum(StudentFeeCollection.remaining_balance).label('balance')
        )
        if user_role != 'superadmin':
            financial_query = financial_query.filter(StudentFeeCollection.school_id == s_id)
            if b_id: financial_query = financial_query.filter(StudentFeeCollection.branch_id == b_id)
        
        fin = financial_query.first()

        # 2. Gender
        gender_query = db.session.query(Student.gender, func.count(Student.id)).group_by(Student.gender)
        if user_role != 'superadmin':
            gender_query = gender_query.filter(Student.school_id == s_id)
            if b_id: gender_query = gender_query.filter(Student.branch_id == b_id)
        
        g_data = dict(gender_query.all())
        # Hubi Case-sensitivity ('Male' vs 'male')
        m_count = g_data.get('Male', 0) + g_data.get('male', 0)
        f_count = g_data.get('Female', 0) + g_data.get('female', 0)

        # 3. Attendance
        attendance_query = StudentAttendance.query
        if user_role != 'superadmin':
            attendance_query = attendance_query.filter_by(school_id=s_id)
            if b_id: attendance_query = attendance_query.filter_by(branch_id=b_id)
        
        present = attendance_query.filter_by(status='present').count()
        absent = attendance_query.filter_by(status='absent').count()

        return jsonify({
            "gender": [int(m_count), int(f_count)],
            "attendance": [int(present), int(absent)],
            "financial": [float(fin.paid or 0), float(fin.balance or 0)]
        })
    except Exception as e:
        print(f"Dashboard Error: {e}") # Kani wuxuu ku tusi doonaa terminal-ka haddii qalad jiro
        return jsonify({"error": str(e)}), 500


@bp.route("/user/toggle-status/<int:user_id>", methods=["POST"])
@login_required
def toggle_user_status(user_id):

    if current_user.role.value != 'superadmin':
        flash("Permission denied!", "danger")
        return redirect(request.referrer)

    target_user = User.query.get_or_404(user_id)

    try:
        new_status = not bool(target_user.status)
        target_user.status = new_status

        sid = target_user.school_id
        bid = target_user.branch_id

        print("DEBUG:", target_user.id, sid, bid, "NEW:", new_status)

        # ==================================================
        # 🔴 DEACTIVATE CASE
        # ==================================================
        if not new_status:

            if sid is not None:

                User.query.filter(User.school_id == sid).update(
                    {User.status: False},
                    synchronize_session=False
                )

                Branch.query.filter(Branch.school_id == sid).update(
                    {Branch.status: 'inactive'},
                    synchronize_session=False
                )

                School.query.filter(School.id == sid).update(
                    {School.status: 'inactive'},
                    synchronize_session=False
                )

                flash("School + branches + users waa la deactivate gareeyay.", "warning")

            elif bid is not None:

                User.query.filter(User.branch_id == bid).update(
                    {User.status: False},
                    synchronize_session=False
                )

                Branch.query.filter(Branch.id == bid).update(
                    {Branch.status: 'inactive'},
                    synchronize_session=False
                )

                flash("Branch + users waa la deactivate gareeyay.", "warning")

        # ==================================================
        # 🟢 ACTIVATE CASE (NEW ADDITION)
        # ==================================================
        else:

            if sid is not None:

                # activate school users
                User.query.filter(User.school_id == sid).update(
                    {User.status: True},
                    synchronize_session=False
                )

                # activate branches
                Branch.query.filter(Branch.school_id == sid).update(
                    {Branch.status: 'active'},
                    synchronize_session=False
                )

                # activate school
                School.query.filter(School.id == sid).update(
                    {School.status: 'active'},
                    synchronize_session=False
                )

                flash("School + branches + users waa la activate gareeyay.", "success")

            elif bid is not None:

                User.query.filter(User.branch_id == bid).update(
                    {User.status: True},
                    synchronize_session=False
                )

                Branch.query.filter(Branch.id == bid).update(
                    {Branch.status: 'active'},
                    synchronize_session=False
                )

                flash("Branch + users waa la activate gareeyay.", "success")

        db.session.commit()
        db.session.expire_all()

    except Exception as e:
        db.session.rollback()
        flash(f"Cillad: {str(e)}", "danger")

    return redirect(request.referrer or url_for('main.users_list'))


#--------------- Notificstion Contacts


#-------------------End Notification Contacts

#-----------------------------------------------------
#---- Function: 4 | Country Codes -=> Phone Numbers --
#-----------------------------------------------------
def get_country_phone_data():
    countries = []

    for country in pycountry.countries:
        region = country.alpha_2
        try:
            code = phonenumbers.country_code_for_region(region)
            metadata = PhoneMetadata.metadata_for_region(region, None)

            if metadata and code > 0:
                possible_lengths = metadata.general_desc.possible_length
                if possible_lengths:
                    min_len = min(possible_lengths)
                    max_len = max(possible_lengths)
                else:
                    min_len = max_len = ''

                countries.append({
                    'name': country.name,
                    'code': str(code),
                    'alpha2': region,
                    'min_length': min_len,
                    'max_length': max_len,
                })
        except:
            continue

    return sorted(countries, key=lambda x: x['name'])


#-------------------------------------------------------------
#---- Function: 5 | Phone SomalSplit Code country and Number -
#-------------------------------------------------------------
def split_phone_number(e164_phone):
    try:
        phone_obj = parse(e164_phone, None)
        country_code = str(phone_obj.country_code)
        national_number = str(phone_obj.national_number)
        return country_code, national_number
    except NumberParseException:
        return None, None
    

#----------------------------------------------------------------
#---- Function: 6 | Helper: Convert locations to dict Somalia ---
#----------------------------------------------------------------
def group_somalia_regions(locations):
    regions_dict = {}
    for loc in locations:
        region = loc.region
        district = loc.district
        regions_dict.setdefault(region, []).append(district)
    return regions_dict


#-------------------------------------------
#---- Function: 18 | Validation Fullname  --
#-------------------------------------------
def is_valid_fullname(name: str) -> bool:
    """
    Validate full name to allow letters, spaces, dots, apostrophes, and hyphens only.
    """
    pattern = r"^[A-Za-z\s.'-]+$"
    return bool(re.match(pattern, name.strip()))

#---------------------------------------------
#---- Function: 19 | Validation Username    --
#---------------------------------------------
def is_valid_username(username: str) -> bool:
    """
    Validate username rules:
    - Starts with a letter
    - Can contain letters, numbers, hyphens
    - No consecutive hyphens
    - No hyphen at start or end
    - No spaces allowed
    """
    pattern = r"^[A-Za-z](?!.*--)[A-Za-z0-9-]*[A-Za-z0-9]$"
    return bool(re.match(pattern, username.strip()))


#---------------------------------------------
#---- Function: 20 | Validation Password  ----
#---------------------------------------------
def is_valid_password(password: str) -> bool:
    """
    Validate password with these rules:
    - Minimum 8 characters
    - At least one uppercase letter
    - At least one lowercase letter
    - At least one digit
    - At least one special character
    """
    if not password or len(password) < 8:
        return False
    if not re.search(r"[A-Z]", password):
        return False
    if not re.search(r"[a-z]", password):
        return False
    if not re.search(r"\d", password):
        return False
    if not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):
        return False
    return True
 

@bp.route('/update/profile', methods=['GET', 'POST'])
@login_required
def profile():
    # ----- Role-based access -----
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))


    # ----- Countries and phone parsing -----
    countries = get_country_phone_data()
    phone = current_user.phone
    country_code, phone_without_code = split_phone_number(phone or '')

    # ----- Initialize form -----
    form = UserProfileForm(obj=current_user)
    form.country.choices = [(c['name'], c['name']) for c in countries]
    form.role.choices = [(role.value, role.name.replace("_", " ").title()) for role in UserRole]

    # ----- Timezone setup -----
    eat = pytz.timezone("Africa/Nairobi")
    now = datetime.now(eat)  # always tz-aware now

    # ----- Form submission -----
    if form.validate_on_submit():
        # Basic info
        current_user.fullname = (form.fullname.data or '').strip()
        current_user.username = (form.username.data or '').strip()
        current_user.email = (form.email.data or '').strip()
        current_user.address = (form.address.data or '').strip()
        current_user.bio = (form.bio.data or '').strip()
        current_user.pob = (form.pob.data or '').strip()

        # Validation
        if not is_valid_fullname(current_user.fullname):
            flash('Invalid full name.', 'danger')
            return redirect(url_for('main.profile'))
        if not is_valid_username(current_user.username):
            flash('Invalid username.', 'danger')
            return redirect(url_for('main.profile'))

        # Social links
        for field in ['facebook','twitter','google','whatsapp','instagram','linkedin','skype','github']:
            setattr(current_user, field, (getattr(form, field).data or '').strip())

        # GitHub ID
        github_id_input = (form.github_id.data or '').strip()
        current_user.github_id = github_id_input if github_id_input else None

        # Photo visibility, gender, dob
        current_user.photo_visibility = (form.photo_visibility.data or '').strip()
        current_user.gender = (form.gender.data or '').strip()
        current_user.dob = form.dob.data  # can be None

        # Country & Phone
        country_name = (form.country.data or '').strip()
        current_user.country = country_name
        try:
            country_obj = pycountry.countries.get(name=country_name)
            if not country_obj:
                raise ValueError()
        except Exception:
            flash("Invalid country selected.", "danger")
            return redirect(url_for('main.profile'))

        region_code = country_obj.alpha_2
        raw_phone = (form.phone.data or '').strip()
        if raw_phone:
            try:
                parsed_phone = parse(raw_phone, region_code)
                if is_valid_number(parsed_phone):
                    current_user.phone = format_number(parsed_phone, PhoneNumberFormat.E164)
                else:
                    flash("Invalid phone number format.", "danger")
                    return redirect(url_for('main.profile'))
            except NumberParseException:
                flash("Invalid phone number input.", "danger")
                return redirect(url_for('main.profile'))

        # State & City
        if country_name.lower() == 'somalia':
            current_user.state = (request.form.get('state') or '').strip()
            city = request.form.get('city')
            current_user.city = (request.form.get('city_text') or '').strip() if city == 'Other' or not city else city.strip()
        else:
            current_user.state = (request.form.get('state_text') or '').strip()
            current_user.city = (request.form.get('city_text') or '').strip()

        # Password update
        if form.password.data:
            current_user.password = generate_password_hash(form.password.data)

        # Profile photo upload
        image = request.files.get('profile_photo')
        if image and image.filename:
            ext = os.path.splitext(image.filename)[1]
            slug = current_user.username.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            filename = f"{slug}-{unique_id}{ext}"

            upload_folder = os.path.join('static/backend/uploads/users', str(current_user.id))
            os.makedirs(upload_folder, exist_ok=True)
            image_path = os.path.join(upload_folder, filename)
            image.save(image_path)

            # Delete old photo
            if current_user.photo:
                old_path = os.path.join('static', current_user.photo)
                if os.path.exists(old_path):
                    os.remove(old_path)

            current_user.photo = os.path.relpath(image_path, 'static').replace("\\", "/")

        # Update timestamp
        current_user.updated_at = now

        # Commit safely
        try:
            db.session.add(current_user)
            db.session.add(UserLog(
                user_id=current_user.id,
                action=f"Updated profile at {current_user.updated_at.strftime('%Y-%m-%d %H:%M:%S')}",
                timestamp=current_user.updated_at,
                is_read=False
            ))
            db.session.commit()
            flash("Profile updated successfully!", "success")
            return redirect(url_for('main.dashboard'))
        except IntegrityError:
            db.session.rollback()
            flash("Error updating profile: username, email, or GitHub ID may already exist.", "danger")
            return redirect(url_for('main.profile'))

    # ----- Prepare data for template -----
    somalia_locations = SomaliaLocation.query.all()
    somaliaRegions = group_somalia_regions(somalia_locations)

    # Fetch active sessions
    user_sessions = UserSession.query.filter_by(user_id=current_user.id).order_by(UserSession.last_activity.desc()).all()

    # Ensure sessions are timezone-aware
    for s in user_sessions:
        if s.last_activity and s.last_activity.tzinfo is None:
            s.last_activity = eat.localize(s.last_activity)

    # Safe last_active handling
    last_active = None
    if current_user.last_active:
        last_active = current_user.last_active
        if last_active.tzinfo is None:
            last_active = eat.localize(last_active)
        else:
            last_active = last_active.astimezone(eat)

    # Online status & humanized time
    current_user.is_online = False
    current_user.last_seen_ago = "Never"
    if last_active:
        diff_seconds = (now - last_active).total_seconds()
        current_user.is_online = current_user.auth_status == 'login' and current_user.session_token and diff_seconds < 30
        current_user.last_seen_ago = humanize_time_diff(last_active, now)

    # Pre-format last_active safely
    formatted_last_active = last_active.strftime('%d %b %Y, %I:%M %p') if last_active else "Unknown"

    return render_template(
        '/backend/pages/components/users/profile_view.html',
        form=form,
        user_country_code=country_code or 'XXX',
        phone_without_code=phone_without_code or '',
        somalia_regions=somaliaRegions,
        all_countries=countries,
        user=current_user,
        formatted_last_active=formatted_last_active,
        user_sessions=user_sessions,
        datetime=datetime,
        now=now,  # tz-aware now for template
    )


#-------------------------------------------------------------------------
##------------------------------------------------------------------------
#---- Route: 45 |   Delete My Account with in Profile Section for Users --
#-------------------------------------------------------------------------
@bp.route('/delete-my-account', methods=['POST'])
@login_required
def delete_my_account():
    try:
        if getattr(current_user, 'status', 1) == 0:
            flash("Account-kaagu ma shaqeynayo.", "danger")
            return redirect(url_for("main.dashboard"))

        allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
        if current_user.role.value not in allowed_roles:
            flash("Ma haysatid ogolaansho.", "danger")
            return redirect(url_for("main.dashboard"))

        user = current_user

        # ==============================
        # 🔥 SUPERADMIN
        # ==============================
        if user.role.value == 'superadmin':
            # Optional: delete EVERYTHING in system 😅
            User.query.delete()
            Branch.query.delete()
            School.query.delete()

        # ==============================
        # 🔥 SCHOOL ADMIN
        # ==============================
        elif user.role.value == 'school_admin':
            school_id = user.school_id

            # Delete all school data (reuse your logic)
            branches = Branch.query.filter_by(school_id=school_id).all()

            for branch in branches:
                Section.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Class.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                ClassLevel.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Student.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Teacher.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Subject.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Parent.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                SchoolSiteSettings.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                User.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)

            # Delete school-level data
            Student.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Teacher.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Subject.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Parent.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Class.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            ClassLevel.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Section.query.filter_by(school_id=school_id).delete(synchronize_session=False)

            Branch.query.filter_by(school_id=school_id).delete(synchronize_session=False)

            school = School.query.get(school_id)
            if school:
                db.session.delete(school)

        # ==============================
        # 🔥 BRANCH ADMIN
        # ==============================
        elif user.role.value == 'branch_admin':
            branch_id = user.branch_id

            Section.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Class.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            ClassLevel.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Student.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Teacher.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Subject.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Parent.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            SchoolSiteSettings.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            User.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

            branch = Branch.query.get(branch_id)
            if branch:
                db.session.delete(branch)

        # ==============================
        # 🔥 DELETE CURRENT USER
        # ==============================
        db.session.delete(user)

        db.session.commit()

        flash("🔥 Your account and ALL related data deleted successfully.", "success")
        return redirect(url_for('main.index'))

    except Exception as e:
        db.session.rollback()
        flash(f"Error: {str(e)}", "danger")
        return redirect(url_for('main.index'))


#---------------------------------------------------------------------
##--------------------------------------------------------------------
#---- Route: 46 |  Update Profile Change Password  Section for Users -
#---------------------------------------------------------------------
@bp.route('/change/password', methods=['GET', 'POST'])
@login_required
def settings():
    # ----- Role-based access control -----
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ----- Form -----
    form = ChangePasswordForm()

    if form.validate_on_submit():
        current_pw = form.current_password.data
        new_pw = form.new_password.data

        # Validate current password
        if not check_password_hash(current_user.password, current_pw):
            flash('❌ Current password is incorrect.', 'danger')
            return redirect(url_for('main.settings'))

        # Update password
        africa_time = datetime.now(pytz.timezone("Africa/Nairobi"))
        current_user.password = generate_password_hash(new_pw)
        current_user.updated_at = africa_time

        # Log action
        db.session.add(UserLog(
            user_id=current_user.id,
            action=f"{current_user.username} changed password at {africa_time.strftime('%Y-%m-%d %H:%M:%S')}",
            timestamp=africa_time,
            is_read=False
        ))

        db.session.add(current_user)
        db.session.commit()

        flash('✅ Password changed successfully!', 'success')
        return redirect(url_for('main.dashboard'))

    # ----- Fetch active sessions -----
    eat = pytz.timezone("Africa/Nairobi")
    user_sessions = UserSession.query.filter_by(user_id=current_user.id).order_by(UserSession.last_activity.desc()).all()

    # Make sessions timezone-aware
    for s in user_sessions:
        if s.last_activity and s.last_activity.tzinfo is None:
            s.last_activity = eat.localize(s.last_activity)

    # Current user last active & online status
    now = datetime.now(eat)
    last_active = current_user.last_active
    if last_active:
        if last_active.tzinfo is None:
            last_active = eat.localize(last_active)
        else:
            last_active = last_active.astimezone(eat)

    current_user.is_online = False
    current_user.last_seen_ago = "Never"
    if last_active:
        diff_seconds = (now - last_active).total_seconds()
        current_user.is_online = current_user.auth_status == 'login' and diff_seconds < 30
        current_user.last_seen_ago = humanize_time_diff(last_active, now)

    formatted_last_active = last_active.strftime('%d %b %Y, %I:%M %p') if last_active else "Unknown"

    return render_template(
        '/backend/pages/components/users/settings.html',
        form=form,
        user=current_user,
        formatted_last_active=formatted_last_active,
        user_sessions=user_sessions,
        now=now,
        datetime=datetime
    )



@bp.route('/auth-settings', methods=['GET', 'POST'])
@login_required
def auth_settings():
    # ----- Role-based access control -----
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ----- Forms -----
    password_form = ChangePasswordForm()
    twofactor_form = TwoFactorForm()

    # ----- Handle Password Change -----
    if password_form.validate_on_submit() and request.form.get('new_password'):
        current_pw = password_form.current_password.data
        new_pw = password_form.new_password.data

        if not check_password_hash(current_user.password, current_pw):
            flash('❌ Current password is incorrect.', 'danger')
            return redirect(url_for('main.auth_settings'))

        africa_time = datetime.now(pytz.timezone("Africa/Nairobi"))
        current_user.password = generate_password_hash(new_pw)
        current_user.updated_at = africa_time

        db.session.add(UserLog(
            user_id=current_user.id,
            action=f"{current_user.username} changed password at {africa_time.strftime('%Y-%m-%d %H:%M:%S')}",
            timestamp=africa_time,
            is_read=False
        ))
        db.session.commit()
        flash('✅ Password changed successfully!', 'success')
        return redirect(url_for('main.auth_settings'))

    # ----- Handle 2FA Setup -----
    secret = current_user.two_factor_code
    if not secret:
        secret = pyotp.random_base32()
        current_user.two_factor_code = secret
        db.session.commit()

    # Generate QR code
    otp_uri = pyotp.TOTP(secret).provisioning_uri(
        name=current_user.email,
        issuer_name="MyAppName"
    )
    qr = qrcode.make(otp_uri)
    buf = io.BytesIO()
    qr.save(buf, format='PNG')
    qr_b64 = base64.b64encode(buf.getvalue()).decode('ascii')

    # Handle 2FA form POST
    if twofactor_form.validate_on_submit() and request.form.get('two_factor_code'):
        otp_input = twofactor_form.two_factor_code.data
        totp = pyotp.TOTP(current_user.two_factor_code)
        if totp.verify(otp_input):
            current_user.two_factor_enabled = twofactor_form.two_factor_enabled.data
            if current_user.two_factor_enabled:
                current_user.two_factor_expires_at = now_eat + timedelta(minutes=5)
            db.session.commit()
            flash("✅ Two-factor authentication settings updated successfully!", "success")
        else:
            flash("❌ Invalid two-factor authentication code.", "danger")
        return redirect(url_for('main.auth_settings'))

    # ----- Fetch active sessions -----
    eat = pytz.timezone("Africa/Nairobi")
    user_sessions = UserSession.query.filter_by(user_id=current_user.id)\
        .order_by(UserSession.last_activity.desc()).all()
    for s in user_sessions:
        if s.last_activity and s.last_activity.tzinfo is None:
            s.last_activity = eat.localize(s.last_activity)

    # ----- Last active & online status -----
    now = datetime.now(eat)
    last_active = current_user.last_active
    if last_active:
        last_active = last_active if last_active.tzinfo else eat.localize(last_active)
        last_active = last_active.astimezone(eat)

    current_user.is_online = False
    current_user.last_seen_ago = "Never"
    if last_active:
        diff_seconds = (now - last_active).total_seconds()
        current_user.is_online = current_user.auth_status == 'login' and diff_seconds < 30
        current_user.last_seen_ago = humanize_time_diff(last_active, now)

    formatted_last_active = last_active.strftime('%d %b %Y, %I:%M %p') if last_active else "Unknown"

    return render_template(
        'backend/pages/components/users/auth-settings.html',
        user=current_user,
        password_form=password_form,
        twofactor_form=twofactor_form,
        qr_b64=qr_b64,
        secret=secret,
        user_sessions=user_sessions,
        formatted_last_active=formatted_last_active,
        now=now,
        datetime=datetime
    )


#----------------------------------------------------------
#---- Route: 140 |In Profile View Remove User Photo -------
#----------------------------------------------------------
@bp.route('/remove_photo', methods=['POST'])
@login_required
def remove_photo():
    # Role-based access control
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin","parent","teacher","student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

 
    if not current_user.photo:
        flash('No profile photo to remove.', 'warning')
        return redirect(url_for('main.profile'))  # adjust to your profile route

    # Construct full path to the photo file
    photo_path = os.path.join(current_app.root_path, 'static', 'users', current_user.photo)

    # Try deleting the photo file
    if os.path.exists(photo_path):
        try:
            os.remove(photo_path)
        except Exception as e:
            flash(f'Failed to delete photo file: {e}', 'danger')
            return redirect(url_for('main.profile'))

    # Remove photo filename from DB record
    current_user.photo = None
    try:
        db.session.commit()
        flash('Profile photo removed successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Failed to update database: {e}', 'danger')

    return redirect(url_for('main.profile'))



#-------------- School All

@bp.route('/all/schools')
@login_required
def all_schools():
    # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Fetch all schools ordered by creation date
    schools = School.query.order_by(School.created_at.desc()).all()

    return render_template(
        'backend/pages/components/schools/all_schools.html',
        schools=schools,
        user=current_user
    )


# ------------------- Add School -------------------
@bp.route('/school/add', methods=['GET', 'POST'])
@login_required
def add_school():
   # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = SchoolForm()
    if form.validate_on_submit():
        school = School(
            name=form.name.data,
            title=form.title.data,
            email=form.email.data,
            phone=form.phone.data,
            status=form.status.data
        )
        db.session.add(school)
        db.session.commit()
        flash("School added successfully.", "success")
        return redirect(url_for('main.all_schools'))

    return render_template('backend/pages/components/schools/add_school.html',
                            user=current_user,
                            form=form, action="Add")


# ------------------- Edit School -------------------
@bp.route('/school/edit/<int:school_id>', methods=['GET', 'POST'])
@login_required
def edit_school(school_id):
    # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    school = School.query.get_or_404(school_id)
    form = SchoolForm(obj=school)
    if form.validate_on_submit():
        school.name = form.name.data
        school.title = form.title.data
        school.email = form.email.data
        school.phone = form.phone.data
        school.status = form.status.data
        db.session.commit()
        flash("School updated successfully.", "success")
        return redirect(url_for('main.all_schools'))

    return render_template('backend/pages/components/schools/edit_school.html',
                            user=current_user,
                              form=form, action="Edit")



# ------------------- Delete School -------------------
@bp.route('/school/delete/<int:school_id>', methods=['POST'])
@login_required
def delete_school(school_id):
   # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    school = School.query.get_or_404(school_id)

    try:
        # 🔥 Delete all branch-related data
        branches = Branch.query.filter_by(school_id=school_id).all()

        for branch in branches:
            Section.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            Class.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            ClassLevel.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            Student.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            Teacher.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            Subject.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            Parent.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            SchoolSiteSettings.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
            User.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)

        # 🔥 Delete school-level data (direct FK)
        Student.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        Teacher.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        Subject.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        Parent.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        User.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        Class.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        ClassLevel.query.filter_by(school_id=school_id).delete(synchronize_session=False)
        Section.query.filter_by(school_id=school_id).delete(synchronize_session=False)

        # 🔥 Delete branches
        Branch.query.filter_by(school_id=school_id).delete(synchronize_session=False)

        # 🔥 Finally delete school
        db.session.delete(school)

        db.session.commit()

        flash("🔥 School and EVERYTHING related deleted successfully.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error: {str(e)}", "danger")

    return redirect(url_for('main.all_schools'))



# ------------------- All Branches -------------------
@bp.route('/all/branches')
@login_required
def all_branches():
  # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    branches = Branch.query.order_by(Branch.created_at.desc()).all()
    return render_template('backend/pages/components/branches/all_branches.html',
                            user=current_user,
                              branches=branches)

# ------------------- Add Branch -------------------
@bp.route('/branch/add', methods=['GET','POST'])
@login_required
def add_branch():
  # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = BranchForm()
    form.school_id.choices = [(s.id, s.name) for s in School.query.order_by(School.name).all()]

    if form.validate_on_submit():
        branch = Branch(
            school_id=form.school_id.data,
            name=form.name.data,
            title=form.title.data,
            address=form.address.data,
            phone=form.phone.data,
            status=form.status.data
        )
        db.session.add(branch)
        db.session.commit()
        flash("Branch added successfully.", "success")
        return redirect(url_for('main.all_branches'))

    return render_template('backend/pages/components/branches/add_branch.html',
                            user=current_user,
                              form=form, action="Add")

# ------------------- Edit Branch -------------------
@bp.route('/branch/edit/<int:branch_id>', methods=['GET','POST'])
@login_required
def edit_branch(branch_id):
  # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    branch = Branch.query.get_or_404(branch_id)
    form = BranchForm(obj=branch)
    form.school_id.choices = [(s.id, s.name) for s in School.query.order_by(School.name).all()]

    if form.validate_on_submit():
        branch.school_id = form.school_id.data
        branch.name = form.name.data
        branch.title = form.title.data
        branch.address = form.address.data
        branch.phone = form.phone.data
        branch.status = form.status.data
        db.session.commit()
        flash("Branch updated successfully.", "success")
        return redirect(url_for('main.all_branches'))

    return render_template('backend/pages/components/branches/edit_branch.html',\
                            user=current_user,
                            form=form, action="Edit")

# ------------------- Delete Branch -------------------
@bp.route('/branch/delete/<int:branch_id>', methods=['POST'])
@login_required
def delete_branch(branch_id):
   # Only superadmins can access this page
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["superadmin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    branch = Branch.query.get_or_404(branch_id)

    try:
        # 🧹 Delete all related data first

        # Sections
        Section.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Classes
        Class.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Class Levels
        ClassLevel.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Students
        Student.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Teachers
        Teacher.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Subjects
        Subject.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Parents
        Parent.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Settings
        SchoolSiteSettings.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # Users
        User.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)

        # 🔥 Finally delete branch
        db.session.delete(branch)

        db.session.commit()

        flash("Branch and all related data deleted successfully.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting branch: {str(e)}", "danger")

    return redirect(url_for('main.all_branches'))



#---------------------------------------------------
#---- Route: 5 | All Users - Backend Template ------
#---------------------------------------------------
@bp.route('/all/users')
@login_required
def all_users():
    # Role-based access control
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- SUPER ADMIN -------------------
    if current_user.role == UserRole.superadmin:
        # Show only school_admin and branch_admin
        users = User.query.filter(
            User.role.in_([UserRole.school_admin, UserRole.branch_admin])
        ).order_by(User.created_at.desc()).all()

    # ------------------- SCHOOL ADMIN -------------------
    elif current_user.role == UserRole.school_admin:
        users = User.query.filter(
            User.school_id == current_user.school_id,
            User.id != current_user.id,  # Do not show self
            User.role.in_([UserRole.teacher, UserRole.student, UserRole.parent])
        ).order_by(User.created_at.desc()).all()

    # ------------------- BRANCH ADMIN -------------------
    elif current_user.role == UserRole.branch_admin:
        users = User.query.filter(
            User.school_id == current_user.school_id,
            User.branch_id == current_user.branch_id,
            User.id != current_user.id,  # Do not show self
            User.role.in_([UserRole.teacher, UserRole.student, UserRole.parent])
        ).order_by(User.created_at.desc()).all()

    # ------------------- NO ACCESS -------------------
    else:
        flash("You are not authorized to view this page.", "danger")
        return redirect(url_for('main.dashboard'))

    return render_template(
        'backend/pages/components/users/all_users.html',
        users=users,
        user=current_user
    )


#----------------------------------------------------
#---- Route: 6 | View User - Backend Template -------
#----------------------------------------------------
@bp.route('/single/user/view/<int:user_id>')
@login_required
def user_view(user_id):

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    user = User.query.get_or_404(user_id)

    return render_template(
        'backend/pages/components/users/user_view.html',
        user=user
    )



#--------------------------------------------------
#---- Route: 7 | Add User - Backend Template ------
#--------------------------------------------------
#---------------------------
# Route: Add User
#---------------------------
@bp.route('/add/user', methods=['GET', 'POST'])
@login_required
def add_user():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))


    # 🔒 Superadmin limit
    existing_superadmin = User.query.filter_by(role=UserRole.superadmin).first()
    if existing_superadmin and current_user.role != UserRole.superadmin:
        flash("There is already a superadmin.", "danger")
        return redirect(url_for('main.add_user'))

    # 🔒 inactive user
    if not current_user.status:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.dashboard'))

    # 🎯 roles
    if current_user.role == UserRole.school_admin:
        roles = Role.query.filter(Role.name != 'superadmin').all()
    elif current_user.role == UserRole.superadmin:
        roles = Role.query.all()
    else:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    countries = get_country_phone_data()

    form = UserForm()

    if request.method == 'POST':

        # 🧾 GET DATA (IMPORTANT → use request.form)
        username = request.form.get('username')
        fullname = request.form.get('fullname')
        email = request.form.get('email')
        role = request.form.get('role')
        phone = request.form.get('phone')
        country = request.form.get('country')
        state = request.form.get('state') or request.form.get('state_text')
        city = request.form.get('city') or request.form.get('city_text')
        status = request.form.get('status')
        password = request.form.get('new_password')

        school_id = request.form.get('school_id')
        branch_id = request.form.get('branch_id')

        # ❌ validation
        if not username or not email or not password:
            flash("Required fields missing", "danger")
            return redirect(url_for('main.add_user'))

        # 🔁 unique check
        if User.query.filter_by(username=username).first():
            flash("Username exists", "danger")
            return redirect(url_for('main.add_user'))

        if User.query.filter_by(email=email).first():
            flash("Email exists", "danger")
            return redirect(url_for('main.add_user'))

        # 🔐 hash password
        hashed_password = generate_password_hash(password)

        # 📞 phone validation (optional skip if error)
        phone_number = phone

        try:
            country_data = next((c for c in countries if c['name'] == country), None)
            if country_data:
                parsed = parse(phone, country_data['alpha2'])
                if is_valid_number(parsed):
                    phone_number = format_number(parsed, PhoneNumberFormat.E164)
        except:
            pass

        # 🕒 time
        current_time = datetime.now(pytz.timezone("Africa/Nairobi"))

        # ✅ CREATE USER
        new_user = User(
            username=username,
            fullname=fullname,
            email=email,
            phone=phone_number,
            country=country,
            state=state,
            city=city,
            role=UserRole(role),
            status=True if status == '1' else False,
            password=hashed_password,
            school_id=school_id,
            branch_id=branch_id if branch_id else None,
            created_at=current_time,
            updated_at=current_time
        )

        db.session.add(new_user)
        db.session.commit()

        # 📸 IMAGE
     
        # ✅ 10. Handle image upload
        image = request.files.get('photo')
        if not image or not image.filename:
            flash("Image is required.", "danger")
            return redirect(request.url)

        ext = os.path.splitext(image.filename)[1]
        slug = username.lower().replace(' ', '-').replace('_', '-')
        unique_id = uuid.uuid4().hex[:8]
        safe_filename = f"{slug}-{unique_id}{ext}"
        user_subfolder = os.path.join('static/backend/uploads/users', str(new_user.id))
        os.makedirs(user_subfolder, exist_ok=True)
        image_path = os.path.join(user_subfolder, safe_filename)
        image.save(image_path)
        saved_filename = os.path.relpath(image_path, 'static').replace("\\", "/")

        new_user.photo = saved_filename
        db.session.commit()

        flash("User created successfully", "success")
        return redirect(url_for('main.all_users'))

    # 🔄 GET
    site_data = SettingsData.query.first()
    somalia_locations = SomaliaLocation.query.all()
    somaliaRegions = group_somalia_regions(somalia_locations)

    return render_template(
        'backend/pages/components/users/add_users.html',
        form=form,
        somalia_regions=somaliaRegions,
        all_countries=countries,
        roles=roles,
        schools=School.query.all(),
        branches=Branch.query.all(),
        user=current_user,
        site_data=site_data
    )



#---------------------------------------------------
#---- Route: 8 | Edit User - Backend Template ------
#---------------------------------------------------
@bp.route("/edit/user/<int:user_id>", methods=["GET", "POST"])
@login_required
def edit_user(user_id):

    # ------------------- CHECK ACTIVE -------------------
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- ROLE ACCESS -------------------
    if current_user.role not in [UserRole.superadmin, UserRole.school_admin, UserRole.branch_admin]:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    user = User.query.get_or_404(user_id)

    # ------------------- PREVENT SELF-EDIT -------------------
    if user.id == current_user.id:
        flash("You cannot edit your own account here.", "warning")
        return redirect(url_for('main.all_users'))

    # ------------------- ROLE-BASED ACCESS -------------------
    if current_user.role == UserRole.school_admin:
        # School admin can only edit teachers, students, parents in their school
        if user.school_id != current_user.school_id or user.role in [UserRole.school_admin, UserRole.branch_admin, UserRole.superadmin]:
            flash("Unauthorized to edit this user", "danger")
            return redirect(url_for('main.all_users'))

    elif current_user.role == UserRole.branch_admin:
        # Branch admin can only edit users in their branch
        if user.school_id != current_user.school_id or user.branch_id != current_user.branch_id or user.role in [UserRole.school_admin, UserRole.superadmin]:
            flash("Unauthorized to edit this user", "danger")
            return redirect(url_for('main.all_users'))

    # ------------------- FORM -------------------
    form = UserForm(obj=user)
    countries = get_country_phone_data()
    country_code, phone_without_code = split_phone_number(user.phone or '')

    # Populate select fields
    form.role.choices = [(r.name, r.name.capitalize()) for r in Role.query.all()]
    form.school_id.choices = [(0, "Select School")] + [(s.id, s.name) for s in School.query.all()]
    form.branch_id.choices = [(0, "Select Branch")] + [(b.id, b.name) for b in Branch.query.all()]
    form.country.choices = [(c['name'], c['name']) for c in countries]

    # ------------------- POST -------------------
    if request.method == "POST":

        username = request.form.get("username", "").strip()
        fullname = request.form.get("fullname", "").strip()
        email = request.form.get("email", "").strip()
        role = request.form.get("role")
        status = request.form.get("status")
        school_id = request.form.get("school_id")
        branch_id = request.form.get("branch_id")
        country_name = request.form.get("country", "").strip()
        password = request.form.get("new_password")
        confirm_password = request.form.get("confirm_password")
        state = request.form.get("state") or request.form.get("state_text")
        city = request.form.get("city") or request.form.get("city_text")
        phone = request.form.get("phone", "")

        # ------------------- VALIDATION -------------------
        if not username or not email or not role:
            flash("Required fields missing", "danger")
            return redirect(request.url)

        if User.query.filter(User.username == username, User.id != user.id).first():
            flash("Username already exists", "danger")
            return redirect(request.url)

        if User.query.filter(User.email == email, User.id != user.id).first():
            flash("Email already exists", "danger")
            return redirect(request.url)

        if password:
            if not is_valid_password(password):
                flash("Password must be at least 8 characters with uppercase, lowercase, number, and special character.", "warning")
                return redirect(request.url)
            if password != confirm_password:
                flash("Passwords do not match!", "danger")
                return redirect(request.url)
            user.password = generate_password_hash(password)

        # ------------------- UPDATE BASIC INFO -------------------
        user.username = username
        user.fullname = fullname
        user.email = email
        user.role = UserRole[role]
        user.status = True if status == '1' else False
        user.updated_at = datetime.now(pytz.timezone("Africa/Nairobi"))

        # ------------------- SCHOOL / BRANCH SECURITY -------------------
        if current_user.role == UserRole.superadmin:
            user.school_id = int(school_id) if school_id and school_id != "0" else None
            user.branch_id = int(branch_id) if branch_id and branch_id != "0" else None
        elif current_user.role == UserRole.school_admin:
            user.school_id = current_user.school_id
            user.branch_id = int(branch_id) if branch_id and branch_id != "0" else None
        elif current_user.role == UserRole.branch_admin:
            user.school_id = current_user.school_id
            user.branch_id = current_user.branch_id

        # ------------------- COUNTRY + PHONE -------------------
        user.country = country_name
        try:
            country_obj = pycountry.countries.get(name=country_name)
            if not country_obj:
                raise ValueError()
        except:
            flash("Invalid country selected.", "danger")
            return redirect(request.url)

        region_code = country_obj.alpha_2
        try:
            parsed_phone = parse(phone, region_code)
            if is_valid_number(parsed_phone):
                user.phone = format_number(parsed_phone, PhoneNumberFormat.E164)
            else:
                flash("Invalid phone number.", "danger")
                return redirect(request.url)
        except:
            flash("Invalid phone input.", "danger")
            return redirect(request.url)

        # ------------------- STATE / CITY -------------------
        if country_name.lower() == 'somalia':
            user.state = request.form.get('state', '').strip()
            city = request.form.get('city')
            user.city = request.form.get('city_text', '').strip() if city == 'Other' or not city else city
        else:
            user.state = request.form.get('state_text', '').strip()
            user.city = request.form.get('city_text', '').strip()

        # ------------------- IMAGE -------------------
        image = request.files.get('photo')
        if image and image.filename:
            ext = os.path.splitext(image.filename)[1]
            slug = user.username.lower().replace(' ', '-')
            filename = f"{slug}-{uuid.uuid4().hex[:8]}{ext}"
            folder = os.path.join('static/backend/uploads/users', str(user.id))
            os.makedirs(folder, exist_ok=True)
            path = os.path.join(folder, filename)
            image.save(path)
            # Delete old
            if user.photo:
                old_path = os.path.join('static', user.photo)
                if os.path.exists(old_path):
                    os.remove(old_path)
            user.photo = os.path.relpath(path, 'static').replace("\\", "/")

        # ------------------- SAVE -------------------
        try:
            db.session.commit()
            flash("User updated successfully", "success")
            return redirect(url_for("main.all_users"))
        except Exception as e:
            db.session.rollback()
            print("ERROR:", e)
            flash("Error updating user", "danger")

    # ------------------- GET DATA -------------------
    site_data = SettingsData.query.first()
    somaliaRegions = group_somalia_regions(SomaliaLocation.query.all())

    return render_template(
        "backend/pages/components/users/edit_user.html",
        form=form,
        user=user,
        roles=Role.query.all(),
        schools=School.query.all(),
        branches=Branch.query.all(),
        user_country_code=country_code or '',
        phone_without_code=phone_without_code or '',
        all_countries=countries,
        somalia_regions=somaliaRegions,
        site_data=site_data,
        UserRole=UserRole 
    )


#---------------------------------------------------
#---- Route: 9 | Delete User - Backend Template ------
#---------------------------------------------------
@bp.route("/delete/user/<int:user_id>", methods=["POST"])
@login_required
def delete_user(user_id):

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","superadmin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    user = User.query.get_or_404(user_id)

    # ❌ Prevent deleting self
    if user.id == current_user.id:
        flash("You cannot delete your own account!", "danger")
        return redirect(url_for("main.all_users"))

    try:
        # ==============================
        # 🔥 SCHOOL ADMIN → delete school
        # ==============================
        if user.role.value == 'school_admin' and user.school_id:
            school_id = user.school_id

            branches = Branch.query.filter_by(school_id=school_id).all()

            for branch in branches:
                Section.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Class.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                ClassLevel.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Student.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Teacher.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Subject.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                Parent.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                SchoolSiteSettings.query.filter_by(branch_id=branch.id).delete(synchronize_session=False)
                User.query.filter(
                    User.branch_id == branch.id,
                    User.id != user.id
                ).delete(synchronize_session=False)

            # school-level delete
            Student.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Teacher.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Subject.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Parent.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Class.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            ClassLevel.query.filter_by(school_id=school_id).delete(synchronize_session=False)
            Section.query.filter_by(school_id=school_id).delete(synchronize_session=False)

            Branch.query.filter_by(school_id=school_id).delete(synchronize_session=False)

            school = School.query.get(school_id)
            if school:
                db.session.delete(school)

        # ==============================
        # 🔥 BRANCH ADMIN → delete branch
        # ==============================
        elif user.role.value == 'branch_admin' and user.branch_id:
            branch_id = user.branch_id

            Section.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Class.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            ClassLevel.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Student.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Teacher.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Subject.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            Parent.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            SchoolSiteSettings.query.filter_by(branch_id=branch_id).delete(synchronize_session=False)
            User.query.filter(
                User.branch_id == branch_id,
                User.id != user.id
            ).delete(synchronize_session=False)

            branch = Branch.query.get(branch_id)
            if branch:
                db.session.delete(branch)

        # ==============================
        # 🔥 NORMAL USER → just delete
        # ==============================
        db.session.delete(user)

        db.session.commit()

        flash(f"🔥 User {user.fullname} and related data deleted successfully.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error: {str(e)}", "danger")

    return redirect(url_for("main.all_users"))

# ==================== ALL CLASS LEVELS ====================
@bp.route("/all/class/levels")
@login_required
def all_class_levels():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # -------------------------
    # Role-based filtering
    # -------------------------
    if current_user.role.value == 'school_admin':
        # School admin → only branch-less class levels
        class_levels = ClassLevel.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(ClassLevel.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Branch admin → only class levels of this branch
        class_levels = ClassLevel.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(ClassLevel.created_at.desc()).all()

    return render_template(
        "backend/pages/components/class_levels/all_class_levels.html",
        class_levels=class_levels,
        user=current_user
    )


# ==================== ADD CLASS LEVEL ====================
@bp.route("/add/class/level", methods=["GET", "POST"])
@login_required
def add_class_level():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = ClassLevelForm()

    # Automatically set default values
    if current_user.role.value == 'school_admin':
        form.school_id.data = current_user.school_id
        if form.branch_id.choices:
            form.branch_id.data = form.branch_id.choices[0][0]
    elif current_user.role.value == 'branch_admin':
        form.school_id.data = current_user.school_id
        form.branch_id.data = current_user.branch_id

    if form.validate_on_submit():
        # Convert empty branch_id to None
        branch_id = form.branch_id.data
        if not branch_id:
            branch_id = None
        else:
            # Verify branch exists
            from app.modal import Branch
            branch = Branch.query.get(branch_id)
            if not branch:
                flash("Selected branch does not exist.", "danger")
                return redirect(url_for('main.add_class_level'))

        # Check if class level already exists
        existing = ClassLevel.query.filter_by(
            name=form.name.data.strip(),
            school_id=form.school_id.data,
            branch_id=branch_id
        ).first()
        if existing:
            flash("This class level already exists!", "warning")
            return redirect(url_for('main.add_class_level'))

        # Create new class level
        new_level = ClassLevel(
            name=form.name.data.strip(),
            price=form.price.data,
            school_id=form.school_id.data,
            branch_id=branch_id
        )
        db.session.add(new_level)

        try:
            db.session.commit()
            flash(f"Class Level '{form.name.data}' added successfully!", "success")
        except Exception as e:
            db.session.rollback()  # clear session after failure
            flash(f"Error adding class level: {str(e)}", "danger")
            return redirect(url_for('main.add_class_level'))

        return redirect(url_for('main.all_class_levels'))

    return render_template(
        "backend/pages/components/class_levels/add_class_level.html",
        form=form,
        user=current_user
    )

# ==================== EDIT CLASS LEVEL ====================
@bp.route("/edit/class/level/<int:id>", methods=["GET", "POST"])
@login_required
def edit_class_level(id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    level = ClassLevel.query.get_or_404(id)

    # Restrict editing to own school/branch
    if current_user.role.value == 'school_admin' and level.school_id != current_user.school_id:
        flash("You cannot edit class levels from another school.", "danger")
        return redirect(url_for('main.all_class_levels'))
    if current_user.role.value == 'branch_admin' and level.branch_id != current_user.branch_id:
        flash("You cannot edit class levels from another branch.", "danger")
        return redirect(url_for('main.all_class_levels'))

    form = ClassLevelForm(obj=level)

    # Auto-fill school/branch
    form.school_id.data = level.school_id
    form.branch_id.data = level.branch_id

    if form.validate_on_submit():
        # Prevent duplicates after edit
        existing = ClassLevel.query.filter(
            ClassLevel.id != id,
            ClassLevel.name == form.name.data.strip(),
            ClassLevel.school_id == form.school_id.data,
            ClassLevel.branch_id == form.branch_id.data
        ).first()
        if existing:
            flash("Another class level with the same name exists!", "warning")
            return redirect(url_for('main.edit_class_level', id=id))

        level.name = form.name.data.strip()
        level.price = form.price.data
        level.school_id = form.school_id.data
        level.branch_id = form.branch_id.data

        db.session.commit()
        flash(f"Class Level '{level.name}' updated successfully!", "success")
        return redirect(url_for('main.all_class_levels'))

    return render_template(
        "backend/pages/components/class_levels/edit_class_level.html",
        form=form,
        user=current_user
    )



# ==================== DELETE CLASS LEVEL ====================
@bp.route("/delete/class/levels/<int:id>", methods=["POST"])
@login_required
def delete_class_level(id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    level = ClassLevel.query.get_or_404(id)

    # Security checks
    if current_user.role.value == 'school_admin' and level.school_id != current_user.school_id:
        flash("You cannot delete class levels from another school.", "danger")
        return redirect(url_for('main.all_class_levels'))
    if current_user.role.value == 'branch_admin' and level.branch_id != current_user.branch_id:
        flash("You cannot delete class levels from another branch.", "danger")
        return redirect(url_for('main.all_class_levels'))

    try:
        # ---------------- Get all classes under this level ----------------
        classes = Class.query.filter_by(level_id=level.id).all()
        total_students = 0
        total_sections = 0
        total_classes = len(classes)

        for cls in classes:
            # Delete students in the class
            students_in_class = Student.query.filter_by(class_id=cls.id).all()
            for student in students_in_class:
                db.session.delete(student)
            total_students += len(students_in_class)

            # Delete sections in the class
            sections_in_class = Section.query.filter_by(class_id=cls.id).all()
            for section in sections_in_class:
                # Delete students in section
                students_in_section = Student.query.filter_by(section_id=section.id).all()
                for student in students_in_section:
                    db.session.delete(student)
                total_students += len(students_in_section)

                db.session.delete(section)
                total_sections += 1

            # Delete the class itself
            db.session.delete(cls)

        # Delete the class level
        db.session.delete(level)
        db.session.commit()

        flash(f"Class Level '{level.name}' deleted successfully! "
              f"Deleted {total_classes} class(es), {total_sections} section(s), and {total_students} student(s).", "success")

    except Exception as e:
        db.session.rollback()
        flash("Error deleting class level or its related classes, sections, and students.", "danger")
        print("DB ERROR:", e)

    return redirect(url_for('main.all_class_levels'))



#----------------------
#---------- Classes 
#----------------------
@bp.route("/classes")
@login_required
def all_classes():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # -------------------------
    # Role-based filtering
    # -------------------------
    if current_user.role.value == 'school_admin':
        # School admin → only branch-less classes
        classes = Class.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Class.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Branch admin → only classes of this branch
        classes = Class.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Class.created_at.desc()).all()

    return render_template(
        "backend/pages/components/classes/all_classes.html",
        classes=classes,
        user=current_user
    )



@bp.route('/export-classes-full')
@login_required
def export_classes():
    import io, csv
    from flask import Response

    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        classes = Class.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None  # only classes without branch
        ).all()
    elif current_user.role.value == 'branch_admin':
        classes = Class.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).all()
    else:
        return Response("Unauthorized", status=403)

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "ID", "Class Name", "Level", "Shift", "Capacity", "Status",
        "School", "Branch", "Created At", "Updated At"
    ])

    # ---------------- CSV Data ----------------
    for c in classes:
        writer.writerow([
            c.id,
            c.name,
            c.level.name if c.level else "",
            c.shift,
            c.capacity or 0,
            c.status,
            c.school.name if c.school else "",
            c.branch.name if c.branch else "",
            c.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            c.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        ])

    output.seek(0)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=classes_full.csv"}
    )


@bp.route('/export-classes-xlsx')
@login_required
def export_classes_xlsx():
    import pandas as pd
    import io
    from flask import send_file, Response

    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    filters = {'school_id': current_user.school_id}
    if current_user.role.value == 'branch_admin':
        filters['branch_id'] = current_user.branch_id
    else:
        filters['branch_id'] = None

    classes = Class.query.filter_by(**filters).order_by(Class.name).all()
    levels = ClassLevel.query.filter_by(**filters).order_by(ClassLevel.name).all()

    # ---------------- Build Data ----------------
    classes_data = []
    for c in classes:
        classes_data.append({
            "ID": c.id,
            "Name": c.name,
            "Level": c.level.name if c.level else "",
            "Shift": c.shift,
            "Capacity": c.capacity or 0,
            "Status": c.status or "",
            "School": c.school.name if c.school else "",
            "Branch": c.branch.name if c.branch else ""
        })

    levels_data = []
    for l in levels:
        levels_data.append({
            "ID": l.id,
            "Name": l.name,
            "Price": float(l.price or 0)
        })

    # ---------------- Excel Multi-Sheet ----------------
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        pd.DataFrame(classes_data).to_excel(writer, sheet_name='Classes', index=False)
        pd.DataFrame(levels_data).to_excel(writer, sheet_name='Levels', index=False)

    output.seek(0)
    return send_file(
        output,
        download_name="classes_data.xlsx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )


@bp.route('/import-classes-full', methods=['POST'])
@login_required
def import_classes():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.all_classes'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    file = request.files.get('file')
    if not file:
        flash("No file selected", "danger")
        return redirect(url_for('main.all_classes'))

    import io, csv

    reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='latin-1'))

    for row in reader:
        # ---------------- School & Branch ----------------
        school_id = current_user.school_id
        branch_id = None

        if current_user.role.value == 'branch_admin':
            branch_id = current_user.branch_id
        else:
            # Only assign branch if provided
            branch_name = row.get('Branch')
            if branch_name:
                branch_obj = Branch.query.filter_by(
                    name=branch_name,
                    school_id=current_user.school_id
                ).first()
                branch_id = branch_obj.id if branch_obj else None

        # ---------------- Level Lookup ----------------
        level_name = row.get('Level')
        level_obj = ClassLevel.query.filter_by(
            name=level_name,
            school_id=current_user.school_id
        ).first()
        if not level_obj:
            flash(f"Level '{level_name}' not found for class '{row.get('Class Name')}'. Skipping.", "warning")
            continue

        # ---------------- Duplicate Check ----------------
        existing_class = Class.query.filter_by(
            name=row.get('Class Name').strip(),
            level_id=level_obj.id,
            school_id=school_id,
            branch_id=branch_id,
            shift=row.get('Shift')
        ).first()
        if existing_class:
            flash(f"Class '{row.get('Class Name')}' with shift '{row.get('Shift')}' already exists.", "warning")
            continue

        # ---------------- Create Class ----------------
        new_class = Class(
            name=row.get('Class Name').strip(),
            level_id=level_obj.id,
            shift=row.get('Shift') or 'morning',
            capacity=int(row.get('Capacity') or 0),
            status=row.get('Status') or 'active',
            school_id=school_id,
            branch_id=branch_id
        )
        db.session.add(new_class)

    try:
        db.session.commit()
        flash("Classes imported successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error importing classes: {str(e)}", "danger")

    return redirect(url_for('main.all_classes'))



@bp.route("/add/class", methods=["GET", "POST"])
@login_required
def add_class():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = ClassForm()

    # =================== Set defaults ===================
    if current_user.role.value == 'school_admin':
        form.school_id.data = current_user.school_id
        # branch default empty (optional)
        if not form.branch_id.data:
            form.branch_id.data = ''
    elif current_user.role.value == 'branch_admin':
        form.school_id.data = current_user.school_id
        form.branch_id.data = current_user.branch_id

    # =================== On Submit ===================
    if form.validate_on_submit():
        # Coerce branch_id to None if empty
        branch_id = form.branch_id.data if form.branch_id.data else None

        # Prevent duplicate
        existing = Class.query.filter_by(
            name=form.name.data.strip(),
            school_id=form.school_id.data,
            branch_id=branch_id,
            level_id=form.level_id.data,
            shift=form.shift.data   # ✅ NEW
        ).first()
        if existing:
            flash("This class already exists!", "warning")
            return redirect(url_for('main.add_class'))

        # Create class
        new_class = Class(
            name=form.name.data.strip(),
            capacity=form.capacity.data or 0,
            status=form.status.data,
            school_id=form.school_id.data,
            branch_id=branch_id,
            level_id=form.level_id.data,
            shift=form.shift.data   # ✅ NEW
        )
        try:
            db.session.add(new_class)
            db.session.commit()
            flash(f"Class '{form.name.data}' added successfully!", "success")
            return redirect(url_for('main.all_classes'))
        except Exception as e:
            db.session.rollback()
            flash("Error saving class. Check console.", "danger")
            print("DB ERROR:", e)

    else:
        # Show WTForms errors in console
        if form.errors:
            print("FORM ERRORS:", form.errors)

    return render_template(
        "backend/pages/components/classes/add_class.html",
        form=form,
        user=current_user
    )

 

@bp.route("/classes/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_class(id):
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    cls = Class.query.get_or_404(id)

    # Security check
    if current_user.role.value == 'school_admin' and cls.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_classes'))

    if current_user.role.value == 'branch_admin' and cls.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_classes'))

    form = ClassForm(obj=cls)

    # Set values
    form.school_id.data = cls.school_id
    form.branch_id.data = cls.branch_id

    if form.validate_on_submit():
        existing = Class.query.filter(
            Class.id != id,
            Class.name == form.name.data.strip(),
            Class.school_id == form.school_id.data,
            Class.branch_id == form.branch_id.data,
            Class.level_id == form.level_id.data,
            Class.shift == form.shift.data   # ✅ NEW
        ).first()

        if existing:
            flash("Duplicate class exists!", "warning")
            return redirect(url_for('main.edit_class', id=id))

        cls.name = form.name.data.strip()
        cls.capacity = form.capacity.data
        cls.school_id = form.school_id.data
        cls.branch_id = form.branch_id.data
        cls.level_id = form.level_id.data
        cls.status = form.status.data
        cls.shift = form.shift.data   # ✅ NEW


        db.session.commit()

        flash(f"Class '{cls.name}' updated successfully!", "success")
        return redirect(url_for('main.all_classes'))

    return render_template(
        "backend/pages/components/classes/edit_class.html",
        form=form,
        user=current_user
    )



@bp.route("/classes/delete/<int:id>", methods=["POST"])
@login_required
def delete_class(id):
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    cls = Class.query.get_or_404(id)

    # Security check
    if current_user.role.value == 'school_admin' and cls.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_classes'))

    if current_user.role.value == 'branch_admin' and cls.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_classes'))

    try:
        # ------------------ Delete students directly under this class ------------------
        students_in_class = Student.query.filter_by(class_id=cls.id).all()
        for student in students_in_class:
            db.session.delete(student)

        # ------------------ Delete related sections and their students ------------------
        sections = Section.query.filter_by(class_id=cls.id).all()
        for section in sections:
            students_in_section = Student.query.filter_by(section_id=section.id).all()
            for student in students_in_section:
                db.session.delete(student)
            db.session.delete(section)

        # ------------------ Delete class ------------------
        db.session.delete(cls)
        db.session.commit()

        flash(f"Class '{cls.name}', its {len(sections)} section(s), and {len(students_in_class)} student(s) deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash("Error deleting class, its sections, or students. Check console.", "danger")
        print("DB ERROR:", e)

    return redirect(url_for('main.all_classes'))




#--------------------
# -------- Sections
#--------------------

# LIST ALL SECTIONS
@bp.route("/sections")
@login_required
def all_sections():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    if current_user.role.value == 'school_admin':
        # School admin sees only branch-less sections
        sections = Section.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Section.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Branch admin sees only their branch sections
        sections = Section.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Section.created_at.desc()).all()

    return render_template(
        "backend/pages/components/sections/all_sections.html",
        sections=sections,
        user=current_user
    )


# ADD NEW SECTION
@bp.route("/add/section", methods=["GET", "POST"])
@login_required
def add_section():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = SectionForm()
    form.school_id.data = current_user.school_id
    if current_user.role.value == 'branch_admin':
        form.branch_id.data = current_user.branch_id

    # Populate Class dropdown
    classes = Class.query.filter_by(school_id=current_user.school_id).all()
    form.class_id.choices = [(c.id, c.name) for c in classes]  # Only class name

    if form.validate_on_submit():
        branch_id = form.branch_id.data or None
        class_obj = Class.query.get(form.class_id.data)
        if not class_obj:
            flash("Selected class not found.", "danger")
            return redirect(url_for('main.add_section'))

        # Use the shift of the selected class automatically
        shift = class_obj.shift

        # Build full section name
        full_section_name = f"{class_obj.name} - {shift} - {form.name.data.strip()}"

        # Prevent duplicate
        existing = Section.query.filter_by(
            name=full_section_name,
            school_id=form.school_id.data,
            branch_id=branch_id,
            class_id=form.class_id.data,
            shift=shift
        ).first()
        if existing:
            flash("This section already exists!", "warning")
            return redirect(url_for('main.add_section'))

        new_section = Section(
            name=full_section_name,
            capacity=form.capacity.data or 0,
            status=form.status.data,
            school_id=form.school_id.data,
            branch_id=branch_id,
            class_id=form.class_id.data,
            shift=shift  # Automatically taken from class
        )
        try:
            db.session.add(new_section)
            db.session.commit()
            flash(f"Section '{full_section_name}' added successfully!", "success")
            return redirect(url_for('main.all_sections'))
        except Exception as e:
            db.session.rollback()
            flash("Error saving section. Check console.", "danger")
            print("DB ERROR:", e)

    return render_template(
        "backend/pages/components/sections/add_section.html",
        form=form,
        user=current_user,
        classes=classes  # Pass classes for JS usage
    )


# EDIT SECTION
@bp.route("/sections/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_section(id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    section = Section.query.get_or_404(id)

    # Ensure user can edit only their school/branch
    if section.school_id != current_user.school_id:
        flash("Unauthorized access to this section.", "danger")
        return redirect(url_for('main.all_sections'))
    if current_user.role.value == 'branch_admin' and section.branch_id != current_user.branch_id:
        flash("Unauthorized access to this section.", "danger")
        return redirect(url_for('main.all_sections'))

    form = SectionForm(obj=section)  # Populate form with existing data
    form.school_id.data = current_user.school_id
    if current_user.role.value == 'branch_admin':
        form.branch_id.data = current_user.branch_id

    # Populate Class dropdown
    classes = Class.query.filter_by(school_id=current_user.school_id).all()
    form.class_id.choices = [(c.id, c.name) for c in classes]

    if form.validate_on_submit():
        branch_id = form.branch_id.data or None
        class_obj = Class.query.get(form.class_id.data)
        if not class_obj:
            flash("Selected class not found.", "danger")
            return redirect(url_for('main.edit_section', section_id=id))

        # Use the shift of the selected class automatically
        shift = class_obj.shift

        # Build full section name
        full_section_name = f"{class_obj.name} - {shift} - {form.name.data.strip()}"

        # Prevent duplicate (exclude current section)
        existing = Section.query.filter(
            Section.id != section.id,
            Section.name == form.name.data.strip(),
            Section.school_id == form.school_id.data,
            Section.branch_id == branch_id,
            Section.class_id == form.class_id.data,
            Section.shift == shift
        ).first()
        if existing:
            flash("This section already exists!", "warning")
            return redirect(url_for('main.edit_section', section_id=IsADirectoryError))

        # Update section
        section.name = form.name.data.strip()
        section.capacity = form.capacity.data or 0
        section.status = form.status.data
        section.class_id = form.class_id.data
        section.branch_id = branch_id
        section.shift = shift

        try:
            db.session.commit()
            flash(f"Section '{full_section_name}' updated successfully!", "success")
            return redirect(url_for('main.all_sections'))
        except Exception as e:
            db.session.rollback()
            flash("Error updating section. Check console.", "danger")
            print("DB ERROR:", e)

    return render_template(
        "backend/pages/components/sections/edit_section.html",
        form=form,
        user=current_user,
        classes=classes,
        section=section
    )


# DELETE SECTION
@bp.route("/sections/delete/<int:id>", methods=["POST"])
@login_required
def delete_section(id):
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    section = Section.query.get_or_404(id)

    # Security check
    if current_user.role.value == 'school_admin' and section.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_sections'))
    if current_user.role.value == 'branch_admin' and section.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_sections'))

    try:
        # ---------------- Delete students related to this section ----------------
        students_in_section = Student.query.filter_by(section_id=section.id).all()
        for student in students_in_section:
            db.session.delete(student)

        # ---------------- Delete section ----------------
        db.session.delete(section)
        db.session.commit()

        flash(f"Section '{section.name}' and its {len(students_in_section)} student(s) deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash("Error deleting section or its students. Check console.", "danger")
        print("DB ERROR:", e)

    return redirect(url_for('main.all_sections'))


#---------- Subjects
@bp.route("/subjects")
@login_required
def all_subjects():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # -------------------------
    # Role-based filtering
    # -------------------------
    if current_user.role.value == 'school_admin':
        # Only branch-less subjects in the school
        subjects = Subject.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Subject.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Only subjects in this branch
        subjects = Subject.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(Subject.created_at.desc()).all()

    return render_template(
        "backend/pages/components/subjects/all_subjects.html",
        subjects=subjects,
         user=current_user,
    )


# -------------------------
# Generate subject code
# -------------------------
def generate_school_prefix(name):
    if not name:
        return "SCH"

    name = name.upper().strip()
    no_vowels = re.sub(r'[AEIOU]', '', name)

    # fallback if too short
    if len(no_vowels) < 3:
        no_vowels = re.sub(r'[^A-Z]', '', name)

    return (no_vowels[:3]).ljust(3, 'X')


def generate_subject_code(school_id, branch_id=None):
    # Generate prefix from branch name if given, otherwise school name
    if branch_id:
        branch = db.session.get(Branch, branch_id)
        prefix = generate_school_prefix(branch.name if branch else "")
    else:
        school = db.session.get(School, school_id)
        prefix = generate_school_prefix(school.name if school else "")

    # Query subjects in the correct scope
    query = Subject.query.filter(Subject.school_id == school_id)
    if branch_id:
        query = query.filter(Subject.branch_id == branch_id)
    else:
        query = query.filter(Subject.branch_id.is_(None))

    # Get max last 3 digits of existing codes
    max_number = query.with_entities(
        func.coalesce(
            func.max(
                cast(func.right(Subject.code, 3), Integer)
            ),
            0
        )
    ).scalar()

    new_number = max_number + 1
    return f"{prefix}{str(new_number).zfill(3)}"



@bp.route("/add/subject", methods=["GET", "POST"])
@login_required
def add_subject():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = SubjectForm()

    # Default values for GET
    if request.method == "GET":
        if hasattr(current_user, "school_id"):
            form.school_id.data = current_user.school_id
        if hasattr(current_user, "branch_id"):
            form.branch_id.data = current_user.branch_id

    if form.validate_on_submit():
        subject_name = form.name.data.strip()
        school_id = form.school_id.data
        branch_id = form.branch_id.data or None

        # Duplicate check
        existing = Subject.query.filter(
            db.func.lower(Subject.name) == subject_name.lower(),
            Subject.school_id == school_id,
            Subject.branch_id == branch_id
        ).first()
        if existing:
            flash("⚠️ Subject already exists!", "warning")
            return redirect(url_for('main.add_subject'))

        # Generate final code
        code = generate_subject_code(school_id=school_id, branch_id=branch_id)

        subject = Subject(
            name=subject_name,
            code=code,
            school_id=school_id,
            branch_id=branch_id
        )

        try:
            db.session.add(subject)
            db.session.commit()
            flash(f"✅ Subject added successfully! Code: {code}", "success")
            return redirect(url_for('main.all_subjects'))
        except Exception as e:
            db.session.rollback()
            print("ERROR saving subject:", e)
            flash("❌ Error saving subject.", "danger")

    return render_template(
        "backend/pages/components/subjects/add_subject.html",
        form=form,
         user=current_user,
    )


@bp.route("/edit/subject/<int:subject_id>", methods=["GET", "POST"])
@login_required
def edit_subject(subject_id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    subject = Subject.query.get_or_404(subject_id)
    form = SubjectForm(obj=subject)  # pre-fill with existing values

    # ===== Restrict school/branch admin =====
    if current_user.role.value == 'school_admin' and subject.school_id != current_user.school_id:
        flash("You cannot edit subjects from other schools.", "danger")
        return redirect(url_for('main.all_subjects'))

    if current_user.role.value == 'branch_admin' and subject.branch_id != current_user.branch_id:
        flash("You cannot edit subjects from other branches.", "danger")
        return redirect(url_for('main.all_subjects'))

    # ===== GET request: pre-fill hidden values for school/branch admin =====
    if request.method == "GET":
        if current_user.role.value == 'school_admin':
            form.school_id.data = current_user.school_id
        if current_user.role.value == 'branch_admin':
            form.branch_id.data = current_user.branch_id

    if form.validate_on_submit():
        new_name = form.name.data.strip()
        school_id = form.school_id.data
        branch_id = form.branch_id.data or None

        # Prevent duplicate in same school/branch
        existing = Subject.query.filter(
            db.func.lower(Subject.name) == new_name.lower(),
            Subject.school_id == school_id,
            Subject.branch_id == branch_id,
            Subject.id != subject.id
        ).first()
        if existing:
            flash("⚠️ Subject with this name already exists!", "warning")
            return redirect(url_for('main.edit_subject', subject_id=subject.id))

        # Update subject
        subject.name = new_name

        # Regenerate code ONLY if name changed
        if new_name != subject.name:
            subject.code = generate_subject_code(school_id=school_id, branch_id=branch_id)

        subject.school_id = school_id
        subject.branch_id = branch_id

        try:
            db.session.commit()
            flash(f"✅ Subject updated successfully! Code: {subject.code}", "success")
            return redirect(url_for('main.all_subjects'))
        except Exception as e:
            db.session.rollback()
            print("ERROR updating subject:", e)
            flash("❌ Error updating subject.", "danger")

    return render_template(
        "backend/pages/components/subjects/edit_subject.html",
        form=form,
         user=current_user,
        subject=subject
    )



@bp.route("/subjects/delete/<int:id>", methods=["POST"])
@login_required
def delete_subject(id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    subject = Subject.query.get_or_404(id)

    if current_user.role.value == 'school_admin' and subject.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_subjects'))

    if current_user.role.value == 'branch_admin' and subject.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_subjects'))

    db.session.delete(subject)
    db.session.commit()

    flash("Subject deleted successfully!", "success")
    return redirect(url_for('main.all_subjects'))


@bp.route("/export-subjects")
@login_required
def export_subjects():
    import io, csv
    from flask import Response
    
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        subjects = Subject.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Subject.name).all()
    elif current_user.role.value == 'branch_admin':
        subjects = Subject.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Subject.name).all()
    else:
        return Response("Unauthorized", status=403)

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "ID", "Name", "Code", "School", "Branch", "Created At", "Updated At"
    ])

    # ---------------- CSV Data ----------------
    for s in subjects:
        writer.writerow([
            s.id,
            s.name,
            s.code,
            s.school.name if s.school else "",
            s.branch.name if s.branch else "",
            s.created_at.strftime('%Y-%m-%d %H:%M:%S') if s.created_at else "",
            s.updated_at.strftime('%Y-%m-%d %H:%M:%S') if s.updated_at else ""
        ])

    output.seek(0)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=subjects.csv"}
    )

@bp.route("/import-subjects", methods=["POST"])
@login_required
def import_subjects():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.all_subjects'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    file = request.files.get('file')
    if not file:
        flash("No file selected", "danger")
        return redirect(url_for('main.all_subjects'))

    import io, csv

    reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='utf-8'))
    added_count = 0

    for row in reader:
        subject_name = row.get("Name", "").strip()
        school_name = row.get("School", "").strip()
        branch_name = row.get("Branch", "").strip()

        # Determine school_id
        school_id = current_user.school_id
        branch_id = None
        if current_user.role.value == 'branch_admin':
            branch_id = current_user.branch_id
        elif current_user.role.value == 'school_admin' and branch_name:
            branch_obj = Branch.query.filter_by(name=branch_name, school_id=current_user.school_id).first()
            branch_id = branch_obj.id if branch_obj else None

        if not subject_name:
            continue

        # Skip duplicates
        existing = Subject.query.filter(
            db.func.lower(Subject.name) == subject_name.lower(),
            Subject.school_id == school_id,
            Subject.branch_id == branch_id
        ).first()
        if existing:
            continue

        # Generate code
        code = generate_subject_code(school_id=school_id, branch_id=branch_id)

        subject = Subject(
            name=subject_name,
            code=code,
            school_id=school_id,
            branch_id=branch_id
        )
        db.session.add(subject)
        added_count += 1

    try:
        db.session.commit()
        flash(f"✅ Imported {added_count} subjects successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"❌ Error importing subjects: {str(e)}", "danger")

    return redirect(url_for('main.all_subjects'))


# ----------------------
# LIST ALL PARENTS
# ----------------------
@bp.route("/parents")
@login_required
def all_parents():
    # Check user status
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # Role-based filtering
    if current_user.role.value == 'school_admin':
        # Only branch-less parents in the school
        parents = Parent.query.filter_by(school_id=current_user.school_id, branch_id=None)\
                              .order_by(Parent.created_at.desc()).all()
    elif current_user.role.value == 'branch_admin':
        # Only parents in this branch
        parents = Parent.query.filter_by(branch_id=current_user.branch_id)\
                              .order_by(Parent.created_at.desc()).all()

    return render_template(
        "backend/pages/components/parents/all_parents.html",
        parents=parents,
        user=current_user
    )

@bp.route('/export-parents-full')
@login_required
def export_parents_full():
    import io, csv
    from flask import Response
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        # School admin → parents school-ka ku jira oo branch_id=NULL
        parents = Parent.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Parent.created_at.desc()).all()
    elif current_user.role.value == 'branch_admin':
        # Branch admin → parents branch-ka uu leeyahay
        parents = Parent.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Parent.created_at.desc()).all()
    else:
        return Response("Unauthorized", status=403)

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "ID", "Full Name", "Password Hash", "Photo", "Roll No",
        "Gender", "Date of Birth", "Occupation", "Emergency Contact",
        "Phone", "Email", "Address", "National ID", "Relationship",
        "Status", "School", "Branch", "Created At", "Updated At"
    ])

    # ---------------- CSV Data ----------------
    for p in parents:
        writer.writerow([
            p.id,
            p.full_name,
            p.password_hash,
            p.photo or "",
            p.roll_no or "",
            p.gender or "",
            p.date_of_birth.strftime('%Y-%m-%d') if p.date_of_birth else "",
            p.occupation or "",
            p.emergency_contact or "",
            p.phone or "",
            p.email or "",
            p.address or "",
            p.national_id or "",
            p.relationship or "",
            p.status or "",
            p.school.name if p.school else "",
            p.branch.name if p.branch else "",
            p.created_at.strftime('%Y-%m-%d %H:%M:%S') if p.created_at else "",
            p.updated_at.strftime('%Y-%m-%d %H:%M:%S') if p.updated_at else ""
        ])

    output.seek(0)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=parents_full.csv"}
    )


@bp.route('/import-parents-full', methods=['POST'])
@login_required
def import_parents_full():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.all_parents'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    file = request.files.get('file')
    if not file:
        flash("No file selected", "danger")
        return redirect(url_for('main.all_parents'))

    import io
    import csv
    from datetime import datetime

    reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='latin-1'))

    for row in reader:
        school_id = current_user.school_id
        branch_id = current_user.branch_id if current_user.role.value == 'branch_admin' else None
        if current_user.role.value == 'school_admin':
            branch_name = row.get('Branch')
            if branch_name:
                branch_obj = Branch.query.filter_by(name=branch_name, school_id=current_user.school_id).first()
                branch_id = branch_obj.id if branch_obj else None

        # Parse DOB
        dob = None
        dob_str = row.get('Date of Birth')
        if dob_str:
            for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
                try:
                    dob = datetime.strptime(dob_str, fmt).date()
                    break
                except ValueError:
                    continue

        # Check if parent exists
        parent = Parent.query.filter_by(full_name=row.get('Full Name'), school_id=current_user.school_id).first()

        if not parent:
            parent = Parent(
                full_name=row.get('Full Name'),
                password_hash=row.get('Password Hash') or generate_password_hash("Default123"),
                photo=row.get('Photo'),
                roll_no=row.get('Roll No'),
                gender=row.get('Gender'),
                date_of_birth=dob,
                occupation=row.get('Occupation'),
                emergency_contact=row.get('Emergency Contact'),
                phone=row.get('Phone'),
                email=row.get('Email'),
                address=row.get('Address'),
                national_id=row.get('National ID'),
                relationship=row.get('Relationship'),
                status=row.get('Status', 'active'),
                school_id=school_id,
                branch_id=branch_id
            )
            db.session.add(parent)
        else:
            # update existing parent
            parent.password_hash = row.get('Password Hash') or parent.password_hash
            parent.photo = row.get('Photo')
            parent.roll_no = row.get('Roll No')
            parent.gender = row.get('Gender')
            parent.date_of_birth = dob
            parent.occupation = row.get('Occupation')
            parent.emergency_contact = row.get('Emergency Contact')
            parent.phone = row.get('Phone')
            parent.email = row.get('Email')
            parent.address = row.get('Address')
            parent.national_id = row.get('National ID')
            parent.relationship = row.get('Relationship')
            parent.status = row.get('Status', 'active')
            parent.branch_id = branch_id

    try:
        db.session.commit()
        flash("Parents imported successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error importing parents: {str(e)}", "danger")

    return redirect(url_for('main.all_parents'))




# ----------------------
# Helper Functions
# ----------------------
#-------------------------------
# Generate parent password
#-------------------------------
def generate_parent_password():
    """
    Generate a 6-digit password starting with current year last 2 digits.
    Example: if 2026 -> '26' + 4 random digits
    """
    year_prefix = str(datetime.now().year)[-2:]
    random_suffix = ''.join(str(random.randint(0, 9)) for _ in range(4))
    return year_prefix + random_suffix

#-------------------------------
# Generate parent roll number
#-------------------------------
def generate_parent_roll_no(school_id, branch_id=None):
    """
    Generate roll_no like YYNN:
    - YY: last 2 digits of current year
    - NN: sequence number of parents in school or branch
    branch_id is optional; if provided, count parents in that branch
    """
    year_prefix = str(datetime.now().year)[-2:]

    if branch_id:  # Branch-specific
        count = Parent.query.filter_by(branch_id=branch_id).count()
    else:  # School-wide, branch-less parents only
        count = Parent.query.filter_by(school_id=school_id, branch_id=None).count()

    seq_number = count + 1
    seq_number_str = str(seq_number).zfill(2)
    return f"{year_prefix}{seq_number_str}"

# -------------------- Route --------------------
# ----------------------
# ADD PARENT
# ----------------------
@bp.route("/add/parent", methods=["GET", "POST"])
@login_required
def add_parent():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = ParentForm()

    if request.method == "GET":
        branch_id = form.branch_id.data or None
        form.roll_no.data = generate_parent_roll_no(form.school_id.data, branch_id)
        form.password.data = generate_parent_password()

    if form.validate_on_submit():
        branch_id = form.branch_id.data or None
        full_name_clean = (form.full_name.data or '').strip()
        roll_no_clean = (form.roll_no.data or '').strip()
        raw_password = (form.password.data or generate_parent_password()).strip()

        # Duplicate email check
        if form.email.data:
            existing_parent = Parent.query.filter_by(email=form.email.data.strip()).first()
            if existing_parent:
                flash("A parent with this email already exists!", "warning")
                return redirect(url_for('main.add_parent'))

        # Create Parent instance
        new_parent = Parent(
            full_name=full_name_clean,
            email=form.email.data.strip() if form.email.data else None,
            phone=form.phone.data.strip() if form.phone.data else None,
            school_id=form.school_id.data,
            branch_id=branch_id,
            roll_no=roll_no_clean,
            gender=form.gender.data or None,
            date_of_birth=form.date_of_birth.data,
            occupation=form.occupation.data.strip() if form.occupation.data else None,
            emergency_contact=form.emergency_contact.data.strip() if form.emergency_contact.data else None,
            address=form.address.data.strip() if form.address.data else None,
            national_id=form.national_id.data.strip() if form.national_id.data else None,
            relationship=form.relationship.data if form.relationship.data else None,
            status=form.status.data
        )

        new_parent.set_password(raw_password)

        # Create User
        first_name = full_name_clean.split()[0].lower() if full_name_clean else 'user'
        email = f"{first_name}{roll_no_clean}@gmail.com" if roll_no_clean else f"{first_name}@gmail.com"
        username = f"{first_name}{roll_no_clean}" if roll_no_clean else first_name

        existing_user = User.query.filter_by(email=email).first()
        if not existing_user:
            user = User(
                username=username,
                fullname=full_name_clean,
                email=email,
                school_id=form.school_id.data,
                branch_id=branch_id,
                role=UserRole.parent,
                auth_status='logout',
                status=True,
                password=generate_password_hash(raw_password)
            )
            db.session.add(user)
            db.session.flush()
            new_parent.user_id = user.id
        else:
            new_parent.user_id = existing_user.id

        # Save Parent
        try:
            db.session.add(new_parent)
            db.session.commit()
            flash(f"✅ Parent '{new_parent.full_name}' added successfully! Roll No: {new_parent.roll_no}, Password: {raw_password}", "success")
            return redirect(url_for('main.all_parents'))
        except Exception as e:
            db.session.rollback()
            flash("❌ Error saving parent. Check console.", "danger")
            print("DB ERROR:", e)

    return render_template(
        "backend/pages/components/parents/add_parent.html",
        form=form,
        user=current_user
    )


# ----------------------
# EDIT PARENT
# ----------------------
@bp.route("/edit/parent/<int:parent_id>", methods=["GET", "POST"])
@login_required
def edit_parent(parent_id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    parent = Parent.query.get_or_404(parent_id)

    # Access control
    if current_user.role.value == 'branch_admin' and parent.branch_id != current_user.branch_id:
        flash("You cannot edit this parent.", "danger")
        return redirect(url_for('main.all_parents'))
    if current_user.role.value == 'school_admin' and parent.school_id != current_user.school_id:
        flash("You cannot edit this parent.", "danger")
        return redirect(url_for('main.all_parents'))

    form = ParentForm(obj=parent)

    if request.method == "GET":
        form.roll_no.data = parent.roll_no
        form.password.data = ""  # do not show current password

    if form.validate_on_submit():
        branch_id = form.branch_id.data or None
        full_name_clean = (form.full_name.data or '').strip()
        roll_no_clean = (form.roll_no.data or '').strip()
        raw_password = (form.password.data or '').strip()

        # Duplicate email check
        if form.email.data:
            existing = Parent.query.filter(Parent.email == form.email.data.strip(), Parent.id != parent.id).first()
            if existing:
                flash("A parent with this email already exists!", "warning")
                return redirect(url_for('main.edit_parent', parent_id=parent.id))

        # Update Parent
        parent.full_name = full_name_clean
        parent.email = form.email.data.strip() if form.email.data else None
        parent.phone = form.phone.data.strip() if form.phone.data else None
        parent.school_id = form.school_id.data
        parent.branch_id = branch_id
        parent.gender = form.gender.data or None
        parent.date_of_birth = form.date_of_birth.data
        parent.occupation = form.occupation.data.strip() if form.occupation.data else None
        parent.emergency_contact = form.emergency_contact.data.strip() if form.emergency_contact.data else None
        parent.address = form.address.data.strip() if form.address.data else None
        parent.national_id = form.national_id.data.strip() if form.national_id.data else None
        parent.relationship = form.relationship.data if form.relationship.data else None
        parent.status = form.status.data

        if raw_password:
            parent.set_password(raw_password)

        # Sync User
        user = User.query.filter_by(email=parent.email).first()
        if user:
            user.username = full_name_clean.split()[0].lower() + (roll_no_clean or '')
            user.fullname = full_name_clean
            user.email = parent.email
            user.school_id = parent.school_id
            user.branch_id = parent.branch_id
            if raw_password:
                user.password = generate_password_hash(raw_password)
            parent.user_id = user.id
        else:
            new_user = User(
                username=full_name_clean.split()[0].lower() + (roll_no_clean or ''),
                fullname=full_name_clean,
                email=parent.email,
                school_id=parent.school_id,
                branch_id=parent.branch_id,
                role=UserRole.parent,
                auth_status='logout',
                status=True,
                password=generate_password_hash(raw_password or generate_parent_password())
            )
            db.session.add(new_user)
            db.session.flush()
            parent.user_id = new_user.id

        # Commit changes
        try:
            db.session.commit()
            flash(f"✅ Parent '{parent.full_name}' updated successfully!", "success")
            return redirect(url_for('main.all_parents'))
        except Exception as e:
            db.session.rollback()
            flash("❌ Error updating parent. Check console.", "danger")
            print("DB ERROR:", e)

    return render_template(
        "backend/pages/components/parents/edit_parent.html",
        form=form,
        user=current_user,
        parent=parent
    )


# ----------------------
# DELETE PARENT
# ----------------------
@bp.route("/parents/delete/<int:parent_id>", methods=["POST"])
@login_required
def delete_parent(parent_id):
    # 🔒 Check if user inactive
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Role check
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # 🔎 Fetch Parent
    parent = Parent.query.get_or_404(parent_id)

    # 🔒 Access control
    if current_user.role.value == 'school_admin' and parent.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_parents'))
    if current_user.role.value == 'branch_admin' and parent.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_parents'))

    try:
        # ------------------ DELETE RELATED STUDENTS ------------------
        related_students = Student.query.filter_by(parent_id=parent.id).all()
        for student in related_students:
            # Delete student photo
            if student.photo:
                photo_path = os.path.join('static', student.photo)
                if os.path.exists(photo_path):
                    os.remove(photo_path)
                student_folder = os.path.dirname(photo_path)
                if os.path.exists(student_folder) and not os.listdir(student_folder):
                    os.rmdir(student_folder)

            # Delete related student User
            if student.full_name and student.roll_no:
                first_name = student.full_name.split()[0].lower()
                roll_no_clean = str(student.roll_no).strip()
                expected_email = f"{first_name}{roll_no_clean}@gmail.com"
                user = User.query.filter_by(email=expected_email).first()
                if user:
                    db.session.delete(user)

            # Delete student
            db.session.delete(student)

        # ------------------ DELETE PARENT USER ------------------
        # 1️⃣ Try email match
        parent_user = None
        if parent.email:
            parent_user = User.query.filter_by(email=parent.email).first()
        
        # 2️⃣ Fallback: try username = first_name + roll_no
        if not parent_user and parent.full_name and parent.roll_no:
            first_name = parent.full_name.split()[0].lower()
            roll_no_clean = str(parent.roll_no).strip()
            username = f"{first_name}{roll_no_clean}"
            parent_user = User.query.filter_by(username=username).first()

        if parent_user:
            db.session.delete(parent_user)

        # ------------------ DELETE PARENT ------------------
        db.session.delete(parent)
        db.session.commit()

        flash(f"Parent '{parent.full_name}' and their {len(related_students)} student(s) including related users deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash("Error deleting parent, related students, or users. Check console.", "danger")
        print("DB ERROR:", e)

    return redirect(url_for('main.all_parents'))


#========================
#------------ Students
#======================

@bp.route("/students")
@login_required
def all_students():
    # Check if current user is active
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin can access
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------------------
    # Role-based student filtering
    # ---------------------------
    if current_user.role.value == 'school_admin':
        # Only students in the school with no branch assigned
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Student.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Only students in the branch
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Student.created_at.desc()).all()

    return render_template(
        "backend/pages/components/students/all_students.html",
        students=students,
        user=current_user
    )


@bp.route('/export-students')
@login_required
def export_students():
    import io, csv
    from flask import Response, flash, redirect, url_for
    from datetime import datetime
    
    # 1. Hubi Permissions-ka
    if current_user.status == 0:
        flash("Account-kaagu ma shaqeynayo. Fadlan la xiriir maamulka.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for('main.dashboard'))

    # 2. Role-based Filter (Base Filter)
    base_filter = {"school_id": current_user.school_id}
    if current_user.role.value == 'branch_admin':
        base_filter["branch_id"] = current_user.branch_id
    else:
        # school_admin wuxuu arkaa kuwa aan branch lahayn
        base_filter["branch_id"] = None

    students = Student.query.filter_by(**base_filter).order_by(Student.created_at.desc()).all()

    # 3. Diyaarinta CSV-ga
    output = io.StringIO()
    writer = csv.writer(output)

    # --- CSV Header ---
    writer.writerow([
        "ID", "Full Name", "Gender", "Shift",
        "Parent", "School", "Branch",
        "Class", "Level", "Section",
        "Password",
        "Date of Birth", "Place of Birth", "Photo",
        "Academic Year", "Roll No",
        "Price", "Registration Fee", "Total",
        "Status"
    ])

    # --- CSV Data ---
    if students:
        for s in students:
            # SAXIDDA: class_obj ayaan isticmaalnay si looga fogaado AttributeError
            class_name = s.class_obj.name if hasattr(s, 'class_obj') and s.class_obj else (s.class_id if s.class_id else "")
            
            writer.writerow([
                s.id,
                s.full_name,
                s.gender or "",
                s.shift or "",
                s.parent.full_name if s.parent else "",
                s.school.name if s.school else "",
                s.branch.name if s.branch else "",
                class_name, # Halkan ayaa la saxay
                s.level.name if s.level else "",
                s.section_rel.name if hasattr(s, 'section_rel') and s.section_rel else "",
                "",  # Password waa in la dhaafaa (Amni ahaan)
                s.date_of_birth.strftime('%Y-%m-%d') if s.date_of_birth else "",
                s.place_of_birth or "",
                s.photo or "",
                s.academic_year or "",
                s.roll_no or "",
                float(s.price or 0),
                float(s.registration_fee or 0),
                float(s.total or 0),
                s.status or ""
            ])
    else:
        # Haddii aysan arday jirin, sii Sample data
        school_name = current_user.school.name if current_user.school else "Your School"
        branch_name = current_user.branch.name if (hasattr(current_user, 'branch') and current_user.branch) else ""
        
        sample_students = [
            ["", "Ahmed Ali", "male", "Morning", "Parent 1", school_name, branch_name, "Class 1", "Primary", "A", "", "2015-05-10", "Mogadishu", "", "2025-2026", "STD-001", 50, 5, 55, "active"],
            ["", "Ayaan Hassan", "female", "Morning", "Parent 2", school_name, branch_name, "Class 1", "Primary", "A", "", "2016-03-15", "Mogadishu", "", "2025-2026", "STD-002", 50, 5, 55, "active"],
        ]
        for row in sample_students:
            writer.writerow(row)

    output.seek(0)
    
    # Magaca faylka oo wata taariikhda maanta
    filename = f"students_export_{datetime.now().strftime('%Y%m%d')}.csv"
    
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )



@bp.route('/export-students-xlsx')
@login_required
def export_students_xlsx():
    import pandas as pd
    import io
    from flask import send_file, Response, flash, redirect, url_for, current_app
    
    if current_user.status == 0:
        flash("Account-kaagu ma shaqeynayo. Fadlan la xiriir maamulka.", "danger")
        return redirect(url_for('main.index'))

    # Kaliya school_admin ama branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    # Waxaan u isticmaalaynaa .filter() si aan u hubino school_id iyo branch_id si sax ah
    base_filter = {"school_id": current_user.school_id}
    if current_user.role.value == 'branch_admin':
        base_filter["branch_id"] = current_user.branch_id
    else:
        # school_admin wuxuu arkaa kuwa aan branch lahayn (None)
        base_filter["branch_id"] = None

    students = Student.query.filter_by(**base_filter).order_by(Student.created_at.desc()).all()
    parents = Parent.query.filter_by(**base_filter).order_by(Parent.created_at.desc()).all()
    classes = Class.query.filter_by(**base_filter).order_by(Class.name).all()
    levels = ClassLevel.query.filter_by(**base_filter).order_by(ClassLevel.name).all()
    sections = Section.query.filter_by(**base_filter).order_by(Section.name).all()

    # ---------------- Build Data ----------------
    students_data = []
    for s in students:
        # SAXIDDA: Waxaan isticmaalaynaa xiriirka saxda ah (level ama class_obj)
        # Haddii model-kaaga uu leeyahay 'class_id', SQLAlchemy badanaaba wuxuu u bixiyaa xiriirka 'level' ama 'class'
        students_data.append({
            "ID": s.id,
            "Full Name": s.full_name,
            "Gender": s.gender,
            "Shift": s.shift,
            "Parent": s.parent.full_name if s.parent else "",
            "School": s.school.name if s.school else "",
            "Branch": s.branch.name if s.branch else "",
            # HALKAN AYUU KHALADKU AHAA - Waxaan u beddelnay 'level' maadaama Class-ku uu ku dhex jiro Level
            "Class": s.class_obj.name if hasattr(s, 'class_obj') and s.class_obj else (s.class_id if s.class_id else ""),
            "Level": s.level.name if s.level else "",
            "Section": s.section_rel.name if hasattr(s, 'section_rel') and s.section_rel else "",
            "Date of Birth": s.date_of_birth.strftime('%Y-%m-%d') if s.date_of_birth else "",
            "Place of Birth": s.place_of_birth or "",
            "Academic Year": s.academic_year or "",
            "Roll No": s.roll_no or "",
            "Price": float(s.price or 0),
            "Registration Fee": float(s.registration_fee or 0),
            "Total": float(s.total or 0),
            "Status": s.status
        })

    parents_data = []
    for p in parents:
        parents_data.append({
            "ID": p.id,
            "Full Name": p.full_name,
            "Gender": p.gender or "",
            "Phone": p.phone or "",
            "Email": p.email or "",
            "Status": p.status or ""
        })

    classes_data = []
    for c in classes:
        classes_data.append({
            "ID": c.id,
            "Name": c.name,
            "Level": c.level.name if c.level else "",
            "Capacity": c.capacity or 0,
            "Status": c.status or ""
        })

    levels_data = []
    for l in levels:
        levels_data.append({
            "ID": l.id,
            "Name": l.name,
            "Price": float(l.price or 0)
        })

    sections_data = []
    for sec in sections:
        sections_data.append({
            "ID": sec.id,
            "Name": sec.name,
            # Hubi haddii xiriirka Section loo bixiyey 'klass' ama 'class_obj'
            "Class": sec.klass.name if hasattr(sec, 'klass') and sec.klass else "",
            "Capacity": sec.capacity or 0,
            "Status": sec.status or ""
        })

    # ---------------- Excel Multi-Sheet ----------------
    try:
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # Haddii liisku maran yahay, DataFrame-ka ha laga dhigo mid banaan laakiin leh columns
            pd.DataFrame(students_data).to_excel(writer, sheet_name='Students', index=False)
            pd.DataFrame(parents_data).to_excel(writer, sheet_name='Parents', index=False)
            pd.DataFrame(classes_data).to_excel(writer, sheet_name='Classes', index=False)
            pd.DataFrame(levels_data).to_excel(writer, sheet_name='Levels', index=False)
            pd.DataFrame(sections_data).to_excel(writer, sheet_name='Sections', index=False)

        output.seek(0)
        return send_file(
            output,
            download_name=f"School_Data_{datetime.now().strftime('%Y%m%d')}.xlsx",
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        flash(f"Excel error: {str(e)}", "danger")
        return redirect(url_for('main.all_students'))
    


# ===================== HELPERS =====================
def clean(value):
    """Nadiifinta qoraalka iyo ka saarista NaN"""
    if pd.isna(value) or str(value).strip().lower() in ['nan', 'none', 'null', '']:
        return None
    return str(value).strip()

def parse_float(value):
    """U bedelida lambar (Decimal) si ammaan ah"""
    if pd.isna(value) or value is None:
        return 0.0
    try:
        # Ka saar wixii aan lambar ahayn
        clean_val = str(value).replace('$', '').replace(',', '').strip()
        return float(clean_val)
    except (ValueError, TypeError):
        return 0.0


@bp.route('/import-students', methods=['POST'])
@login_required
def import_students():
    # 1. Hubi Permissions-ka
    if current_user.status == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Ma haysatid ogolaansho aad xog ku soo dhoweysid.", "danger")
        return redirect(url_for('main.dashboard'))

    # 2. Hel Faylka
    file = request.files.get('file')
    if not file or file.filename == '':
        flash("Fadlan dooro fayl (Excel/CSV).", "danger")
        return redirect(url_for('main.all_students'))

    # 3. Aqri Faylka
    try:
        filename = file.filename.lower()
        if filename.endswith('.csv'):
            df = pd.read_csv(file)
        elif filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file)
        else:
            flash("Kaliya Excel ama CSV ayaa la oggol yahay.", "danger")
            return redirect(url_for('main.all_students'))

        # Nadiifi xogta (Iska saar safafka madhan)
        df = df.where(pd.notnull(df), None)
        rows = df.to_dict(orient='records')
    except Exception as e:
        flash(f"Khalad faylka ah: {str(e)}", "danger")
        return redirect(url_for('main.all_students'))

    default_password = "Student123"
    added_count = 0
    skipped_count = 0

    # 4. LOOP-ka Xogta
    for index, row in enumerate(rows, start=1):
        full_name = clean(row.get('Full Name'))
        roll_no = clean(row.get('Roll No'))
        
        # Haddii magaca ama roll_no ay maqan yihiin
        if not full_name:
            continue

        try:
            # ✅ DUPLICATE CHECK
            existing_student = Student.query.filter_by(
                full_name=full_name,
                roll_no=roll_no,
                school_id=current_user.school_id
            ).first()

            if existing_student:
                skipped_count += 1
                continue 

            # Hel Fasalka iyo Heerka (Class & Level)
            # Hubi in magacyada Excel-ka ku qoran ay la mid yihiin kuwa Database-ka
            class_obj = Class.query.filter_by(name=clean(row.get('Class')), school_id=current_user.school_id).first()
            level_obj = ClassLevel.query.filter_by(name=clean(row.get('Level')), school_id=current_user.school_id).first()

            # Haddii la waayo Fasalka ama Heerka, ka bood safkan
            if not class_obj:
                print(f"Row {index}: Fadhiga '{row.get('Class')}' lama helin.")
                continue
            
            # Haddii level_obj la waayo, isku day in laga soo qaado class_obj
            final_level_id = level_obj.id if level_obj else class_obj.level_id

            # Section & Parent
            section_name = clean(row.get('Section'))
            section_obj = Section.query.filter_by(name=section_name, school_id=current_user.school_id).first() if section_name else None
            
            parent_name = clean(row.get('Parent'))
            parent_obj = Parent.query.filter_by(full_name=parent_name, school_id=current_user.school_id).first() if parent_name else None

            # Date of Birth conversion
            dob = None
            dob_str = str(row.get('Date of Birth')).strip() if row.get('Date of Birth') else None
            if dob_str and dob_str != 'None':
                for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d %H:%M:%S'):
                    try:
                        dob = datetime.strptime(dob_str.split(' ')[0], fmt).date()
                        break
                    except: continue

            # 5. CREATE STUDENT
            student = Student(
                full_name=full_name,
                gender=clean(row.get('Gender')) or 'Other',
                parent_id=parent_obj.id if parent_obj else None,
                school_id=current_user.school_id,
                # Haddii uu yahay branch_admin, branch-kiisa ku xir, haddii kale xogta Excel ka eeg
                branch_id=current_user.branch_id if current_user.role.value == 'branch_admin' else getattr(current_user, 'branch_id', None),
                class_id=class_obj.id,
                level_id=final_level_id,
                section_id=section_obj.id if section_obj else None,
                shift=clean(row.get('Shift')) or 'Morning',
                roll_no=roll_no,
                academic_year=clean(row.get('Academic Year')),
                # Qiimaha haddii uu Excel-ka ka maqan yahay, ka qaado Level-ka
                price=parse_float(row.get('Price')) or (level_obj.price if level_obj else 0),
                registration_fee=parse_float(row.get('Registration Fee')) or 0,
                status='active',
                date_of_birth=dob
            )

            student.set_password(default_password)
            db.session.add(student)
            db.session.flush() # Soo saar student.id si loogu isticmaalo FeeCollection

            # 6. CREATE FEE RECORD
            # Wadarta lacagta laga rabo (Price + Registration)
            total_due = float(student.price or 0) + float(student.registration_fee or 0)
            
            fee = StudentFeeCollection(
                student_id=student.id,
                class_id=student.class_id,
                section_id=student.section_id,
                school_id=student.school_id,
                branch_id=student.branch_id,
                amount_due=total_due,
                amount_paid=0.0,
                payment_status='Pending',
                remarks='Si otomaatig ah ayaa loo abuuray intii lagu guda jiray Import'
            )
            db.session.add(fee)
            added_count += 1

        except Exception as e:
            db.session.rollback()
            print(f"Khalad safka {index}: {str(e)}")
            continue

    # 7. COMMIT & FINISH
    try:
        db.session.commit()
        flash(f"✅ Shaqadu way dhammaatay! La daray: {added_count}, La dhaafay (Hore u jiray): {skipped_count}", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Khalad ka dhacay kaydinta database-ka: {str(e)}", "danger")

    return redirect(url_for('main.all_students'))

# ----------------------
# Helper Functions
# ----------------------
# ----------------------

# ----------------------
# Generate Student Password
# ----------------------
def generate_student_password():
    """
    Generate a 6-digit password starting with last 2 digits of current year.
    Example: 2026 -> '26' + 4 random digits
    """
    year_prefix = str(datetime.now().year)[-2:]
    random_suffix = ''.join(str(random.randint(0, 9)) for _ in range(4))
    return year_prefix + random_suffix

#-------------------------------
# Generate student roll number
#-------------------------------
def generate_student_roll_no(school_id, branch_id=None):
    """
    Generate roll_no like YYNN:
    - YY: last 2 digits of current year
    - NN: sequence number of students
    branch_id is optional:
      - Branch admin → count students in that branch
      - School admin → count students in school with branch_id=NULL
    """
    year_prefix = str(datetime.now().year)[-2:]

    if branch_id:  # Branch-specific
        count = Student.query.filter_by(branch_id=branch_id).count()
    else:  # School-wide, branch-less students only
        count = Student.query.filter_by(school_id=school_id, branch_id=None).count()

    seq_number = count + 1
    seq_number_str = str(seq_number).zfill(2)
    return f"{year_prefix}{seq_number_str}"

# ===========================
# ROUTES: STUDENT DYNAMIC DATA
# ===========================

# Get classes for a level (school + branch isolation)


@bp.route('/get-academic-year/<int:class_id>')
@login_required
def get_academic_year_for_class(class_id):
    if current_user.status == 0:
        return jsonify({"error": "Inactive user"}), 403

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        return jsonify({"error": "Unauthorized"}), 403

    # Get the class
    school_class = Class.query.get(class_id)
    if not school_class:
        return jsonify({"error": "Class not found"}), 404

    # Current year
    start_year = now_eat().year

    # -------------------- Extract numeric grade from class.name --------------------
    import re
    match = re.search(r'\d+', school_class.name)
    if match:
        current_grade = int(match.group())  # e.g., "8aad" -> 8
    else:
        current_grade = 1  # fallback default

    final_grade = 12  # maximum class in school system

    # Calculate remaining years
    remaining_years = final_grade - current_grade + 2  # include current year
    end_year = start_year + remaining_years - 1  # last year

    academic_year = f"{start_year} - {end_year}"
    return jsonify({"academic_year": academic_year})



@bp.route('/get-academic-year')
@login_required
def get_academic_year_route():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    return jsonify({"academic_year": get_academic_year()})
  
# --------------------------
# Get all levels for current school/branch
# --------------------------
@bp.route('/get-levels')
@login_required
def get_levels():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))


    query = ClassLevel.query.filter_by(school_id=current_user.school_id)

    # ==============================
    # 🔥 BRANCH ADMIN → only own branch
    # ==============================
    if current_user.role.value == 'branch_admin':
        if not current_user.branch_id:
            return jsonify([])

        query = query.filter(ClassLevel.branch_id == current_user.branch_id)

    # ==============================
    # 🔥 SCHOOL ADMIN → only school (no branch)
    # ==============================
    elif current_user.role.value == 'school_admin':
        query = query.filter(ClassLevel.branch_id == None)

    # ==============================
    # 🔥 SUPERADMIN (optional)
    # ==============================
    elif current_user.role.value == 'superadmin':
        pass  # sees all

    levels = query.order_by(ClassLevel.name).all()

    return jsonify([{
        "id": l.id,
        "name": l.name
    } for l in levels])



# --------------------------
# Get single level details (price)
# --------------------------
@bp.route('/get-level/<int:level_id>')
@login_required
def get_level(level_id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    query = ClassLevel.query.filter_by(
        id=level_id,
        school_id=current_user.school_id
    )

    # ==============================
    # 🔥 BRANCH ADMIN → only own branch
    # ==============================
    if current_user.role.value == 'branch_admin':
        query = query.filter(ClassLevel.branch_id == current_user.branch_id)

    # ==============================
    # 🔥 SCHOOL ADMIN → only school (no branch)
    # ==============================
    elif current_user.role.value == 'school_admin':
        query = query.filter(ClassLevel.branch_id == None)

    level = query.first_or_404()

    return jsonify({
        "id": level.id,
        "name": level.name,
        "price": float(level.price)
    })




# --------------------------
# Get classes for a level
# --------------------------
@bp.route('/get-classes/<int:level_id>')
@login_required
def get_classes(level_id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    query = Class.query.filter_by(level_id=level_id, school_id=current_user.school_id, status='active')

    if current_user.role.value == 'branch_admin' and current_user.branch_id:
        query = query.filter(
            (Class.branch_id == current_user.branch_id) | (Class.branch_id == None)
        )
    elif current_user.role.value == 'school_admin':
        query = query.filter(Class.branch_id == None)

    classes = query.order_by(Class.name).all()
    return jsonify([{"id": c.id, "name": c.name, "shift": c.shift} for c in classes])


# --------------------------
# Get sections for a class
# --------------------------
@bp.route('/get-sections/<int:class_id>')
@login_required
def get_sections(class_id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    query = Section.query.filter_by(class_id=class_id, school_id=current_user.school_id)

    if current_user.role.value == 'branch_admin' and current_user.branch_id:
        query = query.filter(
            (Section.branch_id == current_user.branch_id) | (Section.branch_id == None)
        )
    elif current_user.role.value == 'school_admin':
        query = query.filter(Section.branch_id == None)

    sections = query.order_by(Section.name).all()
    return jsonify([{"id": s.id, "name": s.name, "shift": s.shift} for s in sections])



# ===========================
# ROUTE: ADD STUDENT
# ===========================
@bp.route("/add/student", methods=["GET", "POST"])
@login_required
def add_student():
    # --------------------- PERMISSIONS ---------------------
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = StudentForm()

    # --------------------- PRE-FILL DEFAULTS ---------------------
    if request.method == "GET":
        form.school_id.data = current_user.school_id
        form.branch_id.data = current_user.branch_id if current_user.role.value == "branch_admin" else None

        branch_id = form.branch_id.data or None
        form.roll_no.data = generate_student_roll_no(form.school_id.data, branch_id)
        form.password.data = generate_student_password()

    if form.validate_on_submit():
        school_id = form.school_id.data or current_user.school_id
        branch_id = form.branch_id.data or None

        # --------------------- VALIDATE BRANCH ---------------------
        if branch_id:
            branch = Branch.query.get(branch_id)
            if not branch or branch.school_id != school_id:
                flash("Selected branch does not belong to your school!", "danger")
                return redirect(url_for('main.add_student'))

        # --------------------- PREVENT DUPLICATES ---------------------
        roll_no_clean = str(form.roll_no.data).strip()
        full_name_clean = form.full_name.data.strip()

        existing_roll = Student.query.filter_by(
            roll_no=roll_no_clean,
            school_id=school_id,
            branch_id=branch_id
        ).first()
        if existing_roll:
            flash("A student with this roll number already exists in this school/branch!", "warning")
            return redirect(url_for('main.add_student'))

        existing_name = Student.query.filter_by(
            full_name=full_name_clean,
            school_id=school_id,
            branch_id=branch_id
        ).first()
        if existing_name:
            flash("A student with this full name already exists in this school/branch!", "warning")
            return redirect(url_for('main.add_student'))

        # --------------------- CREATE STUDENT INSTANCE ---------------------
        new_student = Student(
            full_name=full_name_clean,
            gender=form.gender.data or None,
            parent_id=form.parent_id.data or None,
            school_id=school_id,
            branch_id=branch_id,
            class_id=form.class_id.data or None,
            level_id=form.level_id.data or None,
            section_id=form.section_id.data or None,
            shift=form.shift.data,   # ✅ NEW
            date_of_birth=form.date_of_birth.data,
            place_of_birth=(form.place_of_birth.data or "").strip() or None,
            photo=(form.photo.data or "").strip() or None,
            academic_year=(form.academic_year.data or "").strip() or None,
            roll_no=roll_no_clean,
            price=form.price.data or 0,
            registration_fee=form.registration_fee.data or 0,
            status=form.status.data
        )

        # --------------------- SET PASSWORD ---------------------
        raw_password = form.password.data.strip() if form.password.data else generate_student_password()
        new_student.set_password(raw_password)

        # --------------------- CREATE STUDENT USER ---------------------
        first_name = full_name_clean.split()[0].lower()
        email = f"{first_name}{roll_no_clean}@gmail.com"
        username = f"{first_name}{roll_no_clean}"

        existing_user = User.query.filter_by(email=email).first()
        if not existing_user:
            user = User(
                username=username,
                fullname=full_name_clean,
                email=email,
                school_id=school_id,
                branch_id=branch_id,
                role=UserRole.student,
                auth_status='logout',
                status=True
            )
            user.password = generate_password_hash(raw_password)
            db.session.add(user)
            db.session.flush()  # get user.id
            new_student.user = user  # ✅ LINK to student
        else:
            new_student.user_id = existing_user.id  # ✅ LINK if user already exists

        # --------------------- SAVE STUDENT AND CREATE FEE COLLECTION ---------------------
        try:
            db.session.add(new_student)
            db.session.flush()  # get new_student.id without committing

            # Automatically create initial fee collection
            total_fee = float(new_student.price or 0) + float(new_student.registration_fee or 0)
            fee_collection = StudentFeeCollection(
                student_id=new_student.id,
                class_id=new_student.class_id,
                section_id=new_student.section_id,
                school_id=new_student.school_id,
                branch_id=new_student.branch_id,
                amount_due=total_fee,
                amount_paid=0.0,
                payment_status='Pending',
                payment_date=None,
                remarks='Initial fee collection created automatically'
            )
            db.session.add(fee_collection)

            db.session.commit()
            flash(
                f"✅ Student '{new_student.full_name}' added successfully! "
                f"Roll No: {new_student.roll_no}, Password: {raw_password}",
                "success"
            )
            return redirect(url_for('main.all_students'))

        except Exception as e:
            db.session.rollback()
            print("DB ERROR:", e)
            flash("❌ Error saving student and fee collection. Check console.", "danger")

    return render_template(
        "backend/pages/components/students/add_student.html",
        form=form,
        user=current_user,
        get_academic_year=get_academic_year
    )

@bp.route("/edit/student/<int:student_id>", methods=["GET", "POST"])
@login_required
def edit_student(student_id):
    student = Student.query.get_or_404(student_id)

    # ------------------------- ACCESS CONTROL -------------------------
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = StudentForm(obj=student)

    # Populate select fields
    form.parent_id.choices = [(p.id, p.full_name) for p in Parent.query.order_by(Parent.full_name).all()]
    form.level_id.choices = [(l.id, l.name) for l in ClassLevel.query.order_by(ClassLevel.name).all()]
    form.class_id.choices = [(c.id, c.name) for c in Class.query.order_by(Class.name).all()]
    form.section_id.choices = [(s.id, s.name) for s in Section.query.order_by(Section.name).all()]

    # ------------------------- PRE-FILL DATA FOR EDIT -------------------------
    if request.method == "GET":
        form.academic_year.data = student.academic_year or get_academic_year()
        form.shift.data = student.shift or ''

    # ------------------------- FORM SUBMISSION -------------------------
    if form.validate_on_submit():
        # --------------------- PREVIOUS TOTAL FEES ---------------------
        old_total = float(student.price or 0) + float(student.registration_fee or 0)

        # --------------------- UPDATE STUDENT FIELDS ---------------------
        student.full_name = form.full_name.data.strip()
        student.gender = form.gender.data
        student.parent_id = form.parent_id.data or None
        student.level_id = form.level_id.data
        student.class_id = form.class_id.data or None
        student.section_id = form.section_id.data or None
        student.shift = form.shift.data
        student.academic_year = (form.academic_year.data or "").strip()

        student.price = float(form.price.data or 0)
        student.registration_fee = float(form.registration_fee.data or 0)
        new_total = student.price + student.registration_fee
        student.total = new_total

        student.status = form.status.data
        student.updated_at = now_eat()

        # --------------------- PASSWORD ---------------------
        raw_password = None
        if form.password.data:
            raw_password = form.password.data.strip()
            student.set_password(raw_password)

        # --------------------- IMAGE HANDLING ---------------------
        image = request.files.get('photo')
        if image and image.filename:
            import uuid, os
            ext = os.path.splitext(image.filename)[1]
            slug = student.full_name.lower().replace(' ', '-')
            unique_id = uuid.uuid4().hex[:8]
            filename = f"{slug}-{unique_id}{ext}"

            upload_folder = os.path.join('static/backend/uploads/students', str(student.id))
            os.makedirs(upload_folder, exist_ok=True)
            image_path = os.path.join(upload_folder, filename)
            image.save(image_path)

            student.photo = os.path.relpath(image_path, 'static').replace("\\", "/")

        # --------------------- UPDATE FEE COLLECTION IF TOTAL CHANGED ---------------------
        if old_total != new_total:
            fee_collections = StudentFeeCollection.query.filter_by(student_id=student.id).all()
            for fc in fee_collections:
                fc.amount_due = new_total
                if fc.amount_paid > fc.amount_due:
                    fc.amount_paid = fc.amount_due
                    fc.payment_status = "Paid"
                elif fc.amount_paid == 0:
                    fc.payment_status = "Pending"
                elif fc.amount_paid < fc.amount_due:
                    fc.payment_status = "Partial"

        # --------------------- SYNC STUDENT USER ---------------------
        first_name = student.full_name.split()[0].lower()
        roll_no_clean = str(form.roll_no.data).strip() or str(student.roll_no)
        username = f"{first_name}{roll_no_clean}"
        email = f"{username}@gmail.com"

        existing_user = User.query.filter_by(email=email).first()

        if existing_user:
            # Update existing user
            existing_user.username = username
            existing_user.fullname = student.full_name
            existing_user.school_id = student.school_id
            existing_user.branch_id = student.branch_id
            if raw_password:
                existing_user.password = generate_password_hash(raw_password)

            student.user_id = existing_user.id
        else:
            # Create new user if not exists
            new_user = User(
                username=username,
                fullname=student.full_name,
                email=email,
                school_id=student.school_id,
                branch_id=student.branch_id,
                role=UserRole.student,
                auth_status='logout',
                status=True,
                password=generate_password_hash(raw_password or generate_student_password())
            )
            db.session.add(new_user)
            db.session.flush()
            student.user_id = new_user.id

        # --------------------- COMMIT ---------------------
        try:
            db.session.commit()
            flash(f"✅ Student '{student.full_name}' updated successfully!", "success")
            return redirect(url_for('main.all_students'))
        except Exception as e:
            db.session.rollback()
            print("DB ERROR:", e)
            flash("❌ Error updating student. Check console.", "danger")

    return render_template(
        "backend/pages/components/students/edit_student.html",
        form=form,
        student=student,
        get_academic_year=get_academic_year,
        user=current_user,
    )



@bp.route('/delete/student/<int:student_id>', methods=['POST'])
@login_required
def delete_student(student_id):
    # 1. Soo qaado ardayga ama soo tuur 404
    student = Student.query.get_or_404(student_id)

    # 2. 🔒 Amniga: Inactive user ma tirtiri karo waxba
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # 3. 🔒 Amniga: Role-based access control (RBAC)
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized: You do not have permission to delete students.", "danger")
        return redirect(url_for('main.dashboard'))

    # 4. 🔒 Amniga: Hubi in admin-ku uusan tirtirin arday ka baxsan dugsigiisa
    if current_user.school_id != student.school_id:
        flash("Unauthorized: You can only delete students from your own school.", "danger")
        return redirect(url_for('main.all_students'))

    try:
        # ------------------------- DELETE PHOTO & FOLDER -------------------------
        if student.photo:
            # Ka taxadar in 'static' aysan ku dhex jirin variable-ka student.photo
            photo_path = os.path.join(current_app.root_path, 'static', student.photo)
            
            if os.path.exists(photo_path) and os.path.isfile(photo_path):
                os.remove(photo_path)

                # Tirtir folder-ka ardayga haddii uu madnaado (tusaale: static/uploads/student_123/)
                student_folder = os.path.dirname(photo_path)
                # Hubi inaanan tirtirin folder-ka weyn ee static ama uploads
                if os.path.exists(student_folder) and os.path.isdir(student_folder):
                    if not os.listdir(student_folder):
                        os.rmdir(student_folder)

        # ------------------------- DELETE RELATED USER -------------------------
        # Saadaasha email-ka waa khatar haddii email-ka la bedelo. 
        # Waxaan raadinaynaa User-ka leh email-ka ardayga ee hadda diiwaanka ugu jira.
        # Haddii Student model-kaagu leeyahay 'user_id', taas isticmaal (Best Practice).
        
        # Halkan waxaan u isticmaalaynaa email-ka ku jira Student record-ka (Haddii uu jiro field-kaas)
        # Haddii kale, qaabkaaga email-ka ee f"{first_name}{roll_no_clean}@gmail.com" halkaan ku hay.
        
        first_name = student.full_name.split()[0].lower() if student.full_name else ""
        roll_no_clean = str(student.roll_no).strip() if student.roll_no else ""
        expected_email = f"{first_name}{roll_no_clean}@gmail.com"

        related_user = User.query.filter_by(email=expected_email).first()
        if related_user:
            db.session.delete(related_user)

        # ------------------------- DELETE STUDENT -------------------------
        # Maadaama aad haysato cascade='all, delete-orphan' model-kaaga:
        # Waxaa si toos ah u tirmaya: Attendance, Marks, Results, Tickets, iyo Hall Assignments.
        db.session.delete(student)
        
        db.session.commit()
        flash(f"✅ Student '{student.full_name}' and their account deleted successfully.", "success")

    except Exception as e:
        db.session.rollback()
        # Log error-ka (isticmaal logging halkii aad ka isticmaali lahayd print)
        print(f"CRITICAL DB ERROR: {str(e)}")
        flash("❌ An error occurred during deletion. No data was removed.", "danger")

    return redirect(url_for('main.all_students'))


@bp.route('/delete-repeated-students', methods=['POST'])
@login_required
def delete_repeated_students():
    # 1. Amniga: Kaliya Admin-ka ayaa tirtiri kara
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    try:
        # 2. Raadi magacyada soo noqnoqday ee isku iskuulka ah
        # Waxaan kooxaynaynaa (Group By) magaca iyo fasalka
        repeated_list = db.session.query(
            Student.full_name, 
            Student.class_id, 
            db.func.count(Student.id).label('count')
        ).filter(Student.school_id == current_user.school_id)\
         .group_by(Student.full_name, Student.class_id)\
         .having(db.func.count(Student.id) > 1).all()

        deleted_count = 0

        for entry in repeated_list:
            # Soo qabo dhamaan ardayda magacaas iyo fasalkaas wadaaga
            duplicates = Student.query.filter_by(
                full_name=entry.full_name, 
                class_id=entry.class_id,
                school_id=current_user.school_id
            ).order_by(Student.created_at.asc()).all()

            # Reeb ardayga ugu horeeya (duplicates[0]), tirtir inta kale
            to_delete = duplicates[1:] 

            for s in to_delete:
                # Tirtir sawirka haddii uu jiro
                if s.photo:
                    photo_path = os.path.join(current_app.root_path, 'static', s.photo)
                    if os.path.exists(photo_path):
                        os.remove(photo_path)
                
                # Tirtir ardayga (Cascade ayaa tirtiri doona Marks/Attendance)
                db.session.delete(s)
                deleted_count += 1

        db.session.commit()
        flash(f"✅ Nadiifinta waa dhammaatay. Waxaa la tirtiray {deleted_count} arday oo labalaab ahaa.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"❌ Khalad ayaa dhacay: {str(e)}", "danger")

    return redirect(url_for('main.all_students'))



@bp.route('/delete-all-students', methods=['POST'])
@login_required
def delete_all_students():
    # 🔒 Amniga: Kaliya school_admin ayaa tirtiri kara dhamaan ardayda iskuulka
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    try:
        # 1. Soo qaado dhamaan ardayda iskuulkan
        students = Student.query.filter_by(school_id=current_user.school_id).all()
        student_ids = [s.id for s in students]

        if not student_ids:
            flash("Ma jiraan arday la tirtiro.", "info")
            return redirect(url_for('main.all_students'))

        # 2. 📁 Tirtir Faylasha Sawirada (Physical Photos)
        for s in students:
            if s.photo:
                photo_path = os.path.join(current_app.root_path, 'static', s.photo)
                if os.path.exists(photo_path) and os.path.isfile(photo_path):
                    os.remove(photo_path)
            
            # 🔒 Tirtir User Account-ka haddii uu u jiro email-kaas
            first_name = s.full_name.split()[0].lower() if s.full_name else ""
            roll_no_clean = str(s.roll_no).strip() if s.roll_no else ""
            expected_email = f"{first_name}{roll_no_clean}@gmail.com"
            User.query.filter_by(email=expected_email, school_id=current_user.school_id).delete()

        # 3. 🧹 Tirtir dhamaan xogta xiriirka la leh (Related Data)
        # In kasta oo model-kaagu leeyahay cascade, halkan waa 'Manual Safety Net'
        
        # A. Tirtir Maaliyadda (Invoices & Fee Collections)
        # Invoices-ka ayaa hore loo tirtiraa waayo waxay ku xiran yihiin collections
        invoices_to_delete = FeeInvoice.query.filter(FeeInvoice.student_fee_id.in_(
            db.session.query(StudentFeeCollection.id).filter(StudentFeeCollection.student_id.in_(student_ids))
        )).delete(synchronize_session=False)
        
        StudentFeeCollection.query.filter(StudentFeeCollection.student_id.in_(student_ids)).delete(synchronize_session=False)

        # B. Tirtir Imtixaanada & Natiijooyinka (Marks & Results)
        StudentExamMark.query.filter(StudentExamMark.student_id.in_(student_ids)).delete(synchronize_session=False)
        StudentExamResult.query.filter(StudentExamResult.student_id.in_(student_ids)).delete(synchronize_session=False)
        ExamTicket.query.filter(ExamTicket.student_id.in_(student_ids)).delete(synchronize_session=False)
        ExamHallAssignment.query.filter(ExamHallAssignment.student_id.in_(student_ids)).delete(synchronize_session=False)

        # C. Tirtir Attendance & Promotion History
        StudentAttendance.query.filter(StudentAttendance.student_id.in_(student_ids)).delete(synchronize_session=False)
        StudentPromotion.query.filter(StudentPromotion.student_id.in_(student_ids)).delete(synchronize_session=False)

        # D. Ugu dambeyntii tirtir Ardayda nafsadooda
        Student.query.filter(Student.id.in_(student_ids)).delete(synchronize_session=False)

        db.session.commit()
        flash(f"✅ Si guul leh ayaa loo tirtiray {len(student_ids)} arday iyo dhammaan xogtoodii maaliyadda, imtixaannada, iyo joogitaanka.", "success")

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting all students: {str(e)}")
        flash(f"❌ Khalad ayaa dhacay: {str(e)}", "danger")

    return redirect(url_for('main.all_students'))



#------------------------------
# -------- Teacher 
#------------------------------
@bp.route("/all/teachers")
@login_required
def all_teachers():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # -------------------------
    # Role-based filtering
    # -------------------------
    if current_user.role.value == 'school_admin':
        # School admin sees only teachers in the school without branch
        teachers = Teacher.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Teacher.created_at.desc()).all()

    elif current_user.role.value == 'branch_admin':
        # Branch admin sees only teachers in their branch
        teachers = Teacher.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(Teacher.created_at.desc()).all()

    # -------------------------
    # Render template
    # -------------------------
    return render_template(
        "backend/pages/components/teachers/all_teachers.html",
        teachers=teachers,
        user=current_user
    )


@bp.route("/update-teacher-profile", methods=["POST"])
@login_required
def update_teacher_profile():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    teacher_id = request.form.get("teacher_id")

    teacher = Teacher.query.get_or_404(teacher_id)

    # 🔒 SECURITY → only allow same user
    if teacher.user_id != current_user.id:
        flash("Unauthorized!", "danger")
        return redirect(url_for('main.all_teachers'))

    # UPDATE
    teacher.full_name = request.form.get("full_name")
    teacher.email = request.form.get("email")
    teacher.phone = request.form.get("phone")
    teacher.specialization = request.form.get("specialization")

    try:
        db.session.commit()
        flash("✅ Profile updated successfully", "success")
    except Exception as e:
        db.session.rollback()
        print(e)
        flash("❌ Error updating", "danger")

    return redirect(url_for('main.all_teachers'))


@bp.route('/export-teachers-full')
@login_required
def export_teachers():
    import io, csv
    from flask import Response
    
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))


    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        # School admin → only teachers without branch
        teachers = Teacher.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).all()
    elif current_user.role.value == 'branch_admin':
        # Branch admin → only teachers in their branch
        teachers = Teacher.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).all()
    else:
        return Response("Unauthorized", status=403)

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "ID", "Full Name", "Specialization",
        "Email", "Phone", "Emergency",
        "Gender", "Date of Birth", "Address",
        "Photo", "Password Hash", "Roll No",
        "Status", "School", "Branch",
        "Created At", "Updated At"
    ])

    # ---------------- CSV Data ----------------
    for t in teachers:
        writer.writerow([
            t.id,
            t.full_name,
            t.specialization,
            t.email,
            t.phone or "",
            t.emergency or "",
            t.gender or "",
            t.date_of_birth.strftime('%Y-%m-%d') if t.date_of_birth else "",
            t.address or "",
            t.photo or "",
            t.password_hash,
            t.roll_no or "",
            t.status,
            t.school.name if t.school else "",
            t.branch.name if t.branch else "",
            t.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            t.updated_at.strftime('%Y-%m-%d %H:%M:%S')
        ])

    output.seek(0)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=teachers_full.csv"}
    )



@bp.route('/import-teachers-full', methods=['POST'])
@login_required
def import_teachers():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.all_teachers'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    file = request.files.get('file')
    if not file:
        flash("No file selected", "danger")
        return redirect(url_for('main.all_teachers'))

    import io, csv
    from datetime import datetime
    from werkzeug.security import generate_password_hash

    reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='latin-1'))

    for row in reader:

        # ---------------- School & Branch ----------------
        school_id = current_user.school_id
        branch_id = None

        if current_user.role.value == 'branch_admin':
            branch_id = current_user.branch_id
        else:
            branch_name = row.get('Branch')
            if branch_name:
                branch_obj = Branch.query.filter_by(
                    name=branch_name,
                    school_id=current_user.school_id
                ).first()
                branch_id = branch_obj.id if branch_obj else None

        # ---------------- DOB Parse ----------------
        dob = None
        dob_str = row.get('Date of Birth')
        if dob_str:
            for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
                try:
                    dob = datetime.strptime(dob_str, fmt).date()
                    break
                except ValueError:
                    continue

        # ---------------- Check Existing ----------------
        teacher = Teacher.query.filter_by(
            email=row.get('Email'),
            school_id=current_user.school_id
        ).first()

        if not teacher:
            teacher = Teacher(
                full_name=row.get('Full Name'),
                specialization=row.get('Specialization'),
                email=row.get('Email'),
                phone=row.get('Phone'),
                emergency=row.get('Emergency'),
                gender=row.get('Gender'),
                date_of_birth=dob,
                address=row.get('Address'),
                photo=row.get('Photo'),
                roll_no=row.get('Roll No'),
                status=row.get('Status', 'active'),
                school_id=school_id,
                branch_id=branch_id,
                password_hash=row.get('Password Hash') or generate_password_hash("Teacher123")
            )
            db.session.add(teacher)

        else:
            # ---------------- Update ----------------
            teacher.full_name = row.get('Full Name')
            teacher.specialization = row.get('Specialization')
            teacher.phone = row.get('Phone')
            teacher.emergency = row.get('Emergency')
            teacher.gender = row.get('Gender')
            teacher.date_of_birth = dob
            teacher.address = row.get('Address')
            teacher.photo = row.get('Photo')
            teacher.roll_no = row.get('Roll No')
            teacher.status = row.get('Status', 'active')
            teacher.branch_id = branch_id

            if row.get('Password Hash'):
                teacher.password_hash = row.get('Password Hash')

    try:
        db.session.commit()
        flash("Teachers imported successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error importing teachers: {str(e)}", "danger")

    return redirect(url_for('main.all_teachers'))



def generate_teacher_password():
    """
    Example: 26XXXX (year + 4 digits)
    """
    year_prefix = str(datetime.now().year)[-2:]
    random_suffix = ''.join(str(random.randint(0, 9)) for _ in range(4))
    return year_prefix + random_suffix


def generate_teacher_roll_no(school_id, branch_id=None):
    """
    Generate next roll number: YYNN format.
    - YY: last two digits of year
    - NN: sequential number per school/branch
    """
    year_prefix = str(datetime.now().year)[-2:]

    # Query last teacher in this school/branch
    if branch_id:
        last_teacher = (
            Teacher.query
            .filter_by(school_id=school_id, branch_id=branch_id)
            .order_by(Teacher.id.desc())
            .first()
        )
    else:
        last_teacher = (
            Teacher.query
            .filter_by(school_id=school_id, branch_id=None)
            .order_by(Teacher.id.desc())
            .first()
        )

    if last_teacher and last_teacher.roll_no and last_teacher.roll_no.startswith(year_prefix):
        # Extract last two digits (sequence)
        try:
            last_seq = int(last_teacher.roll_no[-2:])
        except ValueError:
            last_seq = 0
    else:
        last_seq = 0

    next_seq = last_seq + 1
    seq_str = str(next_seq).zfill(2)

    return f"{year_prefix}{seq_str}"



@bp.route("/generate-teacher-roll", methods=["GET"])
@login_required
def generate_teacher_roll():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    try:
        school_id = int(request.args.get('school_id', 0))
    except (TypeError, ValueError):
        return jsonify({"roll_no": ""})

    try:
        branch_id_raw = request.args.get('branch_id', None)
        branch_id = int(branch_id_raw) if branch_id_raw else None
    except (TypeError, ValueError):
        branch_id = None

    if not school_id:
        return jsonify({"roll_no": ""})

    roll_no = generate_teacher_roll_no(school_id, branch_id)
    return jsonify({"roll_no": roll_no})


@bp.route("/add/teacher", methods=["GET", "POST"])
@login_required
def add_teacher():
    form = TeacherForm()
    
    # 🔒 Check active user
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Role access
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- GET ----------------
    if request.method == "GET":
        if current_user.role.value == 'branch_admin':
            form.school_id.data = current_user.school_id
            form.branch_id.data = current_user.branch_id
        else:
            form.school_id.data = current_user.school_id
            form.branch_id.data = None

        form.roll_no.data = generate_teacher_roll_no(current_user.school_id, form.branch_id.data)
        form.password.data = generate_teacher_password()

    # ---------------- POST ----------------
    if form.validate_on_submit():
        school_id = form.school_id.data or current_user.school_id
        branch_id = form.branch_id.data or None

        # Branch validation
        if branch_id:
            branch = Branch.query.get(branch_id)
            if not branch or branch.school_id != school_id:
                flash("Selected branch does not belong to your school!", "danger")
                return redirect(url_for('main.add_teacher'))
        elif current_user.role.value == 'branch_admin':
            branch_id = current_user.branch_id

        # Roll no
        roll_no = form.roll_no.data or f"T-{int(now_eat().timestamp())}"

        # Duplicate checks
        if Teacher.query.filter_by(email=form.email.data.strip()).first():
            flash("A teacher with this email already exists!", "warning")
            return redirect(url_for('main.add_teacher'))

        if Teacher.query.filter_by(
            full_name=form.full_name.data.strip(),
            school_id=school_id,
            branch_id=branch_id
        ).first():
            flash("A teacher with this full name already exists!", "warning")
            return redirect(url_for('main.add_teacher'))

        # ---------------- CREATE TEACHER ----------------
        teacher = Teacher(
            full_name=form.full_name.data.strip(),
            specialization=form.specialization.data.strip(),
            email=form.email.data.strip(),
            phone=form.phone.data.strip() if form.phone.data else None,
            emergency=form.emergency.data.strip() if form.emergency.data else None,
            gender=form.gender.data if form.gender.data in ['male','female','other'] else None,
            address=form.address.data.strip() if form.address.data else None,
            date_of_birth=form.date_of_birth.data,
            school_id=school_id,
            branch_id=branch_id,
            roll_no=roll_no,
            status=form.status.data if form.status.data in ['active','inactive','blocked'] else 'active'
        )

        raw_password = form.password.data or generate_teacher_password()
        teacher.set_password(raw_password)

        # ---------------- IMAGE ----------------
        image = request.files.get('photo')
        if image and image.filename:
            ext = os.path.splitext(image.filename)[1]
            slug = teacher.full_name.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            safe_filename = f"{slug}-{unique_id}{ext}"

            folder = os.path.join('static/backend/uploads/teachers', str(teacher.roll_no))
            os.makedirs(folder, exist_ok=True)

            image_path = os.path.join(folder, safe_filename)
            image.save(image_path)

            teacher.photo = os.path.relpath(image_path, 'static').replace("\\", "/")

        try:
            db.session.add(teacher)
            db.session.flush()  # ✅ get teacher.id

            # ---------------- AUTO CREATE USER ----------------
            existing_user = User.query.filter_by(email=teacher.email).first()

            if not existing_user:
                user = User(
                    username=teacher.email,
                    fullname=teacher.full_name,
                    email=teacher.email,
                    school_id=teacher.school_id,
                    branch_id=teacher.branch_id,
                    role=UserRole.teacher,
                    status=True
                )
                user.password = generate_password_hash(raw_password)

                db.session.add(user)
                db.session.flush()  # ✅ get user.id

                teacher.user = user   # ✅ LINK
            else:
                teacher.user_id = existing_user.id  # ✅ LINK haddii hore u jiray

            db.session.commit()

            flash(f"✅ Teacher '{teacher.full_name}' added! Roll: {teacher.roll_no} | Password: {raw_password}", "success")
            return redirect(url_for('main.all_teachers'))

        except Exception as e:
            db.session.rollback()
            print("DB ERROR:", e)
            flash("❌ Error saving teacher.", "danger")

    # ---------------- RENDER ----------------
    if current_user.role.value == 'branch_admin':
        branches = Branch.query.filter_by(id=current_user.branch_id).all()
    else:
        branches = Branch.query.filter_by(school_id=current_user.school_id).all()

    return render_template(
        "backend/pages/components/teachers/add_teacher.html",
        form=form,
        user=current_user,
        branches=branches
    )



@bp.route("/edit/teacher/<int:teacher_id>", methods=["GET", "POST"])
@login_required
def edit_teacher(teacher_id):
    # 🔒 Check active
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Role access
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    teacher = Teacher.query.get_or_404(teacher_id)

    # 🔒 Scope check
    if current_user.role.value == 'school_admin' and teacher.school_id != current_user.school_id:
        flash("You cannot edit teachers outside your school!", "danger")
        return redirect(url_for('main.all_teachers'))
    if current_user.role.value == 'branch_admin' and teacher.branch_id != current_user.branch_id:
        flash("You cannot edit teachers outside your branch!", "danger")
        return redirect(url_for('main.all_teachers'))

    form = TeacherForm(obj=teacher)

    # ---------------- GET ----------------
    if request.method == "GET":
        if current_user.role.value == 'branch_admin':
            form.school_id.data = current_user.school_id
            form.branch_id.data = current_user.branch_id
        else:
            form.school_id.data = teacher.school_id
            form.branch_id.data = teacher.branch_id

    # ---------------- POST ----------------
    if form.validate_on_submit():
        school_id = form.school_id.data or current_user.school_id
        branch_id = form.branch_id.data or (current_user.branch_id if current_user.role.value == 'branch_admin' else None)

        # ---------------- VALIDATE BRANCH ----------------
        if branch_id:
            branch = Branch.query.get(branch_id)
            if not branch or branch.school_id != school_id:
                flash("Selected branch is invalid!", "danger")
                return redirect(request.url)

        # ---------------- DUPLICATE EMAIL CHECK ----------------
        existing_teacher_email = Teacher.query.filter(
            Teacher.email==form.email.data.strip(),
            Teacher.id != teacher.id
        ).first()
        if existing_teacher_email:
            flash("A teacher with this email already exists!", "warning")
            return redirect(request.url)

        # ---------------- UPDATE TEACHER ----------------
        teacher.full_name = form.full_name.data.strip()
        teacher.specialization = form.specialization.data.strip()
        teacher.email = form.email.data.strip()
        teacher.phone = form.phone.data.strip() if form.phone.data else None
        teacher.emergency = form.emergency.data.strip() if form.emergency.data else None
        teacher.gender = form.gender.data if form.gender.data in ['male','female','other'] else teacher.gender
        teacher.address = form.address.data.strip() if form.address.data else teacher.address
        teacher.date_of_birth = form.date_of_birth.data
        teacher.school_id = school_id
        teacher.branch_id = branch_id
        teacher.roll_no = form.roll_no.data or teacher.roll_no
        teacher.status = form.status.data if form.status.data in ['active','inactive','blocked'] else teacher.status

        # ---------------- PASSWORD ----------------
        raw_password = None
        if form.password.data:
            raw_password = form.password.data
            teacher.set_password(raw_password)

        # ---------------- IMAGE HANDLING ----------------
        if 'remove_photo' in request.form and teacher.photo:
            old_path = os.path.join('static', teacher.photo)
            if os.path.exists(old_path):
                os.remove(old_path)
            teacher.photo = None

        image = request.files.get('photo')
        if image and image.filename:
            ext = os.path.splitext(image.filename)[1]
            slug = teacher.full_name.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            safe_filename = f"{slug}-{unique_id}{ext}"

            folder = os.path.join('static/backend/uploads/teachers', str(teacher.id))
            os.makedirs(folder, exist_ok=True)

            image_path = os.path.join(folder, safe_filename)
            image.save(image_path)
            teacher.photo = os.path.relpath(image_path, 'static').replace("\\", "/")

        # ---------------- SYNC USER & UPDATE user_id ----------------
        try:
            user = User.query.filter_by(email=teacher.email).first()

            if user:
                # Update existing user
                user.fullname = teacher.full_name
                user.email = teacher.email
                user.school_id = teacher.school_id
                user.branch_id = teacher.branch_id

                if raw_password:
                    user.password = generate_password_hash(raw_password)

                # Link teacher to this existing user
                teacher.user_id = user.id
            else:
                # Create new user if not exists
                user = User(
                    username=teacher.email,
                    fullname=teacher.full_name,
                    email=teacher.email,
                    school_id=teacher.school_id,
                    branch_id=teacher.branch_id,
                    role=UserRole.teacher,
                    auth_status='logout',
                    status=True
                )
                user.password = generate_password_hash(raw_password or generate_teacher_password())
                db.session.add(user)
                db.session.flush()  # Get user.id
                teacher.user_id = user.id  # Link teacher to new user

            db.session.commit()
            flash(f"✅ Teacher '{teacher.full_name}' updated successfully!", "success")
            return redirect(url_for('main.all_teachers'))

        except Exception as e:
            db.session.rollback()
            print("DB ERROR:", e)
            flash("❌ Error updating teacher.", "danger")

    # ---------------- RENDER ----------------
    branches = Branch.query.filter_by(
        id=current_user.branch_id
    ).all() if current_user.role.value == 'branch_admin' else Branch.query.filter_by(
        school_id=current_user.school_id
    ).all()

    return render_template(
        "backend/pages/components/teachers/edit_teacher.html",
        form=form,
        teacher=teacher,
        user=current_user,
        branches=branches
    )


@bp.route("/teachers/delete/<int:teacher_id>", methods=["POST"])
@login_required
def delete_teacher(teacher_id):
    teacher = Teacher.query.get_or_404(teacher_id)

    # ---------------- Security ----------------
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.dashboard'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # Check scope
    if current_user.role.value == 'school_admin' and teacher.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_teachers'))

    if current_user.role.value == 'branch_admin' and teacher.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_teachers'))

    try:
        # ---------------- Delete photo file if exists ----------------
        if teacher.photo:
            photo_path = os.path.join('static', teacher.photo)
            if os.path.exists(photo_path):
                os.remove(photo_path)

            # Remove teacher's folder if empty
            teacher_folder = os.path.dirname(photo_path)
            if os.path.exists(teacher_folder) and not os.listdir(teacher_folder):
                os.rmdir(teacher_folder)

        # ---------------- Delete linked User ----------------
        user = User.query.filter_by(email=teacher.email).first()
        if user:
            db.session.delete(user)

        # ---------------- Delete teacher ----------------
        db.session.delete(teacher)
        db.session.commit()

        flash(f"✅ Teacher '{teacher.full_name}' and related user deleted successfully!", "success")

    except Exception as e:
        db.session.rollback()
        print("DB ERROR:", e)
        flash("❌ Error deleting teacher", "danger")

    return redirect(url_for('main.all_teachers'))


#------------------------------------
#------------- Class Subjects
#------------------------------------

@bp.route("/class-subjects")
@login_required
def all_class_subjects():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ================= SCHOOL ADMIN =================
    if current_user.role.value == 'school_admin':
        # Only branch-less classes (school-level)
        class_subjects = ClassSubject.query.join(Class).filter(
            Class.school_id == current_user.school_id,
            Class.branch_id == None  # only school-level classes
        ).order_by(ClassSubject.created_at.desc()).all()

    # ================= BRANCH ADMIN =================
    elif current_user.role.value == 'branch_admin':
        # Only subjects for this branch
        class_subjects = ClassSubject.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(ClassSubject.created_at.desc()).all()

    return render_template(
        "backend/pages/components/subjects/all_class_subjects.html",
        class_subjects=class_subjects,
         user=current_user,
    )



@bp.route("/class-subjects/add", methods=["GET", "POST"])
@login_required
def add_class_subject():
    form = ClassSubjectForm()
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    if form.validate_on_submit():
        class_ids = form.class_ids.data
        subject_ids = form.subject_ids.data
        school_id = current_user.school_id
        branch_id = form.branch_id.data or None

        added_count = 0
        for class_id in class_ids:
            for subject_id in subject_ids:
                # Check role-based ownership
                cls = Class.query.get(class_id)
                subj = Subject.query.get(subject_id)
                
                if current_user.role.value == 'school_admin':
                    if cls.school_id != current_user.school_id or cls.branch_id is not None:
                        continue  # skip classes not owned by school or branch-level
                    if subj.school_id != current_user.school_id or subj.branch_id is not None:
                        continue  # skip subjects not owned by school or branch-level
                elif current_user.role.value == 'branch_admin':
                    if cls.branch_id != current_user.branch_id:
                        continue
                    if subj.branch_id != current_user.branch_id:
                        continue

                # Prevent duplicates
                existing = ClassSubject.query.filter_by(
                    class_id=class_id,
                    subject_id=subject_id,
                    branch_id=branch_id
                ).first()
                if existing:
                    continue

                cs = ClassSubject(
                    class_id=class_id,
                    subject_id=subject_id,
                    school_id=school_id,
                    branch_id=branch_id
                )
                db.session.add(cs)
                added_count += 1

        try:
            db.session.commit()
            flash(f"✅ {added_count} assignments created successfully!", "success")
            return redirect(url_for('main.all_class_subjects'))
        except Exception as e:
            db.session.rollback()
            print("ERROR saving ClassSubjects:", e)
            flash("❌ Error saving assignments.", "danger")

    return render_template(
        "backend/pages/components/subjects/add_class_subject.html",
        form=form,
        user=current_user
    )


@bp.route("/class-subjects/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_class_subject(id):
    cs = ClassSubject.query.get_or_404(id)
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # Restrict access
    if current_user.role.value == 'school_admin' and cs.school_id != current_user.school_id:
        flash("Not allowed to edit assignments from other schools.", "danger")
        return redirect(url_for('main.all_class_subjects'))
    if current_user.role.value == 'branch_admin' and cs.branch_id != current_user.branch_id:
        flash("Not allowed to edit assignments from other branches.", "danger")
        return redirect(url_for('main.all_class_subjects'))

    # Pre-fill existing assignment as initial data
    form = ClassSubjectForm(
        class_ids=[cs.class_id],
        subject_ids=[cs.subject_id]
    )

    if form.validate_on_submit():
        class_ids = form.class_ids.data
        subject_ids = form.subject_ids.data
        school_id = current_user.school_id
        branch_id = form.branch_id.data or None

        added_count = 0

        # Delete the current assignment (we will replace with selected ones)
        db.session.delete(cs)
        db.session.flush()  # flush so we can add new assignments

        for class_id in class_ids:
            for subject_id in subject_ids:
                cls = Class.query.get(class_id)
                subj = Subject.query.get(subject_id)

                # Role-based restrictions
                if current_user.role.value == 'school_admin':
                    if cls.school_id != current_user.school_id or cls.branch_id is not None:
                        continue
                    if subj.school_id != current_user.school_id or subj.branch_id is not None:
                        continue
                elif current_user.role.value == 'branch_admin':
                    if cls.branch_id != current_user.branch_id:
                        continue
                    if subj.branch_id != current_user.branch_id:
                        continue

                # Prevent duplicates
                existing = ClassSubject.query.filter_by(
                    class_id=class_id,
                    subject_id=subject_id,
                    branch_id=branch_id
                ).first()
                if existing:
                    continue

                cs_new = ClassSubject(
                    class_id=class_id,
                    subject_id=subject_id,
                    school_id=school_id,
                    branch_id=branch_id
                )
                db.session.add(cs_new)
                added_count += 1

        try:
            db.session.commit()
            flash(f"✅ {added_count} assignments updated successfully!", "success")
            return redirect(url_for('main.all_class_subjects'))
        except Exception as e:
            db.session.rollback()
            print("ERROR updating ClassSubjects:", e)
            flash("❌ Error updating assignments.", "danger")

    return render_template(
        "backend/pages/components/subjects/edit_class_subject.html",
        form=form,
        class_subject=cs,
        user=current_user
    )



@bp.route("/class-subjects/delete/<int:id>", methods=["POST"])
@login_required
def delete_class_subject(id):
    cs = ClassSubject.query.get_or_404(id)
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # Access control
    if current_user.role.value == 'school_admin' and cs.school_id != current_user.school_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_class_subjects'))

    if current_user.role.value == 'branch_admin' and cs.branch_id != current_user.branch_id:
        flash("Not allowed.", "danger")
        return redirect(url_for('main.all_class_subjects'))

    try:
        db.session.delete(cs)
        db.session.commit()
        flash("✅ Assignment deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        print("ERROR deleting ClassSubject:", e)
        flash("❌ Error deleting assignment.", "danger")

    return redirect(url_for('main.all_class_subjects'))


# ================= EXPORT =================
@bp.route("/class-subjects/export")
@login_required
def export_class_subjects():
    import io, csv
    from flask import Response
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # Only school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        class_subjects = ClassSubject.query.join(Class).filter(
            Class.school_id == current_user.school_id,
            Class.branch_id == None
        ).all()
    elif current_user.role.value == 'branch_admin':
        class_subjects = ClassSubject.query.join(Class).filter(
            Class.branch_id == current_user.branch_id
        ).all()
    else:
        return Response("Unauthorized", status=403)

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "Class Name", "Subject Name", "School", "Branch", "Created At"
    ])

    # ---------------- CSV Data ----------------
    for cs in class_subjects:
        writer.writerow([
            cs.class_obj.name if cs.class_obj else "",
            cs.subject.name if cs.subject else "",
            cs.class_obj.school.name if cs.class_obj and cs.class_obj.school else "",
            cs.class_obj.branch.name if cs.class_obj and cs.class_obj.branch else "",
            cs.created_at.strftime('%Y-%m-%d %H:%M:%S') if cs.created_at else ""
        ])

    output.seek(0)
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=class_subjects.csv"}
    )


@bp.route("/class-subjects/import", methods=["GET", "POST"])
@login_required
def import_class_subjects():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.all_class_subjects'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    if request.method == "POST":
        file = request.files.get('file')
        if not file:
            flash("No file selected", "danger")
            return redirect(url_for('main.all_class_subjects'))

        import io, csv
        reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='utf-8'))

        added_count = 0
        for row in reader:
            class_name = row.get("Class Name")
            subject_name = row.get("Subject Name")

            # Role-based ownership filter
            cls_query = Class.query.filter_by(name=class_name, school_id=current_user.school_id)
            subj_query = Subject.query.filter_by(name=subject_name, school_id=current_user.school_id)

            if current_user.role.value == 'branch_admin':
                cls_query = cls_query.filter_by(branch_id=current_user.branch_id)
                subj_query = subj_query.filter_by(branch_id=current_user.branch_id)
            else:
                cls_query = cls_query.filter_by(branch_id=None)
                subj_query = subj_query.filter_by(branch_id=None)

            cls = cls_query.first()
            subj = subj_query.first()
            if not cls or not subj:
                continue  # Skip invalid entries

            # Prevent duplicates
            existing = ClassSubject.query.filter_by(
                class_id=cls.id, subject_id=subj.id, branch_id=cls.branch_id
            ).first()
            if existing:
                continue

            cs = ClassSubject(
                class_id=cls.id,
                subject_id=subj.id,
                school_id=current_user.school_id,
                branch_id=cls.branch_id
            )
            db.session.add(cs)
            added_count += 1

        try:
            db.session.commit()
            flash(f"✅ Imported {added_count} class-subject assignments successfully!", "success")
        except Exception as e:
            db.session.rollback()
            flash(f"❌ Error importing CSV: {str(e)}", "danger")

        return redirect(url_for('main.all_class_subjects'))

    # GET method → optional: redirect to all_class_subjects (inline import does not need separate page)
    return redirect(url_for('main.all_class_subjects'))




# --------------------------------
#------------ Teacher Assignment 
#---------------------------------

@bp.route("/teacher-assignments")
@login_required
def all_teacher_assignments():
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ================= SCHOOL ADMIN =================
    if current_user.role.value == 'school_admin':
        teacher_assignments = TeacherAssignment.query.join(Class).filter(
            TeacherAssignment.school_id == current_user.school_id,
            TeacherAssignment.branch_id == None
        ).order_by(TeacherAssignment.created_at.desc()).all()

    # ================= BRANCH ADMIN =================
    elif current_user.role.value == 'branch_admin':
        teacher_assignments = TeacherAssignment.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(TeacherAssignment.created_at.desc()).all()

    # Preload subjects for each assignment to avoid querying in template
    assignment_subjects = {}
    for ta in teacher_assignments:
        subjects = []
        if ta.subject_ids:
            subjects = Subject.query.filter(Subject.id.in_(ta.subject_ids)).all()
        assignment_subjects[ta.id] = subjects

    return render_template(
        "backend/pages/components/teachers/all_teacher_assignments.html",
        teacher_assignments=teacher_assignments,
        assignment_subjects=assignment_subjects,  # Pass subjects dict
        user=current_user
    )


# ================= EXPORT =================
@bp.route("/teacher-assignments/export")
@login_required
def export_teacher_assignments():
    import io, csv, json
    from flask import Response, flash, redirect, url_for

    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ---------------- Role-based Filter ----------------
    if current_user.role.value == 'school_admin':
        assignments = TeacherAssignment.query.join(Class).filter(
            Class.school_id == current_user.school_id,
            Class.branch_id == None
        ).order_by(TeacherAssignment.created_at.desc()).all()
    else:  # branch_admin
        assignments = TeacherAssignment.query.join(Class).filter(
            Class.branch_id == current_user.branch_id
        ).order_by(TeacherAssignment.created_at.desc()).all()

    output = io.StringIO()
    writer = csv.writer(output)

    # ---------------- CSV Header ----------------
    writer.writerow([
        "S/N", "Teacher", "Class", "Section", "Subjects", "School", "Branch", "Created At"
    ])

    # ---------------- CSV Data ----------------
    for idx, ta in enumerate(assignments, start=1):
        # Handle subjects
        subjects_str = ""
        if ta.subject_ids:
            # Ensure subject_ids is a list (sometimes stored as JSON string)
            subject_ids = ta.subject_ids
            if isinstance(subject_ids, str):
                try:
                    subject_ids = json.loads(subject_ids)
                except Exception:
                    subject_ids = []
            subject_names = [Subject.query.get(sid).name for sid in subject_ids if Subject.query.get(sid)]
            subjects_str = ", ".join(subject_names)

        writer.writerow([
            idx,
            ta.teacher.full_name if hasattr(ta, 'teacher') and ta.teacher else "",
            ta.class_obj.name if hasattr(ta, 'class_obj') and ta.class_obj else "",
            ta.section.name if hasattr(ta, 'section') and ta.section else "Full Class",
            subjects_str,
            ta.school.name if hasattr(ta, 'school') and ta.school else "",
            ta.branch.name if hasattr(ta, 'branch') and ta.branch else "N/A",
            ta.created_at.strftime('%Y-%m-%d %H:%M:%S') if ta.created_at else ""
        ])

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=teacher_assignments.csv"}
    )


# ================= IMPORT =================
@bp.route("/teacher-assignments/import", methods=["POST"])
@login_required
def import_teacher_assignments():
    import io, csv
    from sqlalchemy.exc import SQLAlchemyError

    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for("main.index"))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    file = request.files.get('file')
    if not file:
        flash("No file selected", "danger")
        return redirect(url_for('main.all_teacher_assignments'))

    reader = csv.DictReader(io.TextIOWrapper(file.stream, encoding='utf-8'))
    total_added = 0

    for row in reader:
        teacher_name = row.get("Teacher")
        class_name = row.get("Class")
        section_name = row.get("Section")
        subjects_str = row.get("Subjects")

        if not teacher_name or not class_name or not subjects_str:
            continue

        try:
            with db.session.no_autoflush:
                # --- Teacher ---
                teacher = Teacher.query.filter_by(full_name=teacher_name, school_id=current_user.school_id).first()
                if not teacher:
                    continue

                # --- Class ---
                cls_query = Class.query.filter_by(name=class_name, school_id=current_user.school_id)
                if current_user.role.value == 'branch_admin':
                    cls_query = cls_query.filter_by(branch_id=current_user.branch_id)
                else:
                    cls_query = cls_query.filter_by(branch_id=None)
                cls = cls_query.first()
                if not cls:
                    continue

                # --- Section ---
                section = None
                if section_name and section_name.lower() != "full class":
                    section_query = Section.query.filter_by(name=section_name, class_id=cls.id, school_id=current_user.school_id)
                    if current_user.role.value == 'branch_admin':
                        section_query = section_query.filter_by(branch_id=current_user.branch_id)
                    else:
                        section_query = section_query.filter_by(branch_id=None)
                    section = section_query.first()

                # --- Subjects ---
                subject_names = [s.strip() for s in subjects_str.split(",") if s.strip()]
                subject_objects = []
                for sub_name in subject_names:
                    subj_query = Subject.query.filter_by(name=sub_name, school_id=current_user.school_id)
                    if current_user.role.value == 'branch_admin':
                        subj_query = subj_query.filter_by(branch_id=current_user.branch_id)
                    else:
                        subj_query = subj_query.filter_by(branch_id=None)
                    subj = subj_query.first()
                    if subj:
                        subject_objects.append({"id": subj.id, "name": subj.name})

                if not subject_objects:
                    continue

                # --- Check existing teacher assignment ---
                assignment = TeacherAssignment.query.filter_by(
                    teacher_id=teacher.id,
                    class_id=cls.id,
                    section_id=section.id if section else None,
                    school_id=current_user.school_id,
                    branch_id=cls.branch_id
                ).first()

                if assignment:
                    # Merge subjects without duplicates
                    existing_ids = {s['id'] for s in assignment.subjects}
                    for obj in subject_objects:
                        if obj['id'] not in existing_ids:
                            assignment.subjects.append(obj)
                            assignment.subject_ids.append(obj['id'])
                            total_added += 1
                else:
                    # Create new assignment
                    new_assignment = TeacherAssignment(
                        teacher_id=teacher.id,
                        class_id=cls.id,
                        section_id=section.id if section else None,
                        subject_ids=[s['id'] for s in subject_objects],
                        subjects=subject_objects,
                        school_id=current_user.school_id,
                        branch_id=cls.branch_id
                    )
                    db.session.add(new_assignment)
                    total_added += len(subject_objects)

        except SQLAlchemyError as e:
            db.session.rollback()
            flash(f"❌ Error processing row: {row}. Error: {str(e)}", "danger")
            continue

    try:
        db.session.commit()
        flash(f"✅ Imported {total_added} teacher assignment subjects successfully!", "success")
    except SQLAlchemyError as e:
        db.session.rollback()
        flash(f"❌ Error committing import: {str(e)}", "danger")

    return redirect(url_for('main.all_teacher_assignments'))


@bp.route("/teacher-assignments/add", methods=["GET", "POST"])
@login_required
def add_teacher_assignment():
    form = TeacherAssignmentForm()

    # User inactive
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for("main.index"))

    # Only admin roles
    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ================= GET CLASSES & SUBJECTS FOR TEMPLATE =================
    classes = []
    subjects = []

    if current_user.role.value == "school_admin":
        classes = Class.query.filter_by(
            school_id=current_user.school_id, branch_id=None
        ).order_by(Class.name).all()
        subjects = Subject.query.filter_by(
            school_id=current_user.school_id, branch_id=None
        ).order_by(Subject.name).all()
    elif current_user.role.value == "branch_admin":
        classes = Class.query.filter_by(
            school_id=current_user.school_id, branch_id=current_user.branch_id
        ).order_by(Class.name).all()
        subjects = Subject.query.filter_by(
            school_id=current_user.school_id, branch_id=current_user.branch_id
        ).order_by(Subject.name).all()

    # ================= HANDLE POST =================
    if form.validate_on_submit():
        teacher_id = form.teacher_id.data
        school_id = current_user.school_id
        branch_id = form.branch_id.data or None
        added_count = 0

        # Parse dynamic pairs
        pairs = []
        for key in request.form:
            if key.startswith("pairs[") and "class_id" in key:
                index = key.split("[")[1].split("]")[0]
                class_id = int(request.form.get(f"pairs[{index}][class_id]"))
                section_id_raw = request.form.get(f"pairs[{index}][section_id]")
                section_id = int(section_id_raw) if section_id_raw else None
                subject_ids = request.form.getlist(f"pairs[{index}][subject_ids][]")
                subject_ids = [int(s) for s in subject_ids]
                pairs.append({"class_id": class_id, "section_id": section_id, "subject_ids": subject_ids})

        if not pairs:
            flash("❌ Please add at least one Class + Subject pair.", "danger")
            return render_template(
                "backend/pages/components/teachers/add_teacher_assignment.html",
                form=form,
                user=current_user,
                classes=classes,
                subjects=subjects
            )

        # Save assignments
        for pair in pairs:
            cls = Class.query.get(pair["class_id"])
            if not cls:
                continue

            # Role ownership check
            if current_user.role.value == "school_admin" and (cls.school_id != school_id or cls.branch_id is not None):
                continue
            if current_user.role.value == "branch_admin" and cls.branch_id != current_user.branch_id:
                continue

            section_id = pair["section_id"]  # None if full class
            existing = TeacherAssignment.query.filter_by(
                teacher_id=teacher_id,
                class_id=cls.id,
                section_id=section_id,
                school_id=school_id,
                branch_id=branch_id
            ).first()

            if existing:
                merged_subjects = list(set(existing.subject_ids + pair["subject_ids"]))
                if merged_subjects != existing.subject_ids:
                    existing.subject_ids = merged_subjects
                    existing.subjects = merged_subjects
                    db.session.commit()
                    added_count += len(pair["subject_ids"])
                continue

            assignment = TeacherAssignment(
                teacher_id=teacher_id,
                class_id=cls.id,
                section_id=section_id,
                subject_ids=pair["subject_ids"],
                subjects=pair["subject_ids"],
                school_id=school_id,
                branch_id=branch_id
            )
            db.session.add(assignment)
            added_count += len(pair["subject_ids"])

        try:
            db.session.commit()
            flash(f"✅ {added_count} teacher assignment(s) created/updated successfully!", "success")
            return redirect(url_for("main.all_teacher_assignments"))
        except Exception as e:
            db.session.rollback()
            print("ERROR saving TeacherAssignments:", e)
            flash("❌ Error saving assignments. Check console.", "danger")

    # ================= RENDER TEMPLATE =================
    return render_template(
        "backend/pages/components/teachers/add_teacher_assignment.html",
        form=form,
        user=current_user,
        classes=classes,
        subjects=subjects
    )


@bp.route("/teacher-assignments/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_teacher_assignment(id):
    # 1. Soo qaado assignment-ka hadda la rabo in la edit gareeyo
    assignment = TeacherAssignment.query.get_or_404(id)

    # Permission checks
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Soo qaado form-ka (Teacher dropdown-ka wuxuu ahaanayaa sidii uu ahaa)
    form = TeacherAssignmentForm(obj=assignment)
    
    # 2. Classes & Subjects Filter (Sida uu qofku u arko dropdown-ka)
    branch_id = current_user.branch_id if current_user.role.value == "branch_admin" else None
    classes = Class.query.filter_by(school_id=current_user.school_id, branch_id=branch_id).order_by(Class.name).all()
    subjects = Subject.query.filter_by(school_id=current_user.school_id, branch_id=branch_id).order_by(Subject.name).all()

    # ---------------------------------------------------------
    # 3. POST LOGIC: Marka la taabto "Update Changes"
    # ---------------------------------------------------------
    if form.validate_on_submit():
        try:
            # A. Sifee/Tirtir dhamaan assignments-ka macalinkan ee dugsigan/branch-kan
            # Kani waa qaybta "Clear kii hore"
            TeacherAssignment.query.filter_by(
                teacher_id=assignment.teacher_id,
                school_id=current_user.school_id,
                branch_id=branch_id
            ).delete()

            # B. Parse garee "pairs"-ka cusub ee ka yimid foomka dynamic-ga ah
            added_count = 0
            # Waxaan dhex mushaaxaynaa request.form si aan u helno pairs-ka
            # Qaabka: pairs[0][class_id], pairs[0][subject_ids][]
            indices = set()
            for key in request.form.keys():
                if key.startswith("pairs["):
                    index = key.split("[")[1].split("]")[0]
                    indices.add(index)

            for i in indices:
                class_id = request.form.get(f"pairs[{i}][class_id]")
                section_id_raw = request.form.get(f"pairs[{i}][section_id]")
                subject_ids = request.form.getlist(f"pairs[{i}][subject_ids][]")

                if class_id and subject_ids:
                    # Abuuro assignment cusub
                    new_assign = TeacherAssignment(
                        teacher_id=assignment.teacher_id,
                        class_id=int(class_id),
                        section_id=int(section_id_raw) if section_id_raw else None,
                        subject_ids=[int(sid) for sid in subject_ids],
                        subjects=[int(sid) for sid in subject_ids], # Haddii aad labada field isticmaasho
                        school_id=current_user.school_id,
                        branch_id=branch_id
                    )
                    db.session.add(new_assign)
                    added_count += 1

            # C. Commit changes
            db.session.commit()
            flash(f"✅ Si guul leh ayaa loo update gareeyay {added_count} class assignment.", "success")
            return redirect(url_for('main.all_teacher_assignments'))

        except Exception as e:
            db.session.rollback()
            print("Update Error:", e)
            flash(f"❌ Khalad ayaa dhacay intii lagu guda jiray update-ka: {str(e)}", "danger")

    # ---------------------------------------------------------
    # 4. GET LOGIC: Soo aruuri xogtii hore si loogu muujiyo Form-ka
    # ---------------------------------------------------------
    existing_assignments = TeacherAssignment.query.filter_by(
        teacher_id=assignment.teacher_id,
        school_id=assignment.school_id,
        branch_id=branch_id
    ).all()

    existing_pairs = []
    for ea in existing_assignments:
        existing_pairs.append({
            "class_id": ea.class_id,
            "section_id": ea.section_id,
            "subject_ids": [int(sid) for sid in ea.subject_ids] if ea.subject_ids else []
        })

    return render_template(
        "backend/pages/components/teachers/edit_teacher_assignment.html",
        form=form,
        classes=classes,
        subjects=subjects,
        user=current_user,
        existing_pairs_json=json.dumps(existing_pairs)
    )


@bp.route("/teacher-assignments/delete/<int:id>", methods=["POST"])
@login_required
def delete_teacher_assignment(id):
    assignment = TeacherAssignment.query.get_or_404(id)

    # Role-based access
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    if current_user.role.value == 'school_admin':
        if assignment.school_id != current_user.school_id:
            flash("You cannot delete this assignment.", "danger")
            return redirect(url_for('main.all_teacher_assignments'))
    elif current_user.role.value == 'branch_admin':
        if assignment.branch_id != current_user.branch_id:
            flash("You cannot delete this assignment.", "danger")
            return redirect(url_for('main.all_teacher_assignments'))

    try:
        db.session.delete(assignment)
        db.session.commit()
        flash("✅ Teacher assignment deleted successfully.", "success")
    except Exception as e:
        db.session.rollback()
        print("ERROR deleting TeacherAssignment:", e)
        flash("❌ Error deleting assignment.", "danger")

    return redirect(url_for('main.all_teacher_assignments'))



#------------------------------------------
#--------------- Students Fee Collections
#------------------------------------------



@bp.route("/all-fee-collections")
@login_required
def all_fee_collections():
    # Check user status
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    # Only allow school_admin or branch_admin
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ================= SCHOOL ADMIN =================
    if current_user.role.value == 'school_admin':
        fee_collections = StudentFeeCollection.query.filter(
            StudentFeeCollection.school_id == current_user.school_id,
            StudentFeeCollection.branch_id == None
        ).order_by(StudentFeeCollection.created_at.desc()).all()

    # ================= BRANCH ADMIN =================
    elif current_user.role.value == 'branch_admin':
        fee_collections = StudentFeeCollection.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(StudentFeeCollection.created_at.desc()).all()

    # Preload related entities to avoid querying in template
    collection_details = {}
    for fc in fee_collections:
        collection_details[fc.id] = {
            "student": fc.student,
            "class": fc.class_obj,
            "section": fc.section,
            "school": fc.school,
            "branch": fc.branch
        }

    return render_template(
        "backend/pages/components/fee_collections/all_fee_collections.html",
        fee_collections=fee_collections,
        collection_details=collection_details,
        user=current_user
    )


# ---------------- INVOICE NUMBER (DYNAMIC WITH YEAR) ----------------
def generate_invoice_number(school_id, branch_id=None):
    # 1. Soo qaado Sanadka hadda taagan (Tusaale: 2026 -> 26)
    # %y waxay soo saartaay labada nambar ee u dambeeya sanadka (26)
    current_year_short = datetime.now().strftime("%y")

    # 2. Go'aami Magaca loo isticmaalayo Prefix-ka koowaad
    if branch_id:
        branch = db.session.get(Branch, branch_id)
        name_to_use = branch.name if branch else "BR"
    else:
        school = db.session.get(School, school_id)
        name_to_use = school.name if school else "SCH"

    # 3. Dhish Prefix-ka (Tusaale: JC)
    base_prefix = generate_school_prefix(name_to_use)
    
    # Isku gee Prefix-ka iyo Sanadka (Tusaale: JC + 26 + - -> JC26-)
    full_prefix = f"{base_prefix}{current_year_short}-"

    # 4. Soo qaado Invoice-kii ugu dambeeyay ee isla prefix-kan wata (JC26-)
    last_invoice = FeeInvoice.query.filter(
        FeeInvoice.invoice_number.like(f"{full_prefix}%")
    ).order_by(FeeInvoice.id.desc()).first()

    if last_invoice and last_invoice.invoice_number:
        try:
            # Ka jar JC26-, ka dibna u beddel lambarka haray 'int'
            parts = last_invoice.invoice_number.split("-")
            if len(parts) > 1:
                last_number = int(parts[-1])
            else:
                last_number = 0
        except (ValueError, IndexError):
            last_number = 0
    else:
        # Haddii sanad cusub la galo ama invoice kii u horeeyay yahay
        last_number = 0

    # 5. Kordhi lambarka
    new_number = last_number + 1

    # 6. Soo celi qaabka: JC26-00001
    return f"{full_prefix}{str(new_number).zfill(5)}"


# ---------------- DESCRIPTION ----------------
def add_generate_description():
    now = now_eat
    month_name = now.strftime("%B")
    year = now.year
    return f"{month_name} {year} - {year + 1}"

# ---------------- DESCRIPTION ----------------
def add_generate_description():
    # Halkan haddii now_eat uu yahay function, waa in loo wacaa ()
    now = now_eat() if callable(now_eat) else now_eat
    month_name = now.strftime("%B")
    year = now.year
    return f"{month_name} {year} - {year + 1}"

def generate_description(text=""):
    now = now_eat() if callable(now_eat) else now_eat
    month_name = now.strftime("%B")
    year = now.year
    if text:
        return f"{text} - {month_name} {year}-{year + 1}"
    return f"{month_name} {year}-{year + 1}"



# ================= ROUTE =================
@bp.route("/fee-collections/add", methods=["GET", "POST"])
@login_required
def add_fee_collection():
    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    form = StudentFeeCollectionForm()

    # ------------------- STUDENTS INFO -------------------
    students_info = {}
    students = Student.query.filter_by(
        school_id=current_user.school_id,
        branch_id=(current_user.branch_id if current_user.role.value == "branch_admin" else None)
    ).order_by(Student.full_name).all()

    for s in students:
        reg_f = float(s.registration_fee or 0)
        tui_f = float(s.price or 0)
        students_info[s.id] = {
            "roll_no": s.roll_no,
            "registration_fee": reg_f,
            "tuition_fee": tui_f,
            "total_fee": reg_f + tui_f
        }

    # ------------------- PROCESS FORM -------------------
    if form.validate_on_submit():
        try:
            student = Student.query.get_or_404(form.student_id.data)
            registration_fee = float(student.registration_fee or 0)
            tuition_fee = float(student.price or 0)
            input_amount = float(form.amount_paid.data or 0)
            
            today = now_eat() if callable(now_eat) else now_eat

            # 1️⃣ HEL BAAQIGII U DAMBEEYAY (Previous Balance)
            last_invoice = FeeInvoice.query.join(StudentFeeCollection)\
                .filter(StudentFeeCollection.student_id == student.id)\
                .order_by(FeeInvoice.id.desc()).first()
            
            running_balance = float(last_invoice.balance if last_invoice else 0.0)

            # 2️⃣ GET AMA CREATE FEE RECORD (Summary Record for Current Month)
            fee_record = StudentFeeCollection.query.filter(
                StudentFeeCollection.student_id == student.id,
                extract("year", StudentFeeCollection.payment_date) == today.year,
                extract("month", StudentFeeCollection.payment_date) == today.month
            ).first()

            is_new_record = False
            if not fee_record:
                is_new_record = True
                fee_record = StudentFeeCollection(
                    student_id=student.id,
                    class_id=student.class_id,
                    section_id=student.section_id,
                    school_id=current_user.school_id,
                    branch_id=(current_user.branch_id if current_user.role.value == "branch_admin" else None),
                    amount_due=(registration_fee + tuition_fee),
                    amount_paid=0,
                    payment_status="Pending",
                    payment_date=form.payment_date.data or today
                )
                db.session.add(fee_record)
                db.session.flush() # Flush si aan u helno fee_record.id

            # ---------------------------------------------------------
            # 3️⃣ TRANSACTIONS: (Receipt First -> Then Charges)
            # ---------------------------------------------------------

            # A. Safka Payment-ka (Receipt)
            # Formula: Balance = Previous - Receipt
            if input_amount > 0:
                running_balance -= input_amount
                
                inv_pay = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD

                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Payment Memo', 
                    amount_due=0,             # Charge column (Eber)
                    amount_paid=input_amount,  # Receipt column
                    balance=running_balance,   # Current Balance
                    description=f"Payment Received - {add_generate_description()}"
                )
                db.session.add(inv_pay)

                # --- UPDATE STUDENT TABLE: REGISTRATION FEE TO 0 ---
                # Haddii uu jiro registration fee lana bixiyay lacag, u badal 0 si uusan bisha dambe u deynaysan
                if registration_fee > 0:
                    student.registration_fee = 0
                    # Waxaan sidoo kale update gareynaynaa 'total' column-ka Student table haddii uu jiro
                    student.total = float(student.price or 0) 
                
                db.session.flush()

            # B. Safka Registration-ka (Charge)
            # Formula: Balance = Previous + Charge
            if is_new_record and registration_fee > 0:
                running_balance += registration_fee 
                
                inv_reg = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD
    
                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Registration',
                    amount_due=registration_fee, # Charge column
                    amount_paid=0,               # Receipt column
                    balance=running_balance,     # Current Balance
                    description=generate_description("Registration Fee")
                )
                db.session.add(inv_reg)
                db.session.flush()

            # C. Safka Tuition-ka (Charge)
            # Formula: Balance = Previous + Charge
            if is_new_record and tuition_fee > 0:
                running_balance += tuition_fee
                
                inv_tui = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD

                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Tuition',
                    amount_due=tuition_fee,   # Charge column
                    amount_paid=0,            # Receipt column
                    balance=running_balance,  # Current Balance
                    description=generate_description("Tuition Fee")
                )
                db.session.add(inv_tui)

            # ---------------------------------------------------------
            # 4️⃣ UPDATE SUMMARY (Fee Collection Record)
            # ---------------------------------------------------------
            fee_record.amount_paid += input_amount
            fee_record.remaining_balance = running_balance 

            # Cusboonaysii Payment Status
            if fee_record.amount_paid >= fee_record.amount_due:
                fee_record.payment_status = "Paid"
            elif fee_record.amount_paid > 0:
                fee_record.payment_status = "Partial"
            else:
                fee_record.payment_status = "Pending"

            db.session.commit()
            flash(f"✅ Transaction processed. Registration updated. Final Balance: {running_balance}", "success")
            return redirect(url_for("main.all_fee_collections"))

        except Exception as e:
            db.session.rollback()
            print(f"Error: {e}")
            flash(f"❌ Error: {str(e)}", "danger")

    return render_template(
        "backend/pages/components/fee_collections/add_fee_collection.html",
        form=form,
        students_info=students_info,
        user=current_user,
        current_date = now_eat().date().isoformat(),
    )

    

@bp.route("/fee-collections/edit/<int:fee_id>", methods=["GET", "POST"])
@login_required
def edit_fee_collection(fee_id):
    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # Soo qaado record-ka la edit-gareynayo
    fee_record = StudentFeeCollection.query.get_or_404(fee_id)
    student = fee_record.student

    # Isticmaal Form-ka (Fill-garee xogta jirta)
    form = StudentFeeCollectionForm(obj=fee_record)

    # ------------------- FIXED LOGIC -------------------
    # Halkan waxaa lagu saxay TypeError-ka: float() ayaa lagu daray labada dhinac
    # Waxaan ka dhigaynaa float() si aan loogu dhicin khilaafka 'float' and 'decimal.Decimal'
    
    current_amount_due = float(fee_record.amount_due or 0)
    current_student_price = float(student.price or 0)
    
    original_reg_fee = current_amount_due - current_student_price
    
    if original_reg_fee < 0: 
        original_reg_fee = 0.0
    
    tuition_fee = current_student_price
    total_fee = original_reg_fee + tuition_fee

    if request.method == "GET":
        form.student_id.data = fee_record.student_id
        form.amount_paid.data = fee_record.amount_paid
        form.payment_date.data = fee_record.payment_date
        form.remarks.data = fee_record.remarks

    # Macluumaadka template-ka loo dirayo
    students_info = {
        student.id: {
            "roll_no": student.roll_no,
            "registration_fee": original_reg_fee,
            "tuition_fee": tuition_fee,
            "total_fee": total_fee
        }
    }

    invoices = FeeInvoice.query.filter_by(student_fee_id=fee_record.id).all()

    if form.validate_on_submit():
        try:
            input_amount = float(form.amount_paid.data or 0)

            # 1️⃣ HEL BAAQIGII HORE (Previous Balance)
            last_invoice_before = FeeInvoice.query.join(StudentFeeCollection)\
                .filter(
                    StudentFeeCollection.student_id == student.id,
                    StudentFeeCollection.id < fee_record.id
                ).order_by(FeeInvoice.id.desc()).first()
            
            running_balance = float(last_invoice_before.balance if last_invoice_before else 0.0)

            # 2️⃣ TIRTIR INVOICES-KII HORE
            old_invoices = FeeInvoice.query.filter_by(student_fee_id=fee_record.id).all()
            for inv in old_invoices:
                db.session.delete(inv)

            db.session.flush()

           
            # A. Receipt First
            if input_amount > 0:
                running_balance -= input_amount
                inv_pay = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD
                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Payment Memo', 
                    amount_due=0, 
                    amount_paid=input_amount,
                    balance=running_balance,
                    description=f"Payment Received - {add_generate_description()}",
                    extra_info=""
                )
                db.session.add(inv_pay)
                
                # UPDATE STUDENT REGISTRATION TO 0
                if float(student.registration_fee or 0) > 0:
                    student.registration_fee = 0
                    student.total = float(student.price or 0)

            # B. Registration Charge
            if original_reg_fee > 0:
                running_balance += original_reg_fee
                inv_reg = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD
                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Registration',
                    amount_due=original_reg_fee,
                    amount_paid=0,
                    balance=running_balance,
                    description=generate_description("Registration Fee"),
                    extra_info="Charge for Registration"
                )
                db.session.add(inv_reg)

            # C. Tuition Charge
            if tuition_fee > 0:
                running_balance += tuition_fee
                inv_tui = FeeInvoice(
                    student_fee_id=fee_record.id,
                    school_id=fee_record.school_id,      # ✅ ADD
                    branch_id=fee_record.branch_id,      # ✅ ADD
                    invoice_number=generate_invoice_number(current_user.school_id, current_user.branch_id),
                    type='Tuition',
                    amount_due=tuition_fee,
                    amount_paid=0,
                    balance=running_balance,
                    description=generate_description("Tuition Fee"),
                    extra_info="Charge for Tuition"
                )
                db.session.add(inv_tui)

            # 4️⃣ UPDATE SUMMARY RECORD
            fee_record.amount_due = total_fee
            fee_record.amount_paid = input_amount
            fee_record.remaining_balance = running_balance 
            
            if fee_record.amount_paid >= fee_record.amount_due:
                fee_record.payment_status = 'Paid'
            elif fee_record.amount_paid > 0:
                fee_record.payment_status = 'Partial'
            else:
                fee_record.payment_status = 'Pending'

            fee_record.payment_date = form.payment_date.data
            fee_record.remarks = form.remarks.data

            db.session.commit()
            flash(f"✅ Collection updated successfully. New Balance: {running_balance}", "success")
            return redirect(url_for("main.all_fee_collections"))

        except Exception as e:
            db.session.rollback()
            flash(f"❌ Error updating fee collection: {str(e)}", "danger")

    return render_template(
        "backend/pages/components/fee_collections/edit_fee_collection.html",
        form=form,
        fee_record=fee_record,
        students_info=students_info,
        invoices=invoices,
        user=current_user,
        current_date = now_eat().date().isoformat()

    )


@bp.route("/delete/fee-collection/<int:fee_id>", methods=["POST"])
@login_required
def delete_fee_collection(fee_id):
    fee_record = StudentFeeCollection.query.get_or_404(fee_id)

    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ✅ Save student name before deletion
    student_name = fee_record.student.full_name if fee_record.student else "Unknown Student"

    try:
        with db.session.no_autoflush:
            db.session.delete(fee_record)
            db.session.commit()
        flash(
            f"✅ Fee collection for {student_name} deleted successfully, along with all related invoices.",
            "success"
        )
    except Exception as e:
        db.session.rollback()
        print("Delete ERROR:", e)
        flash("❌ Error deleting fee collection. Check console.", "danger")

    return redirect(url_for("main.all_fee_collections"))


@bp.route("/fee-collections/print/<int:fee_id>")
@login_required
def print_student_fee_collection(fee_id):
    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- FETCH FEE COLLECTION -------------------
    fee_record = StudentFeeCollection.query.get_or_404(fee_id)
    student = fee_record.student
    invoices = FeeInvoice.query.filter_by(student_fee_id=fee_record.id).all()

    last_invoice_before = FeeInvoice.query.join(StudentFeeCollection)\
        .filter(
            StudentFeeCollection.student_id == student.id,
            StudentFeeCollection.id < fee_record.id
        ).order_by(FeeInvoice.id.desc()).first()
    
    previous_balance = float(last_invoice_before.balance if last_invoice_before else 0.0)

    # 2️⃣ XISAABI TOTALS-KA SI SAX AH (Fixing Float/Decimal issues)
    # Total Fee-ga bishan: Registration + Tuition
    current_month_total = float(fee_record.amount_due or 0)
    
    # Lacagta bishan la bixiyay
    amount_paid = float(fee_record.amount_paid or 0)
    
    # Baaqiga hadda u soo haray ardayga (Final Balance)
    # Waxaan ka soo qaadanaynaa invoice-ka u dambeeyay ee record-kan
    last_invoice_in_record = invoices[-1] if invoices else None
    current_balance = float(last_invoice_in_record.balance if last_invoice_in_record else previous_balance)

    # Total fees
    total_fee = float(student.registration_fee or 0) + float(student.price or 0)
    
    # Remaining balance (what student still owes)
    remaining_balance = max(0.0, total_fee - float(fee_record.amount_paid or 0))
    
    # Overpayment / Receipt (if student paid more than total fee)
    receipt_amount = max(0.0, float(fee_record.amount_paid or 0) - total_fee)

    return render_template(
        "backend/pages/components/fee_collections/print_fee_collection.html",
        fee_record=fee_record,
        student=student,
        invoices=invoices,
        total_fee=total_fee,
        remaining_balance=remaining_balance,
        receipt_amount=receipt_amount,
        user=current_user,
        previous_balance=previous_balance,
        current_month_total=current_month_total,
        amount_paid=amount_paid,
        current_balance=current_balance,
    )


@bp.route("/fee-collections/export-all")
@login_required
def export_fee_collections_excel():
    import io
    import pandas as pd
    from flask import Response

    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.dashboard'))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ------------------- ROLE-BASED FILTER -------------------
    if current_user.role.value == 'school_admin':
        fee_records = StudentFeeCollection.query.filter(
            StudentFeeCollection.school_id == current_user.school_id,
            StudentFeeCollection.branch_id == None
        ).order_by(StudentFeeCollection.payment_date.desc()).all()
    else:  # branch_admin
        fee_records = StudentFeeCollection.query.filter(
            StudentFeeCollection.branch_id == current_user.branch_id
        ).order_by(StudentFeeCollection.payment_date.desc()).all()

    # ------------------- FEE COLLECTION SHEET -------------------
    fee_collection_data = []
    for idx, record in enumerate(fee_records, start=1):
        fee_collection_data.append({
            "S/N": idx,
            "Student": record.student.full_name if record.student else "",
            "Class": record.class_obj.name if record.class_obj else "",
            "Section": record.section.name if record.section else "Full Class",
            "Amount Due": f"{record.amount_due:.2f}" if record.amount_due else "0.00",
            "Amount Paid": f"{record.amount_paid:.2f}" if record.amount_paid else "0.00",
            "Remaining Balance": f"{record.remaining_balance:.2f}" if record.remaining_balance else "0.00",
            "Payment Status": record.payment_status or "",
            "Payment Date": record.payment_date.strftime('%Y-%m-%d %H:%M:%S') if record.payment_date else "",
            "Remarks": record.remarks or ""
        })

    df_fee_collection = pd.DataFrame(fee_collection_data)

    # ------------------- FEE INVOICE SHEET -------------------
    fee_invoice_data = []
    idx = 1
    for record in fee_records:
        invoices = FeeInvoice.query.filter_by(student_fee_id=record.id).all()
        for inv in invoices:
            fee_invoice_data.append({
                "S/N": idx,
                "Student": record.student.full_name if record.student else "",
                "Class": record.class_obj.name if record.class_obj else "",
                "Section": record.section.name if record.section else "Full Class",
                "Invoice Type": inv.type or "",
                "Amount Due": f"{inv.amount_due:.2f}" if inv.amount_due else "0.00",
                "Amount Paid": f"{inv.amount_paid:.2f}" if inv.amount_paid else "0.00",
                "Total Amount": f"{inv.total_amount:.2f}" if inv.total_amount else "0.00",
                "Remaining Balance": f"{inv.remaining_balance:.2f}" if inv.remaining_balance else "0.00",
                "Description": inv.description or "",
                "Extra Info": inv.extra_info or "",
                "Payment Date": record.payment_date.strftime('%Y-%m-%d %H:%M:%S') if record.payment_date else "",
                "Payment Status": record.payment_status or ""
            })
            idx += 1

    df_fee_invoice = pd.DataFrame(fee_invoice_data)

    # ------------------- WRITE TO EXCEL IN MEMORY -------------------
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_fee_collection.to_excel(writer, sheet_name='Fee Collection', index=False)
        df_fee_invoice.to_excel(writer, sheet_name='Fee Invoice', index=False)
    output.seek(0)

    return Response(
        output.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=fee_collections_full.xlsx"}
    )


@bp.route("/fee-collections/import", methods=["POST"])
@login_required
def import_fee_collections_excel():
    import pandas as pd
    from werkzeug.utils import secure_filename
    import numpy as np

    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for("main.all_fee_collections"))

    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.all_fee_collections"))

    # ------------------- GET FILE -------------------
    file = request.files.get("file")
    if not file:
        flash("No file uploaded.", "danger")
        return redirect(url_for("main.all_fee_collections"))

    filename = secure_filename(file.filename)
    if not filename.endswith((".xlsx", ".xls")):
        flash("Please upload a valid Excel file (.xlsx/.xls).", "danger")
        return redirect(url_for("main.all_fee_collections"))

    try:
        # Read Excel file
        xls = pd.ExcelFile(file)

        # ------------------- FEE COLLECTION SHEET -------------------
        if "Fee Collection" in xls.sheet_names:
            df_fee_collection = pd.read_excel(xls, sheet_name="Fee Collection")
            df_fee_collection.fillna("", inplace=True)

            for _, row in df_fee_collection.iterrows():
                student_name = str(row.get("Student")).strip()
                if not student_name:
                    continue

                student = Student.query.filter_by(full_name=student_name).first()
                if not student:
                    continue

                fee_record = StudentFeeCollection.query.filter_by(
                    student_id=student.id,
                    class_id=student.class_id,
                    section_id=student.section_id,
                    school_id=current_user.school_id,
                    branch_id=current_user.branch_id if current_user.role.value == "branch_admin" else None
                ).first()

                if not fee_record:
                    fee_record = StudentFeeCollection(
                        student_id=student.id,
                        class_id=student.class_id,
                        section_id=student.section_id,
                        school_id=current_user.school_id,
                        branch_id=current_user.branch_id if current_user.role.value == "branch_admin" else None
                    )
                    db.session.add(fee_record)

                fee_record.amount_due = float(row.get("Amount Due") or 0)
                fee_record.amount_paid = float(row.get("Amount Paid") or 0)
                fee_record.remaining_balance = float(row.get("Remaining Balance") or 0)
                fee_record.payment_status = str(row.get("Payment Status") or "Pending")
                payment_date = row.get("Payment Date")
                fee_record.payment_date = pd.to_datetime(payment_date) if pd.notna(payment_date) else None
                remarks = row.get("Remarks")
                fee_record.remarks = str(remarks) if pd.notna(remarks) else ""

        # ------------------- FEE INVOICE SHEET -------------------
        if "Fee Invoice" in xls.sheet_names:
            df_fee_invoice = pd.read_excel(xls, sheet_name="Fee Invoice")
            df_fee_invoice.fillna("", inplace=True)

            for _, row in df_fee_invoice.iterrows():
                student_name = str(row.get("Student")).strip()
                if not student_name:
                    continue

                student = Student.query.filter_by(full_name=student_name).first()
                if not student:
                    continue

                fee_record = StudentFeeCollection.query.filter_by(
                    student_id=student.id,
                    class_id=student.class_id,
                    section_id=student.section_id,
                    school_id=current_user.school_id,
                    branch_id=current_user.branch_id if current_user.role.value == "branch_admin" else None
                ).first()
                if not fee_record:
                    continue

                # ------------------- CREATE INVOICE -------------------
                invoice_number = generate_invoice_number(current_user.school_id,
                                                         current_user.branch_id if current_user.role.value == "branch_admin" else None)

                invoice = FeeInvoice(
                    student_fee_id=fee_record.id,
                    invoice_number=invoice_number,
                    type=str(row.get("Invoice Type") or "Unknown"),
                    amount_due=float(row.get("Amount Due") or 0),
                    amount_paid=float(row.get("Amount Paid") or 0),
                    total_amount=float(row.get("Total Amount") or 0),
                    remaining_balance=float(row.get("Remaining Balance") or 0),
                    description=str(row.get("Description") or ""),
                    extra_info=str(row.get("Extra Info") or "")
                )
                db.session.add(invoice)

        db.session.commit()
        flash("✅ Excel imported successfully.", "success")
    except Exception as e:
        db.session.rollback()
        print("Import ERROR:", e)
        flash(f"❌ Error importing Excel: {str(e)}", "danger")

    return redirect(url_for("main.all_fee_collections"))

@bp.route("/fee-invoices/edit/<int:invoice_id>", methods=["POST"])
@login_required
def edit_invoice(invoice_id):
    invoice = FeeInvoice.query.get_or_404(invoice_id)
    fee_record = invoice.fee_collection  # student's fee collection
    student = fee_record.student
    # get first invoice (or specific one)
    fee_invoice = fee_record.invoices[0] if fee_record.invoices else None

    # Permissions check
    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # Get data from modal form
    new_amount = float(request.form.get("registration_fee", 0))
    remarks = request.form.get("remarks", "")

    # Update **Student's fee collection amount**, not the invoice
    if new_amount != float(student.registration_fee or 0):
        old_amount = float(student.registration_fee or 0)
        student.registration_fee = new_amount

        # Update total in fee collection
        tuition_fee = float(student.tuition_fee or 0)
        fee_record.amount_due = new_amount + tuition_fee

        # Recalculate remaining balance
        already_paid = float(fee_record.amount_paid or 0)
        fee_record.payment_status = (
            "Paid" if already_paid >= fee_record.amount_due
            else "Partial" if already_paid > 0
            else "Pending"
        )
        fee_record.remaining_balance = fee_record.amount_due - already_paid

    # Update invoice remarks
    invoice.description = remarks

    try:
        db.session.commit()
        flash("✅ Invoice & Fee updated successfully", "success")
    except Exception as e:
        db.session.rollback()
        print("ERROR:", e)
        flash("❌ Error updating invoice/fee", "danger")

    return redirect(url_for("main.edit_fee_collection",fee_invoice=fee_invoice, fee_id=fee_record.id))


#---------------------------- Fees Reports

@bp.route("/school-reports_fees")
@login_required
def school_reports_fees():
    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- FETCH FEE COLLECTION -------------------
 
    return render_template(
        "backend/pages/components/reports/school-reports-sales.html",
        user=current_user,
    )

@bp.route("/branches-reports_fees")
@login_required
def branches_reports_fees():
    # ------------------- PERMISSIONS -------------------
    if current_user.status == 0:
        flash("Account inactive", "danger")
        return redirect(url_for("main.dashboard"))

    if current_user.role.value not in ["school_admin", "branch_admin"]:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- FETCH FEE COLLECTION -------------------
 
    return render_template(
        "backend/pages/components/reports/branches-reports-sales.html",
        user=current_user,
    )







#-----------------------------------------
#------------ Attendance Students
#-----------------------------------------

# -----------------------------
# GET SECTIONS (Dynamic AJAX)
# -----------------------------
# -----------------------------
# GET SECTIONS for selected class (AJAX)
# -----------------------------
@bp.route("/get-attendance-sections/<int:class_id>")
@login_required
def get_attendance_sections(class_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if not teacher:
        return jsonify([])

    assignments = TeacherAssignment.query.filter_by(teacher_id=teacher.id, class_id=class_id).all()
    section_ids = set(a.section_id for a in assignments if a.section_id is not None)
    sections = Section.query.filter(Section.id.in_(section_ids)).all()
    return jsonify([{"id": s.id, "name": s.name} for s in sections])

# -----------------------------
# GET STUDENTS (AJAX)
# -----------------------------
@bp.route("/get-students")
@login_required
def get_students():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get("class_id", type=int)
    section_id = request.args.get("section_id", type=int)
    if not class_id:
        return jsonify([])

    query = Student.query.filter_by(class_id=class_id)
    if section_id:
        query = query.filter_by(section_id=section_id)

    students = query.order_by(Student.full_name).all()
    return jsonify([{"id": s.id, "name": s.full_name, "roll_no": s.roll_no} for s in students])


# -----------------------------
# GET SUBJECTS (AJAX)
# -----------------------------
@bp.route("/get-subjects")
@login_required
def get_subjects():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get("class_id", type=int)
    section_id = request.args.get("section_id", type=int)

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if not teacher or not class_id:
        return jsonify([])

    assignments = TeacherAssignment.query.filter_by(teacher_id=teacher.id, class_id=class_id).all()
    subject_ids = set()
    for a in assignments:
        if section_id is None or a.section_id == section_id:
            subject_ids.update(a.subject_ids)

    if not subject_ids:
        return jsonify([])

    subjects = Subject.query.filter(Subject.id.in_(subject_ids)).all()
    return jsonify([{"id": s.id, "name": s.name} for s in subjects])

# -----------------------------
# TAKE ATTENDANCE
# -----------------------------

# -----------------------------
# TAKE ATTENDANCE
# -----------------------------

@bp.route("/attendance/take", methods=["GET", "POST"])
@login_required
def take_attendance():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = AttendanceForm()
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()

    # Security
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for("main.index"))

    if current_user.role.value != "teacher":
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    if not teacher:
        flash("❌ Teacher profile not found.", "danger")
        return redirect(url_for("main.dashboard"))

    if request.method == "POST":
        try:
            # ✅ IMPORTANT FIXES
            class_id = int(request.form.get("class_id"))

            section_id = request.form.get("section_id")
            section_id = int(section_id) if section_id and section_id != "-" else None

            subject_id = int(request.form.get("subject_id"))

            today = date.today()
            school_id = current_user.school_id
            branch_id = current_user.branch_id

            # =========================
            # CHECK (SUBJECT PER DAY ONLY)
            # =========================
            query = StudentAttendance.query.filter(
                StudentAttendance.teacher_id == teacher.id,
                StudentAttendance.subject_id == subject_id,
                StudentAttendance.school_id == school_id,
                StudentAttendance.branch_id == branch_id,
                StudentAttendance.date == today
            )

            if section_id is None:
                query = query.filter(StudentAttendance.section_id.is_(None))
            else:
                query = query.filter(StudentAttendance.section_id == section_id)

            if query.first():
                flash("❌ Attendance already submitted for this subject today.", "warning")
                return redirect(url_for("attendance.take_attendance"))

            # =========================
            # SAVE STUDENTS
            # =========================
            index = 0
            while True:
                student_key = f"students[{index}][student_id]"
                status_key = f"students[{index}][status]"

                if student_key not in request.form:
                    break

                student_id = int(request.form[student_key])
                status = request.form[status_key]

                new_attendance = StudentAttendance(
                    student_id=student_id,
                    class_id=class_id,
                    section_id=section_id,
                    subject_id=subject_id,
                    teacher_id=teacher.id,
                    school_id=school_id,
                    branch_id=branch_id,
                    status=status,
                    date=today
                )

                db.session.add(new_attendance)
                index += 1

            db.session.commit()

            flash("✅ Attendance saved successfully!", "success")
            return redirect(url_for("attendance.take_attendance"))

        except Exception as e:
            db.session.rollback()
            print("ERROR:", e)
            flash("❌ Error saving attendance.", "danger")

    return render_template(
        "backend/pages/components/attendance/take_attendance.html",
        form=form,
        user=current_user
    )


# -----------------------------
# AJAX: check attendance per SUBJECT
# -----------------------------
@bp.route("/check-attendance")
@login_required
def check_attendance():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin","teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get("class_id", type=int)
    section_id = request.args.get("section_id")
    subject_id = request.args.get("subject_id", type=int)
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    today = date.today()

    if not teacher or not subject_id:
        return jsonify({"submitted": False})

    section_id = int(section_id) if section_id and section_id != "-" else None

    existing = StudentAttendance.query.filter(
        StudentAttendance.teacher_id == teacher.id,
        StudentAttendance.subject_id == subject_id,
        StudentAttendance.school_id == current_user.school_id,
        StudentAttendance.branch_id == current_user.branch_id,
        StudentAttendance.date == today
    )
    if section_id is None:
        existing = existing.filter(StudentAttendance.section_id.is_(None))
    else:
        existing = existing.filter(StudentAttendance.section_id == section_id)

    return jsonify({"submitted": bool(existing.first())})


@bp.route("/attendance/all")
@login_required
def all_attendance():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    from sqlalchemy import and_
    
    query = StudentAttendance.query.join(Teacher, StudentAttendance.teacher_id == Teacher.id)

    # ================= TEACHER =================
    if current_user.role.value == "teacher":
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()
        if not teacher:
            flash("❌ Teacher profile not found.", "danger")
            return redirect(url_for("main.dashboard"))
        query = query.filter(StudentAttendance.teacher_id == teacher.id)

    # ================= SCHOOL ADMIN =================
    elif current_user.role.value == "school_admin":
        query = query.filter(StudentAttendance.school_id == current_user.school_id)

    # ================= BRANCH ADMIN =================
    elif current_user.role.value == "branch_admin":
        query = query.filter(
            and_(
                StudentAttendance.school_id == current_user.school_id,
                StudentAttendance.branch_id == current_user.branch_id
            )
        )

    else:
        flash("Unauthorized", "danger")
        return redirect(url_for("main.dashboard"))

    # ✅ SORT BY LAST UPDATED
    attendances = query.order_by(StudentAttendance.updated_at.desc()).all()

    return render_template(
        "backend/pages/components/attendance/all_attendance.html",
        attendances=attendances,
        user=current_user
    )



@bp.route("/attendance/update_status/<int:id>", methods=["POST"])
@login_required
def update_attendance_status(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    from flask import request, jsonify
    try:
        attendance = StudentAttendance.query.get_or_404(id)
        data = request.get_json()
        new_status = data.get('status')

        if new_status not in ['present','absent','late','excused']:
            return jsonify({"success": False, "message": "Invalid status"}), 400

        attendance.status = new_status
        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500




#--------------------------------------
#--------- Exams Sections
#--------------------------------------
@bp.route("/academic-years", methods=["GET"])
@login_required
def all_academic_years():
    # Check if user is active
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Role-based filtering
    query = AcademicYear.query

    if getattr(current_user, "role", None):
        if current_user.role.value == "school_admin":
            query = query.filter_by(school_id=current_user.school_id, branch_id=None)
        elif current_user.role.value == "branch_admin":
            query = query.filter_by(school_id=current_user.school_id, branch_id=current_user.branch_id)
        elif current_user.role.value != "superadmin":
            flash("You are not authorized to view academic years.", "danger")
            return redirect(url_for('main.index'))

    academic_years = query.order_by(AcademicYear.year_name.desc()).all()

    return render_template(
        "backend/pages/components/academic_years/all_academic_years.html",
        academic_years=academic_years,
        user=current_user
    )


@bp.route("/add/academic-year", methods=["GET", "POST"])
@login_required
def add_academic_year():
    # Check if user is active
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only allow school or branch admins
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("You are not authorized to add academic years.", "danger")
        return redirect(url_for('main.all_academic_years'))

    form = AcademicYearForm()

    if form.validate_on_submit():
        # Normalize inputs
        year_name = form.year_name.data.strip()
        school_id = form.school_id.data
        branch_id = form.branch_id.data or None

        # Prevent duplicate academic years for the same school/branch
        existing = AcademicYear.query.filter_by(
            school_id=school_id,
            branch_id=branch_id,
            year_name=year_name
        ).first()
        if existing:
            flash(f"Academic year '{year_name}' already exists for this school/branch!", "warning")
            return redirect(url_for('main.add_academic_year'))

        # Create new academic year
        new_year = AcademicYear(
            school_id=school_id,
            branch_id=branch_id,
            year_name=year_name,
            is_active=form.is_active.data
        )
        db.session.add(new_year)
        db.session.commit()

        flash(f"Academic year '{year_name}' added successfully!", "success")
        return redirect(url_for('main.all_academic_years'))

    return render_template(
        "backend/pages/components/academic_years/add_academic_year.html",
        form=form,
        user=current_user
    )


@bp.route("/edit/academic-year/<int:id>", methods=["GET", "POST"])
@login_required
def edit_academic_year(id):
    # ------------------- Check user status -------------------
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ------------------- Load academic year -------------------
    academic_year = AcademicYear.query.get_or_404(id)

    # ------------------- Role-based access -------------------
    if current_user.role.value == "school_admin":
        if academic_year.school_id != current_user.school_id or academic_year.branch_id is not None:
            flash("You are not authorized to edit this academic year.", "danger")
            return redirect(url_for('main.all_academic_years'))

    elif current_user.role.value == "branch_admin":
        if (academic_year.school_id != current_user.school_id or 
            academic_year.branch_id != current_user.branch_id):
            flash("You are not authorized to edit this academic year.", "danger")
            return redirect(url_for('main.all_academic_years'))

    elif current_user.role.value != "superadmin":
        flash("You are not authorized to edit academic years.", "danger")
        return redirect(url_for('main.all_academic_years'))

    # ------------------- Prepopulate form -------------------
    form = AcademicYearForm(obj=academic_year)

    if form.validate_on_submit():
        # Normalize inputs
        year_name = form.year_name.data.strip()
        school_id = form.school_id.data
        branch_id = form.branch_id.data or None

        # Prevent duplicate academic years for same school/branch
        existing = AcademicYear.query.filter(
            AcademicYear.school_id == school_id,
            AcademicYear.branch_id == branch_id,
            AcademicYear.year_name == year_name,
            AcademicYear.id != id  # exclude current record
        ).first()

        if existing:
            flash(f"Academic year '{year_name}' already exists for this school/branch!", "warning")
            return redirect(url_for('main.edit_academic_year', id=id))

        # Update academic year
        academic_year.year_name = year_name
        academic_year.is_active = form.is_active.data
        db.session.commit()

        flash(f"Academic year '{year_name}' updated successfully!", "success")
        return redirect(url_for('main.all_academic_years'))

    return render_template(
        "backend/pages/components/academic_years/edit_academic_year.html",
        form=form,
        user=current_user,
        edit=True  # optional flag to indicate editing
    )


@bp.route("/delete/academic-year/<int:id>", methods=["POST"])
@login_required
def delete_academic_year(id):
    academic_year = AcademicYear.query.get_or_404(id)

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    role = getattr(current_user, "role", None)
    if role:
        if role.value == "school_admin":
            if academic_year.school_id != current_user.school_id or academic_year.branch_id is not None:
                flash("You are not authorized to delete this academic year.", "danger")
                return redirect(url_for('main.all_academic_years'))

        elif role.value == "branch_admin":
            if academic_year.school_id != current_user.school_id or academic_year.branch_id != current_user.branch_id:
                flash("You are not authorized to delete this academic year.", "danger")
                return redirect(url_for('main.all_academic_years'))

        elif role.value != "superadmin":
            flash("You are not authorized to delete academic years.", "danger")
            return redirect(url_for('main.all_academic_years'))

    try:
        # ✅ STEP 1: delete dependent records first
        StudentPromotion.query.filter_by(from_academic_year_id=id).delete()
        StudentPromotion.query.filter_by(to_academic_year_id=id).delete()  # if exists

        # ⚠️ commit child deletions first
        db.session.flush()

        # ✅ STEP 2: delete parent
        db.session.delete(academic_year)
        db.session.commit()

        flash(f"Academic year '{academic_year.year_name}' deleted successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting academic year: {str(e)}", "danger")

    return redirect(url_for('main.all_academic_years'))


# -------------------------------------
#------------ Terms
#---------------------------------------
# ================= All Terms =================
@bp.route("/terms", methods=["GET"])
@login_required
def all_terms():
    # Check if user is active
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Start query
    query = Term.query.join(AcademicYear, Term.academic_year_id == AcademicYear.id)

    # Role-based filtering
    if getattr(current_user, "role", None):
        if current_user.role.value == "school_admin":
            query = query.filter(Term.school_id == current_user.school_id, Term.branch_id.is_(None))
        elif current_user.role.value == "branch_admin":
            query = query.filter(Term.school_id == current_user.school_id,
                                 Term.branch_id == current_user.branch_id)
        elif current_user.role.value != "superadmin":
            flash("You are not authorized to view terms.", "danger")
            return redirect(url_for('main.index'))

    # Order by most recent academic year, then term name
    terms = query.order_by(AcademicYear.year_name.desc(), Term.term_name.asc()).all()

    return render_template(
        "backend/pages/components/terms/all_terms.html",
        terms=terms,
        user=current_user
    )


@bp.route("/add/term", methods=["GET", "POST"])
@login_required
def add_term():
    # Check if user is active
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only allow school or branch admins
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("You are not authorized to add terms.", "danger")
        return redirect(url_for('main.all_terms'))

    form = TermForm()

    if form.validate_on_submit():
        term_name = form.term_name.data.strip()
        school_id = form.school_id.data
        branch_id = form.branch_id.data or None
        academic_year_id = form.academic_year_id.data
        start_date = form.start_date.data
        end_date = form.end_date.data
        is_active = form.is_active.data

        # Prevent duplicate term names
        existing_name = Term.query.filter_by(
            school_id=school_id,
            branch_id=branch_id,
            academic_year_id=academic_year_id,
            term_name=term_name
        ).first()
        if existing_name:
            flash(f"Term '{term_name}' already exists for this academic year!", "warning")
            return redirect(url_for('main.add_term'))

        # Prevent overlapping terms
        overlapping_term = Term.query.filter(
            Term.school_id == school_id,
            Term.branch_id == branch_id,
            Term.academic_year_id == academic_year_id,
            db.or_(
                db.and_(Term.start_date <= start_date, Term.end_date >= start_date),
                db.and_(Term.start_date <= end_date, Term.end_date >= end_date),
                db.and_(Term.start_date >= start_date, Term.end_date <= end_date)
            )
        ).first()

        if overlapping_term:
            flash(
                f"The date range {start_date} to {end_date} overlaps with existing term '{overlapping_term.term_name}' "
                f"({overlapping_term.start_date} to {overlapping_term.end_date})!",
                "warning"
            )
            return redirect(url_for('main.add_term'))

        # ✅ Deactivate all previously active terms for this school/branch/academic year
        active_terms = Term.query.filter_by(
            school_id=school_id,
            branch_id=branch_id,
            academic_year_id=academic_year_id,
            is_active=True
        ).all()

        for term in active_terms:
            term.is_active = False
            db.session.add(term)

        # Create new Term
        new_term = Term(
            school_id=school_id,
            branch_id=branch_id,
            academic_year_id=academic_year_id,
            term_name=term_name,
            start_date=start_date,
            end_date=end_date,
            is_active=is_active
        )
        db.session.add(new_term)
        db.session.commit()

        flash(f"Term '{term_name}' added successfully! All previous terms have been deactivated.", "success")
        return redirect(url_for('main.all_terms'))

    return render_template(
        "backend/pages/components/terms/add_term.html",
        form=form,
        user=current_user
    )

@bp.route("/edit/term/<int:id>", methods=["GET", "POST"])
@login_required
def edit_term(id):
    term = Term.query.get_or_404(id)

    # Only allow school/branch admins or superadmin
    if current_user.role.value not in ["superadmin", "school_admin", "branch_admin"]:
        flash("You are not authorized to edit terms.", "danger")
        return redirect(url_for("main.all_terms"))

    # Role-based access: restrict school/branch admins to their own terms
    if current_user.role.value == "school_admin" and term.branch_id is not None:
        flash("You are not authorized to edit this term.", "danger")
        return redirect(url_for("main.all_terms"))

    if current_user.role.value == "branch_admin" and term.branch_id != current_user.branch_id:
        flash("You are not authorized to edit this term.", "danger")
        return redirect(url_for("main.all_terms"))

    form = TermForm(obj=term)

    if form.validate_on_submit():
        term_name = form.term_name.data.strip()
        academic_year_id = form.academic_year_id.data
        start_date = form.start_date.data
        end_date = form.end_date.data
        is_active = form.is_active.data

        # Check for duplicate term names (excluding current term)
        existing_name = Term.query.filter_by(
            school_id=term.school_id,
            branch_id=term.branch_id,
            academic_year_id=academic_year_id,
            term_name=term_name
        ).filter(Term.id != term.id).first()

        if existing_name:
            flash(f"Term '{term_name}' already exists for this academic year!", "warning")
            return redirect(url_for("main.edit_term", id=term.id))

        # Check for overlapping date ranges (excluding current term)
        overlapping_term = Term.query.filter(
            Term.school_id == term.school_id,
            Term.branch_id == term.branch_id,
            Term.academic_year_id == academic_year_id,
            Term.id != term.id,
            db.or_(
                db.and_(Term.start_date <= start_date, Term.end_date >= start_date),
                db.and_(Term.start_date <= end_date, Term.end_date >= end_date),
                db.and_(Term.start_date >= start_date, Term.end_date <= end_date)
            )
        ).first()

        if overlapping_term:
            flash(
                f"The date range {start_date} to {end_date} overlaps with existing term "
                f"'{overlapping_term.term_name}' ({overlapping_term.start_date} to {overlapping_term.end_date})!",
                "warning"
            )
            return redirect(url_for("main.edit_term", id=term.id))

        # If this term is set to active, deactivate all other terms for this school/branch/year
        if is_active:
            other_active_terms = Term.query.filter(
                Term.school_id == term.school_id,
                Term.branch_id == term.branch_id,
                Term.academic_year_id == academic_year_id,
                Term.is_active == True,
                Term.id != term.id
            ).all()
            for other in other_active_terms:
                other.is_active = False
                db.session.add(other)

        # Update the current term
        term.term_name = term_name
        term.academic_year_id = academic_year_id
        term.start_date = start_date
        term.end_date = end_date
        term.is_active = is_active

        db.session.commit()
        flash(f"Term '{term.term_name}' updated successfully!", "success")
        return redirect(url_for("main.all_terms"))

    return render_template(
        "backend/pages/components/terms/edit_term.html",
        form=form,
        edit=True,
        user=current_user
    )

@bp.route("/delete/term/<int:id>", methods=["POST"])
@login_required
def delete_term(id):
    term = Term.query.get_or_404(id)

    # ---------------- AUTH CHECK ----------------
    if not current_user.role:
        flash("Unauthorized access.", "danger")
        return redirect(url_for("main.all_terms"))

    if current_user.role.value not in ["superadmin", "school_admin", "branch_admin"]:
        flash("You are not authorized to delete terms.", "danger")
        return redirect(url_for("main.all_terms"))

    if current_user.role.value == "school_admin":
        if term.branch_id is not None:
            flash("You are not authorized to delete this term.", "danger")
            return redirect(url_for("main.all_terms"))

    if current_user.role.value == "branch_admin":
        if term.branch_id != current_user.branch_id:
            flash("You are not authorized to delete this term.", "danger")
            return redirect(url_for("main.all_terms"))

    try:
        # ---------------- IMPORTANT: DELETE CHILD RECORDS FIRST ----------------

        # Delete exam results linked to exams under this term
        exams = Exam.query.filter_by(term_id=id).all()

        for exam in exams:
            StudentExamResult.query.filter_by(exam_id=exam.id).delete()

        # Delete exams under this term
        Exam.query.filter_by(term_id=id).delete()

        # (Optional) delete promotions or other dependent data
        StudentPromotion.query.join(Exam).filter(Exam.term_id == term.id)


        db.session.flush()

        # ---------------- DELETE PARENT ----------------
        db.session.delete(term)
        db.session.commit()

        flash(f"Term '{term.term_name}' has been deleted successfully!", "success")

    except Exception as e:
        db.session.rollback()   # 🔥 VERY IMPORTANT (fixes PendingRollbackError)
        flash(f"Error deleting term: {str(e)}", "danger")

    return redirect(url_for("main.all_terms"))



#---------------------------------
#----------- Exams
#---------------------------------

@bp.route("/all/exams")
@login_required
def all_exams():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 🔐 ROLE-BASED ACCESS
    # =========================
    if current_user.role.value == "superadmin":
        # 🌍 Access to all exams
        exams = Exam.query.order_by(Exam.created_at.desc()).all()

    elif current_user.role.value == "school_admin":
        # 🏫 Only exams for this school with branch_id=None
        exams = Exam.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(Exam.created_at.desc()).all()

    elif current_user.role.value == "branch_admin":
        # 🏢 Only exams for this branch
        exams = Exam.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(Exam.created_at.desc()).all()

    else:
        # ❌ Unauthorized
        flash("You are not authorized to view exams.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 📤 RENDER TEMPLATE
    # =========================
    return render_template(
        "backend/pages/components/exams/all_exams.html",
        exams=exams,
        user=current_user
    )


@bp.route("/add/exam", methods=["GET", "POST"])
@login_required
def add_exam():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = ExamForm()

    # Check if term selected to auto-fill start/end dates
    term_dates = None
    if form.term_id.data:
        selected_term = Term.query.get(form.term_id.data)
        if selected_term:
            term_dates = selected_term  # object with start_date, end_date
            # Auto-fill start/end dates for read-only
            form.start_date.data = selected_term.start_date
            form.end_date.data = selected_term.end_date

    if form.validate_on_submit():
        # Prevent duplicate exam names for same term
        existing_exam = Exam.query.filter_by(
            school_id=form.school_id.data,
            branch_id=form.branch_id.data or None,
            academic_year_id=form.academic_year_id.data,
            term_id=form.term_id.data,
            exam_name=form.exam_name.data.strip()
        ).first()
        if existing_exam:
            flash(f"Exam '{form.exam_name.data}' already exists for this term!", "warning")
            return redirect(url_for("main.add_exam"))

        new_exam = Exam(
            school_id=form.school_id.data,
            branch_id=form.branch_id.data or None,
            academic_year_id=form.academic_year_id.data,
            term_id=form.term_id.data,
            exam_name=form.exam_name.data.strip(),
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            status=form.status.data
        )
        db.session.add(new_exam)
        db.session.commit()

        flash(f"Exam '{new_exam.exam_name}' added successfully!", "success")
        return redirect(url_for("main.all_exams"))

    return render_template(
        "backend/pages/components/exams/add_exam.html",
        form=form,
        user=current_user,
        term_dates=term_dates
    )

@bp.route("/edit/exam/<int:id>", methods=["GET", "POST"])
@login_required
def edit_exam(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    exam = Exam.query.get_or_404(id)
    form = ExamForm(obj=exam)

    # Auto-fill start/end dates from term if selected
    term_dates = None
    if form.term_id.data:
        selected_term = Term.query.get(form.term_id.data)
        if selected_term:
            term_dates = selected_term
            form.start_date.data = selected_term.start_date
            form.end_date.data = selected_term.end_date
            # Make read-only in template
            form.start_date.render_kw = {"readonly": True}
            form.end_date.render_kw = {"readonly": True}

    if form.validate_on_submit():
        # Prevent duplicate exam names for the same term (excluding current exam)
        existing_exam = Exam.query.filter_by(
            school_id=form.school_id.data,
            branch_id=form.branch_id.data or None,
            academic_year_id=form.academic_year_id.data,
            term_id=form.term_id.data,
            exam_name=form.exam_name.data.strip()
        ).filter(Exam.id != exam.id).first()

        if existing_exam:
            flash(f"Exam '{form.exam_name.data}' already exists for this term!", "warning")
            return redirect(url_for("main.edit_exam", id=exam.id))

        # Update exam fields
        exam.exam_name = form.exam_name.data.strip()
        exam.academic_year_id = form.academic_year_id.data
        exam.term_id = form.term_id.data
        exam.start_date = form.start_date.data
        exam.end_date = form.end_date.data
        exam.status = form.status.data

        db.session.commit()
        flash(f"Exam '{exam.exam_name}' updated successfully!", "success")
        return redirect(url_for("main.all_exams"))

    return render_template(
        "backend/pages/components/exams/add_exam.html",
        form=form,
        edit=True,
        user=current_user,
        term_dates=term_dates
    )



# Delete Exam
@bp.route("/delete/exam/<int:id>", methods=["POST"])
@login_required
def delete_exam(id):
    exam = Exam.query.get_or_404(id)

    # Check if user is active
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for("main.all_exams"))

    # Only allow school/branch admins or superadmin
    if current_user.role.value not in ["superadmin", "school_admin", "branch_admin"]:
        flash("You are not authorized to delete exams.", "danger")
        return redirect(url_for("main.all_exams"))

    # Role-based restrictions
    if current_user.role.value == "school_admin" and exam.branch_id is not None:
        flash("You are not authorized to delete this exam.", "danger")
        return redirect(url_for("main.all_exams"))

    if current_user.role.value == "branch_admin" and exam.branch_id != current_user.branch_id:
        flash("You are not authorized to delete this exam.", "danger")
        return redirect(url_for("main.all_exams"))

    # Delete exam
    db.session.delete(exam)
    db.session.commit()
    flash(f"Exam '{exam.exam_name}' deleted successfully!", "success")
    return redirect(url_for("main.all_exams"))


@bp.route("/exams/update_status/<int:id>", methods=["POST"])
@login_required
def update_exam_status(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    exam = Exam.query.get_or_404(id)
    if current_user.status == 0:
        return jsonify({"success": False, "message": "Your account is inactive."}), 403

    data = request.get_json()
    status = data.get("status")
    if status not in ["draft", "published", "closed"]:
        return jsonify({"success": False, "message": "Invalid status"}), 400

    exam.status = status
    db.session.commit()
    return jsonify({"success": True})


#-----------------------------------
#--------- 
#-----------------------------------

@bp.route("/get-exam-subjects/<int:level_id>")
@login_required
def get_exam_subjects(level_id):
    """Return subjects assigned to all classes under a specific class level."""
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1️⃣ Get all classes for this level
    classes_query = Class.query.filter_by(
        level_id=level_id,
        school_id=current_user.school_id
    )
    
    if current_user.role.value == "branch_admin":
        classes_query = classes_query.filter(
            or_(
                Class.branch_id == current_user.branch_id,
                Class.branch_id.is_(None)
            )
        )
    
    classes = classes_query.all()
    
    # 2️⃣ Collect all subjects from ClassSubject for these classes
    subjects_set = set()
    for klass in classes:
        for cs in klass.class_subjects:
            # Include only subjects in the same school/branch (branch may be NULL)
            if cs.school_id == current_user.school_id and (
                cs.branch_id == current_user.branch_id or cs.branch_id is None
            ):
                subjects_set.add((cs.subject.id, cs.subject.name))
    
    # 3️⃣ Convert set to list of dicts
    subjects_list = [{"id": sid, "name": name} for sid, name in subjects_set]
    subjects_list.sort(key=lambda x: x["name"])  # Optional: sort alphabetically
    
    return {"subjects": subjects_list}


@bp.route("/add/exam-timetable", methods=["GET", "POST"])
@login_required
def add_exam_timetable():

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = ExamTimetableForm()

    # ================= SCHOOL / BRANCH =================
    school_id = current_user.school_id
    branch_id = current_user.branch_id if current_user.role.value == "branch_admin" else None

    form.school_id.data = school_id
    form.branch_id.data = branch_id

    # ================= EXAMS =================
    exam_query = Exam.query.filter_by(school_id=school_id)
    if branch_id:
        exam_query = exam_query.filter_by(branch_id=branch_id)

    exams = exam_query.all()
    form.exam_id.choices = [(0, "-- Select Exam --")] + [(e.id, e.exam_name) for e in exams]

    # ================= LEVELS =================
    level_query = ClassLevel.query.filter_by(school_id=school_id)
    if branch_id:
        level_query = level_query.filter_by(branch_id=branch_id)

    levels = level_query.all()
    form.level_id.choices = [(0, "-- Select Level --")] + [(l.id, l.name) for l in levels]

    # ⚠️ IMPORTANT: allow dynamic subject
    form.subject_id.choices = [(0, "-- Select Subject --")]

    # ================= FIX SUBJECT VALIDATION =================
    if request.method == "POST":
        subject_id = request.form.get("subject_id")
        if subject_id:
            form.subject_id.choices.append((int(subject_id), "Selected"))

    # ================= SUBMIT =================
    if form.validate_on_submit():

        if form.exam_id.data == 0 or form.level_id.data == 0 or form.subject_id.data == 0:
            flash("Please select valid Exam, Level, and Subject.", "warning")
            return redirect(url_for("main.add_exam_timetable"))

        existing = ExamTimetable.query.filter_by(
            exam_id=form.exam_id.data,
            level_id=form.level_id.data,
            subject_id=form.subject_id.data,
            date=form.date.data,
            school_id=school_id,
            branch_id=branch_id
        ).first()

        if existing:
            flash("Timetable already exists.", "warning")
            return redirect(url_for("main.add_exam_timetable"))

        new_tt = ExamTimetable(
            exam_id=form.exam_id.data,
            level_id=form.level_id.data,
            subject_id=form.subject_id.data,
            date=form.date.data,
            start_time=form.start_time.data,
            end_time=form.end_time.data,
            school_id=school_id,
            branch_id=branch_id
        )

        db.session.add(new_tt)
        db.session.commit()

        flash("Exam timetable added successfully!", "success")
        return redirect(url_for("main.all_exam_timetables"))

    return render_template(
        "backend/pages/components/exams/add_exam_timetable.html",
        form=form,
        user=current_user
    )


# ---------------- VIEW ALL TIMETABLES ----------------
@bp.route("/all/exam-timetables")
@login_required
def all_exam_timetables():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))


    # ================= TIMETABLES =================
    if current_user.role.value == "superadmin":
        timetables = ExamTimetable.query.order_by(
            ExamTimetable.date,
            ExamTimetable.start_time
        ).all()

        levels = ClassLevel.query.all()

    elif current_user.role.value == "school_admin":
        timetables = ExamTimetable.query.filter_by(
            school_id=current_user.school_id
        ).order_by(
            ExamTimetable.date,
            ExamTimetable.start_time
        ).all()

        levels = ClassLevel.query.filter_by(
            school_id=current_user.school_id
        ).all()

    elif current_user.role.value == "branch_admin":
        timetables = ExamTimetable.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(
            ExamTimetable.date,
            ExamTimetable.start_time
        ).all()

        levels = ClassLevel.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).all()

    else:
        flash("You are not authorized to view exam timetables.", "danger")
        return redirect(url_for("main.dashboard"))

    # ================= SEND TO TEMPLATE =================
    return render_template(
        "backend/pages/components/exams/all_exam_timetables.html",
        timetables=timetables,
        levels=levels,   # ✅ IMPORTANT
        user=current_user
    )


@bp.route("/exam-timetable/print")
@login_required
def print_exam_timetable():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    level_id = request.args.get("level_id", type=int)

    if not level_id:
        flash("Please select a level first.", "warning")
        return redirect(url_for("main.all_exam_timetables"))

    # ================= QUERY =================
    query = ExamTimetable.query.filter_by(level_id=level_id)

    if current_user.role.value == "school_admin":
        query = query.filter_by(school_id=current_user.school_id)

    elif current_user.role.value == "branch_admin":
        query = query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        )

    timetables = query.order_by(
        ExamTimetable.date,
        ExamTimetable.start_time
    ).all()

    # ================= GROUP BY DATE =================
    grouped = defaultdict(list)

    for tt in timetables:
        grouped[tt.date].append(tt)

    # ================= SPLIT (2 PER ROW) =================
    final_rows = []

    for date in sorted(grouped.keys()):
        exams = grouped[date]

        for i in range(0, len(exams), 2):
            first = exams[i]
            second = exams[i+1] if i+1 < len(exams) else None

            final_rows.append({
                "date": date,
                "first": first,
                "second": second
            })

    # ================= LEVEL =================
    level = ClassLevel.query.get(level_id)

    # ================= SCHOOL & BRANCH =================
    school = current_user.school if current_user.school_id else None
    branch = current_user.branch if current_user.branch_id else None

    # ================= RETURN =================
    return render_template(
        "backend/pages/components/exams/print_exam_timetable.html",
        rows=final_rows,
        level=level,
        school=school,   # 👈 muhiim
        branch=branch    # 👈 muhiim
    )


@bp.route('/download-exam-timetable-word/<int:level_id>')
@login_required
def download_exam_timetable_word(level_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Soo qaado xogta
    level = ClassLevel.query.get_or_404(level_id)
    
    query = ExamTimetable.query.filter_by(level_id=level_id)
    if current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)
    else:
        query = query.filter_by(school_id=current_user.school_id)
        
    timetables = query.order_by(ExamTimetable.date, ExamTimetable.start_time).all()

    # Grouping data (2 exams per row)
    grouped = defaultdict(list)
    for tt in timetables:
        grouped[tt.date].append(tt)
    
    final_rows = []
    for date in sorted(grouped.keys()):
        exams = grouped[date]
        for i in range(0, len(exams), 2):
            final_rows.append({
                "date": date,
                "first": exams[i],
                "second": exams[i+1] if i+1 < len(exams) else None
            })

    # 2. Create Word Document
    doc = Document()
    
    # Page Setup (Portrait)
    section = doc.sections[-1]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    # Logo Path
    logo_path = None
    if g.get('site_logo') and g.site_logo.get('main_logo'):
        logo_path = os.path.join(current_app.root_path, 'static', g.site_logo['main_logo'])

    # --- 1. HEADER (Logo) ---
    if logo_path and os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(18.0)) 
    else:
        title = doc.add_heading(current_user.school.name.upper(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 2. BLUE BANNER ---
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_banner = p_banner.add_run(f"JADWALKA IMTIXAANKA - {level.name.upper()}")
    run_banner.bold = True
    run_banner.font.size = Pt(16)
    run_banner.font.color.rgb = RGBColor(255, 255, 255)
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '0000FF') 
    p_banner._element.get_or_add_pPr().append(shd)

    doc.add_paragraph() # Space

    # --- 3. TIMETABLE TABLE ---
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.width = Cm(18.0)

    # Header Row
    hdr_cells = table.rows[0].cells
    headers = ["DATE", "DAY", "FIRST EXAM (Time)", "SECOND EXAM (Time)"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        # Set Header Background (Light Blue)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '5DADE2')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading)

    # Fill Data
    for row_data in final_rows:
        row = table.add_row().cells
        row[0].text = row_data['date'].strftime('%d/%m/%Y')
        row[1].text = row_data['date'].strftime('%A')
        
        # First Exam
        f = row_data['first']
        row[2].text = f"{f.subject.name}\n({f.start_time.strftime('%I:%M %p')} - {f.end_time.strftime('%I:%M %p')})"
        
        # Second Exam
        s = row_data['second']
        if s:
            row[3].text = f"{s.subject.name}\n({s.start_time.strftime('%I:%M %p')} - {s.end_time.strftime('%I:%M %p')})"
        else:
            row[3].text = "-"

    # Center all cells
    for r in table.rows:
        for cell in r.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 4. FOOTER ---
    doc.add_paragraph("\n" + "_" * 85)
    footer = doc.add_paragraph()
    f_run = footer.add_run(f"{current_user.school.name.upper()}")
    f_run.bold = True
    
    branch_phone = current_user.branch.phone if (hasattr(current_user, 'branch') and current_user.branch) else current_user.school.phone
    footer.add_run(f"             TEL: {branch_phone}").bold = True

    # 3. Save & Send
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)

    return send_file(
        target, 
        as_attachment=True, 
        download_name=f"Timetable_{level.name}_{now_eat().strftime('%Y%m%d')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@bp.route('/download-exam-timetable-png/<int:level_id>')
@login_required
def download_exam_timetable_png(level_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Xogta Database-ka
    level = ClassLevel.query.get_or_404(level_id)
    school = current_user.school
    branch = current_user.branch if hasattr(current_user, 'branch') else None
    
    query = ExamTimetable.query.filter_by(level_id=level_id)
    if current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)
    else:
        query = query.filter_by(school_id=current_user.school_id)
        
    timetables = query.order_by(ExamTimetable.date, ExamTimetable.start_time).all()

    # 2. Grouping Data (2 exams per row)
    grouped = defaultdict(list)
    for tt in timetables:
        grouped[tt.date].append(tt)
    
    final_rows = []
    for date in sorted(grouped.keys()):
        exams = grouped[date]
        for i in range(0, len(exams), 2):
            final_rows.append({
                "date": date,
                "first": exams[i],
                "second": exams[i+1] if i+1 < len(exams) else None
            })

    # 3. Logada Saxda ah (Isticmaalka get_user_site_logo)
    site_logos = get_user_site_logo(current_user)
    main_logo = site_logos.get("main_logo")
    
    if main_logo:
        logo_path = url_for('static', filename=main_logo)
    else:
        logo_path = url_for('static', filename='backend/uploads/images/no_image.jpg')
    
    # URL Dhamaystiran oo wkhtmltoimage akhrisan karo
    full_logo_url = f"{request.url_root.rstrip('/')}{logo_path}"

    # 4. Render HTML
    html_content = render_template(
        'backend/pages/components/exams/print_exam_timetable-png.html', 
        rows=final_rows,
        level=level,
        school=school,
        branch=branch,
        logo_url=full_logo_url,
        now_eat=now_eat
    )

    # 5. Path-ka wkhtmltoimage
    path_to_wk = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltoimage.exe'
    if not os.path.exists(path_to_wk):
        return f"Khalad: Ma jiro faylka: {path_to_wk}", 500

    config = imgkit.config(wkhtmltoimage=path_to_wk)

    options = {
        'format': 'png',
        'encoding': "UTF-8",
        'width': 1250, 
        'enable-local-file-access': '', 
        'quiet': ''
    }

    try:
        img_bytes = imgkit.from_string(html_content, False, options=options, config=config)
        response = make_response(img_bytes)
        response.headers['Content-Type'] = 'image/png'
        response.headers['Content-Disposition'] = f'attachment; filename=Jadwalka_{level.name}.png'
        return response
    except Exception as e:
        return f"Error: {str(e)}", 500


# ---------------- EDIT TIMETABLE ----------------
@bp.route("/edit/exam-timetable/<int:id>", methods=["GET", "POST"])
@login_required
def edit_exam_timetable(id):

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    timetable = ExamTimetable.query.get_or_404(id)

    # ================= SECURITY (school/branch check) =================
    if timetable.school_id != current_user.school_id:
        flash("Unauthorized access.", "danger")
        return redirect(url_for("main.all_exam_timetables"))

    if current_user.role.value == "branch_admin":
        if timetable.branch_id != current_user.branch_id:
            flash("Unauthorized branch access.", "danger")
            return redirect(url_for("main.all_exam_timetables"))

    form = ExamTimetableForm(obj=timetable)

    # ================= SCHOOL / BRANCH =================
    school_id = current_user.school_id
    branch_id = current_user.branch_id if current_user.role.value == "branch_admin" else None

    form.school_id.data = school_id
    form.branch_id.data = branch_id

    # ================= EXAMS =================
    exam_query = Exam.query.filter_by(school_id=school_id)
    if branch_id:
        exam_query = exam_query.filter_by(branch_id=branch_id)

    exams = exam_query.all()
    form.exam_id.choices = [(0, "-- Select Exam --")] + [(e.id, e.exam_name) for e in exams]

    # ================= LEVELS =================
    level_query = ClassLevel.query.filter_by(school_id=school_id)
    if branch_id:
        level_query = level_query.filter_by(branch_id=branch_id)

    levels = level_query.all()
    form.level_id.choices = [(0, "-- Select Level --")] + [(l.id, l.name) for l in levels]

    # ================= SUBJECTS (IMPORTANT) =================
    # Load subjects based on current level (for edit page)
    classes = Class.query.filter_by(level_id=timetable.level_id, school_id=school_id)
    if branch_id:
        classes = classes.filter_by(branch_id=branch_id)

    class_ids = [c.id for c in classes.all()]

    subject_query = ClassSubject.query.filter(
        ClassSubject.class_id.in_(class_ids),
        ClassSubject.school_id == school_id
    )
    if branch_id:
        subject_query = subject_query.filter(ClassSubject.branch_id == branch_id)

    subjects = subject_query.all()

    form.subject_id.choices = [(0, "-- Select Subject --")] + [
        (s.subject.id, s.subject.name) for s in subjects if s.subject
    ]

    # ================= FIX SUBJECT VALIDATION =================
    if request.method == "POST":
        subject_id = request.form.get("subject_id")
        if subject_id:
            form.subject_id.choices.append((int(subject_id), "Selected"))

    # ================= SUBMIT =================
    if form.validate_on_submit():

        if form.exam_id.data == 0 or form.level_id.data == 0 or form.subject_id.data == 0:
            flash("Please select valid Exam, Level, and Subject.", "warning")
            return redirect(url_for("main.edit_exam_timetable", id=id))

        existing = ExamTimetable.query.filter(
            ExamTimetable.exam_id == form.exam_id.data,
            ExamTimetable.level_id == form.level_id.data,
            ExamTimetable.subject_id == form.subject_id.data,
            ExamTimetable.date == form.date.data,
            ExamTimetable.school_id == school_id,
            ExamTimetable.branch_id == branch_id,
            ExamTimetable.id != id  # 🔥 IMPORTANT (exclude current)
        ).first()

        if existing:
            flash("Timetable already exists.", "warning")
            return redirect(url_for("main.edit_exam_timetable", id=id))

        # ================= UPDATE =================
        timetable.exam_id = form.exam_id.data
        timetable.level_id = form.level_id.data
        timetable.subject_id = form.subject_id.data
        timetable.date = form.date.data
        timetable.start_time = form.start_time.data
        timetable.end_time = form.end_time.data
        timetable.school_id = school_id
        timetable.branch_id = branch_id

        db.session.commit()

        flash("Exam timetable updated successfully!", "success")
        return redirect(url_for("main.all_exam_timetables"))

    return render_template(
        "backend/pages/components/exams/edit_exam_timetable.html",  # reuse same template
        form=form,
        timetable=timetable,
         user=current_user
    )


# ---------------- DELETE TIMETABLE ----------------
@bp.route("/delete/exam-timetable/<int:id>", methods=["POST", "GET"])
@login_required
def delete_exam_timetable(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    timetable = ExamTimetable.query.get_or_404(id)

    # Only allow deletion if the timetable belongs to user's school/branch
    if timetable.school_id != current_user.school_id or (
        current_user.role.value == "branch_admin" and timetable.branch_id != current_user.branch_id
    ):
        flash("You do not have permission to delete this timetable.", "danger")
        return redirect(url_for("main.all_exam_timetables"))

    db.session.delete(timetable)
    db.session.commit()
    flash("Exam timetable deleted successfully!", "success")
    return redirect(url_for("main.all_exam_timetables"))



#----------------------------------
#------------ Exam Halls
#----------------------------------

@bp.route("/all/exam-halls")
@login_required
def all_exam_halls():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 🔐 ROLE-BASED ACCESS
    # =========================
    if current_user.role.value == "superadmin":
        # 🌍 Full access
        halls = ExamHall.query.order_by(ExamHall.created_at.desc()).all()

    elif current_user.role.value == "school_admin":
        # 🏫 Only halls in school with branch_id=None
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(ExamHall.created_at.desc()).all()

    elif current_user.role.value == "branch_admin":
        # 🏢 Only halls in this branch
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(ExamHall.created_at.desc()).all()

    else:
        # ❌ Unauthorized
        flash("You are not authorized to view exam halls.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 📤 RENDER TEMPLATE
    # =========================
    return render_template(
        "backend/pages/components/exam_halls/all_exam_halls.html",
        halls=halls,
        user=current_user
    )


@bp.route("/exam-halls/add", methods=["GET", "POST"])
@login_required
def add_exam_hall():
    # 🔒 Check user status
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Check permission
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("You are not authorized to edit exam halls.", "danger")
        return redirect(url_for('main.dashboard'))

    form = ExamHallForm()

    # 👉 Set hidden values (important)
    form.school_id.data = current_user.school_id
    form.branch_id.data = current_user.branch_id

    if form.validate_on_submit():
        name = form.name.data.strip()

        # 🚫 Prevent duplicate hall name per school/branch
        existing = ExamHall.query.filter_by(
            school_id=form.school_id.data,
            branch_id=form.branch_id.data,
            name=name
        ).first()

        if existing:
            flash("This exam hall already exists!", "warning")
            return redirect(url_for('main.add_exam_hall'))

        # ✅ Create hall
        hall = ExamHall(
            name=name,
            capacity=form.capacity.data,
            school_id=form.school_id.data,
            branch_id=form.branch_id.data
        )

        db.session.add(hall)
        db.session.commit()

        flash(f"Exam hall '{name}' added successfully!", "success")
        return redirect(url_for('main.all_exam_halls'))

    return render_template(
        "backend/pages/components/exam_halls/add-exam-hall.html",
        form=form,
        user=current_user
    )




@bp.route("/exam-halls/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_exam_hall(id):
    # 🔒 Check user status
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Check permission
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("You are not authorized to edit exam halls.", "danger")
        return redirect(url_for('main.dashboard'))

    # 🔍 Get hall
    hall = ExamHall.query.get_or_404(id)

    # 🔒 Extra security (school/branch isolation)
    if current_user.role.value != "superadmin":
        if hall.school_id != current_user.school_id:
            flash("Unauthorized access.", "danger")
            return redirect(url_for('main.all_exam_halls'))

        if current_user.role.value == "branch_admin" and hall.branch_id != current_user.branch_id:
            flash("Unauthorized access.", "danger")
            return redirect(url_for('main.all_exam_halls'))

    form = ExamHallForm(obj=hall)

    # 👉 Set hidden fields
    form.school_id.data = hall.school_id
    form.branch_id.data = hall.branch_id

    if form.validate_on_submit():
        name = form.name.data.strip()

        # 🚫 Prevent duplicate (exclude current record)
        existing = ExamHall.query.filter(
            ExamHall.school_id == form.school_id.data,
            ExamHall.branch_id == form.branch_id.data,
            ExamHall.name == name,
            ExamHall.id != hall.id
        ).first()

        if existing:
            flash("Another exam hall with this name already exists!", "warning")
            return redirect(url_for('main.edit_exam_hall', id=hall.id))

        # ✅ Update
        hall.name = name
        hall.capacity = form.capacity.data

        db.session.commit()

        flash(f"Exam hall '{name}' updated successfully!", "success")
        return redirect(url_for('main.all_exam_halls'))

    return render_template(
        "backend/pages/components/exam_halls/edit_exam_hall.html",
        form=form,
        hall=hall,
        user=current_user
    )


@bp.route("/exam-halls/delete/<int:id>", methods=["POST"])
@login_required
def delete_exam_hall(id):
    # 🔒 Check user status
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # 🔒 Check permission
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("You are not authorized to edit exam halls.", "danger")
        return redirect(url_for('main.dashboard'))

    # 🔍 Get hall
    hall = ExamHall.query.get_or_404(id)

    # 🔒 Extra security (school isolation)
    if current_user.role.value != "superadmin":
        if hall.school_id != current_user.school_id:
            flash("Unauthorized access.", "danger")
            return redirect(url_for('main.all_exam_halls'))

    try:
        db.session.delete(hall)
        db.session.commit()
        flash(f"Exam hall '{hall.name}' deleted successfully!", "success")

    except Exception as e:
        db.session.rollback()
        flash("Cannot delete hall. It may be assigned to exams.", "danger")

    return redirect(url_for('main.all_exam_halls'))



#-----------------------------------------
#--------- Exam Hall Assignments
#-----------------------------------------
@bp.route("/all/exam-hall-assignments")
@login_required
def all_exam_hall_assignments():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # ROLE-BASED FILTERING
    # =========================
    if current_user.role.value == "superadmin":
        # 🌍 Full access
        halls = ExamHall.query.order_by(ExamHall.created_at.desc()).all()
        total_students = ExamHallAssignment.query.count()  # all students in all schools

    elif current_user.role.value == "school_admin":
        # 🏫 Only halls in school (branch_id=None)
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(ExamHall.created_at.desc()).all()

        # 🏫 Total students in this school, only branch_id=None
        total_students = ExamHallAssignment.query.join(
            Student, ExamHallAssignment.student_id == Student.id
        ).filter(
            Student.school_id == current_user.school_id,
            Student.branch_id.is_(None)
        ).count()

    elif current_user.role.value == "branch_admin":
        # 🏢 Only halls in this branch
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(ExamHall.created_at.desc()).all()

        # 🏢 Total students in this branch
        total_students = ExamHallAssignment.query.join(
            Student, ExamHallAssignment.student_id == Student.id
        ).filter(
            Student.school_id == current_user.school_id,
            Student.branch_id == current_user.branch_id
        ).count()

    else:
        # ❌ Unauthorized
        flash("You are not authorized to view exam halls.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # SEND TO TEMPLATE
    # =========================
    return render_template(
        "backend/pages/components/exam_halls/all_exam_hall_assignments.html",
        halls=halls,
        total_students=total_students,
        user=current_user
    )


@bp.route("/delete/all-exam-hall-assignments", methods=["POST"])
@login_required
def delete_all_exam_hall_assignments():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # ROLE-BASED DELETION
    # =========================
    if current_user.role.value == "superadmin":
        # 🌍 Delete all assignments
        ExamHallAssignment.query.delete()

    elif current_user.role.value == "school_admin":
        # 🏫 Get all student IDs in this school, branch=None
        student_ids = [s.id for s in Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).all()]
        if student_ids:
            ExamHallAssignment.query.filter(ExamHallAssignment.student_id.in_(student_ids)).delete(synchronize_session=False)

    elif current_user.role.value == "branch_admin":
        # 🏢 Get all student IDs in this branch
        student_ids = [s.id for s in Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).all()]
        if student_ids:
            ExamHallAssignment.query.filter(ExamHallAssignment.student_id.in_(student_ids)).delete(synchronize_session=False)

    db.session.commit()
    flash("✅ All relevant exam hall assignments have been deleted.", "success")
    return redirect(url_for("main.all_exam_hall_assignments"))


@bp.route("/generate/exam-hall-assignments")
@login_required
def generate_exam_hall_assignments():

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    import random
    from collections import defaultdict

    assignments = []
    unassigned_students = []

    # =========================
    # HELPER FUNCTION – SMART DISTRIBUTION
    # =========================
    def generate_for_students(students, halls):
        random.shuffle(halls)
        hall_remaining_capacity = {h.id: h.capacity for h in halls}
        hall_ids = [h.id for h in halls]
        hall_index = 0
        hall_count = len(halls)

        # Avoid same hall
        current_assignments = {a.student_id: a.hall_id for a in ExamHallAssignment.query.all()}

        # Group students by class
        class_groups = defaultdict(list)
        for s in students:
            if not s.class_id:
                continue
            class_groups[s.class_id].append(s)

        # Shuffle classes
        class_ids = list(class_groups.keys())
        random.shuffle(class_ids)

        total_students = len(students)
        hall_fill_order = hall_ids.copy()
        if hall_count > total_students:
            hall_fill_order = hall_ids[:total_students]

        for class_id in class_ids:
            cls_students = class_groups[class_id]
            random.shuffle(cls_students)

            for student in cls_students:
                assigned = False
                start_index = hall_index % len(hall_fill_order)
                attempts = 0

                while attempts < len(hall_fill_order):
                    hall_id = hall_fill_order[start_index]

                    if hall_remaining_capacity[hall_id] > 0 and hall_id != current_assignments.get(student.id):
                        assignments.append(
                            ExamHallAssignment(
                                school_id=student.school_id,
                                branch_id=student.branch_id,
                                hall_id=hall_id,
                                student_id=student.id,
                                class_id=student.class_id
                            )
                        )
                        hall_remaining_capacity[hall_id] -= 1
                        assigned = True
                        hall_index = (start_index + 1) % len(hall_fill_order)
                        break
                    else:
                        start_index = (start_index + 1) % len(hall_fill_order)
                        attempts += 1

                if not assigned:
                    unassigned_students.append(student)

    # =========================
    # VALIDATE DATA
    # =========================
    def validate_data(students, halls):
        if not students:
            flash("❌ No students found.", "danger")
            return False
        if not halls:
            flash("❌ No exam halls found.", "danger")
            return False
        valid_students = [s for s in students if s.class_id]
        if not valid_students:
            flash("❌ Students have no classes assigned.", "danger")
            return False
        return True

    # =========================
    # ROLE-BASED DATA SELECTION
    # =========================
    if current_user.role.value == "superadmin":
        halls = ExamHall.query.order_by(ExamHall.capacity.asc()).all()
        students = Student.query.filter_by(shift=shift).order_by(Student.class_id).all()
        if not validate_data(students, halls):
            return redirect(url_for("main.all_exam_hall_assignments"))
        ExamHallAssignment.query.delete()
        db.session.commit()
        generate_for_students(students, halls)

    elif current_user.role.value == "school_admin":
        # ✅ Only halls branch_id=None
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(ExamHall.capacity.asc()).all()

        # ✅ Only students branch_id=None
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None,
            shift=shift
        ).order_by(Student.class_id).all()

        if not validate_data(students, halls):
            return redirect(url_for("main.all_exam_hall_assignments"))

        # Delete previous assignments in school branch_id=None
        ExamHallAssignment.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).delete()
        db.session.commit()
        generate_for_students(students, halls)

    else:  # branch_admin
        # ✅ Only halls in branch
        halls = ExamHall.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(ExamHall.capacity.asc()).all()

        # ✅ Only students in this branch
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id,
            shift=shift
        ).order_by(Student.class_id).all()

        if not validate_data(students, halls):
            return redirect(url_for("main.all_exam_hall_assignments"))

        # Delete previous assignments in this branch
        ExamHallAssignment.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).delete()
        db.session.commit()
        generate_for_students(students, halls)

    # =========================
    # SAVE ASSIGNMENTS
    # =========================
    if assignments:
        db.session.bulk_save_objects(assignments)
        db.session.commit()

    flash(f"✅ Assigned: {len(assignments)} | ❌ Not Assigned: {len(unassigned_students)} (Shift: {shift})", "success")
    return redirect(url_for("main.all_exam_hall_assignments"))

@bp.route('/print-exam-halls')
@login_required
def print_exam_halls():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ✅ SHIFT (QASAB)
    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    from collections import defaultdict

    # =========================
    # 🔐 SECURITY FILTER
    # =========================
    query = ExamHall.query.filter_by(school_id=current_user.school_id)

    if current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)
        current_branch = current_user.branch
    else:
        current_branch = None

    # =========================
    # ✅ HALLS WITH SHIFT FILTER (FROM STUDENT)
    # =========================
    halls = []

    for hall in query.order_by(ExamHall.branch_id, ExamHall.name).all():

        # ✅ FILTER by student.shift (NOT assignment.shift)
        filtered_assignments = [
            a for a in hall.hall_assignments
            if a.student and a.student.shift == shift
        ]

        if filtered_assignments:
            # 🔥 temp attach
            hall.filtered_assignments = filtered_assignments
            halls.append(hall)

    # =========================
    # 📅 TIMETABLE
    # =========================
    timetables = ExamTimetable.query.filter_by(
        school_id=current_user.school_id
    ).order_by(
        ExamTimetable.date, ExamTimetable.start_time
    ).all()

    timetable_map = defaultdict(lambda: defaultdict(list))
    all_dates = sorted(list(set([tt.date for tt in timetables])))

    for tt in timetables:
        timetable_map[tt.level_id][tt.date].append(tt.subject.name)

    # =========================
    # 📤 SEND TO TEMPLATE
    # =========================
    return render_template(
        'backend/pages/components/exam_halls/print_exam_halls.html',
        halls=halls,
        all_dates=all_dates,
        timetable_map=timetable_map,
        school=current_user.school,
        branch=current_branch,
        shift=shift,
        now_eat=now_eat
    )


@bp.route('/download-exam-halls-word')
@login_required
def download_exam_halls_word():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ✅ SHIFT (QASAB)
    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    from collections import defaultdict
    import io

    # =========================
    # 🔐 SECURITY FILTER
    # =========================
    query = ExamHall.query.filter_by(school_id=current_user.school_id)

    if current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)

    # =========================
    # ✅ FILTER HALLS BY SHIFT (FROM STUDENT)
    # =========================
    halls = []

    for hall in query.order_by(ExamHall.branch_id, ExamHall.name).all():

        filtered_assignments = [
            a for a in hall.hall_assignments
            if a.student and a.student.shift == shift
        ]

        if filtered_assignments:
            hall.filtered_assignments = filtered_assignments
            halls.append(hall)

    # =========================
    # 📅 TIMETABLE
    # =========================
    timetables = ExamTimetable.query.filter_by(
        school_id=current_user.school_id
    ).order_by(ExamTimetable.date).all()

    timetable_map = defaultdict(lambda: defaultdict(list))
    all_dates = sorted(list(set([tt.date for tt in timetables])))

    for tt in timetables:
        timetable_map[tt.level_id][tt.date].append(tt.subject.name)

    # =========================
    # 📝 WORD DOCUMENT
    # =========================
    doc = Document()

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.left_margin = section.right_margin = Cm(1)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(1)

    # LOGO
    logo_path = None
    if g.get('site_logo') and g.site_logo.get('main_logo'):
        logo_path = os.path.join(current_app.root_path, 'static', g.site_logo['main_logo'])

    for index, hall in enumerate(halls):
        if index > 0:
            doc.add_page_break()

        # =========================
        # HEADER (Dynamic Branch / School)
        # =========================
        branch_obj = hall.branch if hall.branch else None
        header_title = branch_obj.name.upper() if branch_obj else current_user.school.name.upper()

        if logo_path and os.path.exists(logo_path):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Cm(27.5))
        else:
            title = doc.add_heading(header_title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ✅ SHIFT TITLE
        p_shift = doc.add_paragraph()
        p_shift.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_shift = p_shift.add_run(f"SHIFT: {shift.upper()}")
        run_shift.bold = True

        # =========================
        # HALL INFO
        # =========================
        info_table = doc.add_table(rows=1, cols=2)

        hall_cell = info_table.cell(0, 0).paragraphs[0]
        hall_cell.add_run(f"HALL: {hall.name.upper()}").bold = True

        date_cell = info_table.cell(0, 1).paragraphs[0]
        date_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell.add_run(f"DATE: {now_eat().strftime('%d/%m/%Y')}").bold = True

        doc.add_paragraph()

        # =========================
        # TABLE
        # =========================
        num_cols = 3 + (len(all_dates) * 2)
        table = doc.add_table(rows=3, cols=num_cols)
        table.style = 'Table Grid'

        table.cell(0, 0).merge(table.cell(2, 0)).text = "NO"
        table.cell(0, 1).merge(table.cell(2, 1)).text = "STUDENT NAME"
        table.cell(0, 2).merge(table.cell(2, 2)).text = "CLASS"

        col_offset = 3
        for d in all_dates:

            d_cell = table.cell(0, col_offset).merge(table.cell(0, col_offset + 1))
            d_cell.text = d.strftime('%d-%b')

            first_assign = hall.filtered_assignments[0] if hall.filtered_assignments else None

            subjects = timetable_map[first_assign.class_obj.level_id][d] if first_assign else []

            table.cell(1, col_offset).text = subjects[0] if len(subjects) >= 1 else "-"
            table.cell(1, col_offset + 1).text = subjects[1] if len(subjects) >= 2 else "-"

            table.cell(2, col_offset).text = "SIGN"
            table.cell(2, col_offset + 1).text = "SIGN"

            col_offset += 2

        # =========================
        # DATA ROWS
        # =========================
        for i, assign in enumerate(hall.filtered_assignments, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = assign.student.full_name.title()
            row[2].text = assign.class_obj.name

            for c_idx in range(3, num_cols):
                row[c_idx].paragraphs[0].add_run().add_break()

        # =========================
        # FOOTER
        # =========================
        doc.add_paragraph("\n" + "_" * 100)
        footer_name = hall.branch.name.upper() if hall.branch else current_user.school.name.upper()
        footer = doc.add_paragraph()
        footer.add_run(footer_name).bold = True

    # =========================
    # SAVE
    # =========================
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)

    return send_file(
        target,
        as_attachment=True,
        download_name=f"Exam_Attendance_{shift}_{now_eat().strftime('%Y%m%d')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@bp.route('/download-exam-halls-png')
@login_required
def download_exam_halls_png():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ✅ SHIFT (QASAB)
    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    from collections import defaultdict

    # =========================
    # 📊 BASIC DATA
    # =========================
    school = current_user.school
    branch = current_user.branch if hasattr(current_user, 'branch') else None

    query = ExamHall.query.filter_by(school_id=current_user.school_id)

    if current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)

    # =========================
    # ✅ FILTER HALLS BY SHIFT (FROM STUDENT)
    # =========================
    halls = []

    for hall in query.order_by(ExamHall.branch_id, ExamHall.name).all():

        filtered_assignments = [
            a for a in hall.hall_assignments
            if a.student and a.student.shift == shift
        ]

        if filtered_assignments:
            hall.filtered_assignments = filtered_assignments
            halls.append(hall)

    # =========================
    # 📅 TIMETABLE
    # =========================
    timetables = ExamTimetable.query.filter_by(
        school_id=current_user.school_id
    ).order_by(ExamTimetable.date).all()

    timetable_map = defaultdict(lambda: defaultdict(list))
    all_dates = sorted(list(set([tt.date for tt in timetables])))

    for tt in timetables:
        timetable_map[tt.level_id][tt.date].append(tt.subject.name)

    # =========================
    # 🖼️ RENDER HTML
    # =========================
    html_content = render_template(
        'backend/pages/components/exam_halls/print_exam_halls-png.html',
        halls=halls,
        all_dates=all_dates,
        timetable_map=timetable_map,
        school=school,
        branch=branch,
        shift=shift,
        now_eat=now_eat
    )

    # =========================
    # ⚙️ WKHTMLTOIMAGE CONFIG
    # =========================
    path_to_wk = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltoimage.exe'

    if not os.path.exists(path_to_wk):
        return f"Khalad: Ma jiro faylka: {path_to_wk}", 500

    config = imgkit.config(wkhtmltoimage=path_to_wk)

    options = {
        'format': 'png',
        'encoding': "UTF-8",
        'width': 1300,
        'enable-local-file-access': '',
        'quiet': ''
    }

    try:
        img_bytes = imgkit.from_string(
            html_content,
            False,
            options=options,
            config=config
        )

        response = make_response(img_bytes)
        response.headers['Content-Type'] = 'image/png'
        response.headers['Content-Disposition'] = (
            f'attachment; filename=Exam_Attendance_{shift}_{now_eat().strftime("%Y%m%d")}.png'
        )

        return response

    except Exception as e:
        return f"Khalad ayaa ka dhacay: {str(e)}", 500



#-------------------------------------------
#------------- Exam Ticket
#-------------------------------------------

@bp.route("/all/exam-tickets")
@login_required
def all_exam_tickets():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # ================= ROLE FILTER =================
    if current_user.role.value == "superadmin":
        query = ExamTicket.query

    elif current_user.role.value == "school_admin":
        # ✅ Only his own students with branch_id = None
        query = ExamTicket.query.filter(
            ExamTicket.school_id == current_user.school_id,
            ExamTicket.branch_id.is_(None)  # Only unassigned branch
        )

    elif current_user.role.value == "branch_admin":
        # ✅ Only his branch
        query = ExamTicket.query.filter(
            ExamTicket.school_id == current_user.school_id,
            ExamTicket.branch_id == current_user.branch_id
        )

    else:
        flash("You are not authorized to view exam tickets.", "danger")
        return redirect(url_for("main.dashboard"))

    # ================= FETCH DATA =================
    tickets = query.order_by(ExamTicket.created_at.desc()).all()
    total_tickets = query.count()

    # ================= COUNT PER CLASS =================
    class_counts = {}
    for t in tickets:
        cls = t.class_obj.name
        class_counts[cls] = class_counts.get(cls, 0) + 1

    # ================= RESPONSE =================
    return render_template(
        "backend/pages/components/exam_halls/all_exam_tickets.html",
        tickets=tickets,
        total_tickets=total_tickets,
        class_counts=class_counts,
        user=current_user
    )


def get_shift_prefix(shift: str) -> str:
    """Return dynamic prefix based on shift."""
    shift_map = {
        "morning": "",
        "afternoon": "",
        "evening": ""  # optional
    }
    return shift_map.get(shift.lower(), "")


def generate_ticket_number_no_cache(shift="morning") -> str:
    """
    Generate ticket number dynamically from DB each time.
    No in-memory cache is used.
    Scoped to current user's school and branch.
    """
    # Get current user's school and branch
    school_id = current_user.school_id
    branch_id = getattr(current_user, "branch_id", None)

    # Prefix from shift + branch/school
    shift_prefix = get_shift_prefix(shift)

    if branch_id:
        branch = db.session.get(Branch, branch_id)
        name_prefix = generate_school_prefix(branch.name if branch else "")
    else:
        school = db.session.get(School, school_id)
        name_prefix = generate_school_prefix(school.name if school else "")

    prefix = f"{shift_prefix}{name_prefix}"

    # Query tickets with this prefix to find max number
    query = ExamTicket.query.filter(
        ExamTicket.school_id == school_id,
        ExamTicket.ticket_number.like(f"{prefix}%")
    )

    if branch_id:
        query = query.filter(ExamTicket.branch_id == branch_id)
    else:
        query = query.filter(ExamTicket.branch_id.is_(None))

    # Get max number after prefix
    max_number = query.with_entities(
        func.coalesce(
            func.max(
                cast(func.substring(ExamTicket.ticket_number, len(prefix)+1), Integer)
            ),
            0
        )
    ).scalar()

    new_number = max_number + 1
    return f"{prefix}{str(new_number).zfill(4)}"


@bp.route("/generate/exam-tickets")
@login_required
def generate_exam_tickets():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_tickets"))

    # =========================
    # 🔹 Fetch students based on role
    # =========================
    if current_user.role.value == "superadmin":
        students = Student.query.filter_by(shift=shift).all()
        exams = Exam.query.filter(Exam.status == "draft").all()
        assigned_students = {a.student_id for a in ExamHallAssignment.query.all()}
    elif current_user.role.value == "school_admin":
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None,
            shift=shift
        ).all()
        exams = Exam.query.filter(
            Exam.school_id == current_user.school_id,
            Exam.branch_id.is_(None),
            Exam.status == "draft"
        ).all()
        assigned_students = {
            a.student_id for a in ExamHallAssignment.query.join(Student).filter(
                Student.school_id == current_user.school_id,
                Student.branch_id.is_(None),
                Student.shift == shift
            ).all()
        }
    else:  # branch_admin
        students = Student.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id,
            shift=shift
        ).all()
        exams = Exam.query.filter(
            Exam.school_id == current_user.school_id,
            Exam.branch_id == current_user.branch_id,
            Exam.status == "draft"
        ).all()
        assigned_students = {
            a.student_id for a in ExamHallAssignment.query.join(Student).filter(
                Student.school_id == current_user.school_id,
                Student.branch_id == current_user.branch_id,
                Student.shift == shift
            ).all()
        }

    # =========================
    # 🔹 Check existing tickets
    # =========================
    existing_tickets = {(t.student_id, t.exam_id) for t in ExamTicket.query.all()}

    tickets_created = 0
    tickets_skipped = 0

    # =========================
    # 🔹 Generate tickets
    # =========================
    for student in students:
        if student.id not in assigned_students:
            tickets_skipped += 1
            continue

        for exam in exams:
            if (student.id, exam.id) in existing_tickets:
                tickets_skipped += 1
                continue

            hall_assignment = ExamHallAssignment.query.filter_by(student_id=student.id).first()
            if not hall_assignment:
                tickets_skipped += 1
                continue

            ticket_number = generate_ticket_number_no_cache(shift=shift)

            ticket = ExamTicket(
                school_id=student.school_id,
                branch_id=student.branch_id,
                exam_id=exam.id,
                student_id=student.id,
                class_id=student.class_id,
                hall_id=hall_assignment.hall_id,
                ticket_number=ticket_number,
                status="active"
            )
            db.session.add(ticket)
            tickets_created += 1

    if tickets_created:
        db.session.commit()

    flash(
        f"✅ Tickets created: {tickets_created} | ❌ Skipped (no hall, draft exam, or already exists): {tickets_skipped} (Shift: {shift})",
        "success"
    )
    return redirect(url_for("main.all_exam_tickets"))


@bp.route("/print-exam-ticket/<int:ticket_id>")
@login_required
def print_exam_ticket(ticket_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    ticket = ExamTicket.query.get_or_404(ticket_id)
    return render_template(
        "backend/pages/components/exam_halls/print_exam_tickets.html",
        ticket=ticket
    )


@bp.route("/delete/exam-ticket/<int:ticket_id>", methods=["POST"])
@login_required
def delete_exam_ticket(ticket_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    ticket = ExamTicket.query.get_or_404(ticket_id)

    # ✅ Role-based access check
    if current_user.role.value == "superadmin":
        pass  # full access
    elif current_user.role.value == "school_admin":
        if ticket.school_id != current_user.school_id or ticket.branch_id is not None:
            flash("❌ Unauthorized to delete this ticket.", "danger")
            return redirect(url_for("main.all_exam_tickets"))
    elif current_user.role.value == "branch_admin":
        if ticket.school_id != current_user.school_id or ticket.branch_id != current_user.branch_id:
            flash("❌ Unauthorized to delete this ticket.", "danger")
            return redirect(url_for("main.all_exam_tickets"))
    else:
        flash("❌ Unauthorized.", "danger")
        return redirect(url_for("main.all_exam_tickets"))

    # Delete ticket
    db.session.delete(ticket)
    db.session.commit()
    flash(f"✅ Ticket {ticket.ticket_number} deleted successfully.", "success")
    return redirect(url_for("main.all_exam_tickets"))


@bp.route("/delete/all-exam-tickets", methods=["POST"])
@login_required
def delete_all_exam_tickets():
    # Role-based restriction
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Base query depending on role
    if current_user.role.value == "superadmin":
        query = ExamTicket.query
    elif current_user.role.value == "school_admin":
        query = ExamTicket.query.filter(
            ExamTicket.school_id == current_user.school_id,
            ExamTicket.branch_id.is_(None)
        )
    else:  # branch_admin
        query = ExamTicket.query.filter(
            ExamTicket.school_id == current_user.school_id,
            ExamTicket.branch_id == current_user.branch_id
        )

    count = query.count()
    if count == 0:
        flash("⚠️ No tickets to delete.", "info")
        return redirect(url_for("main.all_exam_tickets"))

    query.delete(synchronize_session=False)
    db.session.commit()
    flash(f"✅ All {count} tickets deleted successfully.", "success")
    return redirect(url_for("main.all_exam_tickets"))


@bp.route("/print-all-exam-tickets")
@login_required
def print_all_exam_tickets():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required (morning/afternoon).", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    # Base query
    query = ExamTicket.query.join(ExamTicket.student)

    # Role-based filtering
    if current_user.role.value == "superadmin":
        current_branch = None
        current_school = None
    elif current_user.role.value == "school_admin":
        query = query.filter(
            Student.school_id == current_user.school_id,
            Student.branch_id.is_(None)
        )
        current_school = current_user.school
        current_branch = None
    elif current_user.role.value == "branch_admin":
        query = query.filter(
            Student.school_id == current_user.school_id,
            Student.branch_id == current_user.branch_id
        )
        current_school = current_user.school
        current_branch = current_user.branch
    else:
        flash("❌ Unauthorized access.", "danger")
        return redirect(url_for("main.dashboard"))

    # Shift filter
    query = query.filter(Student.shift == shift)

    # Range filter (optional)
    min_id = request.args.get("min", type=int)
    max_id = request.args.get("max", type=int)
    if min_id is not None and max_id is not None:
        query = query.filter(ExamTicket.id.between(min_id, max_id))

    tickets = query.order_by(ExamTicket.id).all()

    return render_template(
        "backend/pages/components/exam_halls/print_all_exam_tickets.html",
        tickets=tickets,
        shift=shift,
        school=current_school,
        branch=current_branch,
        now_eat=now_eat
    )


@bp.route('/download-exam-tickets-png')
@login_required
def download_exam_tickets_png():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    shift = request.args.get("shift")
    if not shift:
        flash("❌ Shift is required.", "danger")
        return redirect(url_for("main.all_exam_hall_assignments"))

    # 🔐 Query sax ah oo leh Join
    query = ExamTicket.query.join(Student, ExamTicket.student_id == Student.id)

    if current_user.role.value == "school_admin":
        query = query.filter(Student.school_id == current_user.school_id)
    elif current_user.role.value == "branch_admin":
        query = query.filter(
            Student.school_id == current_user.school_id,
            Student.branch_id == current_user.branch_id
        )
    
    query = query.filter(Student.shift == shift)
    tickets = query.all()

    # 🖼️ Render Template
    html = render_template(
        "backend/pages/components/exam_halls/print_exam_tickets_png.html",
        tickets=tickets
    )

    # ⚙️ Config
    config = imgkit.config(
        wkhtmltoimage=r"C:\Program Files\wkhtmltopdf\bin\wkhtmltoimage.exe"
    )

    options = {
        'format': 'png',
        'enable-local-file-access': '', 
        'width': 2480,  
        'quality': '100',
        'encoding': "UTF-8",
        'quiet': ''
    }

    try:
        # Halkan 'base_url' waa laga saaray si uusan error u bixin
        img = imgkit.from_string(html, False, config=config, options=options)

        response = make_response(img)
        response.headers['Content-Type'] = 'image/png'
        response.headers['Content-Disposition'] = f'attachment; filename=exam_tickets_{shift}.png'
        return response

    except Exception as e:
        return f"Error: {str(e)}", 500



#------------------------------------------------
#-------------------- Exam Subjects
#------------------------------------------------


@bp.route("/all/exam/subjects")
@login_required
def all_exam_subjects():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 🔐 ROLE-BASED ACCESS
    # =========================
    if current_user.role.value == "superadmin":
        # 🌍 Access to all exam subjects
        exam_subjects = ExamSubject.query.order_by(ExamSubject.created_at.desc()).all()

    elif current_user.role.value == "school_admin":
        # 🏫 Only exam subjects for this school with branch_id=None
        exam_subjects = ExamSubject.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(ExamSubject.created_at.desc()).all()

    elif current_user.role.value == "branch_admin":
        # 🏢 Only exam subjects for this branch
        exam_subjects = ExamSubject.query.filter_by(
            school_id=current_user.school_id,
            branch_id=current_user.branch_id
        ).order_by(ExamSubject.created_at.desc()).all()

    else:
        # ❌ Unauthorized
        flash("You are not authorized to view exam subjects.", "danger")
        return redirect(url_for("main.dashboard"))

    # =========================
    # 📤 RENDER TEMPLATE
    # =========================
    return render_template(
        "backend/pages/components/exam_subjects/all_exam_subjects.html",
        exam_subjects=exam_subjects,
        user=current_user
    )


# =========================
# Route for Creating Exam Subject
# =========================
@bp.route("/add/exam-subject", methods=["GET", "POST"])
@login_required
def add_exam_subject():
    # 1. Hubi status-ka user-ka
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = ExamSubjectForm()

    # 2. ROLE-BASED DATA
    school_id = current_user.school_id
    # Haddii uu yahay branch_admin, isticmaal branch-kiisa, haddii kale waa None (Main School)
    branch_id = current_user.branch_id if (hasattr(current_user.role, 'value') and current_user.role.value == "branch_admin") else None

    # 3. DYNAMIC CHOICES
    if request.method == 'POST':
        # Soo qaado sanadka si loogu sifeeyo imtixaannada
        selected_year_id = request.form.get('academic_year_id', type=int)
        if selected_year_id:
            exams = Exam.query.filter_by(
                school_id=school_id, 
                branch_id=branch_id, 
                academic_year_id=selected_year_id
            ).all()
            form.exam_id.choices = [(e.id, e.exam_name) for e in exams]
    else:
        # GET request: Deji xogta bilowga ah ee hidden fields
        form.school_id.data = school_id
        form.branch_id.data = branch_id

    # 4. VALIDATION & SAVE (Multi-Class & Multi-Subject Logic)
    if form.validate_on_submit():
        exam_id = form.exam_id.data
        academic_year_id = form.academic_year_id.data
        selected_classes = form.class_ids.data    # Liiska fasallada (CUSUB)
        selected_subjects = form.subject_ids.data  # Liiska maaddooyinka
        
        added_count = 0
        skipped_count = 0

        try:
            # Hubi imtixaanka amniga dartiis
            exam_check = Exam.query.filter_by(id=exam_id, academic_year_id=academic_year_id).first()
            if not exam_check:
                flash("Imtixaanka la doortay kama tirsana sannad-dugsiyeedka aad dooratay.", "danger")
                return render_template("backend/pages/components/exam_subjects/add_exam_subject.html", form=form)

            # LOOP-KA 1AAD: Mari fasal kasta oo la doortay
            for c_id in selected_classes:
                # LOOP-KA 2AAD: Mari maaddo kasta oo la doortay
                for sub_id in selected_subjects:
                    
                    # Duplicate Check: Hubi in xiriirkan (Exam + Class + Subject) uu jiro
                    existing = ExamSubject.query.filter_by(
                        exam_id=exam_id,
                        class_id=c_id,
                        subject_id=sub_id,
                        academic_year_id=academic_year_id
                    ).first()

                    if not existing:
                        new_es = ExamSubject(
                            school_id=school_id,
                            branch_id=branch_id,
                            academic_year_id=academic_year_id,
                            exam_id=exam_id,
                            class_id=c_id,    # Keydi fasalka hadda loop-ka ku jira
                            subject_id=sub_id, # Keydi maaddada hadda loop-ka ku jirta
                            total_marks=form.total_marks.data,
                            pass_marks=form.pass_marks.data
                        )
                        db.session.add(new_es)
                        added_count += 1
                    else:
                        skipped_count += 1

            # 5. COMMIT & FEEDBACK
            if added_count > 0:
                db.session.commit()
                flash(f"Si guul leh ayaa loo daray {added_count} xiriir (Fasallo & Maaddooyin)!", "success")
            
            if skipped_count > 0:
                flash(f"{skipped_count} xog hore ayey u jireen, waa laga booday.", "info")

            return redirect(url_for("main.all_exam_subjects"))

        except Exception as e:
            db.session.rollback()
            flash(f"Khalad ayaa dhacay intii lagu guda jiray kaydinta: {str(e)}", "danger")

    return render_template(
        "backend/pages/components/exam_subjects/add_exam_subject.html",
        form=form,
        user=current_user
    )

    

@bp.route("/get-exams/<int:year_id>")
@login_required
def get_exams_by_year(year_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Aqoonsiga dugsiga iyo laanta qofka soo galay
    school_id = current_user.school_id
    # Haddii uu yahay branch_admin, geli branch_id, haddii kale (school_admin) waa None
    branch_id = current_user.branch_id if current_user.role.value == "branch_admin" else None

    # 2. Query-ga imtixaannada iyadoo la raacayo shuruucda:
    # - Waa inay ahaadaan dugsiga user-ka (school_id)
    # - Waa inay ahaadaan sannad-dugsiyeedka la doortay (year_id)
    query = Exam.query.filter_by(school_id=school_id, academic_year_id=year_id)

    # 3. Haddii uu yahay branch_admin, kaliya u soo saar imtixaannada laantiisa
    if branch_id:
        query = query.filter_by(branch_id=branch_id)
    
    exams = query.order_by(Exam.exam_name).all()

    # 4. U soo celi xogta qaab JSON ah si JS-ka u akhriyo
    return jsonify({
        "exams": [{"id": e.id, "name": e.exam_name} for e in exams]
    })

@bp.route("/edit/exam-subject/<int:id>", methods=["GET", "POST"])
@login_required
def edit_exam_subject(id):
    # 1. Hubi xaaladda isticmaalaha
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 2. Soo jiid xogta hadda jirta
    exam_subject = ExamSubject.query.get_or_404(id)

    # 3. Bilow foomka (obj=exam_subject waxay si otomaatig ah u pre-fill gareysaa field-yada caadiga ah)
    form = ExamSubjectForm(obj=exam_subject)

    if form.validate_on_submit():
        # Soo qabo subject_id-ga la doortay (maadaama foomku yahay SelectMultipleField)
        selected_subject_id = form.subject_ids.data[0] if isinstance(form.subject_ids.data, list) and form.subject_ids.data else form.subject_ids.data
        
        # 4. Duplicate Check (Si looga hortago in hal maaddo laba jeer loo diwaangaliyo hal imtixaan/sannad gudahiis)
        existing = ExamSubject.query.filter(
            ExamSubject.id != id,  # Iska reeb kan hadda la edit-gareynayo
            ExamSubject.exam_id == form.exam_id.data,
            ExamSubject.subject_id == selected_subject_id,
            ExamSubject.academic_year_id == form.academic_year_id.data,
            ExamSubject.branch_id == (form.branch_id.data or None)
        ).first()

        if existing:
            flash("Habayntan (Imtixaanka, Maaddada, iyo Sannadka) horey ayay u jirtay!", "warning")
            return redirect(url_for("main.edit_exam_subject", id=id))

        try:
            # 5. Cusboonaysii xogta
            exam_subject.academic_year_id = form.academic_year_id.data
            exam_subject.exam_id = form.exam_id.data
            exam_subject.subject_id = selected_subject_id
            exam_subject.total_marks = form.total_marks.data
            exam_subject.pass_marks = form.pass_marks.data

            db.session.commit()
            flash("Si guul leh ayaa loo cusboonaysiiyay xogta!", "success")
            return redirect(url_for("main.all_exam_subjects"))

        except Exception as e:
            db.session.rollback()
            flash(f"Khalad ayaa dhacay xilli la kaydinayay: {str(e)}", "danger")

    # 6. Pre-populate logic loogu talagalay GET Request
    elif request.method == 'GET':
        # Field-yada 'Select' waxay mararka qaarkood u baahan yihiin in gacanta lagu sii xaqiijiyo xogtooda
        form.academic_year_id.data = exam_subject.academic_year_id
        form.exam_id.data = exam_subject.exam_id
        # Sababta subject_ids loogu shubayo list [] waa maadaama uu foomku yahay SelectMultipleField
        form.subject_ids.data = [exam_subject.subject_id]

    return render_template(
        "backend/pages/components/exam_subjects/edit_exam_subject.html",
        form=form,
        user=current_user,
        exam_subject=exam_subject
    )


@bp.route("/delete/exam-subject/<int:id>", methods=["POST"])
@login_required
def delete_exam_subject(id):
    # Hubi in user-ka uu firfircoon yahay
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Hel ExamSubject-ka la tirtirayo
    exam_subject = ExamSubject.query.get_or_404(id)

    # 🔐 ROLE-BASED ACCESS
    if current_user.role.value == "superadmin":
        # superadmin wuu tirtiri karaa
        pass
    elif current_user.role.value == "school_admin":
        if exam_subject.school_id != current_user.school_id:
            flash("You are not authorized to delete this exam subject.", "danger")
            return redirect(url_for("main.all_exam_subjects"))
    elif current_user.role.value == "branch_admin":
        if exam_subject.branch_id != current_user.branch_id:
            flash("You are not authorized to delete this exam subject.", "danger")
            return redirect(url_for("main.all_exam_subjects"))
    else:
        flash("You are not authorized to perform this action.", "danger")
        return redirect(url_for("main.all_exam_subjects"))

    # Tirtir record-ka
    try:
        db.session.delete(exam_subject)
        db.session.commit()
        flash("Exam subject deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting exam subject: {str(e)}", "danger")

    return redirect(url_for("main.all_exam_subjects"))



#------------------------
#----------- Exam Paper
#--------------------------


@bp.route("/exam/multi-publish", methods=["GET", "POST"])
@login_required
def multi_publish_exam():

    # =====================
    # 1. SECURITY CHECK
    # =====================
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = ExamMultiPublishForm()

    s_id = current_user.school_id
    b_id = getattr(current_user, 'branch_id', None)
    role = current_user.role.value

    is_branch = (role == "branch_admin" and b_id is not None)

    # =====================
    # STRICT ROLE SCOPE FIX
    # =====================
    def scope_query(query):
        query = query.filter_by(school_id=s_id)

        if role == "school_admin":
            # ONLY global school-level data
            return query.filter_by(branch_id=None)

        elif is_branch:
            # ONLY this branch data
            return query.filter_by(branch_id=b_id)

        return query

    # =====================
    # 2. EXAMS
    # =====================
    exam_query = scope_query(
        Exam.query.filter_by(status='draft')
    )

    form.exam_id.choices = [
        (e.id, e.exam_name)
        for e in exam_query.order_by(Exam.id.desc()).all()
    ]

    # =====================
    # 3. TEACHERS
    # =====================
    assign_query = scope_query(TeacherAssignment.query)

    assigned_teacher_ids = [
        t.teacher_id
        for t in assign_query.with_entities(TeacherAssignment.teacher_id)
        .distinct()
        .all()
    ]

    teachers = (
        Teacher.query.filter(Teacher.id.in_(assigned_teacher_ids)).all()
        if assigned_teacher_ids else []
    )

    # =====================
    # 4. FORM SUBMIT
    # =====================
    if form.validate_on_submit():

        selected_exam_id = form.exam_id.data
        duration = form.duration_minutes.data

        try:
            raw_data = form.teacher_ids_hidden.data

            if not raw_data:
                flash("Fadlan dooro ugu yaraan hal macalin.", "warning")
                return redirect(url_for('main.multi_publish_exam'))

            selected_teacher_ids = json.loads(raw_data)

            papers_created = 0

            # ExamSubject map
            exam_subjects = ExamSubject.query.filter_by(
                exam_id=selected_exam_id
            ).all()

            exam_map = {
                (int(es.subject_id), int(es.class_id)): es.id
                for es in exam_subjects
            }

            for t_id in selected_teacher_ids:

                t_assign_query = scope_query(
                    TeacherAssignment.query.filter_by(teacher_id=t_id)
                )

                assignments = t_assign_query.all()

                for assign in assignments:

                    s_ids = assign.subject_ids

                    if isinstance(s_ids, str):
                        try:
                            s_ids = json.loads(s_ids)
                        except:
                            continue

                    if not isinstance(s_ids, list):
                        continue

                    for sub_id in s_ids:

                        key = (int(sub_id), int(assign.class_id))

                        if key not in exam_map:
                            continue

                        es_id = exam_map[key]

                        existing = ExamPaper.query.filter_by(
                            exam_subject_id=es_id,
                            teacher_id=t_id,
                            school_id=s_id
                        ).first()

                        if existing:
                            continue

                        db.session.add(ExamPaper(
                            school_id=s_id,
                            branch_id=(b_id if is_branch else None),
                            exam_subject_id=es_id,
                            teacher_id=t_id,
                            duration_minutes=duration,
                            status='draft'
                        ))

                        papers_created += 1

            if papers_created:
                db.session.commit()
                flash(f"✅ {papers_created} warqadood ayaa la sameeyay.", "success")
                return redirect(url_for('main.all_exam_papers'))
            else:
                flash("⚠️ Warqado cusub lama helin.", "info")

        except Exception as e:
            db.session.rollback()
            flash(f"❌ Khalad: {str(e)}", "danger")

    # =====================
    # 5. GET DEFAULT VALUES
    # =====================
    if request.method == "GET":
        form.school_id.data = s_id
        form.branch_id.data = b_id if is_branch else None
        form.duration_minutes.data = 60

    # =====================
    # 6. RECENT PAPERS
    # =====================
    recent_query = scope_query(ExamPaper.query)

    recent_papers = recent_query.order_by(
        ExamPaper.id.desc()
    ).limit(5).all()

    return render_template(
        "backend/pages/components/exams/multi_publish.html",
        form=form,
        teachers=teachers,
        user=current_user,
        recent_papers=recent_papers
    )

@bp.route("/all/exam/papers")
@login_required
def all_exam_papers():

    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin",  "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    role = current_user.role.value
    s_id = current_user.school_id
    b_id = getattr(current_user, 'branch_id', None)

    query = ExamPaper.query.options(
        joinedload(ExamPaper.teacher),
        joinedload(ExamPaper.exam_subject).joinedload(ExamSubject.subject)
    )

    # ================= SUPERADMIN =================
    if role == "superadmin":
        exam_papers = query.order_by(ExamPaper.id.desc()).all()

    # ================= SCHOOL ADMIN =================
    elif role == "school_admin":
        exam_papers = query.filter(
            ExamPaper.school_id == s_id
        ).order_by(ExamPaper.id.desc()).all()

    # ================= BRANCH ADMIN =================
    elif role == "branch_admin":
        exam_papers = query.filter(
            ExamPaper.school_id == s_id,
            ExamPaper.branch_id == b_id
        ).order_by(ExamPaper.id.desc()).all()

    # ================= TEACHER =================
    elif role == "teacher":

        teacher = current_user.teacher

        if not teacher:
            flash("Teacher profile lama helin.", "danger")
            return redirect(url_for("main.dashboard"))

        # 🔥 FIX LIST ISSUE
        if isinstance(teacher, list):
            teacher = teacher[0]

        exam_papers = query \
            .join(ExamSubject, ExamPaper.exam_subject_id == ExamSubject.id) \
            .join(Exam, ExamSubject.exam_id == Exam.id) \
            .join(AcademicYear, Exam.academic_year_id == AcademicYear.id) \
            .filter(
                ExamPaper.school_id == s_id,
                ExamPaper.teacher_id == teacher.id,
                ExamPaper.status == 'draft',
                AcademicYear.is_active.is_(True)
            ) \
            .order_by(ExamPaper.id.desc()) \
            .all()


    else:
        flash("Access denied.", "danger")
        return redirect(url_for("main.dashboard"))

    return render_template(
        "backend/pages/components/exams/all_exam_papers.html",
        exam_papers=exam_papers,
        user=current_user
    )


@bp.route("/exams/paper/update_status/<int:id>", methods=["POST"])
@login_required
def update_exam_paper_status(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))


    # 🔥 JSON SAFETY
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"success": False, "message": "Invalid JSON"}), 400

    new_status = data.get("status")

    # 🔥 CORRECT MODEL
    paper = ExamPaper.query.get_or_404(id)

    # Permissions
    if current_user.role.value not in ['superadmin', 'school_admin']:
        return jsonify({"success": False, "message": "No permission"}), 403

    # Validate status
    valid_statuses = ['draft', 'published', 'closed']
    if new_status not in valid_statuses:
        return jsonify({"success": False, "message": "Invalid status"}), 400

    try:
        paper.status = new_status
        db.session.commit()

        return jsonify({
            "success": True,
            "message": f"Paper status changed to {new_status}"
        })

    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500


@bp.route("/delete/exam-paper/<int:id>", methods=["POST"])
@login_required
def delete_exam_paper(id):
    # 1. Hubi in account-ku active yahay
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 2. Hel Warqadda (Halkii aad ka raadin lahayd ExamSubject)
    paper = ExamPaper.query.get_or_404(id)
    user_role = current_user.role.value
    can_delete = False

    # 3. 🔐 ROLE-BASED ACCESS CONTROL
    if user_role == "superadmin":
        can_delete = True
    elif user_role == "school_admin":
        if paper.school_id == current_user.school_id:
            can_delete = True
    elif user_role == "branch_admin":
        if current_user.branch_id is not None and paper.branch_id == current_user.branch_id:
            can_delete = True
    elif user_role == "teacher":
        # Macalinku wuxuu tirtiri karaa kaliya warqadiisa haddii aysan 'published' ahayn
        if paper.teacher_id == current_user.id and paper.status != 'published':
            can_delete = True

    if not can_delete:
        flash("Ma haysatid ogolaansho aad ku tirtirto warqaddan.", "danger")
        return redirect(url_for("main.all_exam_papers"))

    # 4. 🧨 TIRtiridda XOGTA LA XIRIIRTA (Questions)
    # Haddii aadan database-ka ku dhex qeexin 'ondelete=CASCADE'
    try:
        # Tirtir dhamaan su'aalaha warqaddan ku xiran ka hor intaan warqadda la tirtirin
        ExamQuestion.query.filter_by(paper_id=id).delete()
        
        # 5. Tirtir Warqadda lafteeda
        db.session.delete(paper)
        db.session.commit()
        
        flash("Warqaddii imtixaanka iyo su'aalihii ku xirnaa waa la tirtiray!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Khalad ayaa dhacay xiligii tirtirista: {str(e)}", "danger")

    return redirect(url_for("main.all_exam_papers"))


@bp.route("/delete-all/exam-papers", methods=["POST"])
@login_required
def delete_all_exam_papers():
    # 1. Hubi in account-ku active yahay
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    user_role = current_user.role.value
    
    # 2. 🔐 ROLE-BASED ACCESS CONTROL
    # Kaliya heerarka sare ayaa oggol in ay 'Bulk Delete' sameeyaan
    allowed_roles = ["superadmin", "school_admin", "branch_admin"]
    if user_role not in allowed_roles:
        flash("Ma haysatid ogolaansho aad ku tirtirto dhammaan warqadaha.", "danger")
        return redirect(url_for("main.all_exam_papers"))

    try:
        # 3. 🎯 GO'AAMI XOGTA LA TIRTIRAYO (Filtering)
        query = ExamPaper.query

        if user_role == "superadmin":
            # Superadmin: Wax kasta oo nidaamka ku jira
            papers_to_delete = query.all()
            
        elif user_role == "school_admin":
            # School Admin: Dhammaan warqadaha dugsiga (Branch kasta ha ahaadeen)
            papers_to_delete = query.filter_by(school_id=current_user.school_id).all()
            
        elif user_role == "branch_admin":
            # Branch Admin: Kaliya warqadaha laankiisa (Branch)
            papers_to_delete = query.filter_by(
                school_id=current_user.school_id, 
                branch_id=current_user.branch_id
            ).all()

        if not papers_to_delete:
            flash("Ma jiraan warqado imtixaan oo halkan laga tirtiro.", "info")
            return redirect(url_for("main.all_exam_papers"))

        # 4. 🧨 TIRtiridda XOGTA (Bulk Deletion)
        deleted_count = 0
        for paper in papers_to_delete:
            # A. Tirtir su'aalaha (Questions)
            ExamQuestion.query.filter_by(paper_id=paper.id).delete()
            
            # B. Tirtir Natiijooyinka (Haddii aad leedahay StudentResult)
            # StudentResult.query.filter_by(paper_id=paper.id).delete()
            
            # C. Tirtir Warqadda (The Paper itself)
            db.session.delete(paper)
            deleted_count += 1

        # 5. COMMIT & FINALIZE
        db.session.commit()
        
        flash(f"Guul! Waxaa la sifeeyay {deleted_count} warqadood iyo xogtooda.", "success")

    except Exception as e:
        db.session.rollback()
        flash(f"Khalad ayaa dhacay: {str(e)}", "danger")

    return redirect(url_for("main.all_exam_papers"))




#----------------------------------------------
#------------------- Exam Results
#----------------------------------------------
# ===========================


@bp.route("/exam-paper/<int:paper_id>/add-question", methods=["GET", "POST"])
@login_required
def add_exam_question(paper_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["teacher"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    paper = ExamPaper.query.get_or_404(paper_id)
    # Soo qaad su'aalihii u dambeeyay (oo loo habeeyay revindex-ka template-ka)
    recent_questions = ExamQuestion.query.filter_by(paper_id=paper.id)\
                                         .order_by(ExamQuestion.created_at.desc()).all()
    questions_count = ExamQuestion.query.filter_by(paper_id=paper.id).count()
    total_marks = db.session.query(db.func.sum(ExamQuestion.marks))\
                            .filter(ExamQuestion.paper_id == paper.id).scalar() or 0

    # Security check
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if current_user.role.value == 'teacher':
        if not teacher or paper.teacher_id != teacher.id:
            flash("Ma haysatid ogolaansho aad wax kaga beddesho warqaddan.", "danger")
            return redirect(url_for('main.all_exam_papers'))

    form = ExamQuestionForm()

    if request.method == "GET":
        form.paper_id.data = paper.id
        form.school_id.data = paper.school_id
        form.branch_id.data = paper.branch_id
        form.sort_order.data = questions_count + 1

    if form.validate_on_submit():
        try:
            q_id = request.form.get('question_id')
            question = ExamQuestion.query.get(q_id) if q_id and q_id != "" else None

            # --- HANDLE IMAGE ---
            image_file = request.files.get('question_image_file')
            image_rel_path = question.question_image if question else None

            if image_file and image_file.filename:
                ext = os.path.splitext(image_file.filename)[1]
                unique_id = uuid.uuid4().hex[:8]
                safe_filename = f"q-{paper.id}-{unique_id}{ext}"
                user_subfolder = os.path.join(current_app.root_path, 'static/backend/uploads/questions', str(paper.school_id))
                os.makedirs(user_subfolder, exist_ok=True)
                full_path = os.path.join(user_subfolder, safe_filename)
                image_file.save(full_path)

                if question and question.question_image:
                    old_path = os.path.join(current_app.root_path, 'static', question.question_image)
                    if os.path.exists(old_path): os.remove(old_path)

                image_rel_path = f"backend/uploads/questions/{paper.school_id}/{safe_filename}"

            # --- MCQ OPTIONS ---
            options_data = None
            if form.question_type.data == 'mcq':
                options_list = request.form.getlist('mcq_options[]')
                options_data = {chr(65 + i): opt.strip() for i, opt in enumerate(options_list) if opt.strip()}

            # --- SAVE OR UPDATE ---
            if question:
                question.question_type = form.question_type.data
                question.question_text = form.question_text.data
                question.question_image = image_rel_path
                question.marks = form.marks.data
                question.options = options_data
                question.correct_answer = form.correct_answer.data.upper() if form.correct_answer.data else None
                question.explanation = form.explanation.data
                question.shuffle_options = form.shuffle_options.data
                msg = "Su'aasha waa la cusboonaysiiyay!"
            else:
                question = ExamQuestion(
                    paper_id=paper.id, school_id=paper.school_id, branch_id=paper.branch_id,
                    question_type=form.question_type.data, question_text=form.question_text.data,
                    question_image=image_rel_path, marks=form.marks.data, options=options_data,
                    correct_answer=form.correct_answer.data.upper() if form.correct_answer.data else None,
                    explanation=form.explanation.data, sort_order=questions_count + 1,
                    shuffle_options=form.shuffle_options.data
                )
                db.session.add(question)
                msg = "Su'aasha waa la kaydiyay!"

            db.session.commit()

            # AJAX Response (Kani waa qaybta muhiimka ah)
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'success': True,
                    'message': msg,
                    'question': {
                        'id': question.id,
                        'text': question.question_text,
                        'type': question.question_type,
                        'marks': question.marks,
                        'correct': question.correct_answer,
                        'options': question.options
                    }
                })

            flash(msg, "success")
            return redirect(url_for('main.add_exam_question', paper_id=paper.id))

        except Exception as e:
            db.session.rollback()
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': str(e)}), 500
            flash(f"Khalad: {str(e)}", "danger")

    return render_template(
        "backend/pages/components/exams/add_question.html",
        form=form, paper=paper, questions_count=questions_count,
        total_marks=total_marks, recent_questions=recent_questions,
        user=current_user
    )


# --- DELETE ROUTE ---

@bp.route("/delete-question/<int:id>", methods=["POST"])
@login_required
def delete_question(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["teacher"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    question = ExamQuestion.query.get_or_404(id)
    try:
        # Tirtir sawirka haddii uu jiro
        if question.question_image:
            img_path = os.path.join(current_app.root_path, 'static', question.question_image.lstrip('/'))
            if os.path.exists(img_path):
                os.remove(img_path)
        
        db.session.delete(question)
        db.session.commit()
        return jsonify({'success': True, 'message': 'Su\'aasha waa la tirtiray'})
    except Exception as e:
        db.session.rollback()
        # Aad bay muhiim u tahay inaan halkan ku soo celino JSON si uusan JS u dhiman
        return jsonify({'success': False, 'message': str(e)}), 500
    

@bp.route("/exam-paper/<int:id>/view", methods=["GET"])
@login_required
def view_exam_paper(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Soo qaad warqadda imtixaanka
    paper = ExamPaper.query.get_or_404(id)
    
    # 2. Soo qaad dhammaan su'aalaha u dhexaysa (loo habeeyay sort_order)
    questions = ExamQuestion.query.filter_by(paper_id=id).order_by(ExamQuestion.sort_order.asc()).all()
    
    # 3. Xisaabi wadarta dhibcaha (Total Marks)
    total_marks = db.session.query(db.func.sum(ExamQuestion.marks))\
                            .filter(ExamQuestion.paper_id == id).scalar() or 0
    
    # 4. Security: Hubi in macalinku isagu leeyahay warqadda ama uu yahay Admin
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if current_user.role.value == 'teacher':
        if not teacher or paper.teacher_id != teacher.id:
            flash("Ma haysatid ogolaansho aad ku aragto warqaddan.", "danger")
            return redirect(url_for('main.all_exam_papers'))

    return render_template(
        "backend/pages/components/exams/view_paper.html",
        paper=paper,
        questions=questions,
        total_marks=total_marks,
        user=current_user
    )

# ===========================
# TEACHER ADD EXAM MARKS
# ===========================








def calculate_student_performance(student_id, exam_id, active_year_id):
    student = Student.query.get(student_id)
    if not student: 
        return False

    # 1. Soo xaqiiji dhammaan dhibcaha maaddooyinka (Marks) ee ardaygu helay imtixaankan
    marks_records = StudentExamMark.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id,
        academic_year_id=active_year_id
    ).all()
    
    # Isku gee dhibcaha (Total Obtained)
    total_obtained = sum(Decimal(str(m.marks_obtained or 0)) for m in marks_records)
    
    # 2. Hel tirada maaddooyinka imtixaanka ee fasalkaas (si Average-ka loo helo)
    exam_subjects = ExamSubject.query.filter_by(
        exam_id=exam_id,
        class_id=student.class_id
    ).all()
    
    total_required_subjects = len(exam_subjects)
    subjects_entered = len(marks_records)

    if total_required_subjects == 0: 
        return False

    # 3. Xisaabi Celceliska (Average)
    average = (total_obtained / Decimal(total_required_subjects))
    
    # Grading Logic
    if average >= 90: grade, decision = "A+", "Pass"
    elif average >= 80: grade, decision = "A", "Pass"
    elif average >= 70: grade, decision = "B", "Pass"
    elif average >= 60: grade, decision = "C", "Pass"
    elif average >= 50: grade, decision = "D", "Pass"
    else: grade, decision = "F", "Fail"

    # 4. UPDATE ama INSERT Natiijada Guud (StudentExamResult)
    result = StudentExamResult.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id
    ).first()
    
    count_display = f"{subjects_entered} of {total_required_subjects}"

    if not result:
        # Create New Result Record
        result = StudentExamResult(
            school_id=student.school_id, 
            branch_id=student.branch_id,
            academic_year_id=active_year_id, 
            exam_id=exam_id, 
            student_id=student_id, 
            class_id=student.class_id,
            section_id=student.section_id, 
            total_marks=total_obtained, 
            average=average, 
            grade=grade, 
            decision=decision,
            count_subjects=count_display, 
            published='pending'
        )
        db.session.add(result)
    else:
        # Update Existing Result Record
        result.total_marks = total_obtained
        result.average = average
        result.grade = grade
        result.decision = decision
        result.count_subjects = count_display
        result.class_id = student.class_id
        result.section_id = student.section_id
        result.academic_year_id = active_year_id

    return True




@bp.route("/teacher/exam/marks", methods=["GET", "POST"])
@login_required
def teacher_add_exam_marks():
    # 1. Hubinta Status-ka User-ka
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    # 2. Hubinta Role-ka
    allowed_roles = ["school_admin", "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 3. Hel Macalinka iyo Sanad Dugsiyeedka Firfircoon
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    active_year = AcademicYear.query.filter_by(school_id=current_user.school_id, is_active=True).first()

    if not teacher or not active_year:
        msg = "Cillad: Profile macalin ama Sanad Dugsiyeed firfircoon lama helin."
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({"success": False, "error": msg}), 400
        flash(msg, "danger")
        return redirect(url_for('main.dashboard'))

    form = StudentExamMarkForm()

    if request.method == "POST":
        try:
            # Hel xogta asaasiga ah ee foomka
            ex_sub_id = request.form.get("exam_subject_id")
            exam_id = request.form.get("exam_id")
            
            if not ex_sub_id or not exam_id:
                return jsonify({"success": False, "error": "Fadlan dooro imtixaanka iyo maaddada."}), 400

            exam_subject = ExamSubject.query.get_or_404(int(ex_sub_id))
            updated_student_ids = set()

            # 4. Keydi dhibcaha maaddada iyo Update-garee Xogta Ardayga
            for key, value in request.form.items():
                if key.startswith("marks["):
                    # Soo saar Student ID
                    s_id = int(key.split("[")[1].split("]")[0])
                    
                    # Nadiifi oo u beddel dhibicda Decimal
                    val_raw = value.strip()
                    val = Decimal(val_raw) if val_raw else Decimal('0')
                    
                    # Hubi inaan dhibicdu ka badnaan karin total-ka maaddada
                    if val > exam_subject.total_marks: 
                        val = exam_subject.total_marks

                    # Update ama Insert dhibicda maaddada
                    mark_rec = StudentExamMark.query.filter_by(
                        exam_id=int(exam_id), 
                        exam_subject_id=exam_subject.id, 
                        student_id=s_id
                    ).first()

                    if mark_rec:
                        mark_rec.marks_obtained = val
                        mark_rec.academic_year_id = active_year.id
                    else:
                        db.session.add(StudentExamMark(
                            school_id=current_user.school_id, 
                            branch_id=current_user.branch_id,
                            academic_year_id=active_year.id,
                            exam_id=int(exam_id), 
                            exam_subject_id=exam_subject.id,
                            student_id=s_id, 
                            marks_obtained=val
                        ))
                    
                    # --- CUSBOONAYSIINTA SANADKA ARDAYGA (Sync Logic) ---
                    student = Student.query.get(s_id)
                    if student:
                        student.academic_year_id = active_year.id
                        student.year_name_str = active_year.year_name # String field-ka oo lagu shubay name-ka
                    
                    updated_student_ids.add(s_id)

            # Flush si dhibcaha loo arko ka hor intaan xisaabinta la samayn
            db.session.flush()

            # 5. Dib u xisaabi Performance-ka iyo Natiijada Guud ee arday kasta
            for s_id in updated_student_ids:
                calculate_student_performance(s_id, int(exam_id), active_year.id)

            db.session.flush()

            # 6. Dib u xisaabi Kaalmaha Fasalka (Class Positions)
            students = Student.query.filter(Student.id.in_(list(updated_student_ids))).all()
            affected_classes = set([s.class_id for s in students])
            
            for c_id in affected_classes:
                # Soo qaado dhammaan natiijooyinka fasalkaas ee imtixaankan
                all_res = StudentExamResult.query.filter_by(
                    exam_id=int(exam_id), 
                    class_id=c_id
                ).order_by(StudentExamResult.total_marks.desc()).all()
                
                # Sii kaalmaha (Ranking)
                for index, res in enumerate(all_res):
                    res.class_position = index + 1

            db.session.commit()
            return jsonify({"success": True, "message": "Dhibcaha, Sanadka Ardayga, iyo kaalmaha waa la kaydiyey!"})

        except Exception as e:
            db.session.rollback()
            return jsonify({"success": False, "error": str(e)}), 500

    # 7. GET Logic: Buuxinta Dropdowns
    assignments = TeacherAssignment.query.filter_by(teacher_id=teacher.id).all()
    allowed_class_ids = list(set([a.class_id for a in assignments if a.class_id]))
    
    form.class_id.choices = [(0, "-- Dooro Fasal --")] + \
                            [(c.id, c.name) for c in Class.query.filter(Class.id.in_(allowed_class_ids)).all()]
    
    form.exam_id.choices = [(0, "-- Dooro Imtixaan --")] + \
                           [(e.id, e.exam_name) for e in Exam.query.filter_by(
                                school_id=current_user.school_id, 
                                academic_year_id=active_year.id,
                                status='draft'
                           ).all()]

    return render_template(
        "backend/pages/components/exam_marks/add_exam_marks.html", 
        form=form, 
        user=current_user,
        active_year=active_year
    )


@bp.route("/get-teacher-sections/<int:class_id>")
@login_required
def get_teacher_sections(class_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if not teacher:
        return jsonify([])

    # Waxaan soo saaraynaa sections-ka macalinka loogu talagalay fasalkan
    assignments = TeacherAssignment.query.filter_by(
        teacher_id=teacher.id, 
        class_id=class_id
    ).all()
    
    sections = []
    seen_ids = set()
    
    for a in assignments:
        if a.section_id and a.section_id not in seen_ids:
            sec = Section.query.get(a.section_id)
            if sec:
                # Hubi magaca column-ka (section_name ama name)
                name = getattr(sec, 'section_name', getattr(sec, 'name', 'N/A'))
                sections.append({'id': sec.id, 'name': name})
                seen_ids.add(a.section_id)
                
    return jsonify(sections)


@bp.route("/get-exams-subjects")
@login_required
def get_exams_subjects():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get('class_id', type=int)
    exam_id = request.args.get('exam_id', type=int)
    
    teacher = Teacher.query.filter_by(user_id=current_user.id).first()
    if not teacher or not class_id or not exam_id:
        return jsonify([])

    # 1. Soo saar ID-yada maaddooyinka macalinka loo xilsaaray
    assignments = TeacherAssignment.query.filter_by(
        teacher_id=teacher.id, 
        class_id=class_id
    ).all()
    
    assigned_sub_ids = []
    for a in assignments:
        if a.subject_ids:
            try:
                import json
                ids = json.loads(a.subject_ids) if isinstance(a.subject_ids, str) else a.subject_ids
                if isinstance(ids, list):
                    assigned_sub_ids.extend([int(i) for i in ids])
                else:
                    assigned_sub_ids.append(int(ids))
            except: continue

    unique_sub_ids = list(set(assigned_sub_ids))

    if not unique_sub_ids:
        return jsonify([])

    # 2. HEL MAADDOOYINKA IMTIXAANKA (HAL KASTA MAR)
    # Waxaan isticmaalaynaa .group_by(ExamSubject.subject_id) si uusan u soo noqnoqon
    exam_subjects = ExamSubject.query.filter(
        ExamSubject.exam_id == exam_id,
        ExamSubject.subject_id.in_(unique_sub_ids)
    ).group_by(ExamSubject.subject_id).all() # <--- Halkan ayaa muhiim ah

    result = []
    for es in exam_subjects:
        if es.subject:
            result.append({
                'id': es.id, # Ama es.subject_id haddii aad u baahan tahay ID-ga maaddada asalka ah
                'name': es.subject.name
            })

    return jsonify(result)




@bp.route("/get-assigned-students")
@login_required
def get_assigned_students():
    # 1. Hubinta Status-ka
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    # 2. Hubinta Role-ka
    allowed_roles = ["school_admin", "teacher", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 3. Qaadashada Parameter-ada
    exam_subject_id = request.args.get('subject_id', type=int)
    class_id = request.args.get('class_id', type=int)
    section_id = request.args.get('section_id', type=int, default=0)
    exam_id = request.args.get('exam_id', type=int)

    if not exam_subject_id or not class_id or not exam_id:
        return jsonify([])

    ex_sub = ExamSubject.query.get_or_404(exam_subject_id)
    
    # 4. Sifee ardayda fasalka
    s_query = Student.query.filter(
        Student.class_id == class_id,
        Student.school_id == current_user.school_id,
        or_(Student.status == 'active', Student.status.is_(None))
    )
    
    if section_id and section_id != 0:
        s_query = s_query.filter(Student.section_id == section_id)
    
    students = s_query.order_by(Student.full_name.asc()).all()
    
    results = []
    for s in students:
        # Hubi haddii ardaygu uu horey u leeyahay marks
        existing = StudentExamMark.query.filter_by(
            exam_id=exam_id,
            exam_subject_id=ex_sub.id, 
            student_id=s.id
        ).first()
        
        # --- ISBEDELKA Halkan: ---
        # Haddii ardaygu uu horey u lahaa marks, markaas 'continue' ayaan dhihi
        # si looga boodo ardaygan oo aan liiska loogu darin.
        if existing and existing.marks_obtained is not None:
            continue
        # -------------------------
        
        results.append({
            'id': s.id,
            'full_name': s.full_name,
            'roll_no': getattr(s, 'roll_no', 'N/A') or "N/A",
            'total_marks': float(ex_sub.total_marks),
            'existing_mark': float(existing.marks_obtained) if existing else ""
        })
        
    return jsonify(results)


@bp.route("/exam/results/all")
@login_required
def all_exam_results():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Soo qaad sanadka uu qofku doortay (haddii uu jiro)
    selected_year = request.args.get('academic_year_id', type=int)
    
    # 2. Base Query
    query = Exam.query.filter_by(school_id=current_user.school_id)
    
    # 3. Haddii sanad la doortay, ku sifee query-ga
    if selected_year:
        query = query.filter(Exam.academic_year_id == selected_year)
    
    # 🔐 ROLE-BASED LOGIC
    if current_user.role.value == "school_admin":
        query = query.filter(Exam.branch_id == None)

    elif current_user.role.value == "branch_admin":
        query = query.filter_by(branch_id=current_user.branch_id)

    elif current_user.role.value == "teacher":
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()
        if teacher:
            assignments = TeacherAssignment.query.filter_by(teacher_id=teacher.id).all()
            class_ids = [a.class_id for a in assignments]
            query = query.filter(exists().where(
                (StudentExamResult.exam_id == Exam.id) & 
                (StudentExamResult.class_id.in_(class_ids))
            ))
        else:
            return render_template("backend/pages/components/exam_marks/view_results.html", exams=[], years=[], selected_year=None)

    # Soo saar natiijada ugu dambaysa
    exams = query.order_by(Exam.created_at.desc()).all()
    
    # Soo qaad dhamaan sanadaha dugsiga si loogu soo bandhigo Select-ka
    years = AcademicYear.query.filter_by(school_id=current_user.school_id).order_by(AcademicYear.year_name.desc()).all()

    return render_template(
        "backend/pages/components/exam_marks/view_results.html",
        exams=exams,
        years=years,
        selected_year=selected_year,
        user=current_user
    )


    
@bp.route("/exam/results/view/<int:exam_id>")
@login_required
def view_exam_results(exam_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    exam = Exam.query.get_or_404(exam_id)
    
    # 1. Base Query
    results_query = StudentExamResult.query.filter_by(
        exam_id=exam_id, 
        school_id=current_user.school_id
    )

    # 2. Role-based Filtering
    user_role = current_user.role.value if hasattr(current_user.role, 'value') else str(current_user.role)
    if user_role == "branch_admin":
        results_query = results_query.filter_by(branch_id=current_user.branch_id)
    elif user_role == "teacher":
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()
        if teacher:
            class_ids = [a.class_id for a in TeacherAssignment.query.filter_by(teacher_id=teacher.id).all() if a.class_id]
            results_query = results_query.filter(StudentExamResult.class_id.in_(class_ids))
        else:
            results_query = results_query.filter(StudentExamResult.id == 0)

    results = results_query.order_by(StudentExamResult.total_marks.desc()).all()

    # 3. Grading Logic (Sida sawirkaaga)
    def calculate_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    # 4. Habaynta Max Marks & Mapping
    formatted_results = []
    max_marks_map = {}
    active_class_ids = set([r.class_id for r in results])
    
    for c_id in active_class_ids:
        # Xisaabi wadarta dhibcaha guud ee imtixaanka ee fasalkaas
        total_exam_marks = db.session.query(db.func.sum(ExamSubject.total_marks))\
                             .filter(ExamSubject.exam_id == exam_id, ExamSubject.class_id == c_id).scalar() or 0
        max_marks_map[c_id] = float(total_exam_marks)

    # 5. Dhisidda formatted_results (Dynamic Decision lagu daray)
    for index, res in enumerate(results):
        m_marks = max_marks_map.get(res.class_id, 0)
        obtained = float(res.total_marks or 0)
        percentage = (obtained / m_marks * 100) if m_marks > 0 else 0
        
        # 🟢 DYNAMIC DECISION: Haddii uu 50 ka sareeyo waa PASS
        decision = "PASS" if percentage >= 50 else "FAIL"
        
        formatted_results.append({
            'obj': res,
            'student': res.student,
            'student_id': res.student_id,
            'class_obj': res.class_rel,
            'class_id': res.class_id,
            'total_marks': res.total_marks,
            'max_marks': m_marks,
            'percentage': round(percentage, 2),
            'average': round(percentage, 2),
            'grade': calculate_grade(percentage),
            'decision': decision,  # ✅ Halkan ayaa laga saxay FAIL-kii badnaa
            'position': index + 1
        })

    return render_template(
        "backend/pages/components/exam_marks/view_result.html",
        exam=exam,
        results=formatted_results,
        max_marks_map=max_marks_map,
        user=current_user
    )


@bp.route('/student/exam-report/<int:student_id>/<int:exam_id>')
@login_required
def student_exam_report(student_id, exam_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    student = Student.query.get_or_404(student_id)
    exam = Exam.query.get_or_404(exam_id)
    
    # 1. Hel natiijada summary-ga ah (Summary Result)
    summary_result = StudentExamResult.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id
    ).first_or_404()
    
    # 2. Hel dhibcaha maaddooyinka (Subject Marks)
    subject_marks = StudentExamMark.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id
    ).all()
    
    # 3. 🟢 SAXID: Xisaabi Grand Total (Wadarta dhibcaha imtixaanka ee fasalkaas)
    # Tani waxay soo ururinaysaa dhamaan total_marks ee maaddooyinka imtixaankan loogu talagalay
    total_possible = db.session.query(db.func.sum(ExamSubject.total_marks))\
                        .filter(
                            ExamSubject.exam_id == exam_id, 
                            ExamSubject.class_id == summary_result.class_id
                        ).scalar() or 0
    total_possible = float(total_possible)

    # 4. Grading & Decision Logic
    def calculate_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    # Xisaabi boqolleyda iyo natiijada (Decision)
    obtained_marks = float(summary_result.total_marks or 0)
    percentage = (obtained_marks / total_possible * 100) if total_possible > 0 else 0
    
    # Natiijada: Pass haddii uu 50 ka sareeyo
    current_decision = "PASS" if percentage >= 50 else "FAIL"
    current_grade = calculate_grade(percentage)

    # 5. Soo saar Macallinka Fasalka (Class Teacher)
    teacher_assign = TeacherAssignment.query.options(joinedload(TeacherAssignment.teacher))\
        .filter_by(
            class_id=summary_result.class_id,
            section_id=summary_result.section_id,
            school_id=summary_result.school_id
        ).first()

    # Isku day 2: Haddii la waayo, raadi macallinka fasalka guud ahaan (iyadoon section loo eegin)
    if not teacher_assign:
        teacher_assign = TeacherAssignment.query.options(joinedload(TeacherAssignment.teacher))\
            .filter_by(
                class_id=summary_result.class_id,
                school_id=summary_result.school_id
            ).first()

    # Helitaanka magaca
    if teacher_assign and teacher_assign.teacher:
        class_teacher = teacher_assign.teacher.full_name
    else:
        # Haddii xitaa la waayo, magaca iskuulka macallin ka mid ah ama fariin kale
        class_teacher = "Ma jiro macallin loo xilsaaray"

        
    return render_template(
        'backend/pages/components/exam_marks/student_report.html',
        student=student, 
        exam=exam, 
        result=summary_result, 
        subject_marks=subject_marks,
        total_possible=total_possible,
        percentage=round(percentage, 2),
        grade=current_grade,
        decision=current_decision,
        class_teacher=class_teacher,
        user=current_user,
        now_eat=now_eat() # Hubi in function-kan uu shaqaynayo (now_eat())
    )



@bp.route('/student/exam-report/edit/<int:student_id>/<int:exam_id>', methods=['GET', 'POST'])
@login_required
def edit_student_exam_report(student_id, exam_id):

    # =========================
    # CSRF FORM (ONLY FOR TOKEN)
    # =========================
    form = StudentExamMarkForm()

    # 1. STATUS CHECK
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    # 2. ROLE CHECK
    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 3. FETCH DATA
    student = Student.query.get_or_404(student_id)
    exam = Exam.query.get_or_404(exam_id)

    summary_result = StudentExamResult.query.filter_by(
        student_id=student_id,
        exam_id=exam_id
    ).first_or_404()

    subject_marks = db.session.query(StudentExamMark)\
        .join(ExamSubject, ExamSubject.id == StudentExamMark.exam_subject_id)\
        .filter(
            StudentExamMark.student_id == student_id,
            StudentExamMark.exam_id == exam_id
        ).all()

    total_possible = db.session.query(db.func.sum(ExamSubject.total_marks))\
        .filter(
            ExamSubject.exam_id == exam_id,
            ExamSubject.class_id == summary_result.class_id
        ).scalar() or 0

    total_possible = float(total_possible)

    # =========================
    # GRADE FUNCTION
    # =========================
    def calculate_grade(p):
        if p >= 95: return 'A+'
        elif p >= 90: return 'A'
        elif p >= 85: return 'A-'
        elif p >= 80: return 'B+'
        elif p >= 75: return 'B'
        elif p >= 70: return 'B-'
        elif p >= 65: return 'C+'
        elif p >= 60: return 'C'
        elif p >= 50: return 'C-'
        elif p >= 40: return 'D'
        elif p >= 20: return 'E'
        else: return 'F'

    # =========================
    # 🟡 POST (SAFE FIX ONLY)
    # =========================
    if request.method == 'POST':
        try:
            total_marks = 0
            errors = []

            for mark in subject_marks:
                field_name = f"subject_{mark.exam_subject_id}"
                raw_value = request.form.get(field_name)

                # =========================
                # 🔥 SAFE FLOAT FIX (87.00 PROBLEM SOLVED)
                # =========================
                raw_value = str(raw_value).strip().replace(",", ".")

                if raw_value == "":
                    errors.append(f"{mark.exam_subject.subject.name}: Qiime waa banaan")
                    continue

                try:
                    new_score = float(raw_value)
                except ValueError:
                    errors.append(f"{mark.exam_subject.subject.name}: Qiime sax ah geli")
                    continue

                max_marks = float(mark.exam_subject.total_marks)

                if new_score < 0:
                    errors.append(f"{mark.exam_subject.subject.name}: Kama yaraan karo 0")

                if new_score > max_marks:
                    errors.append(f"{mark.exam_subject.subject.name}: Kama badnaan karo {max_marks}")

                mark.marks_obtained = new_score
                total_marks += new_score

            if errors:
                for e in errors:
                    flash(e, "danger")
                return redirect(request.url)

            if total_marks > total_possible:
                flash("Wadarta guud kama badnaan karto total-ka imtixaanka.", "danger")
                return redirect(request.url)

            # =========================
            # UPDATE RESULT TABLE
            # =========================
            percentage = (total_marks / total_possible * 100) if total_possible > 0 else 0

            summary_result.total_marks = total_marks
            summary_result.percentage = percentage
            summary_result.grade = calculate_grade(percentage)
            summary_result.decision = "PASS" if percentage >= 50 else "FAIL"
            summary_result.updated_at = now_eat()

            db.session.commit()

            flash("Natiijada si guul leh ayaa loo cusbooneysiiyay.", "success")

            return redirect(url_for(
                'main.student_exam_report',
                student_id=student_id,
                exam_id=exam_id
            ))

        except Exception as e:
            db.session.rollback()
            flash(f"Khalad ayaa dhacay: {str(e)}", "danger")

    # =========================
    # GET
    # =========================
    obtained_marks = float(summary_result.total_marks or 0)
    percentage = (obtained_marks / total_possible * 100) if total_possible > 0 else 0

    return render_template(
        'backend/pages/components/exam_marks/edit_exam_result.html',
        form=form,
        student=student,
        exam=exam,
        result=summary_result,
        subject_marks=subject_marks,
        total_possible=total_possible,
        percentage=round(percentage, 2),
        grade=calculate_grade(percentage),
        decision="PASS" if percentage >= 50 else "FAIL",
        user=current_user
    )


@bp.route("/exam/results/cumulative")
@login_required
def cumulative_results():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Hel dhamaan sanad dugsiyeedyada
   
    all_years = AcademicYear.query.filter_by(school_id=current_user.school_id).order_by(AcademicYear.year_name.desc()).all()
    
    # 2. Hel sanadka la doortay ama kan hadda firfircoon
    selected_year_id = request.args.get('academic_year_id', type=int)
    if selected_year_id:
        active_year = AcademicYear.query.get_or_404(selected_year_id)
    else:
        active_year = AcademicYear.query.filter_by(school_id=current_user.school_id, is_active=True).first()

    if not active_year:
        flash("Ma jiro Sannad Dugsiyeed firfircoon.", "warning")
        return redirect(url_for('main.all_exam_results'))

    # 3. Query-ga Natiijooyinka (Kaliya kuwa la daabacay)
    query = StudentExamResult.query.join(Exam, StudentExamResult.exam_id == Exam.id).filter(
        StudentExamResult.school_id == current_user.school_id,
        Exam.academic_year_id == active_year.id,
        Exam.status == 'published'
    )

    # 4. Role-Based Filtering (Admin/Branch/Teacher)
    user_role = current_user.role.value if hasattr(current_user.role, 'value') else str(current_user.role)
    if user_role == "branch_admin":
        query = query.filter(StudentExamResult.branch_id == current_user.branch_id)
    elif user_role == "teacher":
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()
        if teacher:
            class_ids = [a.class_id for a in TeacherAssignment.query.filter_by(teacher_id=teacher.id).all() if a.class_id]
            query = query.filter(StudentExamResult.class_id.in_(class_ids)) if class_ids else query.filter(StudentExamResult.id == 0)

    results = query.all()

    # 5. Habaynta Xogta & Nidaamka Grading-ka (Sida ku cad sawirka)
    student_data = {}
    exam_max_cache = {}

    def get_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    for res in results:
        s_id = res.student_id
        e_id = res.exam_id
        
        # Hel wadarta dhibcaha ugu sareeya ee imtixaankan si dynamic ah
        if e_id not in exam_max_cache:
            total_max = db.session.query(db.func.sum(ExamSubject.total_marks))\
                          .filter(ExamSubject.exam_id == e_id).scalar() or 0
            exam_max_cache[e_id] = float(total_max)

        if s_id not in student_data:
            student_data[s_id] = {
                'student': res.student,
                'class': res.class_rel,
                'grand_total': 0.0,
                'total_out_of': 0.0,
                'exam_count': 0,
                'average': 0.0,
                'grade': 'F'
            }
        
        # Xisaabi dhibcaha la helay (Obtained Marks)
        obtained = float(res.total_marks) if res.total_marks else 0.0
        student_data[s_id]['grand_total'] += obtained
        student_data[s_id]['total_out_of'] += exam_max_cache[e_id]
        student_data[s_id]['exam_count'] += 1

    # 6. Xisaabi Average-ka iyo Grade-ka kama dambaysta ah
    for s_id in student_data:
        total = student_data[s_id]['grand_total']
        out_of = student_data[s_id]['total_out_of']
        
        if out_of > 0:
            avg = (total / out_of) * 100
            student_data[s_id]['average'] = round(avg, 2)
            student_data[s_id]['grade'] = get_grade(avg)
        else:
            student_data[s_id]['average'] = 0.0
            student_data[s_id]['grade'] = 'F'

    # 7. Kala saarista (Ranking) - Marka hore Average-ka ugu sareeya
    sorted_items = sorted(
        student_data.items(), 
        key=lambda item: item[1]['average'], 
        reverse=True
    )
    
    # Dib ugu celi Dict la habeeyay
    sorted_students = {k: v for k, v in sorted_items}
    # Tusaale ahaan, haddii aad u baahan tahay imtixaanka ugu dambeeyay
    active_exam = Exam.query.filter_by(status='published').first()
    active_id = active_exam.id if active_exam else 0

    available_exams = Exam.query.filter_by(
        school_id=current_user.school_id,
        academic_year_id=active_year.id,
        status='published'
    ).all()

   # 4. Role-Based Filtering - Hel fasallada
    user_role = current_user.role.value if hasattr(current_user.role, 'value') else str(current_user.role)
    classes_query = Class.query.filter_by(school_id=current_user.school_id)

    if user_role == "branch_admin":
        classes_query = classes_query.filter_by(branch_id=current_user.branch_id)
    elif user_role == "teacher":
        teacher = Teacher.query.filter_by(user_id=current_user.id).first()
        if teacher:
            assignments = TeacherAssignment.query.filter_by(teacher_id=teacher.id).all()
            class_ids = [a.class_id for a in assignments if a.class_id]
            classes_query = classes_query.filter(Class.id.in_(class_ids))
        else:
            classes_query = classes_query.filter(Class.id == 0)

    classes = classes_query.all()
    
    return render_template(
        "backend/pages/components/exam_marks/cumulative.html",
        student_data=sorted_students,
        active_year=active_year,
        all_years=all_years,
        user=current_user,
        active_exam_id=active_id,
         available_exams=available_exams ,
         classes=classes
    )



#-----------------------------------------------------
#--------------- School Time Table
#-----------------------------------------------------
@bp.route("/time-slots/add", methods=["GET", "POST"])
@login_required
def add_timeslot():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = TimeSlotForm()
    
    # 1. Hubi haddii user-ku uu yahay active
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # 2. Ogolaanshaha (Only school_admin or branch_admin)
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    if form.validate_on_submit():
        school_id = current_user.school_id
        # Haddii uu yahay school_admin, branch_id waa None, haddii kale waa branch-ka uu maamulo
        branch_id = current_user.branch_id if current_user.role.value == 'branch_admin' else None

        # 3. Prevent Duplicates (Hubi haddii slot isku waqti ah uu jiro)
        existing = TimeSlot.query.filter_by(
            school_id=school_id,
            branch_id=branch_id,
            start_time=form.start_time.data,
            shift=form.shift.data
        ).first()

        if existing:
            flash(f"❌ Slot leh waqtigan ({form.start_time.data}) mar hore ayaa loo sameeyay shift-ka {form.shift.data}.", "warning")
        else:
            try:
                new_slot = TimeSlot(
                    school_id=school_id,
                    branch_id=branch_id,
                    label=form.label.data,
                    start_time=form.start_time.data,
                    end_time=form.end_time.data,
                    is_break=form.is_break.data,
                    shift=form.shift.data
                )
                db.session.add(new_slot)
                db.session.commit()
                flash("✅ Time Slot created successfully!", "success")
                return redirect(url_for('main.all_timeslots')) # Bedel haddii route-kaagu magac kale leeyahay
            
            except Exception as e:
                db.session.rollback()
                print("ERROR saving TimeSlot:", e)
                flash("❌ Error saving time slot.", "danger")

    return render_template(
        "backend/pages/components/timetable/add_timeslot.html", # Hubi path-ka template-kaaga
        form=form,
        user=current_user
    )


@bp.route("/time-slots/all")
@login_required
def all_timeslots():
    # 1. Sharciga Account Status (Active/Inactive)
    if current_user.status == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    # 2. Xaqiijinta Role-ka (Authorization)
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    # ================= SHARCIGA SCHOOL ADMIN =================
    if current_user.role.value == 'school_admin':
        # Kaliya wuxuu arkayaa slots-ka dugsiga weyn (branch_id is None)
        # Sida aad ku samaysay ClassSubject-ka
        timeslots = TimeSlot.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).order_by(TimeSlot.shift, TimeSlot.start_time).all()

    # ================= SHARCIGA BRANCH ADMIN =================
    elif current_user.role.value == 'branch_admin':
        # Kaliya wuxuu arkayaa slots-ka u gaarka ah laantiisa
        timeslots = TimeSlot.query.filter_by(
            branch_id=current_user.branch_id
        ).order_by(TimeSlot.shift, TimeSlot.start_time).all()

    return render_template(
        "backend/pages/components/timetable/all_timeslots.html",
        timeslots=timeslots,
        user=current_user
    )

@bp.route("/time-slots/edit/<int:id>", methods=["GET", "POST"])
@login_required
def edit_timeslot(id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Hubi Account Status
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    # 2. Soo saar slot-ka, haddii kale 404
    slot = TimeSlot.query.get_or_404(id)

    # 3. Sharciga Authorization & Ownership
    # School Admin wuxuu bedeli karaa kuwa branch_id-ga leh iyo kuwa aan lahayn ee school-kiisa ah
    # Branch Admin wuxuu bedeli karaa KALIYA kuwa laantiisa
    if current_user.role.value == 'branch_admin':
        if slot.branch_id != current_user.branch_id:
            flash("Unauthorized: Ma bedeli kartid slot aan laantaada ka tirsanayn.", "danger")
            return redirect(url_for('main.all_timeslot_list'))
    elif current_user.role.value == 'school_admin':
        if slot.school_id != current_user.school_id:
            flash("Unauthorized access.", "danger")
            return redirect(url_for('main.all_timeslot_list'))
    else:
        flash("Unauthorized", "danger")
        return redirect(url_for('main.dashboard'))

    form = TimeSlotForm(obj=slot) # Ku shub xogta jirta foomka

    if form.validate_on_submit():
        # 4. Hubi Duplicates (Waqti isku mid ah oo aan ahayn kan hadda la bedelayo)
        existing = TimeSlot.query.filter(
            TimeSlot.id != id, # Ka saar kan hadda la edit-gareynayo
            TimeSlot.school_id == current_user.school_id,
            TimeSlot.branch_id == slot.branch_id,
            TimeSlot.start_time == form.start_time.data,
            TimeSlot.shift == form.shift.data
        ).first()

        if existing:
            flash(f"❌ Error: Waqtigan ({form.start_time.data}) hore ayaa loogu diwaangeliyey {form.shift.data}.", "warning")
        else:
            try:
                # Update xogta
                slot.label = form.label.data
                slot.start_time = form.start_time.data
                slot.end_time = form.end_time.data
                slot.shift = form.shift.data
                slot.is_break = form.is_break.data
                
                # Haddii school_admin wax bedelayo, wuxuu dooran karaa branch
                if current_user.role.value == 'school_admin':
                    slot.branch_id = form.branch_id.data or None

                db.session.commit()
                flash("✅ Time Slot updated successfully!", "success")
                return redirect(url_for('main.all_timeslots'))
            
            except Exception as e:
                db.session.rollback()
                print("ERROR updating TimeSlot:", e)
                flash("❌ Error saving changes.", "danger")

    return render_template(
        "backend/pages/components/timetable/edit_timeslot.html", # Waxaad isticmaali kartaa isla template-ka Add-ka
        form=form,
        user=current_user,
        edit_mode=True, # Si aad ugu muujiso "Update" halkii ay ka ahaan lahayd "Save"
        slot=slot
    )


@bp.route("/time-slots/delete/<int:id>", methods=["POST"])
@login_required
def delete_timeslot(id):
    # 1. Hubi Account Status
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    slot = TimeSlot.query.get_or_404(id)

    # 2. ROLE-BASED OWNERSHIP CHECK
    # Branch Admin: Kaliya laantiisa
    if current_user.role.value == 'branch_admin':
        if slot.branch_id != current_user.branch_id:
            flash("❌ Ma tirtiri kartid waqti aysan laantaadu lahayn.", "danger")
            return redirect(url_for('main.all_timeslots'))
    
    # School Admin: Kaliya dugsigiisa (dhammaan branches-ka hoos yimaada)
    elif current_user.role.value == 'school_admin':
        if slot.school_id != current_user.school_id:
            flash("❌ Ma haysatid ogolaansho aad ku tirtirto xogta iskuul kale.", "danger")
            return redirect(url_for('main.all_timeslots'))
    
    else:
        flash("❌ Unauthorized access.", "danger")
        return redirect(url_for('main.dashboard'))

    # 3. FULL DELETE LOGIC
    try:
        # Marka hore tirtir xiisadaha (Timetable) ku xiran waqtigan
        Timetable.query.filter_by(time_slot_id=id).delete()
        
        # Marka labaad tirtir waqtiga (TimeSlot) laftiisa
        db.session.delete(slot)
        db.session.commit()
        
        flash("✅ Waqtiga iyo dhammaan xiisadihii ku xirnaa waa la tirtiray!", "success")
    except Exception as e:
        db.session.rollback()
        print(f"Error: {e}")
        flash("❌ Cilad baa dhacday intii tirtirista lagu jiray.", "danger")

    return redirect(url_for('main.all_timeslots'))

#---------------------------------
#--------  Time Table Sections
#---------------------------------
# 1. Soo saar maaddooyinka fasal leeyahay
@bp.route("/get-subjects/<int:class_id>")
@login_required
def get_subjects_by_class(class_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # Waxaan ka soo dhex raadinaynaa ClassSubject model-kaaga
    class_subjects = ClassSubject.query.filter_by(class_id=class_id).all()
    
    # Waxaan u dhisaynaa liis (List)
    subjects = []
    for cs in class_subjects:
        subjects.append({
            'id': cs.subject.id,
            'name': cs.subject.name
        })
    return jsonify(subjects)

# 2. Soo saar macallimiinta maaddo bixiya (Optional - haddii ay macallimiintu maaddo ku xidhan yihiin)
@bp.route("/get-teachers-by-subject/<int:subject_id>")
@login_required
def get_teachers_by_subject(subject_id):
    try:
        if getattr(current_user, 'status', 1) == 0:
            flash("Account-kaagu ma shaqeynayo.", "danger")
            return redirect(url_for("main.dashboard"))

        allowed_roles = ["school_admin", "branch_admin"]
        if current_user.role.value not in allowed_roles:
            flash("Ma haysatid ogolaansho.", "danger")
            return redirect(url_for("main.dashboard"))

        # 1. Waxaan raadinaynaa dhammaan assignments-ka
        # Waxaan dhex baaraynaa subject_ids oo ah JSON column
        assignments = TeacherAssignment.query.all()
        
        teacher_list = []
        seen_teacher_ids = set()

        for assign in assignments:
            # Hubi haddii subject_id uu ku jiro subject_ids-ka macallinka
            # subject_ids waa JSON (list), markaa si toos ah ayaan 'in' ugu isticmaali karnaa
            if subject_id in assign.subject_ids:
                if assign.teacher_id not in seen_teacher_ids:
                    # Hubi in macallinka uu jiro magaciisana soo rido
                    if assign.teacher:
                        teacher_list.append({
                            'id': assign.teacher.id,
                            'name': assign.teacher.full_name # Ama assign.teacher.name
                        })
                        seen_teacher_ids.add(assign.teacher_id)

        return jsonify(teacher_list)

    except Exception as e:
        print(f"Cillad baa dhacday: {str(e)}")
        return jsonify([]), 500
    

@bp.route("/timetable/add", methods=["GET", "POST"])
@login_required
def add_timetable():
    # 1. Hubi haddii user-ku uu active yahay
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    form = TimetableForm()

    if form.validate_on_submit():
        school_id = current_user.school_id
        branch_id = current_user.branch_id if current_user.role.value == 'branch_admin' else None
        
        # Section ID handling
        # Waa muhiim inaan hubino haddii section-ka uu yahay 0 ama None
        section_id = form.section_id.data if form.section_id.data and form.section_id.data != 0 else None

        # ---------------------------------------------------------
        # 3. KAHORTAGGA ISKU-DHACA (CONFLICT CHECK)
        # ---------------------------------------------------------

        # A. Class Conflict: Fasalka waqtigaas maaddo kale ma u qorsheysan tahay?
        # Kani wuxuu hubinayaa Day + Time + Class + Section
        class_busy = Timetable.query.filter_by(
            day_of_week=form.day_of_week.data,
            time_slot_id=form.time_slot_id.data,
            class_id=form.class_id.data,
            section_id=section_id,
            school_id=school_id  # Hubi inuu isku iskuul yahay
        ).first()

        if class_busy:
            # Magaca maaddada mashquulka ku ah fasalka
            msg = f"❌ Fasalka {class_busy.klass.name}"
            if class_busy.section:
                msg += f" (Section: {class_busy.section.name})"
            msg += f" waqtigan horey ayaa loogu qoondeeyay Maaddada: {class_busy.subject.name}."
            
            flash(msg, "warning")
            return render_template("backend/pages/components/timetable/add_timetable.html", form=form, user=current_user)

        # B. Teacher Conflict: Macallinka waqtigaas ma mashquul yahay?
        teacher_busy = Timetable.query.filter_by(
            day_of_week=form.day_of_week.data,
            time_slot_id=form.time_slot_id.data,
            teacher_id=form.teacher_id.data,
            school_id=school_id
        ).first()

        if teacher_busy:
            flash(f"❌ Macallinka waqtigan waa mashquul. Wuxuu xiisad kale ugu jiraa Fasalka {teacher_busy.klass.name}.", "warning")
            return render_template("backend/pages/components/timetable/add_timetable.html", form=form, user=current_user)

        # 4. Kaydinta
        try:
            new_entry = Timetable(
                school_id=school_id,
                branch_id=branch_id,
                class_id=form.class_id.data,
                section_id=section_id,
                subject_id=form.subject_id.data,
                teacher_id=form.teacher_id.data,
                time_slot_id=form.time_slot_id.data,
                day_of_week=form.day_of_week.data
            )
            db.session.add(new_entry)
            db.session.commit()
            flash("✅ Jadwalka waa la kaydiyey si guul leh!", "success")
            return redirect(url_for('main.all_timetables')) 

        except Exception as e:
            db.session.rollback()
            print("ERROR saving Timetable:", e)
            flash("❌ Cilad farsamo ayaa dhacday.", "danger")

    return render_template(
        "backend/pages/components/timetable/add_timetable.html",
        form=form,
        user=current_user
    )


@bp.route("/timetables/all")
@login_required
def all_timetables():
    # 1. Sharciga Account Status (Active/Inactive)
    if getattr(current_user, 'status', 1) == 0:
        flash("Your account is inactive.", "danger")
        return redirect(url_for('main.index'))

    # 2. Xaqiijinta Role-ka (Authorization)
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Unauthorized access.", "danger")
        return redirect(url_for('main.dashboard'))

    # ================= SHARCIGA SCHOOL ADMIN =================
    if current_user.role.value == 'school_admin':
        # Wuxuu arkayaa jadwalka dugsiga weyn (branch_id is None)
        timetables = Timetable.query.filter_by(
            school_id=current_user.school_id,
            branch_id=None
        ).join(TimeSlot).order_by(
            Timetable.day_of_week, 
            TimeSlot.start_time
        ).all()

    # ================= SHARCIGA BRANCH ADMIN =================
    elif current_user.role.value == 'branch_admin':
        # Wuxuu arkayaa jadwalka u gaarka ah laantiisa oo kaliya
        timetables = Timetable.query.filter_by(
            branch_id=current_user.branch_id
        ).join(TimeSlot).order_by(
            Timetable.day_of_week, 
            TimeSlot.start_time
        ).all()

    return render_template(
        "backend/pages/components/timetable/all_timetables.html",
        timetables=timetables,
        user=current_user
    )




@bp.route("/timetable/edit/<int:timetable_id>", methods=["GET", "POST"])
@login_required
def edit_timetable(timetable_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Soo saar xogta jirta
    entry = Timetable.query.get_or_404(timetable_id)

    # 2. Ogolaanshaha (Authorization)
    if current_user.role.value not in ['school_admin', 'branch_admin']:
        flash("Ma lihid ogolaansho aad wax ku bedesho.", "danger")
        return redirect(url_for('main.all_timetables'))

    # 3. Ku shub xogta form-ka (Pre-populate)
    form = TimetableForm(obj=entry)

    if form.validate_on_submit():
        school_id = current_user.school_id
        section_id = form.section_id.data if form.section_id.data and form.section_id.data != 0 else None

        # ---------------------------------------------------------
        # 4. KAHORTAGGA ISKU-DHACA (CONFLICT CHECK)
        # ---------------------------------------------------------
        # FIIRO GAAR AH: Waxaan ku darnay '.id != timetable_id' si uusan isaga isku celin

        # A. Class Conflict
        class_busy = Timetable.query.filter(
            Timetable.id != timetable_id, # Iska indha-tir kan aan hadda dhiseyno
            Timetable.day_of_week == form.day_of_week.data,
            Timetable.time_slot_id == form.time_slot_id.data,
            Timetable.class_id == form.class_id.data,
            Timetable.section_id == section_id,
            Timetable.school_id == school_id
        ).first()

        if class_busy:
            flash(f"❌ Isku-dhac: Fasalkan waqtigan waxaa ugu jirta Maaddada {class_busy.subject.name}.", "warning")
            return render_template("backend/pages/components/timetable/edit_timetable.html", form=form, user=current_user, entry=entry)

        # B. Teacher Conflict
        teacher_busy = Timetable.query.filter(
            Timetable.id != timetable_id,
            Timetable.day_of_week == form.day_of_week.data,
            Timetable.time_slot_id == form.time_slot_id.data,
            Timetable.teacher_id == form.teacher_id.data,
            Timetable.school_id == school_id
        ).first()

        if teacher_busy:
            flash(f"❌ Isku-dhac: Macallinka waqtigan wuxuu fasal kale ugu jiraa {teacher_busy.klass.name}.", "warning")
            return render_template("backend/pages/components/timetable/edit_timetable.html", form=form, user=current_user, entry=entry)

        # 5. Cusboonaysiinta (Update)
        try:
            entry.day_of_week = form.day_of_week.data
            entry.time_slot_id = form.time_slot_id.data
            entry.class_id = form.class_id.data
            entry.section_id = section_id
            entry.subject_id = form.subject_id.data
            entry.teacher_id = form.teacher_id.data
            
            db.session.commit()
            flash("✅ Jadwalka si guul leh ayaa loo cusboonaysiiyay!", "success")
            return redirect(url_for('main.all_timetables'))

        except Exception as e:
            db.session.rollback()
            print("Update Error:", e)
            flash("❌ Cilad baa dhacday intii la cusboonaysiinayay.", "danger")

    return render_template(
        "backend/pages/components/timetable/edit_timetable.html",
        form=form,
        user=current_user,
        entry=entry
    )

@bp.route("/timetable/delete/<int:timetable_id>", methods=["POST"])
@login_required
def delete_timetable(timetable_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    entry = Timetable.query.get_or_404(timetable_id)

    # ROLE-BASED OWNERSHIP CHECK
    if current_user.role.value == 'branch_admin':
        if entry.branch_id != current_user.branch_id:
            flash("❌ Ma tirtiri kartid jadwal aysan laantaadu lahayn.", "danger")
            return redirect(url_for('main.all_timetables'))
            
    elif current_user.role.value == 'school_admin':
        if entry.school_id != current_user.school_id:
            flash("❌ Unauthorized access.", "danger")
            return redirect(url_for('main.all_timetables'))

    try:
        db.session.delete(entry)
        db.session.commit()
        flash("✅ Xiisadda waa la tirtiray!", "success")
    except Exception as e:
        db.session.rollback()
        flash("❌ Cilad farsamo ayaa dhacday.", "danger")

    return redirect(url_for('main.all_timetables'))








#----------------------- Students  UI
@bp.route("/student/results/view")
@login_required
def student_cumulative_view():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "student", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Hubi xogta ardayga login-ka ah
    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        flash("Qaybtan waxaa loogu talagalay ardayda kaliya.", "danger")
        return redirect(url_for('main.dashboard'))

    student_current_class = Class.query.get(student.class_id)

    # 2. Hel Sannad Dugsiyeedka hadda socda (Active Year)
    active_year = AcademicYear.query.filter_by(
        school_id=current_user.school_id, 
        is_active=True
    ).first()

    # 3. Dynamic Auto-Select (Fasalkii ugu dambeeyay ee uu natiijo ku leeyahay)
    latest_result = StudentExamResult.query.filter_by(student_id=student.id)\
        .order_by(StudentExamResult.id.desc()).first()
    
    default_id = latest_result.class_id if latest_result else student.class_id
    selected_class_id = request.args.get('class_id', type=int) or default_id
    current_selected_class = Class.query.get_or_404(selected_class_id)

    # 4. --- HEL SUBJECT MARKS (Fasalka la doortay) ---
    marks_query = db.session.query(StudentExamMark, ExamSubject, Subject, Exam).join(
        ExamSubject, StudentExamMark.exam_subject_id == ExamSubject.id
    ).join(
        Subject, ExamSubject.subject_id == Subject.id
    ).join(
        Exam, StudentExamMark.exam_id == Exam.id
    ).join(
        StudentExamResult, (StudentExamResult.exam_id == Exam.id) & (StudentExamResult.student_id == student.id)
    ).filter(
        StudentExamMark.student_id == student.id,
        StudentExamResult.class_id == selected_class_id,
        Exam.status.in_(['published'])
    ).all()

    # Nidaamka Grading-ka ee sawirka ku salaysan
    def get_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    subjects_results = {}
    for mark_record, ex_sub, sub, ex in marks_query:
        s_name = sub.name
        if s_name not in subjects_results:
            subjects_results[s_name] = {'scores': {}, 'row_total': 0.0, 'row_max': 0.0}
        
        attained = float(mark_record.marks_obtained or 0)
        max_of_exam = float(ex_sub.total_marks or 100)
        
        subjects_results[s_name]['scores'][ex.exam_name] = {
            'score': attained,
            'max_score': max_of_exam
        }
        subjects_results[s_name]['row_total'] += attained
        subjects_results[s_name]['row_max'] += max_of_exam

    # 5. --- XISAABINTA GRAND TOTAL ---
    exam_results = StudentExamResult.query.filter_by(
        student_id=student.id,
        class_id=selected_class_id
    ).join(Exam).filter(
        Exam.status.in_(['published'])
    ).all()

    grand_total_obtained = float(sum(res.total_marks or 0 for res in exam_results))
    grand_total_max = sum(data['row_max'] for data in subjects_results.values())

    # 6. --- PERCENT, GRADE & STATUS (DYNAMIC) ---
    percent_gained = (grand_total_obtained / grand_total_max * 100) if grand_total_max > 0 else 0
    
    # Grade-ka wuxuu raacayaa boqolleyda guud (Cumulative Percentage)
    final_grade = get_grade(percent_gained)
    
    # Status-ka wuxuu ku xiran yahay haddii uu ka sareeyo 50% (ama C- ka sareeyo)
    dynamic_status = "Pass" if percent_gained >= 50 else "Fail"
    status_color = "success" if dynamic_status == "Pass" else "danger"

    # Average-ka saxda ah (Total / Maaddooyinka)
    unique_subjects_count = len(subjects_results)
    final_average = (grand_total_obtained / unique_subjects_count) if unique_subjects_count > 0 else 0

    # 7. --- RANKING (Class Position) ---
    class_rank_query = db.session.query(
        StudentExamResult.student_id, 
        db.func.sum(StudentExamResult.total_marks).label('total')
    ).join(Exam).filter(
        StudentExamResult.class_id == selected_class_id,
        Exam.status.in_(['published'])
    ).group_by(StudentExamResult.student_id).order_by(db.desc('total')).all()
    
    class_pos = next((i + 1 for i, r in enumerate(class_rank_query) if r.student_id == student.id), 0)
    total_class_students = len(class_rank_query)

    # 8. --- RANKING (School Position) ---
    school_rank_query = db.session.query(
        StudentExamResult.student_id, 
        db.func.sum(StudentExamResult.total_marks).label('total')
    ).join(Exam).filter(
        StudentExamResult.school_id == current_user.school_id,
        Exam.academic_year_id == (active_year.id if active_year else None),
        Exam.status.in_(['published'])
    ).group_by(StudentExamResult.student_id).order_by(db.desc('total')).all()

    school_pos = next((i + 1 for i, r in enumerate(school_rank_query) if r.student_id == student.id), 0)
    total_school_students = len(school_rank_query)

    # 9. --- EXAM HISTORY ---
    history_classes = db.session.query(Class).join(StudentExamResult, Class.id == StudentExamResult.class_id)\
        .filter(StudentExamResult.student_id == student.id).distinct().order_by(Class.id.desc()).all()

    return render_template(
        "backend/pages/components/exam_marks/student_exam_results.html",
        subjects=subjects_results,
        results=exam_results,
        grand_total=grand_total_obtained,
        grand_max=grand_total_max,
        percent=round(percent_gained, 2),
        average=round(final_average, 2),
        grade=final_grade,
        status=dynamic_status,
        status_color=status_color,
        class_pos=class_pos,
        total_class_students=total_class_students,
        school_pos=school_pos,
        total_school_students=total_school_students,
        current_selected_class=current_selected_class,
        student_current_class=student_current_class,
        active_year=active_year,
        history_classes=history_classes,
        user=current_user,
        student=student
    )

    
@bp.route("/students/promotion", methods=["GET", "POST"])
@login_required
def student_promotion():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get('class_id', type=int)
    academic_year_id = request.args.get('academic_year_id', type=int)
    form = StudentPromotionForm()

    school_id = current_user.school_id
    branch_id_filter = current_user.branch_id 

    if request.method == 'GET':
        if class_id: form.current_class_id.data = class_id
        if academic_year_id: form.current_academic_year_id.data = academic_year_id

    students_with_results = []
    
    if class_id and academic_year_id:
        valid_year = AcademicYear.query.filter_by(id=academic_year_id, school_id=school_id).first()
        if not valid_year:
            flash("Sanad dugsiyeedka la doortay ma jiro.", "danger")
            return redirect(url_for('main.student_promotion'))

        student_query = Student.query.filter_by(
            class_id=class_id, 
            academic_year_id=academic_year_id,
            school_id=school_id
        )
        
        if branch_id_filter is not None:
            student_query = student_query.filter_by(branch_id=branch_id_filter)
            
        all_students = student_query.all()

        for student in all_students:
            total_obtained = db.session.query(func.sum(StudentExamMark.marks_obtained)).filter(
                StudentExamMark.student_id == student.id,
                StudentExamMark.academic_year_id == academic_year_id
            ).scalar() or 0

            total_target_marks = db.session.query(func.sum(ExamSubject.total_marks)).filter(
                ExamSubject.class_id == class_id,
                ExamSubject.academic_year_id == academic_year_id
            ).scalar() or 0

            pass_threshold = float(total_target_marks) * 0.5
            status_decision = 'Pass' if float(total_obtained) >= pass_threshold and total_target_marks > 0 else 'Fail'

            student.yearly_total = float(total_obtained)
            student.total_required = float(total_target_marks)
            student.status_decision = status_decision
            student.yearly_avg = round((float(total_obtained) / float(total_target_marks) * 100), 2) if total_target_marks > 0 else 0
            
            students_with_results.append(student)

    if form.validate_on_submit():
        import json
        from datetime import datetime
        try:
            student_ids_raw = form.student_ids.data
            if not student_ids_raw or student_ids_raw in ['[]', '']:
                flash("Fadlan dooro ugu yaraan hal arday", "warning")
                return redirect(request.url)

            student_ids = json.loads(student_ids_raw)
            to_class_id = form.to_class_id.data
            target_class = Class.query.get(to_class_id) if to_class_id and to_class_id != 0 else None
            new_year_id = form.to_academic_year_id.data
            new_year = AcademicYear.query.filter_by(id=new_year_id, school_id=school_id).first()

            if not new_year:
                flash("Fadlan dooro sanad dugsiyeedka socda/cusub.", "danger")
                return redirect(request.url)

            processed_count = 0
            for s_id in student_ids:
                student = Student.query.filter_by(id=s_id, school_id=school_id).first()
                if not student: continue
                
                if branch_id_filter is not None and student.branch_id != branch_id_filter:
                    continue

                marks = db.session.query(func.sum(StudentExamMark.marks_obtained)).filter(
                    StudentExamMark.student_id == student.id,
                    StudentExamMark.academic_year_id == academic_year_id
                ).scalar() or 0
                
                targets = db.session.query(func.sum(ExamSubject.total_marks)).filter(
                    ExamSubject.class_id == class_id,
                    ExamSubject.academic_year_id == academic_year_id
                ).scalar() or 0

                is_passed = float(marks) >= (float(targets) * 0.5) if targets > 0 else True
                final_promotion_type = 'repeat' if not is_passed else form.promotion_type.data 
                
                # ✅ FIXING DATA FOR MYSQL CONSTRAINTS
                p_to_class_id = None
                p_to_section_id = None
                p_promotion_type = final_promotion_type

                if is_passed and not target_class:
                    # GRADUATED LOGIC
                    student.status = 'graduated'
                    student.class_id = None
                    student.section_id = None
                    p_to_class_id = class_id # Fasalkii ugu dambeeyay
                    p_to_section_id = None   # MySQL waxay ogoshahay NULL halkan (SET NULL)
                    p_promotion_type = 'graduated'
                else:
                    if final_promotion_type == 'repeat':
                        student.status = 'active'
                        p_to_class_id = class_id
                        p_to_section_id = student.section_id
                    elif target_class:
                        student.class_id = target_class.id
                        student.level_id = target_class.level_id
                        student.status = 'active'
                        p_to_class_id = target_class.id
                        
                        # Section-ka cusub haddii la doortay, haddii kale NULL
                        if form.to_section_id.data and form.to_section_id.data != 0:
                            student.section_id = form.to_section_id.data
                            p_to_section_id = form.to_section_id.data
                        else:
                            p_to_section_id = None 

                        if target_class.level:
                            student.price = target_class.level.price
                    
                    student.academic_year = new_year.year_name
                    student.academic_year_id = new_year.id 

                student.updated_at = datetime.now()

                # RECORD PROMOTION HISTORY
                promotion = StudentPromotion(
                    school_id=school_id,
                    branch_id=student.branch_id,
                    student_id=student.id,
                    from_class_id=class_id, 
                    from_section_id=student.section_id if student.section_id else None,
                    from_academic_year_id=academic_year_id,
                    to_class_id=p_to_class_id if p_to_class_id else class_id, 
                    to_section_id=p_to_section_id, # ✅ FIXED: None (NULL) halkii ay ka ahaan lahayd 0
                    to_academic_year_id=new_year.id,
                    promotion_type=p_promotion_type,
                    remarks=f"Obtained: {marks}/{targets}. Result: {'Pass' if is_passed else 'Fail'}",
                    promoted_by=current_user.id
                )
                db.session.add(promotion)

                # INVOICING
                if student.status != 'graduated':
                    existing_fee = StudentFeeCollection.query.filter_by(
                        student_id=student.id, 
                        class_id=student.class_id, 
                        school_id=school_id
                    ).first()

                    new_tuition_price = float(student.price or 0)

                    if not existing_fee:
                        new_fee = StudentFeeCollection(
                            student_id=student.id, class_id=student.class_id,
                            school_id=school_id, branch_id=student.branch_id,
                            amount_due=new_tuition_price, amount_paid=0.0
                        )
                        db.session.add(new_fee)
                        db.session.flush()
                        new_fee.recalculate()
                    else:
                        existing_fee.amount_due = new_tuition_price
                        existing_fee.recalculate()

                processed_count += 1

            db.session.commit()
            flash(f"✅ Si guul ah ayaa loo farsameeyay {processed_count} arday!", "success")
            return redirect(url_for('main.student_promotion', class_id=class_id, academic_year_id=academic_year_id))

        except Exception as e:
            db.session.rollback()
            flash(f"❌ Khalad Database: Hubi in dhamaan xogta (Section/Class) ay sax tahay.", "danger")
            print(f"DEBUG ERROR: {str(e)}") # Tan ka eeg console-ka si aad u aragto error-ka rasmiga ah

    return render_template(
        "backend/pages/components/students/student_promotion.html",
        form=form,
        students=students_with_results,
        current_class_id=class_id,
        current_year_id=academic_year_id,
        user=current_user
    )



#----------------------------------------------------------------
#---------------- Student UI Print for Admins -------------------
#----------------------------------------------------------------
@bp.route('/student/cumulative-report/print/<int:student_id>')
@login_required
def print_cumulative_report(student_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    student = Student.query.get_or_404(student_id)
    school = student.school
    branch = student.branch if student.branch_id else None

    # 1. Helitaanka Sannadka
    selected_year_id = request.args.get('academic_year_id', type=int)
    active_year = AcademicYear.query.get(selected_year_id) if selected_year_id else AcademicYear.query.filter_by(school_id=current_user.school_id, is_active=True).first()

    if not active_year:
        flash("Sannad dugsiyeed firfircoon lama helin.", "warning")
        return redirect(url_for('exam.cumulative_results'))

    # 2. Hel Natiijooyinka Ardayga (Published Exams Only)
    results = StudentExamResult.query.join(Exam).filter(
        StudentExamResult.student_id == student_id,
        Exam.academic_year_id == active_year.id,
        Exam.status == 'published'
    ).order_by(Exam.id.asc()).all()

    if not results:
        flash("Wax natiijo ah lama helin.", "warning")
        return redirect(url_for('exam.cumulative_results'))

    # Sharciga Darajooyinka (New Grade Logic)
    def get_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    exam_reports = []
    subject_matrix = {} 
    exam_names = []
    exam_scores = [] # Trend scores (Line chart)
    grand_total_obtained = 0.0
    grand_total_possible = 0.0

    for res in results:
        exam_names.append(res.exam.exam_name)
        marks = StudentExamMark.query.filter_by(student_id=student_id, exam_id=res.exam_id).all()
        
        e_total, e_max = 0.0, 0.0
        current_subjects = []

        for m in marks:
            max_m = float(m.exam_subject.total_marks if m.exam_subject else 0)
            obtained = float(m.marks_obtained or 0)
            s_name = m.exam_subject.subject.name if m.exam_subject else "Unknown"

            if s_name not in subject_matrix:
                subject_matrix[s_name] = {'marks': {}, 'total': 0, 'max': 0}
            
            subject_matrix[s_name]['marks'][res.exam.exam_name] = obtained
            subject_matrix[s_name]['total'] += obtained
            subject_matrix[s_name]['max'] += max_m

            # Calculate grade for individual subject
            s_percentage = (obtained / max_m * 100) if max_m > 0 else 0
            current_subjects.append({
                'subject': s_name, 
                'marks': obtained, 
                'max': max_m,
                'grade': get_grade(s_percentage)
            })
            e_total += obtained
            e_max += max_m

        e_avg = round((e_total/e_max*100), 2) if e_max > 0 else 0
        exam_scores.append(e_avg) # Key for Trend Analytics

        exam_reports.append({
            'exam_name': res.exam.exam_name,
            'subjects': current_subjects,
            'total': e_total,
            'max': e_max,
            'avg': e_avg,
            'grade': get_grade(e_avg)
        })
        grand_total_obtained += e_total
        grand_total_possible += e_max

    # 3. Dynamic Analytics (Subject vs Class Average)
    subject_labels = list(subject_matrix.keys())
    student_subject_avgs = []
    class_subject_avgs = []

    for s_name in subject_labels:
        s_total = subject_matrix[s_name]['total']
        s_max = subject_matrix[s_name]['max']
        student_subject_avgs.append(round((s_total/s_max*100), 2) if s_max > 0 else 0)

        # Helitaanka celceliska dhamaan fasalka ee maadadan
        avg_query = db.session.query(db.func.avg((StudentExamMark.marks_obtained / ExamSubject.total_marks) * 100))\
            .join(ExamSubject, StudentExamMark.exam_subject_id == ExamSubject.id)\
            .join(Subject, ExamSubject.subject_id == Subject.id)\
            .filter(
                Subject.name == s_name, 
                ExamSubject.academic_year_id == active_year.id, 
                ExamSubject.class_id == student.class_id
            ).scalar()
        
        class_subject_avgs.append(round(float(avg_query or 0), 2))

    final_percentage = round((grand_total_obtained / grand_total_possible * 100), 2) if grand_total_possible > 0 else 0

    return render_template(
        "backend/pages/components/exam_marks/print_cumulative.html",
        student=student, school=school, active_year=active_year, branch=branch,
        exam_reports=exam_reports,
        exam_names=exam_names,
        exam_scores=exam_scores, # Trend Data
        subject_matrix=subject_matrix,
        total_obtained=grand_total_obtained,
        total_possible=grand_total_possible,
        final_percentage=final_percentage,
        final_grade=get_grade(final_percentage),
        subject_labels=subject_labels,
        student_subject_avgs=student_subject_avgs,
        class_subject_avgs=class_subject_avgs,
        now_eat=now_eat()
    )


@bp.route('/student/exam-report/print/<int:student_id>/<int:exam_id>')
@login_required
def print_student_exam_report(student_id, exam_id):
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Soo saar xogta ardayga iyo imtixaanka
    student = Student.query.get_or_404(student_id)
    exam = Exam.query.get_or_404(exam_id)
    
    # --- CUSBOONAYSIIN: Hel xogta dugsiga iyo laanta ---
    # Waxaan ka soo qaadaynaa ardayga dugsiga uu ka tirsan yahay
    school = student.school 
    branch = student.branch if student.branch_id else None
    # --------------------------------------------------

    # 2. Hel summary-ga natiijada
    summary_result = StudentExamResult.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id
    ).first_or_404()
    
    # 3. Hel dhibcaha maaddooyinka
    subject_marks = StudentExamMark.query.filter_by(
        student_id=student_id, 
        exam_id=exam_id
    ).all()
    
    # 4. Xisaabi wadarta guud ee suurtagalka ah
    total_possible = db.session.query(db.func.sum(ExamSubject.total_marks))\
                        .filter(
                            ExamSubject.exam_id == exam_id, 
                            ExamSubject.class_id == summary_result.class_id
                        ).scalar() or 0
    total_possible = float(total_possible)

    # 5. Xisaabi Boqolleyda
    obtained_marks = float(summary_result.total_marks or 0)
    percentage = (obtained_marks / total_possible * 100) if total_possible > 0 else 0

    # 6. Hel Macallinka Fasalka
    teacher_assign = TeacherAssignment.query.filter_by(
        class_id=summary_result.class_id,
        section_id=summary_result.section_id
    ).first()
    
    class_teacher = teacher_assign.teacher.full_name if teacher_assign and teacher_assign.teacher else "________________"

    # 7. Render Template (Hubi inaad raaciso school iyo branch)
    return render_template(
        "backend/pages/components/exam_marks/print_report.html",
        student=student,
        exam=exam,
        result=summary_result,
        subject_marks=subject_marks,
        total_possible=total_possible,
        percentage=round(percentage, 2),
        class_teacher=class_teacher,
        school=school,      # <--- TANI WAA MUHIIM
        branch=branch,      # <--- TANI WAA MUHIIM
        now_eat=now_eat()
    )



@bp.route('/student/class-merit-report/print')
@login_required
def print_class_merit_report():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Helitaanka ID-yada laga soo diray URL-ka
    class_id = request.args.get('class_id', type=int)
    selected_year_id = request.args.get('academic_year_id', type=int)

    if not class_id or not selected_year_id:
        flash("Fadlan dooro fasal iyo sanad dugsiyeed sax ah.", "danger")
        return redirect(url_for('exam.cumulative_results'))

    # 2. Xogta Aasaasiga ah
    active_year = AcademicYear.query.get_or_404(selected_year_id)
    selected_class = Class.query.get_or_404(class_id)
    
    # Qeexitaanka xaaladaha imtixaanka la oggol yahay
    allowed_statuses = ['published', 'closed']

    # 3. Hel Maadooyinka (Subject -> ExamSubject -> Exam)
    # Halkan waxaan ku xiraynaa School iyo Branch si xogtu u noqoto mid sax ah
    class_subjects = db.session.query(Subject).join(ExamSubject).join(Exam).filter(
        ExamSubject.class_id == class_id,
        ExamSubject.academic_year_id == active_year.id,
        Exam.school_id == current_user.school_id,
        Exam.status.in_(allowed_statuses)
    ).group_by(Subject.id).all()

    # 4. Helitaanka dhamaan dhibcaha (Marks) ee fasalkaas iyo sanadkaas
    # MUHIIM: Halkan koodhku ma eegayo 'Student.class_id', wuxuu eegayaa dhibcaha u diwaangashan fasalkan
    all_marks = StudentExamMark.query.join(ExamSubject).join(Exam).filter(
        ExamSubject.class_id == class_id,
        StudentExamMark.academic_year_id == active_year.id,
        Exam.school_id == current_user.school_id,
        Exam.status.in_(allowed_statuses)
    ).all()

    # Hel dhamaan ardayda dhibcaha leh (Unique Students from Marks)
    student_ids = list(set([m.student_id for m in all_marks]))
    report_students = Student.query.filter(Student.id.in_(student_ids)).all()

    # Map dhibcaha: {(student_id, subject_id): sum_of_marks}
    marks_map = {}
    for m in all_marks:
        if m.exam_subject:
            key = (m.student_id, m.exam_subject.subject_id)
            marks_map[key] = marks_map.get(key, 0) + float(m.marks_obtained or 0)

    # Xisaabi wadarta dhibcaha maado kasta laga rabay (Max Marks)
    subject_max_marks = {}
    for sub in class_subjects:
        total = db.session.query(db.func.sum(ExamSubject.total_marks)).join(Exam).filter(
            ExamSubject.class_id == class_id, 
            ExamSubject.subject_id == sub.id,
            ExamSubject.academic_year_id == active_year.id,
            Exam.status.in_(allowed_statuses)
        ).scalar() or 0
        subject_max_marks[sub.id] = float(total)

    # 5. Sharciga Darajooyinka (New Grade Logic)
    def get_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    # 6. Diyaarinta Xogta
    class_report_data = []
    for student in report_students:
        student_total_obtained = 0.0
        student_total_possible = 0.0
        student_subjects_detail = []

        for sub in class_subjects:
            obtained = marks_map.get((student.id, sub.id), 0.0)
            possible = subject_max_marks.get(sub.id, 0.0)
            percentage = (obtained / possible * 100) if possible > 0 else 0
            
            student_subjects_detail.append({
                'subject_name': sub.name,
                'obtained': round(obtained, 1),
                'possible': round(possible, 1),
                'avg': round(percentage, 1),
                'grade': get_grade(percentage)
            })
            student_total_obtained += obtained
            student_total_possible += possible

        if student_total_possible > 0:
            final_avg = round((student_total_obtained / student_total_possible * 100), 2)
            class_report_data.append({
                'student': student,
                'subjects': student_subjects_detail,
                'grand_total': round(student_total_obtained, 1),
                'total_out_of': round(student_total_possible, 1),
                'final_avg': final_avg,
                'final_grade': get_grade(final_avg)
            })

    # Ranking
    class_report_data = sorted(class_report_data, key=lambda x: x['final_avg'], reverse=True)

    return render_template(
        "backend/pages/components/exam_marks/print_class_merit.html",
        class_report_data=class_report_data,
        active_year=active_year,
        selected_class=selected_class,
        school=current_user.school,
        branch=current_user.branch,
        now_eat=now_eat()
    )


@bp.route('/student/class-summarise-merit-report/print')
@login_required
def print_class_summarise_merit_report():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["school_admin", "branch_admin"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    class_id = request.args.get('class_id', type=int)
    selected_year_id = request.args.get('academic_year_id', type=int)

    if not class_id or not selected_year_id:
        flash("Fadlan dooro fasal iyo sanad dugsiyeed sax ah.", "danger")
        return redirect(url_for('exam.cumulative_results'))

    active_year = AcademicYear.query.get_or_404(selected_year_id)
    selected_class = Class.query.get_or_404(class_id)
    allowed_statuses = ['published', 'closed']

    # 1. Hel Maadooyinka leh imtixaan Published ama Closed ah
    class_subjects = db.session.query(Subject).join(ExamSubject).join(Exam).filter(
        ExamSubject.class_id == class_id,
        ExamSubject.academic_year_id == active_year.id,
        Exam.school_id == current_user.school_id,
        Exam.status.in_(allowed_statuses)
    ).group_by(Subject.id).all()

    # 2. Hel dhamaan dhibcaha loo qoray fasalkan (xataa haddii ardaygu class kale u wareegay)
    all_marks = StudentExamMark.query.join(ExamSubject).join(Exam).filter(
        ExamSubject.class_id == class_id,
        StudentExamMark.academic_year_id == active_year.id,
        Exam.school_id == current_user.school_id,
        Exam.status.in_(allowed_statuses)
    ).all()

    # Soo saar list-ka ardayda dhibcaha leh
    student_ids = list(set([m.student_id for m in all_marks]))
    report_students = Student.query.filter(Student.id.in_(student_ids)).all()

    # Map dhibcaha
    marks_map = {}
    for m in all_marks:
        if m.exam_subject:
            key = (m.student_id, m.exam_subject.subject_id)
            marks_map[key] = marks_map.get(key, 0) + float(m.marks_obtained or 0)

    # Max Marks per Subject
    subject_max_marks = {}
    for sub in class_subjects:
        total = db.session.query(db.func.sum(ExamSubject.total_marks)).join(Exam).filter(
            ExamSubject.class_id == class_id, 
            ExamSubject.subject_id == sub.id,
            ExamSubject.academic_year_id == active_year.id,
            Exam.status.in_(allowed_statuses)
        ).scalar() or 0
        subject_max_marks[sub.id] = float(total)

    # Sharciga Darajooyinka (New Grade Scale)
    def get_grade(percentage):
        if percentage >= 95: return 'A+'
        elif percentage >= 90: return 'A'
        elif percentage >= 85: return 'A-'
        elif percentage >= 80: return 'B+'
        elif percentage >= 75: return 'B'
        elif percentage >= 70: return 'B-'
        elif percentage >= 65: return 'C+'
        elif percentage >= 60: return 'C'
        elif percentage >= 50: return 'C-'
        elif percentage >= 40: return 'D'
        elif percentage >= 20: return 'E'
        else: return 'F'

    # 3. Diyaarinta liiska hal-bog ah
    class_report_data = []
    for student in report_students:
        student_total_obtained = 0.0
        student_total_possible = 0.0
        
        for sub in class_subjects:
            obtained = marks_map.get((student.id, sub.id), 0.0)
            possible = subject_max_marks.get(sub.id, 0.0)
            student_total_obtained += obtained
            student_total_possible += possible

        if student_total_possible > 0:
            final_avg = round((student_total_obtained / student_total_possible * 100), 2)
            class_report_data.append({
                'student': student,
                'grand_total': round(student_total_obtained, 1),
                'total_out_of': round(student_total_possible, 1),
                'final_avg': final_avg,
                'final_grade': get_grade(final_avg)
            })

    # Ranking: U kala horree dhibcaha ugu sarreeya
    class_report_data = sorted(class_report_data, key=lambda x: x['final_avg'], reverse=True)

    return render_template(
        "backend/pages/components/exam_marks/print_all_class_merit.html",
        class_report_data=class_report_data,
        active_year=active_year,
        selected_class=selected_class,
        school=current_user.school,
        branch=current_user.branch,
        now_eat=now_eat()
    )


#-------------------------------------
#----------- Student Attendance View |
#-------------------------------------

@bp.route("/student/attendance/view")
@login_required
def student_attendance_view():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["student"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        flash("Qaybtan waxaa loo oggol yahay ardayda oo kaliya.", "danger")
        return redirect(url_for('main.dashboard'))

    active_year = AcademicYear.query.filter_by(
        school_id=current_user.school_id, 
        is_active=True
    ).first()

    selected_subject_id = request.args.get('subject_id', type=int)

    # SAXID: Halkii aan ka isticmaali lahayn student.current_class oo error bixiyay
    # Waxaan si toos ah u soo xulanaa maadooyinka uu ardaygu leeyahay diiwaanka attendance-ka
    # Tani waxay keenaysaa kaliya maadooyinka uu cashar ka qaatay.
    subjects = db.session.query(Subject).join(StudentAttendance).filter(
        StudentAttendance.student_id == student.id
    ).distinct().all()

    # Query-ga Stats-ka
    query = db.session.query(
        db.func.count(StudentAttendance.id).label('total'),
        db.func.sum(db.case((StudentAttendance.status == 'present', 1), else_=0)).label('present'),
        db.func.sum(db.case((StudentAttendance.status == 'absent', 1), else_=0)).label('absent'),
        db.func.sum(db.case((StudentAttendance.status == 'late', 1), else_=0)).label('late'),
        db.func.sum(db.case((StudentAttendance.status == 'excused', 1), else_=0)).label('excused')
    ).filter(
        StudentAttendance.student_id == student.id,
        StudentAttendance.school_id == current_user.school_id
    )

    # Query-ga Logs-ka
    logs_query = StudentAttendance.query.filter_by(student_id=student.id)

    if selected_subject_id:
        query = query.filter(StudentAttendance.subject_id == selected_subject_id)
        logs_query = logs_query.filter(StudentAttendance.subject_id == selected_subject_id)

    attendance_stats = query.first()
    logs = logs_query.order_by(StudentAttendance.date.desc()).limit(50).all()

    total_days = attendance_stats.total or 0
    present_count = (attendance_stats.present or 0) + (attendance_stats.late or 0)
    attendance_percentage = (present_count / total_days * 100) if total_days > 0 else 0

    return render_template(
        "backend/pages/components/students/view_attendance.html",
        stats=attendance_stats,
        percentage=round(attendance_percentage, 1),
        logs=logs,
        subjects=subjects,
        selected_subject_id=selected_subject_id,
        student=student,
        active_year=active_year,
        user=current_user
    )

#--------------------------------------
#-------- Student Financial Statement |
#--------------------------------------
@bp.route("/student/finance/statement")
@login_required
def student_financial_statement():
    if getattr(current_user, 'status', 1) == 0:
        flash("Account-kaagu ma shaqeynayo.", "danger")
        return redirect(url_for("main.dashboard"))

    allowed_roles = ["student"]
    if current_user.role.value not in allowed_roles:
        flash("Ma haysatid ogolaansho.", "danger")
        return redirect(url_for("main.dashboard"))

    # 1. Hubi inuu yahay arday (User relationship to Student)
    student = Student.query.filter_by(user_id=current_user.id).first()
    if not student:
        flash("Arday kaliya ayaa arki kara boggan.", "danger")
        return redirect(url_for('main.dashboard'))

    # 2. Qabashada Filtering-ka URL-ka ka imaanaya
    date_filter = request.args.get('date_filter', 'all')
    from_date_str = request.args.get('from_date')
    to_date_str = request.args.get('to_date')

    # Query-ga aasaasiga ah (Waxaan isku xiraynaa Invoice iyo Collection)
    invoice_query = FeeInvoice.query.join(StudentFeeCollection).filter(
        StudentFeeCollection.student_id == student.id
    )

    start_date = None
    end_date = None

    # Apply Custom Filter
    if date_filter == 'custom' and from_date_str and to_date_str:
        try:
            start_date = datetime.strptime(from_date_str, '%Y-%m-%d').date()
            end_date = datetime.strptime(to_date_str, '%Y-%m-%d').date()
            invoice_query = invoice_query.filter(FeeInvoice.date_issued.between(start_date, end_date))
        except ValueError:
            flash("Fadlan geli taariikh sax ah.", "warning")

    # Invoices-ka u habee si kor u kac ah (ASC) si running balance-ka loogu xisaabiyo
    invoices_list = invoice_query.order_by(FeeInvoice.date_issued.asc()).all()

    # 3. Previous Balance (Wadarta deynta ka horreysa taariikhda la doortay)
    previous_balance = 0.0
    if start_date:
        past_collections = StudentFeeCollection.query.join(FeeInvoice).filter(
            StudentFeeCollection.student_id == student.id,
            FeeInvoice.date_issued < start_date
        ).all()
        previous_balance = sum(fc.remaining_balance for fc in past_collections) or 0.0

    # 4. Running Balance Calculation
    running_balance = previous_balance
    for inv in invoices_list:
        running_balance += (inv.amount_due or 0) - (inv.amount_paid or 0)
        inv.running_balance = running_balance # Ku darista doorsoome kumeel gaar ah

    # 5. Dashboard Summary (Wadarta Guud)
    all_collections = StudentFeeCollection.query.filter_by(student_id=student.id).all()
    total_invoiced = sum(fc.amount_due for fc in all_collections) or 0.0
    total_paid = sum(fc.amount_paid for fc in all_collections) or 0.0
    current_balance = sum(fc.remaining_balance for fc in all_collections) or 0.0

    return render_template(
        "backend/pages/components/students/financial_statement.html",
        invoices=reversed(invoices_list), # Latest First for UI
        total_invoiced=total_invoiced,
        total_paid=total_paid,
        current_balance=current_balance,
        previous_balance=previous_balance,
        date_filter=date_filter,
        from_date=from_date_str,
        to_date=to_date_str,
        user=current_user
    )


















#===============================
#----------- Somali Localtion
#===========================

# -------------------------------
# All Somalia Locations
# -------------------------------
@bp.route("/all/somalia-locations")
@login_required
def all_somalia_locations():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact the admin.", "danger")
        return redirect(url_for('main.index'))

    if getattr(current_user.role, 'value', current_user.role) in ['superadmin']:
        locations = SomaliaLocation.query.order_by(SomaliaLocation.created_at.desc()).all()
    else:
        flash("You are not authorized to view this page.", "danger")
        return redirect(url_for('main.dashboard'))

    return render_template(
        "backend/pages/components/somali_locations/all-somali-locations.html",
        locations=locations,
        user=current_user
    )


# -------------------------------
# Export Locations
# -------------------------------
@bp.route('/somalia-locations/export/<string:file_type>')
@login_required
def export_somalia_locations(file_type):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    locations = SomaliaLocation.query.all()
    
    if file_type == 'csv':
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['ID', 'Region', 'District', 'Created At', 'Updated At'])
        for loc in locations:
            writer.writerow([loc.id, loc.region, loc.district, loc.created_at, loc.updated_at])
        output.seek(0)
        return send_file(io.BytesIO(output.getvalue().encode()), 
                         mimetype='text/csv',
                         as_attachment=True,
                         download_name='somalia_locations.csv')
    
    elif file_type == 'excel':
        df = pd.DataFrame([{
            'ID': loc.id,
            'Region': loc.region,
            'District': loc.district,
            'Created At': loc.created_at,
            'Updated At': loc.updated_at
        } for loc in locations])
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Locations')
        output.seek(0)
        return send_file(output, 
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                         as_attachment=True,
                         download_name='somalia_locations.xlsx')
    else:
        flash("Invalid export format.", "danger")
        return redirect(url_for('main.all_somalia_locations'))


# -------------------------------
# Import Locations
# -------------------------------
@bp.route('/somalia-locations/import', methods=['POST'])
@login_required
def import_somalia_locations():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if 'file' not in request.files:
        flash("No file selected!", "danger")
        return redirect(url_for('main.all_somalia_locations'))
    
    file = request.files['file']
    filename = secure_filename(file.filename)
    
    try:
        if filename.endswith('.csv'):
            df = pd.read_csv(file)
        elif filename.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(file)
        else:
            flash("Unsupported file format!", "danger")
            return redirect(url_for('main.all_somalia_locations'))
        
        added_count = 0
        for _, row in df.iterrows():
            existing = SomaliaLocation.query.filter_by(
                region=row['Region'], district=row['District']
            ).first()
            if not existing:
                loc = SomaliaLocation(region=row['Region'], district=row['District'])
                db.session.add(loc)
                added_count += 1
        db.session.commit()
        flash(f"Imported {added_count} new locations successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error importing file: {str(e)}", "danger")
    
    return redirect(url_for('main.all_somalia_locations'))


# -------------------------------
# Add Somalia Location
# -------------------------------
@bp.route("/add/somalia-location", methods=["GET", "POST"])
@login_required
def add_somalia_location():
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to add locations.", "danger")
        return redirect(url_for('main.all_somalia_locations'))

    form = SomaliaLocationForm()

    if form.validate_on_submit():
        # Prevent duplicates
        existing = SomaliaLocation.query.filter_by(
            region=form.region.data.strip(),
            district=form.district.data.strip()
        ).first()
        if existing:
            flash("This location already exists!", "warning")
            return redirect(url_for('main.add_somalia_location'))

        location = SomaliaLocation(
            region=form.region.data.strip(),
            district=form.district.data.strip()
        )
        db.session.add(location)
        db.session.commit()

        flash(f"Location '{form.region.data} - {form.district.data}' added successfully!", "success")
        return redirect(url_for('main.all_somalia_locations'))

    return render_template(
        "backend/pages/components/somali_locations/add-somali-location.html",
        form=form,
        user=current_user
    )



@bp.route("/edit/somalia-location/<int:id>", methods=["GET", "POST"])
@login_required
def edit_somalia_location(id):
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to edit locations.", "danger")
        return redirect(url_for('main.all_somalia_locations'))

    location = SomaliaLocation.query.get_or_404(id)
    form = SomaliaLocationForm(obj=location)

    if form.validate_on_submit():
        # Prevent duplicates
        existing = SomaliaLocation.query.filter(
            SomaliaLocation.region == form.region.data.strip(),
            SomaliaLocation.district == form.district.data.strip(),
            SomaliaLocation.id != id
        ).first()
        if existing:
            flash("Another location with same Region and District exists!", "warning")
            return redirect(url_for('main.edit_somalia_location', id=id))

        location.region = form.region.data.strip()
        location.district = form.district.data.strip()
        db.session.commit()

        flash(f"Location '{form.region.data} - {form.district.data}' updated successfully!", "success")
        return redirect(url_for('main.all_somalia_locations'))

    return render_template(
        "backend/pages/components/somali_locations/edit-somali-location.html",
        form=form,
        user=current_user
    )



@bp.route("/delete/somalia-location/<int:id>", methods=["POST"])
@login_required
def delete_somalia_location(id):

    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to delete locations.", "danger")
        return redirect(url_for('main.all_somalia_locations'))

    try:
        location = SomaliaLocation.query.get_or_404(id)
        db.session.delete(location)
        db.session.commit()
        flash(f"Location '{location.region} - {location.district}' deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting location: {str(e)}", "danger")

    return redirect(url_for('main.all_somalia_locations'))


# -------------------------------
# Delete All Locations
# -------------------------------
@bp.route("/deleteall/somalia-locations", methods=["POST"])
@login_required
def deleteall_somalia_locations():
    # Check if user is inactive
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only admin/manager allowed
    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to delete all locations.", "danger")
        return redirect(url_for('main.all_somalia_locations'))

    try:
        num_deleted = SomaliaLocation.query.delete()
        db.session.commit()
        flash(f"All Somalia Locations ({num_deleted}) have been deleted successfully!", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"Error deleting all locations: {str(e)}", "danger")

    return redirect(url_for('main.all_somalia_locations'))


@bp.route("/all/school-site-settings")
@login_required
def all_school_site_settings():
    # Active user check
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    # Only superadmin can view
    if getattr(current_user.role, 'value', current_user.role) not in ['superadmin']:
        flash("You are not authorized to view this page.", "danger")
        return redirect(url_for('main.dashboard'))

    settings_list = SchoolSiteSettings.query.order_by(SchoolSiteSettings.created_at.desc()).all()

    return render_template(
        "backend/pages/components/school_site_Settings/all_school_site_Settings.html",
        settings_list=settings_list,
        user=current_user
    )



@bp.route("/add/school-site-settings", methods=["GET", "POST"])
@login_required
def add_school_site_settings():
    # Check user permissions
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to add site settings.", "danger")
        return redirect(url_for('main.dashboard'))

    form = SchoolSiteSettingsForm()
    form.set_choices()  # populate schools and branches

    if form.validate_on_submit():
        branch_id = form.branch_id.data if form.branch_id.data != 0 else None

        # Create new SchoolSiteSettings record
        settings = SchoolSiteSettings(
            school_id=form.school_id.data,
            branch_id=branch_id,
            site_title=form.site_title.data.strip() if form.site_title.data else None,
            short_desc=form.short_desc.data.strip() if form.short_desc.data else None,
            long_desc=form.long_desc.data.strip() if form.long_desc.data else None,
            address=form.address.data.strip() if form.address.data else None,
            phone=form.phone.data.strip() if form.phone.data else None,
            email=form.email.data.strip() if form.email.data else None,
            facebook=form.facebook.data.strip() if form.facebook.data else None,
            twitter=form.twitter.data.strip() if form.twitter.data else None,
            instagram=form.instagram.data.strip() if form.instagram.data else None
        )

        # Save to get an ID for folder structure
        db.session.add(settings)
        db.session.commit()

        # Ensure site_title is not None for slug
        if not settings.site_title:
            flash("Site title is required to generate slug.", "danger")
            db.session.delete(settings)
            db.session.commit()
            return redirect(request.url)

        slug = settings.site_title.lower().replace(' ', '-').replace('_', '-')
        user_subfolder = os.path.join('static/backend/uploads/schools', str(settings.id))
        os.makedirs(user_subfolder, exist_ok=True)

        # Handle main_logo
        main_image = request.files.get('main_logo')
        if not main_image or not main_image.filename:
            flash("Main logo is required.", "danger")
            return redirect(request.url)

        ext = os.path.splitext(main_image.filename)[1]
        unique_id = uuid.uuid4().hex[:8]
        main_filename = f"{slug}-{unique_id}{ext}"
        main_path = os.path.join(user_subfolder, main_filename)
        main_image.save(main_path)
        settings.main_logo = os.path.relpath(main_path, 'static').replace("\\", "/")

        # Handle sub_logo
        sub_image = request.files.get('sub_logo')
        if not sub_image or not sub_image.filename:
            flash("Sub logo is required.", "danger")
            return redirect(request.url)

        ext = os.path.splitext(sub_image.filename)[1]
        unique_id = uuid.uuid4().hex[:8]
        sub_filename = f"{slug}-{unique_id}{ext}"
        sub_path = os.path.join(user_subfolder, sub_filename)
        sub_image.save(sub_path)
        settings.sub_logo = os.path.relpath(sub_path, 'static').replace("\\", "/")

        # Handle sign_logo
        sign_image = request.files.get('sign_logo')
        if not sign_image or not sign_image.filename:
            flash("Sub logo is required.", "danger")
            return redirect(request.url)

        ext = os.path.splitext(sign_image.filename)[1]
        unique_id = uuid.uuid4().hex[:8]
        sub_filename = f"{slug}-{unique_id}{ext}"
        sign_path = os.path.join(user_subfolder, sub_filename)
        sign_image.save(sign_path)
        settings.sign_logo = os.path.relpath(sign_path, 'static').replace("\\", "/")

        # Commit logo paths
        db.session.commit()

        flash("School Site Settings added successfully!", "success")
        return redirect(url_for('main.dashboard'))

    return render_template(
        "backend/pages/components/school_site_Settings/add_school_site_Setting.html",
        form=form,
        user=current_user
    )


@bp.route("/edit/school-site-settings/<int:id>", methods=["GET", "POST"])
@login_required
def edit_school_site_settings(id):
    # ------------------------------
    # Permission check
    # ------------------------------
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if current_user.role.value not in ['superadmin']:
        flash("You are not authorized to edit site settings.", "danger")
        return redirect(url_for('main.dashboard'))

    # ------------------------------
    # Get existing record
    # ------------------------------
    settings = SchoolSiteSettings.query.get_or_404(id)
    form = SchoolSiteSettingsForm(obj=settings)
    form.set_choices()  # populate school and branch dropdowns

    if form.validate_on_submit():
        # ------------------------------
        # Update form fields
        # ------------------------------
        settings.branch_id = form.branch_id.data if form.branch_id.data != 0 else None
        settings.site_title = form.site_title.data.strip() if form.site_title.data else None
        settings.short_desc = form.short_desc.data.strip() if form.short_desc.data else None
        settings.long_desc = form.long_desc.data.strip() if form.long_desc.data else None
        settings.address = form.address.data.strip() if form.address.data else None
        settings.phone = form.phone.data.strip() if form.phone.data else None
        settings.email = form.email.data.strip() if form.email.data else None
        settings.facebook = form.facebook.data.strip() if form.facebook.data else None
        settings.twitter = form.twitter.data.strip() if form.twitter.data else None
        settings.instagram = form.instagram.data.strip() if form.instagram.data else None

        # ------------------------------
        # Folder path for the school (do not create yet)
        # ------------------------------

           # Handle image
        main_image = request.files.get('main_logo')
        if main_image and main_image.filename:
            ext = os.path.splitext(main_image.filename)[1]
            slug = settings.site_title.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            safe_filename = f"{slug}-{unique_id}{ext}"

            user_subfolder = os.path.join('static/backend/uploads/schools', str(settings.id))
            os.makedirs(user_subfolder, exist_ok=True)
            image_path = os.path.join(user_subfolder, safe_filename)
            main_image.save(image_path)

            # Delete old image if exists
            if settings.main_logo:
                old_image_path = os.path.join('static', settings.main_logo)
                if os.path.exists(old_image_path):
                    os.remove(old_image_path)

            settings.main_logo = os.path.relpath(image_path, 'static').replace("\\", "/")
      
           # Handle image
        sub_image = request.files.get('sub_logo')
        if sub_image and sub_image.filename:
            ext = os.path.splitext(sub_image.filename)[1]
            slug = settings.site_title.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            safe_filename = f"{slug}-{unique_id}{ext}"

            user_subfolder = os.path.join('static/backend/uploads/schools', str(settings.id))
            os.makedirs(user_subfolder, exist_ok=True)
            image_path = os.path.join(user_subfolder, safe_filename)
            sub_image.save(image_path)

            # Delete old image if exists
            if settings.sub_logo:
                old_image_path = os.path.join('static', settings.sub_logo)
                if os.path.exists(old_image_path):
                    os.remove(old_image_path)

            settings.sub_logo = os.path.relpath(image_path, 'static').replace("\\", "/")
      
           # Handle image
        sign_image = request.files.get('sign_logo')
        if sign_image and sign_image.filename:
            ext = os.path.splitext(sign_image.filename)[1]
            slug = settings.site_title.lower().replace(' ', '-').replace('_', '-')
            unique_id = uuid.uuid4().hex[:8]
            safe_filename = f"{slug}-{unique_id}{ext}"

            user_subfolder = os.path.join('static/backend/uploads/schools', str(settings.id))
            os.makedirs(user_subfolder, exist_ok=True)
            image_path = os.path.join(user_subfolder, safe_filename)
            sign_image.save(image_path)

            # Delete old image if exists
            if settings.sign_logo:
                old_image_path = os.path.join('static', settings.sign_logo)
                if os.path.exists(old_image_path):
                    os.remove(old_image_path)

            settings.sign_logo = os.path.relpath(image_path, 'static').replace("\\", "/")
      

        # ------------------------------
        # Commit changes
        # ------------------------------
        db.session.commit()
        flash("School Site Settings updated successfully!", "success")
        return redirect(url_for('main.all_school_site_settings'))

    # ------------------------------
    # Render edit template
    # ------------------------------
    return render_template(
        "backend/pages/components/school_site_Settings/edit_school_site_Setting.html",
        form=form,
        settings=settings,
        user=current_user
    )


@bp.route("/delete/school-site-settings/<int:id>", methods=["POST"])
@login_required
def delete_school_site_settings(id):
    # ------------------------------
    # Active user check
    # ------------------------------
    if current_user.status == 0:
        flash("Your account is inactive. Please contact admin.", "danger")
        return redirect(url_for('main.index'))

    if getattr(current_user.role, 'value', current_user.role) not in ['superadmin']:
        flash("You are not authorized to delete site settings.", "danger")
        return redirect(url_for('main.dashboard'))

    # ------------------------------
    # Get the record
    # ------------------------------
    settings = SchoolSiteSettings.query.get_or_404(id)

    # ------------------------------
    # Delete associated images from filesystem
    # ------------------------------
    image_fields = ['main_logo', 'sub_logo', 'sign_logo']
    for field in image_fields:
        image_path = getattr(settings, field)
        if image_path:  # hubi haddii image uu jiro
            full_path = os.path.join('static', image_path)
            if os.path.exists(full_path):
                try:
                    os.remove(full_path)
                except Exception as e:
                    # Log or print error
                    print(f"Failed to delete {full_path}: {e}")

    # ------------------------------
    # Delete the database record
    # ------------------------------
    db.session.delete(settings)
    db.session.commit()

    flash("School Site Settings and associated images deleted successfully!", "success")
    return redirect(url_for('main.all_school_site_settings'))



@bp.route('/settings/site', methods=['GET', 'POST'])
@login_required
def site_settings():
    # Get or create settings
    settings = SettingsData.query.first()
    if not settings:
        settings = SettingsData(group_name="default", address="", phone1="")
        db.session.add(settings)
        db.session.commit()

    form = SettingsDataForm(obj=settings)

    if form.validate_on_submit():
        try:
            # 🔹 Update text fields
            fields = [
                "group_name", "system_name", "address",
                "short_desc", "long_desc", "success_desc",
                "video_url", "phone1", "phone2", "email",
                "facebook", "twitter", "instagram", "dribbble"
            ]

            for field in fields:
                setattr(settings, field, getattr(form, field).data)

            # 🔹 Handle file uploads safely
            upload_fields = ["head_image", "image_success", "about_image", "logo", "logo2"]
            upload_folder = os.path.join("static", "backend", "uploads", "settings")
            os.makedirs(upload_folder, exist_ok=True)

            for field in upload_fields:
                file = getattr(form, field).data

                # ✅ Only process real uploaded files
                if file and isinstance(file, FileStorage) and file.filename:
                    filename = f"{uuid.uuid4().hex}_{secure_filename(file.filename)}"
                    file_path = os.path.join(upload_folder, filename)

                    file.save(file_path)

                    # Save relative path to DB
                    setattr(settings, field, f"backend/uploads/settings/{filename}")

            db.session.commit()
            flash("Site settings updated successfully!", "success")
            return redirect(url_for("settings.site_settings"))

        except IntegrityError as e:
            db.session.rollback()
            flash(f"Database error: {str(e)}", "danger")

        except Exception as e:
            db.session.rollback()
            flash(f"Error: {str(e)}", "danger")

    return render_template(
        "backend/pages/components/settings/site-settings.html",
        form=form,
        settings=settings,
        user=current_user
    )



#---------------------------------------------------
#---- Route: 70 | Dashboard - Backend Template -----
#---------------------------------------------------
@bp.route("/logout")
def logout():
    if current_user.is_authenticated:

        # Log the logout action
        create_user_log(
            user_id=current_user.id,
            action="logout",
            extra_info="User logged out",
            status="info"
        )

        # Only log out from Flask-Login
        logout_user()

        # ✅ Do NOT clear session or delete DB session yet
        # session.clear()  <-- remove this
        # db.session.delete(user_session)  <-- remove this

        # Flash message
        flash("You have been logged out! Your session record remains for inspection.", "success")

    # Clear remember_token cookie to prevent auto-login
    resp = make_response(redirect(url_for("main.login")))
    resp.set_cookie("remember_token", "", expires=0)
    return resp








